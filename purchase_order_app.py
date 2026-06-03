"""
Purchase Order — Unified Flask App (SQLite Version)
Replaces JSON storage with SQLite for performance & safety
Port: 5005
"""

import csv
import os, json, uuid, shutil, base64, re, html, urllib.error, urllib.request, mimetypes, sqlite3, time, traceback, zipfile
from datetime import date, datetime
from flask import Flask, render_template, request, jsonify, send_file, session, redirect, url_for, flash, abort
import io
import hashlib
from functools import wraps
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    WEASYPRINT_AVAILABLE = False
from io import BytesIO

from database import get_db, get_connection, init_db, init_default_settings, init_default_rbac

app = Flask(__name__, template_folder="templates")
app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024  # 20 MB
app.secret_key = "dev_secret_key_replace_in_production_123"

# Initialize database on startup
init_db()
with get_db() as conn:
    init_default_settings(conn)
    init_default_rbac(conn)

DATA_DIR   = os.path.join(os.path.dirname(__file__), "data", "po")
ATTACH_DIR = os.path.join(DATA_DIR, "attachments")
os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(ATTACH_DIR, exist_ok=True)

GEMINI_URL = (
    "https://generativelanguage.googleapis.com/v1beta/"
    "models/{model}:generateContent?key={key}"
)



# ── AUTHENTICATION & RBAC ───────────────────────────────────────────────────

@app.context_processor
def inject_permissions():
    """Inject current user's permissions into all templates"""
    if 'user_id' not in session:
        return dict(current_user=None, user_permissions=[])
        
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id, username, role_id FROM users WHERE id = ?", (session['user_id'],))
        user = cursor.fetchone()
        
        if not user or not user['role_id']:
            return dict(current_user=None, user_permissions=[])
            
        cursor.execute("""
            SELECT p.page_identifier 
            FROM permissions p
            JOIN role_permissions rp ON rp.permission_id = p.id
            WHERE rp.role_id = ?
        """, (user['role_id'],))
        perms = [row['page_identifier'] for row in cursor.fetchall()]
        
        return dict(current_user=dict(user), user_permissions=perms)


def require_permission(perm_name):
    """Decorator to require a specific page-level permission"""
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if 'user_id' not in session:
                return redirect(url_for('login', next=request.url))
                
            with get_db() as conn:
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT 1 FROM users u
                    JOIN role_permissions rp ON u.role_id = rp.role_id
                    JOIN permissions p ON rp.permission_id = p.id
                    WHERE u.id = ? AND p.page_identifier = ?
                """, (session['user_id'], perm_name))
                
                if not cursor.fetchone():
                    abort(403)  # Forbidden
                    
            return f(*args, **kwargs)
        return decorated_function
    return decorator


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        
        hashed_pw = hashlib.sha256(password.encode()).hexdigest()
        
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id FROM users WHERE username = ? AND password_hash = ?", (username, hashed_pw))
            user = cursor.fetchone()
            
            if user:
                session['user_id'] = user['id']
                
                # Smart redirect if they are heading to index but don't have permission
                next_page = request.args.get("next")
                is_root_dest = not next_page or next_page.endswith("/") or next_page.endswith("/po") or next_page.endswith("/po/")
                
                if is_root_dest:
                    cursor.execute("""
                        SELECT p.page_identifier FROM permissions p
                        JOIN role_permissions rp ON rp.permission_id = p.id
                        WHERE rp.role_id = (SELECT role_id FROM users WHERE id = ?)
                    """, (user['id'],))
                    perms = [r['page_identifier'] for r in cursor.fetchall()]
                    
                    if "po_dashboard" not in perms:
                        if "supplier_book" in perms: return redirect("/supplier-book")
                        if "customer_book" in perms: return redirect("/customer-book")
                        if "forwarder_dashboard" in perms: return redirect("/forwarder-dashboard")
                        if "admin_rbac" in perms: return redirect("/admin/rbac")
                
                return redirect(next_page or url_for("index"))
            else:
                flash("Invalid username or password", "error")
                
    return render_template("login.html")

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))

@app.errorhandler(403)
def forbidden_error(error):
    return render_template("403.html"), 403


# ── HELPER FUNCTIONS ──────────────────────────────────────────────────────────

def _call_gemini(api_key, model, prompt, file_data=None, mime_type=None):
    """Call Gemini API using urllib with retry logic for 503 errors"""
    url = GEMINI_URL.format(model=model, key=api_key)
    
    contents = []
    part_prompt = {"text": prompt}
    
    if file_data and mime_type:
        part_file = {
            "inline_data": {
                "mime_type": mime_type,
                "data": base64.b64encode(file_data).decode("utf-8")
            }
        }
        contents.append({"parts": [part_prompt, part_file]})
    else:
        contents.append({"parts": [part_prompt]})

    data = json.dumps({"contents": contents}).encode("utf-8")
    
    max_retries = 3
    retry_delay = 2 # seconds
    
    for attempt in range(max_retries):
        req = urllib.request.Request(
            url,
            data=data,
            headers={"Content-Type": "application/json"}
        )
        
        try:
            with urllib.request.urlopen(req) as response:
                res_data = json.loads(response.read().decode("utf-8"))
                return res_data
        except urllib.error.HTTPError as e:
            error_body = e.read().decode("utf-8")
            # If 429 (Too Many Requests) or 503 (Service Unavailable) and we have retries left, wait and try again
            if e.code in (429, 503) and attempt < max_retries - 1:
                # For 429, we might want to wait longer
                wait_time = retry_delay * (attempt + 1)
                if e.code == 429: wait_time *= 2 
                time.sleep(wait_time)
                continue
            
            # User friendly error message for common codes
            msg = f"Gemini API Error: {e.code}"
            try:
                err_json = json.loads(error_body)
                msg = err_json.get("error", {}).get("message", error_body)
                if e.code == 429:
                    msg = "AI Rate Limit / Quota exceeded. Please wait 60 seconds and try again."
            except: pass
            raise Exception(msg)
        except Exception as e:
            raise Exception(f"Failed to call Gemini: {str(e)}")

def _e(val):
    if val is None: return ""
    return html.escape(str(val))

def fmt_inr(n):
    return "₹" + f"{int(round(float(n))):,}"


def _sync_po_sequence(cursor, po_number):
    """Update po_sequence in settings if the given PO number has a higher sequence.
    This ensures the auto-numbering stays in sync with the latest created or edited PO.
    """
    if not po_number:
        return
    seq_match = re.search(r'(\d+)$', po_number)
    if seq_match:
        new_seq = int(seq_match.group(1))
        cursor.execute("SELECT value FROM settings WHERE key = 'po_sequence'")
        row = cursor.fetchone()
        current_seq = int(row[0]) if row and row[0].isdigit() else 0

        if new_seq > current_seq:
            cursor.execute(
                "INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)",
                ("po_sequence", str(new_seq))
            )

def _sync_inv_sequence(cursor, inv_number):
    """Update inv_sequence in settings if the given Invoice number has a higher sequence."""
    if not inv_number:
        return
    seq_match = re.search(r'(\d+)$', inv_number)
    if seq_match:
        new_seq = int(seq_match.group(1))
        cursor.execute("SELECT value FROM settings WHERE key = 'inv_sequence'")
        row = cursor.fetchone()
        current_seq = int(row[0]) if row and row[0].isdigit() else 0

        if new_seq > current_seq:
            cursor.execute(
                "INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)",
                ("inv_sequence", str(new_seq))
            )

def _log_status_change(cursor, po_id, from_status, to_status, note=None, force=False):
    """Insert one row into po_status_log. Call within an open transaction."""
    import uuid as _uuid
    if from_status == to_status and not force:
        return  # no change, nothing to log
    cursor.execute(
        """INSERT INTO po_status_log (id, po_id, from_status, to_status, note)
           VALUES (?, ?, ?, ?, ?)""",
        (str(_uuid.uuid4()), po_id, from_status, to_status, note)
    )

def _sync_shipment_from_po(cursor, po_id, po_status):
    """
    Automatically updates linked shipment status when a PO status changes.
    Rules:
    - PO 'Shipped' -> Shipment 'Shipped' (if currently 'With Forwarder')
    - PO 'In Transit' -> Shipment 'In Transit' (if currently Shipped/'With Forwarder')
    - PO 'Received' -> Shipment 'Delivered' (if ALL POs in shipment are Received) or 'Arrived'
    """
    try:
        # Find the active shipment linked to this PO
        cursor.execute("""
            SELECT s.id, s.status, s.notes, s.actual_arrival 
            FROM shipments s
            JOIN shipment_po_link spl ON spl.shipment_id = s.id
            WHERE spl.po_id = ? AND s.deleted_at IS NULL
        """, (po_id,))
        ship_rows = cursor.fetchall()
        for ship_row in ship_rows:
            sid = ship_row['id']
            current_ship_status = ship_row['status']
            
            new_ship_status = None
            
            if po_status == 'Shipped':
                if current_ship_status == 'With Forwarder':
                    new_ship_status = 'Shipped'
            
            elif po_status == 'In Transit':
                if current_ship_status in ['With Forwarder', 'Shipped']:
                    new_ship_status = 'In Transit'
            
            elif po_status == 'Received':
                # Check if ALL POs in this shipment are now 'Received'
                cursor.execute("""
                    SELECT COUNT(*) as total, 
                           SUM(CASE WHEN po.status = 'Received' THEN 1 ELSE 0 END) as received
                    FROM shipment_po_link spl
                    JOIN purchase_orders po ON po.id = spl.po_id
                    WHERE spl.shipment_id = ? AND po.deleted_at IS NULL
                """, (sid,))
                stats = cursor.fetchone()
                if stats and stats['total'] == stats['received']:
                    new_ship_status = 'Delivered'
                else:
                    if current_ship_status in ['In Transit', 'Shipped', 'With Forwarder']:
                        new_ship_status = 'Arrived'

            if new_ship_status and new_ship_status != current_ship_status:
                # Add an auto-note to shipment
                notes = json.loads(ship_row['notes'] or '[]')
                notes.append({
                    "date": str(date.today()),
                    "author": "System Sync",
                    "text": f"Status auto-updated to {new_ship_status} (PO status changed to {po_status})"
                })
                
                # Consolidate updates into one call
                actual_arr = ship_row.get("actual_arrival")
                if new_ship_status == "Delivered" and not actual_arr:
                    actual_arr = str(date.today())

                cursor.execute("""
                    UPDATE shipments 
                    SET status = ?, notes = ?, actual_arrival = ?, updated_at = CURRENT_TIMESTAMP 
                    WHERE id = ?
                """, (new_ship_status, json.dumps(notes), actual_arr, sid))
    except Exception as e:
        print(f"Error syncing shipment from PO: {e}")


# ── SUPPLIER BOOKS HELPER FUNCTIONS ───────────────────────────────────────────

def calc_running_balance(entries):
    """Calculate running balance per Tally Prime DR/CR accounting logic.

    CR (INVOICE / CREDIT_NOTE) → payable INCREASES (we owe supplier more)
    DR (PAYMENT / DEBIT_NOTE / ADJUSTMENT) → payable DECREASES (we paid)

    Positive balance = we still OWE the supplier  → display as 'Cr'
    Negative balance = supplier owes us (advance overpaid) → display as 'Dr'

    entries must already be sorted by entry_date ASC before calling this function.
    """
    balance = 0.0
    for e in entries:
        if e['dr_cr'] == 'CR':
            balance += e['amount_usd']
        elif e['dr_cr'] == 'DR':
            balance -= e['amount_usd']
        e['running_balance'] = round(balance, 2)
        # Human-readable suffix for UI display
        e['balance_suffix'] = 'Cr' if balance >= 0 else 'Dr'
    return entries, round(balance, 2)


def calc_due_date(po_date_str, terms):
    """Auto-calculate payment due date from PO date + supplier payment terms.

    terms is a dict with keys: terms_type, balance_trigger, credit_days
    Requires shipment data for AFTER_BL / AFTER_DELIVERY triggers —
    those are handled in the calling code; this covers ON_INVOICE / BEFORE_SHIPMENT.
    """
    from datetime import timedelta
    try:
        base = date.fromisoformat(po_date_str)
    except (ValueError, TypeError):
        base = date.today()

    credit_days = int(terms.get('credit_days', 0))
    trigger = terms.get('balance_trigger', 'ON_INVOICE')

    if trigger == 'ON_INVOICE':
        return (base + timedelta(days=credit_days)).isoformat()
    elif trigger == 'BEFORE_SHIPMENT':
        # Due before shipment — we use po due_date if available, else today
        return base.isoformat()
    else:
        # AFTER_BL / AFTER_DELIVERY → needs shipment data, return None (caller handles)
        return None


def _auto_ledger_on_status_change(cursor, po_id, new_status, po_dict):
    """Auto-create supplier ledger entries when a PO status changes.

    Confirmed  → Insert CR INVOICE entry (once per PO — idempotent)
    Cancelled  → Insert DR DEBIT_NOTE to reverse any existing INVOICE entry
    """
    supplier_id = po_dict.get('supplier_id', '')
    if not supplier_id:
        return

    # Fetch CNY→INR rate from settings (RMB rate is used for Chinese suppliers)
    cursor.execute("SELECT value FROM settings WHERE key = 'default_rmb_rate'")
    rate_row = cursor.fetchone()
    cny_rate = _safe_float(rate_row['value'] if rate_row else None, 11.5)

    # Calculate PO total (in CNY if CNY PO, else treat as USD → convert)
    cursor.execute(
        "SELECT SUM(qty * unit_price) as total FROM po_items WHERE po_id = ?",
        (po_id,)
    )
    total_row = cursor.fetchone()
    raw_total = float(total_row['total'] or 0) if total_row else 0.0

    po_currency = po_dict.get('currency', 'USD')
    if po_currency == 'CNY':
        # Already in CNY
        amount_cny = round(raw_total, 2)
    else:
        # USD PO — convert USD → INR first via USD rate, then INR → CNY via cny_rate
        cursor.execute("SELECT value FROM settings WHERE key = 'default_usd_rate'")
        usd_rate_row = cursor.fetchone()
        usd_rate = _safe_float(usd_rate_row['value'] if usd_rate_row else None, 84.0)
        amount_inr_from_usd = raw_total * usd_rate
        # Convert INR to CNY
        amount_cny = round(amount_inr_from_usd / cny_rate, 2) if cny_rate > 0 else 0.0

    amount_inr = round(amount_cny * cny_rate, 2)
    po_number = po_dict.get('po_number', '')
    today = str(date.today())

    if new_status == 'Confirmed':
        # Idempotent: only create if no INVOICE entry exists for this PO
        cursor.execute(
            """SELECT id FROM supplier_ledger_entries
               WHERE po_id = ? AND entry_type = 'INVOICE' AND deleted_at IS NULL""",
            (po_id,)
        )
        if cursor.fetchone():
            return  # INVOICE already exists, skip

        cursor.execute("""
            INSERT INTO supplier_ledger_entries
                (id, supplier_id, po_id, entry_type, entry_date, ref_number,
                 description, amount_usd, amount_inr, usd_rate, dr_cr,
                 created_by)
            VALUES (?, ?, ?, 'INVOICE', ?, ?, ?, ?, ?, ?, 'CR', 'System')
        """, (
            str(uuid.uuid4()), supplier_id, po_id, today,
            po_number,
            f"Purchase Invoice — {po_number} (auto on Confirmed)",
            amount_cny, amount_inr, cny_rate
        ))

    elif new_status == 'Cancelled':
        # Only reverse if an undeleted INVOICE entry exists
        cursor.execute(
            """SELECT id, amount_usd, amount_inr, usd_rate
               FROM supplier_ledger_entries
               WHERE po_id = ? AND entry_type = 'INVOICE' AND deleted_at IS NULL""",
            (po_id,)
        )
        inv_row = cursor.fetchone()
        if not inv_row:
            return  # No invoice to reverse

        # Check if a DEBIT_NOTE reversal already exists for this PO
        cursor.execute(
            """SELECT id FROM supplier_ledger_entries
               WHERE po_id = ? AND entry_type = 'DEBIT_NOTE' AND deleted_at IS NULL
               AND description LIKE '%Reversal%'""",
            (po_id,)
        )
        if cursor.fetchone():
            return  # Already reversed

        cursor.execute("""
            INSERT INTO supplier_ledger_entries
                (id, supplier_id, po_id, entry_type, entry_date, ref_number,
                 description, amount_usd, amount_inr, usd_rate, dr_cr,
                 created_by)
            VALUES (?, ?, ?, 'DEBIT_NOTE', ?, ?, ?, ?, ?, ?, 'DR', 'System')
        """, (
            str(uuid.uuid4()), supplier_id, po_id, today,
            po_number,
            f"Reversal — PO {po_number} Cancelled",
            float(inv_row['amount_usd']), float(inv_row['amount_inr']),
            float(inv_row['usd_rate'])
        ))




# Statuses that should trigger auto-creation of a shipment
_AUTO_SHIP_STATUSES = {"Confirmed", "Shipped", "In Transit"}


def _auto_create_shipment_for_po(cursor, po_id, po_dict):
    """
    Auto-create a shipment when a PO has a forwarder assigned and reaches
    a meaningful status (Confirmed / Shipped / In Transit) — but ONLY if the
    PO is not already linked to any shipment.

    Shipment defaults:
      - departure_date  = po.due_date  (the "expected shipment date" filled in PO form)
      - expected_arrival = departure_date + 35 days  (typical LCL sea transit China→India)
      - status          = "With Forwarder"
      - description     = "<PO number> — <Supplier company>"
    """
    from datetime import timedelta

    fid = (po_dict.get("forwarder_id") or "").strip()
    if not fid:
        return  # no forwarder assigned

    if po_dict.get("status") not in _AUTO_SHIP_STATUSES:
        return  # not a trigger status

    # Already linked to a shipment?
    cursor.execute("SELECT shipment_id FROM shipment_po_link WHERE po_id = ?", (po_id,))
    if cursor.fetchone():
        return

    # Forwarder must be active
    cursor.execute("SELECT id, name FROM forwarders WHERE id = ? AND active = 1", (fid,))
    fwd = cursor.fetchone()
    if not fwd:
        return

    # Departure date from po.due_date; fallback to today
    raw_due = po_dict.get("due_date") or ""
    try:
        departure = date.fromisoformat(raw_due) if raw_due else date.today()
    except ValueError:
        departure = date.today()

    arrival = departure + timedelta(days=35)  # adjust per your typical transit time

    # Build description label
    snap = po_dict.get("supplier_snapshot") or {}
    if isinstance(snap, str):
        try:
            snap = json.loads(snap)
        except Exception:
            snap = {}
    supplier_company = snap.get("company") or snap.get("name") or ""
    po_number = po_dict.get("po_number", "")
    description = f"{po_number} — {supplier_company}".strip(" —") or po_number

    # Create shipment row
    ship_id = str(uuid.uuid4())
    now_str = str(date.today())
    auto_notes = json.dumps([{
        "date": now_str,
        "author": "System",
        "text": (
            f"Shipment auto-created from PO {po_number} "
            f"(PO status: {po_dict.get('status')}, "
            f"forwarder: {dict(fwd)['name']}). "
            "Update departure & arrival dates via Edit on the Forwarder Dashboard."
        )
    }])

    cursor.execute("""
        INSERT INTO shipments
            (id, forwarder_id, booking_ref, departure_date, expected_arrival,
             actual_arrival, status, description, notes, created_at, updated_at)
        VALUES (?, ?, NULL, ?, ?, NULL, 'With Forwarder', ?, ?, ?, ?)
    """, (
        ship_id, fid,
        departure.isoformat(), arrival.isoformat(),
        description, auto_notes, now_str, now_str,
    ))

    # Link PO → shipment
    cursor.execute("""
        INSERT OR IGNORE INTO shipment_po_link (id, shipment_id, po_id)
        VALUES (?, ?, ?)
    """, (str(uuid.uuid4()), ship_id, po_id))


def calc_landed(items, rate, bank, ship, duty, trans, gst_duty=0, doc_pct=0):
    # Filter out items with zero or negative qty before any calculation
    items = [it for it in items if float(it.get("qty", 0)) > 0]
    if not items:
        return {"items": [], "inv_usd": 0, "inv_inr": 0, "total_addl": 0,
                "grand": 0, "bank": bank, "ship": ship, "duty": duty,
                "trans": trans, "gst_duty": gst_duty, "doc_pct": doc_pct,
                "doc_inr": 0, "rate": rate}

    inv_usd    = sum(float(it["qty"]) * float(it.get("unitPrice", it.get("unit_price", 0))) for it in items)
    inv_inr    = inv_usd * rate
    doc_inr    = inv_inr * (doc_pct / 100.0)
    total_addl = bank + ship + duty + trans + gst_duty + doc_inr
    grand      = inv_inr + total_addl
    result = []
    for it in items:
        qty = float(it.get("qty", 0))
        price = float(it.get("unitPrice", it.get("unit_price", 0)))
        item_inr   = qty * price * rate
        share      = (item_inr / inv_inr) if inv_inr > 0 else 0
        addl_share = total_addl * share
        total_item = item_inr + addl_share
        per_unit   = (total_item / qty) if qty > 0 else 0
        result.append({
            **it,
            "item_inr":   round(item_inr, 2),
            "share":      round(share * 100, 2),
            "addl_share": round(addl_share, 2),
            "total_item": round(total_item, 2),
            "per_unit":   round(per_unit, 2),
            "bank_s":     round(bank * share, 2),
            "ship_s":     round(ship * share, 2),
            "duty_s":     round(duty * share, 2),
            "trans_s":    round(trans * share, 2),
            "gst_s":      round(gst_duty * share, 2),
            "doc_s":      round(doc_inr * share, 2),
        })
    return {
        "items":      result,
        "inv_usd":    round(inv_usd, 2),
        "inv_inr":    round(inv_inr, 2),
        "total_addl": round(total_addl, 2),
        "grand":      round(grand, 2),
        "bank":       round(bank, 2),
        "ship":       round(ship, 2),
        "duty":       round(duty, 2),
        "trans":      round(trans, 2),
        "gst_duty":   round(gst_duty, 2),
        "doc_pct":    round(doc_pct, 2),
        "doc_inr":    round(doc_inr, 2),
        "rate":       rate,
    }


# ── ROUTES ────────────────────────────────────────────────────────────────────

@app.route("/")
@app.route("/po")
@app.route("/po/")
@require_permission("po_dashboard")
def index():
    return render_template("purchase_order.html")


@app.route("/sales-invoice")
@require_permission("customer_books")
def sales_invoice():
    return render_template("sales_invoice.html")



# ═══════════════════════════════════════════════════════════════════════════════
# SUPPLIERS
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/api/suppliers", methods=["GET"])
def get_suppliers():
    """Fetch all suppliers"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM suppliers ORDER BY name")
        suppliers = [dict(row) for row in cursor.fetchall()]
    return jsonify(suppliers)


@app.route("/api/suppliers", methods=["POST"])
@require_permission("supplier_edit")
def add_supplier():
    """Create new supplier"""
    req = request.get_json(silent=True)
    if not req:
        return jsonify({"error": "Invalid JSON"}), 400
    
    s_id = str(uuid.uuid4())
    supplier = {
        "id": s_id,
        "name": req.get("name", "").strip(),
        "company": req.get("company", "").strip(),
        "address": req.get("address", "").strip(),
        "country": req.get("country", "China").strip(),
        "email": req.get("email", "").strip(),
        "phone": req.get("phone", "").strip(),
        "wechat": req.get("wechat", "").strip(),
        "bank_name": req.get("bank_name", "").strip(),
        "bank_account": req.get("bank_account", "").strip(),
        "swift_code": req.get("swift_code", "").strip(),
    }
    
    if not supplier["name"]:
        return jsonify({"error": "Supplier name is required"}), 400
    
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO suppliers (id, name, company, address, country, email, phone, wechat, bank_name, bank_account, swift_code)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                supplier["id"],
                supplier["name"],
                supplier["company"],
                supplier["address"],
                supplier["country"],
                supplier["email"],
                supplier["phone"],
                supplier["wechat"],
                supplier["bank_name"],
                supplier["bank_account"],
                supplier["swift_code"],
            ))
    except sqlite3.IntegrityError:
        return jsonify({"error": f"Supplier '{supplier['name']}' already exists"}), 409
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    
    return jsonify(supplier), 201


@app.route("/api/suppliers/<sid>", methods=["PUT"])
@require_permission("supplier_edit")
def update_supplier(sid):
    """Update supplier"""
    req = request.get_json(silent=True)
    if not req:
        return jsonify({"error": "Invalid JSON"}), 400
    
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            
            # Fetch existing supplier
            cursor.execute("SELECT * FROM suppliers WHERE id = ?", (sid,))
            row = cursor.fetchone()
            if not row:
                return jsonify({"error": "Not found"}), 404
            
            supplier = dict(row)
            
            # Update fields
            for k in ["name", "company", "address", "country", "email", "phone", "wechat", "bank_name", "bank_account", "swift_code"]:
                if k in req:
                    supplier[k] = req[k]
            
            cursor.execute("""
                UPDATE suppliers
                SET name = ?, company = ?, address = ?, country = ?, email = ?, phone = ?, wechat = ?, bank_name = ?, bank_account = ?, swift_code = ?
                WHERE id = ?
            """, (
                supplier["name"], supplier["company"], supplier["address"], supplier["country"],
                supplier["email"], supplier["phone"], supplier["wechat"], supplier["bank_name"],
                supplier["bank_account"], supplier["swift_code"], sid
            ))
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    
    return jsonify(supplier)


@app.route("/api/suppliers/<sid>", methods=["DELETE"])
@require_permission("supplier_delete")
def delete_supplier(sid):
    """Delete supplier (check for usage in POs and Quotations)"""
    with get_db() as conn:
        cursor = conn.cursor()
        blockers = []
        
        # Check if supplier is used in any PO
        cursor.execute("SELECT po_number FROM purchase_orders WHERE supplier_id = ? AND deleted_at IS NULL", (sid,))
        used_pos = [row[0] for row in cursor.fetchall()]
        if used_pos:
            blockers.append(f"POs: {', '.join(used_pos)}")
        
        # O(1) lookup via quotation_refs index (no JSON scanning)
        cursor.execute("""
            SELECT q.quotation_number, q.title, q.id
            FROM quotation_refs r
            JOIN quotations q ON q.id = r.quotation_id
            WHERE r.ref_type = 'supplier' AND r.ref_id = ?
        """, (sid,))
        used_quotes = [
            row["quotation_number"] or row["title"] or row["id"][:8]
            for row in cursor.fetchall()
        ]
        if used_quotes:
            blockers.append(f"Quotations: {', '.join(used_quotes)}")
        
        if blockers:
            return jsonify({"error": f"Cannot delete — used in {' | '.join(blockers)}"}), 409
        
        cursor.execute("DELETE FROM suppliers WHERE id = ?", (sid,))
    
    return jsonify({"ok": True})


# ═══════════════════════════════════════════════════════════════════════════════
# SUPPLIER BOOKS — PAYMENT TERMS & LEDGER API
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/api/suppliers/<sid>/payment-terms", methods=["GET"])
def get_payment_terms(sid):
    """Fetch payment terms for a supplier."""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM supplier_payment_terms WHERE supplier_id = ?", (sid,))
        row = cursor.fetchone()
    if not row:
        return jsonify({})
    return jsonify(dict(row))


@app.route("/api/suppliers/<sid>/payment-terms", methods=["POST"])
@require_permission("supplier_edit")
def save_payment_terms(sid):
    """Create or update payment terms for a supplier (upsert)."""
    req = request.get_json(silent=True) or {}
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM suppliers WHERE id = ?", (sid,))
        if not cursor.fetchone():
            return jsonify({"error": "Supplier not found"}), 404
        cursor.execute("""
            INSERT INTO supplier_payment_terms
                (supplier_id, terms_type, advance_pct, balance_trigger, credit_days, currency, notes, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
            ON CONFLICT(supplier_id) DO UPDATE SET
                terms_type      = excluded.terms_type,
                advance_pct     = excluded.advance_pct,
                balance_trigger = excluded.balance_trigger,
                credit_days     = excluded.credit_days,
                currency        = excluded.currency,
                notes           = excluded.notes,
                updated_at      = CURRENT_TIMESTAMP
        """, (
            sid,
            req.get("terms_type", "NET30"),
            _safe_float(req.get("advance_pct", 30), 30.0),
            req.get("balance_trigger", "ON_INVOICE"),
            _safe_int(req.get("credit_days", 0), 0),
            req.get("currency", "USD"),
            req.get("notes", ""),
        ))
        cursor.execute("SELECT * FROM supplier_payment_terms WHERE supplier_id = ?", (sid,))
        row = cursor.fetchone()
    return jsonify(dict(row)), 200


# ── Outstanding Summary (Dashboard) — must be before <sid>/ledger routes ──────

@app.route("/api/suppliers/summary/outstanding", methods=["GET"])
def outstanding_summary():
    """Return one row per supplier with total billed, paid, and outstanding balance."""
    today = str(date.today())
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id, name, company FROM suppliers ORDER BY name")
        suppliers = [dict(r) for r in cursor.fetchall()]

        cursor.execute("SELECT value FROM settings WHERE key = 'default_rmb_rate'")
        r2 = cursor.fetchone()
        cny_rate = _safe_float(r2["value"] if r2 else None, 11.5)

        cursor.execute("SELECT supplier_id, credit_days, balance_trigger FROM supplier_payment_terms")
        terms_by_sup = {row["supplier_id"]: dict(row) for row in cursor.fetchall()}

        cursor.execute("""
            SELECT 
                supplier_id,
                SUM(CASE WHEN dr_cr = 'CR' THEN amount_usd ELSE 0 END) as total_cr,
                SUM(CASE WHEN dr_cr = 'DR' THEN amount_usd ELSE 0 END) as total_dr,
                SUM(CASE WHEN entry_type = 'PAYMENT' AND dr_cr = 'DR' THEN amount_usd ELSE 0 END) as advance_paid_usd,
                MIN(CASE WHEN entry_type = 'INVOICE' AND dr_cr = 'CR' THEN entry_date END) as oldest_invoice_date
            FROM supplier_ledger_entries
            WHERE deleted_at IS NULL
            GROUP BY supplier_id
        """)
        aggregates = {row["supplier_id"]: dict(row) for row in cursor.fetchall()}

        result = []
        total_outstanding_usd = 0.0
        total_outstanding_inr = 0.0
        overdue_count = 0
        total_advance_paid_usd = 0.0

        for sup in suppliers:
            sid = sup["id"]
            agg = aggregates.get(sid)
            if not agg:
                continue

            total_cr = agg["total_cr"] or 0.0
            total_dr = agg["total_dr"] or 0.0
            balance_usd = round(total_cr - total_dr, 2)
            advance_paid_usd = agg["advance_paid_usd"] or 0.0

            terms = terms_by_sup.get(sid, {"credit_days": 0, "balance_trigger": "ON_INVOICE"})
            
            oldest_overdue_days = 0
            oldest_date = agg["oldest_invoice_date"]
            if oldest_date:
                due = calc_due_date(oldest_date, terms)
                if due and due < today:
                    diff = (date.fromisoformat(today) - date.fromisoformat(due)).days
                    if diff > oldest_overdue_days:
                        oldest_overdue_days = diff

            balance_inr = round(balance_usd * cny_rate, 2)

            if oldest_overdue_days > 90:
                overdue_count += 1

            total_advance_paid_usd += advance_paid_usd
            total_outstanding_usd += balance_usd
            total_outstanding_inr += balance_inr

            if balance_usd != 0:
                result.append({
                    **sup,
                    "total_billed_usd": round(total_cr, 2),
                    "total_paid_usd": round(total_dr, 2),
                    "outstanding_usd": balance_usd,
                    "outstanding_inr": balance_inr,
                    "oldest_overdue_days": oldest_overdue_days,
                    "oldest_date": oldest_date,
                })

    return jsonify({
        "suppliers": result,
        "total_outstanding_usd": round(total_outstanding_usd, 2),
        "total_outstanding_inr": round(total_outstanding_inr, 2),
        "overdue_bills_count": overdue_count,
        "advance_paid_usd": round(total_advance_paid_usd, 2),
    })


# ── Per-Supplier Ledger Entries ───────────────────────────────────────────────

@app.route("/api/suppliers/<sid>/ledger", methods=["GET"])
def get_supplier_ledger(sid):
    """Fetch all active ledger entries for a supplier with running balance."""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM suppliers WHERE id = ?", (sid,))
        if not cursor.fetchone():
            return jsonify({"error": "Supplier not found"}), 404
        cursor.execute("""
            SELECT * FROM supplier_ledger_entries
            WHERE supplier_id = ? AND deleted_at IS NULL
            ORDER BY entry_date ASC, created_at ASC
        """, (sid,))
        entries = [dict(r) for r in cursor.fetchall()]

    entries, closing_balance = calc_running_balance(entries)
    balance_suffix = "Cr" if closing_balance >= 0 else "Dr"
    return jsonify({
        "entries": entries,
        "closing_balance": closing_balance,
        "balance_suffix": balance_suffix,
    })


@app.route("/api/suppliers/<sid>/ledger", methods=["POST"])
@require_permission("supplier_edit")
def add_ledger_entry(sid):
    """Add a new ledger entry for a supplier (CNY / RMB amounts)."""
    req = request.get_json(silent=True) or {}
    VALID_TYPES = {"INVOICE", "PAYMENT", "DEBIT_NOTE", "CREDIT_NOTE", "ADJUSTMENT"}
    entry_type = req.get("entry_type", "").upper()
    if entry_type not in VALID_TYPES:
        return jsonify({"error": f"entry_type must be one of {VALID_TYPES}"}), 400

    # Auto-assign dr_cr if not provided
    dr_cr = req.get("dr_cr", "").upper()
    if not dr_cr:
        if entry_type in ("INVOICE", "CREDIT_NOTE"):
            dr_cr = "CR"
        elif entry_type in ("PAYMENT", "DEBIT_NOTE"):
            dr_cr = "DR"
        else:
            return jsonify({"error": "dr_cr must be DR or CR for ADJUSTMENT"}), 400
    elif dr_cr not in ("DR", "CR"):
        return jsonify({"error": "dr_cr must be DR or CR"}), 400

    # Fetch exchange rate
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM suppliers WHERE id = ?", (sid,))
        if not cursor.fetchone():
            return jsonify({"error": "Supplier not found"}), 404
        cursor.execute("SELECT value FROM settings WHERE key = 'default_rmb_rate'")
        rate_row = cursor.fetchone()
    default_cny_rate = _safe_float(rate_row['value'] if rate_row else None, 11.5)

    cny_rate = _safe_float(req.get("cny_rate") or req.get("usd_rate"), default_cny_rate)

    amount_inr = _safe_float(req.get("amount_inr"), 0)
    amount_cny = _safe_float(req.get("amount_cny") or req.get("amount_usd"), 0)

    if amount_inr > 0 and amount_cny == 0:
        amount_cny = round(amount_inr / cny_rate, 4)
    elif amount_cny > 0 and amount_inr == 0:
        amount_inr = round(amount_cny * cny_rate, 2)
    elif amount_inr == 0 and amount_cny == 0:
        return jsonify({"error": "Either amount_inr or amount_cny must be positive"}), 400

    entry_date = req.get("entry_date", str(date.today()))
    eid = str(uuid.uuid4())
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO supplier_ledger_entries
                (id, supplier_id, po_id, entry_type, entry_date, ref_number,
                 description, amount_usd, amount_inr, usd_rate, dr_cr,
                 payment_mode, bank_ref, attachment_id, notes, created_by)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            eid, sid,
            req.get("po_id") or None,
            entry_type, entry_date,
            req.get("ref_number", ""),
            req.get("description", ""),
            amount_cny, amount_inr, cny_rate, dr_cr,
            req.get("payment_mode", ""),
            req.get("bank_ref", ""),
            req.get("attachment_id") or None,
            req.get("notes", ""),
            req.get("created_by", "User"),
        ))
        cursor.execute("SELECT * FROM supplier_ledger_entries WHERE id = ?", (eid,))
        row = dict(cursor.fetchone())
    return jsonify(row), 201


@app.route("/api/suppliers/<sid>/ledger/<eid>", methods=["PUT"])
@require_permission("supplier_edit")
def update_ledger_entry(sid, eid):
    """Edit narration, notes, bank_ref of an existing entry (amounts locked)."""
    req = request.get_json(silent=True) or {}
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT * FROM supplier_ledger_entries WHERE id = ? AND supplier_id = ? AND deleted_at IS NULL",
            (eid, sid)
        )
        row = cursor.fetchone()
        if not row:
            return jsonify({"error": "Entry not found"}), 404
        cursor.execute("""
            UPDATE supplier_ledger_entries
            SET description = ?, notes = ?, bank_ref = ?, ref_number = ?
            WHERE id = ?
        """, (
            req.get("description", row["description"]),
            req.get("notes", row["notes"]),
            req.get("bank_ref", row["bank_ref"]),
            req.get("ref_number", row["ref_number"]),
            eid,
        ))
        cursor.execute("SELECT * FROM supplier_ledger_entries WHERE id = ?", (eid,))
        updated = dict(cursor.fetchone())
    return jsonify(updated)


@app.route("/api/suppliers/<sid>/ledger/<eid>", methods=["DELETE"])
@require_permission("supplier_delete")
def delete_ledger_entry(sid, eid):
    """Soft-delete a ledger entry (data preserved in DB)."""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT id FROM supplier_ledger_entries WHERE id = ? AND supplier_id = ? AND deleted_at IS NULL",
            (eid, sid)
        )
        if not cursor.fetchone():
            return jsonify({"error": "Entry not found"}), 404
        cursor.execute(
            "UPDATE supplier_ledger_entries SET deleted_at = CURRENT_TIMESTAMP WHERE id = ?",
            (eid,)
        )
    return jsonify({"ok": True})


# ── BATCH PROCESSING ENDPOINTS ───────────────────────────────────────────────

@app.route("/api/suppliers/batch/upload", methods=["POST"])
@require_permission("supplier_edit")
def batch_upload_supplier_ledger():
    """Batch upload ledger entries from CSV file.
    
    Expected CSV columns:
    supplier_id, entry_type, entry_date, ref_number, description, 
    amount_cny, dr_cr, due_date, notes
    
    Returns: {
        "total": number of rows processed,
        "success": number of successful inserts,
        "errors": list of {row_num, error_msg},
        "results": list of inserted entry IDs
    }
    """
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    file = request.files["file"]
    if not file or not file.filename.endswith(".csv"):
        return jsonify({"error": "File must be CSV format"}), 400
    
    try:
        stream = io.StringIO(file.stream.read().decode("utf8"), newline=None)
        reader = csv.DictReader(stream)
        if not reader.fieldnames:
            return jsonify({"error": "CSV file is empty"}), 400
        
        results = {"total": 0, "success": 0, "errors": [], "ids": []}
        
        with get_db() as conn:
            cursor = conn.cursor()
            for row_num, row in enumerate(reader, start=2):  # Start at 2 (header is row 1)
                results["total"] += 1
                try:
                    # Validate required fields
                    supplier_id = row.get("supplier_id", "").strip()
                    entry_type = row.get("entry_type", "").upper()
                    entry_date = row.get("entry_date", str(date.today())).strip()
                    dr_cr = row.get("dr_cr", "").upper()
                    
                    if not supplier_id:
                        raise ValueError("supplier_id is required")
                    if entry_type not in {"INVOICE", "PAYMENT", "DEBIT_NOTE", "CREDIT_NOTE", "ADJUSTMENT"}:
                        raise ValueError(f"Invalid entry_type: {entry_type}")
                    if dr_cr not in ("DR", "CR"):
                        raise ValueError("dr_cr must be DR or CR")
                    
                    amount_cny = _safe_float(row.get("amount_cny"), 0)
                    if amount_cny <= 0:
                        raise ValueError("amount_cny must be positive")
                    
                    # Check supplier exists
                    cursor.execute("SELECT id FROM suppliers WHERE id = ?", (supplier_id,))
                    if not cursor.fetchone():
                        raise ValueError(f"Supplier {supplier_id} not found")
                    
                    # Get exchange rate
                    cny_rate = _safe_float(row.get("cny_rate"), 11.5)
                    amount_inr = round(amount_cny * cny_rate, 2)
                    
                    # Insert entry
                    eid = str(uuid.uuid4())
                    cursor.execute("""
                        INSERT INTO supplier_ledger_entries
                            (id, supplier_id, entry_type, entry_date, ref_number,
                             description, amount_usd, amount_inr, usd_rate, dr_cr,
                             payment_mode, notes, created_by)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """, (
                        eid, supplier_id, entry_type, entry_date,
                        row.get("ref_number", "").strip(),
                        row.get("description", "").strip(),
                        amount_cny, amount_inr, cny_rate, dr_cr,
                        row.get("payment_mode", "").strip(),
                        row.get("notes", "").strip(),
                        "Batch Import",
                    ))
                    results["success"] += 1
                    results["ids"].append(eid)
                
                except Exception as e:
                    results["errors"].append({"row": row_num, "error": str(e)})
        
        return jsonify(results), 201
    
    except Exception as e:
        return jsonify({"error": f"Failed to process file: {str(e)}"}), 400


@app.route("/api/customers/batch/upload", methods=["POST"])
@require_permission("customer_edit")
def batch_upload_customer_ledger():
    """Batch upload customer ledger entries from CSV file.
    
    Expected CSV columns:
    customer_id, entry_type, entry_date, ref_number, description, 
    amount_inr, dr_cr, due_date, notes
    """
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    file = request.files["file"]
    if not file or not file.filename.endswith(".csv"):
        return jsonify({"error": "File must be CSV format"}), 400
    
    try:
        stream = io.StringIO(file.stream.read().decode("utf8"), newline=None)
        reader = csv.DictReader(stream)
        if not reader.fieldnames:
            return jsonify({"error": "CSV file is empty"}), 400
        
        results = {"total": 0, "success": 0, "errors": [], "ids": []}
        
        with get_db() as conn:
            cursor = conn.cursor()
            for row_num, row in enumerate(reader, start=2):
                results["total"] += 1
                try:
                    customer_id = row.get("customer_id", "").strip()
                    entry_type = row.get("entry_type", "").upper()
                    entry_date = row.get("entry_date", str(date.today())).strip()
                    dr_cr = row.get("dr_cr", "").upper()
                    
                    if not customer_id:
                        raise ValueError("customer_id is required")
                    if entry_type not in {"INVOICE", "PAYMENT", "DEBIT_NOTE", "CREDIT_NOTE", "ADJUSTMENT"}:
                        raise ValueError(f"Invalid entry_type: {entry_type}")
                    if dr_cr not in ("DR", "CR"):
                        raise ValueError("dr_cr must be DR or CR")
                    
                    amount_inr = _safe_float(row.get("amount_inr"), 0)
                    if amount_inr <= 0:
                        raise ValueError("amount_inr must be positive")
                    
                    cursor.execute("SELECT id FROM customers WHERE id = ?", (customer_id,))
                    if not cursor.fetchone():
                        raise ValueError(f"Customer {customer_id} not found")
                    
                    eid = str(uuid.uuid4())
                    cursor.execute("""
                        INSERT INTO customer_ledger_entries
                            (id, customer_id, entry_type, entry_date, ref_number,
                             description, amount_inr, dr_cr, payment_mode, 
                             due_date, notes, created_by)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """, (
                        eid, customer_id, entry_type, entry_date,
                        row.get("ref_number", "").strip(),
                        row.get("description", "").strip(),
                        amount_inr, dr_cr,
                        row.get("payment_mode", "").strip(),
                        row.get("due_date", "").strip(),
                        row.get("notes", "").strip(),
                        "Batch Import",
                    ))
                    results["success"] += 1
                    results["ids"].append(eid)
                
                except Exception as e:
                    results["errors"].append({"row": row_num, "error": str(e)})
        
        return jsonify(results), 201
    
    except Exception as e:
        return jsonify({"error": f"Failed to process file: {str(e)}"}), 400


@app.route("/api/suppliers/batch/report", methods=["GET"])
def get_batch_summary():
    """Get summary of batch operations and total entries by type"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT entry_type, dr_cr, COUNT(*) as count, 
                   SUM(amount_inr) as total_amount
            FROM supplier_ledger_entries
            WHERE deleted_at IS NULL
            GROUP BY entry_type, dr_cr
            ORDER BY entry_type, dr_cr
        """)
        entries_by_type = [dict(row) for row in cursor.fetchall()]
        
        cursor.execute("""
            SELECT created_by, COUNT(*) as count, MAX(created_at) as last_created
            FROM supplier_ledger_entries
            GROUP BY created_by
        """)
        entries_by_source = [dict(row) for row in cursor.fetchall()]
    
    return jsonify({
        "by_type": entries_by_type,
        "by_source": entries_by_source
    })


@app.route("/api/suppliers/<sid>/outstanding", methods=["GET"])
def get_supplier_outstanding(sid):
    """Return outstanding (unpaid) invoices grouped by ageing bucket."""
    today_dt = date.today()
    today_str = str(today_dt)
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM suppliers WHERE id = ?", (sid,))
        if not cursor.fetchone():
            return jsonify({"error": "Supplier not found"}), 404

        cursor.execute("SELECT * FROM supplier_payment_terms WHERE supplier_id = ?", (sid,))
        terms_row = cursor.fetchone()
        terms = dict(terms_row) if terms_row else {"credit_days": 0, "balance_trigger": "ON_INVOICE"}

        cursor.execute("SELECT value FROM settings WHERE key = 'default_rmb_rate'")
        rate_row = cursor.fetchone()
        cny_rate = _safe_float(rate_row['value'] if rate_row else None, 11.5)

        cursor.execute("""
            SELECT * FROM supplier_ledger_entries
            WHERE supplier_id = ? AND entry_type = 'INVOICE' AND dr_cr = 'CR' AND deleted_at IS NULL
            ORDER BY entry_date ASC
        """, (sid,))
        invoices = [dict(r) for r in cursor.fetchall()]

        buckets = {"0_30": [], "31_60": [], "61_90": [], "91_plus": []}
        for inv in invoices:
            # Sum payments linked to this PO
            if inv["po_id"]:
                cursor.execute("""
                    SELECT COALESCE(SUM(amount_usd),0) as paid
                    FROM supplier_ledger_entries
                    WHERE po_id = ? AND dr_cr = 'DR' AND deleted_at IS NULL
                """, (inv["po_id"],))
            else:
                cursor.execute("""
                    SELECT COALESCE(SUM(amount_usd),0) as paid
                    FROM supplier_ledger_entries
                    WHERE supplier_id = ? AND ref_number = ? AND dr_cr = 'DR' AND deleted_at IS NULL
                """, (sid, inv["ref_number"] or ""))
            paid_row = cursor.fetchone()
            paid = float(paid_row["paid"])
            pending = round(inv["amount_usd"] - paid, 2)
            if pending <= 0:
                continue

            pending_inr = round(pending * cny_rate, 2)
            due_date = calc_due_date(inv["entry_date"], terms) or inv["entry_date"]
            try:
                days_overdue = (today_dt - date.fromisoformat(due_date)).days
            except Exception:
                days_overdue = 0

            age_label = "0-30 Days"
            bucket_key = "0_30"
            if days_overdue > 90:
                age_label = "91+ Days ⚠"
                bucket_key = "91_plus"
            elif days_overdue > 60:
                age_label = "61–90 Days"
                bucket_key = "61_90"
            elif days_overdue > 30:
                age_label = "31–60 Days"
                bucket_key = "31_60"

            item = {
                "ref": inv["ref_number"],
                "bill_date": inv["entry_date"],
                "due_date": due_date,
                "pending_usd": pending,
                "pending_inr": pending_inr,
                "days_overdue": days_overdue,
                "age_label": age_label,
                "po_id": inv["po_id"],
            }
            buckets[bucket_key].append(item)

    return jsonify({"buckets": buckets, "as_of": today_str})


# ── Supplier Ledger PDF Export ────────────────────────────────────────────────

@app.route("/api/suppliers/<sid>/ledger/pdf", methods=["GET"])
def supplier_ledger_pdf(sid):
    """Generate a formal Account Statement PDF for a supplier."""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM suppliers WHERE id = ?", (sid,))
        sup_row = cursor.fetchone()
        if not sup_row:
            return "Supplier not found", 404
        sup = dict(sup_row)
        cursor.execute("""
            SELECT * FROM supplier_ledger_entries
            WHERE supplier_id = ? AND deleted_at IS NULL
            ORDER BY entry_date ASC, created_at ASC
        """, (sid,))
        entries = [dict(r) for r in cursor.fetchall()]
        cursor.execute("SELECT key, value FROM settings")
        settings = {r["key"]: r["value"] for r in cursor.fetchall()}

    entries, closing_balance = calc_running_balance(entries)
    balance_suffix = "Cr" if closing_balance >= 0 else "Dr"
    company_name = settings.get("company_name", "")
    cur_code = 'USD'
    cur_sym = '$'
    if 'china' in str(sup.get('country', '')).lower():
        cur_code = 'CNY'
        cur_sym = '¥'

    rows_html = ""
    for e in entries:
        dr_amt = f"{cur_sym} {e['amount_usd']:,.2f}" if e["dr_cr"] == "DR" else ""
        cr_amt = f"{cur_sym} {e['amount_usd']:,.2f}" if e["dr_cr"] == "CR" else ""
        bal = e.get("running_balance", 0)
        bal_sfx = e.get("balance_suffix", "Cr")
        row_color = {"INVOICE": "#e8f4fd", "PAYMENT": "#effdf0", "DEBIT_NOTE": "#fff3e0",
                     "CREDIT_NOTE": "#fce4ec", "ADJUSTMENT": "#f3f4f6"}.get(e["entry_type"], "#fff")
        rows_html += f"""<tr style="background:{row_color}">
          <td>{_e(e['entry_date'])}</td>
          <td><span style="font-size:9px;padding:2px 6px;border-radius:3px;background:#ddd">{_e(e['entry_type'])}</span></td>
          <td>{_e(e['ref_number'] or '')}</td>
          <td style="max-width:200px;font-size:10px">{_e(e['description'] or '')}</td>
          <td style="text-align:right;color:#c62828">{_e(dr_amt)}</td>
          <td style="text-align:right;color:#1565c0">{_e(cr_amt)}</td>
          <td style="text-align:right;font-weight:700">{cur_sym} {abs(bal):,.2f} {_e(bal_sfx)}</td>
        </tr>"""

    total_dr = sum(e["amount_usd"] for e in entries if e["dr_cr"] == "DR")
    total_cr = sum(e["amount_usd"] for e in entries if e["dr_cr"] == "CR")

    html = f"""<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8">
<title>Account Statement — {_e(sup.get('company',''))}</title>
<style>
  body{{font-family:'Segoe UI',Arial,sans-serif;font-size:11px;color:#111;margin:0;padding:24px 32px}}
  .hdr{{border-bottom:3px solid #1e3a8a;padding-bottom:12px;margin-bottom:16px;display:flex;justify-content:space-between;align-items:flex-start}}
  .co-name{{font-size:20px;font-weight:800;color:#1e3a8a}}.co-sub{{font-size:9px;color:#6b7280;margin-top:2px}}
  .stmt-title{{font-size:14px;font-weight:700;color:#111;text-align:right}}.stmt-meta{{font-size:10px;color:#6b7280;text-align:right}}
  .parties{{display:grid;grid-template-columns:1fr 1fr;gap:0;border:1px solid #e5e7eb;border-radius:6px;overflow:hidden;margin-bottom:14px}}
  .party{{padding:10px 14px}}.party:first-child{{background:#f9fafb;border-right:1px solid #e5e7eb}}
  .party-label{{font-size:8px;font-weight:700;letter-spacing:.1em;color:#6b7280;text-transform:uppercase;margin-bottom:6px}}
  table{{width:100%;border-collapse:collapse;font-size:10px}}
  thead th{{padding:7px 10px;background:#1e3a8a;color:#fff;font-size:9px;text-align:left;font-weight:600;text-transform:uppercase;letter-spacing:.05em}}
  tbody td{{padding:6px 10px;border-bottom:1px solid #e5e7eb;vertical-align:middle}}
  tfoot td{{padding:7px 10px;background:#1e3a8a;color:#fff;font-weight:700;font-size:11px}}
  .closing{{margin-top:16px;padding:12px 16px;border-radius:8px;border:2px solid #1e3a8a;background:#eff6ff;text-align:right}}
  .closing-lbl{{font-size:10px;color:#6b7280;font-weight:700;text-transform:uppercase;letter-spacing:.06em}}
  .closing-val{{font-size:22px;font-weight:900;color:#1e3a8a;margin-top:4px}}
  .footer{{margin-top:16px;font-size:8px;color:#9ca3af;text-align:center;border-top:1px solid #e5e7eb;padding-top:8px}}
  @media print{{@page{{margin:10mm;size:A4}}body{{padding:0}}}}
</style></head><body>
<div class="hdr">
  <div><div class="co-name">{_e(company_name or 'Purchase Order')}</div>
  <div class="co-sub">Account Statement</div></div>
  <div><div class="stmt-title">SUPPLIER ACCOUNT STATEMENT</div>
  <div class="stmt-meta">As of {date.today().strftime('%d %b %Y')}</div></div>
</div>
<div class="parties">
  <div class="party"><div class="party-label">Statement For</div>
  <div style="font-weight:700;font-size:13px">{_e(sup.get('company',''))}</div>
  <div style="color:#4b5563">{_e(sup.get('name',''))}</div>
  <div style="color:#4b5563">{_e(sup.get('country',''))}</div></div>
  <div class="party"><div class="party-label">Prepared By</div>
  <div style="font-weight:700">{_e(company_name)}</div></div>
</div>
<table><thead><tr>
  <th style="width:80px">Date</th><th style="width:90px">Type</th>
  <th style="width:110px">Reference</th><th>Narration</th>
  <th style="width:100px;text-align:right">Dr ({cur_code})</th>
  <th style="width:100px;text-align:right">Cr ({cur_code})</th>
  <th style="width:110px;text-align:right">Balance</th>
</tr></thead>
<tbody>{rows_html}</tbody>
<tfoot><tr>
  <td colspan="4" style="text-align:right">TOTAL</td>
  <td style="text-align:right">{cur_sym} {total_dr:,.2f}</td>
  <td style="text-align:right">{cur_sym} {total_cr:,.2f}</td>
  <td style="text-align:right">{cur_sym} {abs(closing_balance):,.2f} {_e(balance_suffix)}</td>
</tr></tfoot></table>
<div class="closing">
  <div class="closing-lbl">Closing Balance</div>
  <div class="closing-val">{cur_sym} {abs(closing_balance):,.2f} <span style="font-size:14px">{_e(balance_suffix)}</span></div>
</div>
<div class="footer">Computer-generated Account Statement · {_e(company_name)} · {date.today().strftime('%d %b %Y')}</div>
</body></html>"""

    try:
        if WEASYPRINT_AVAILABLE:
            pdf_bytes = HTML(string=html, base_url=None).write_pdf()
            return send_file(BytesIO(pdf_bytes), mimetype="application/pdf", as_attachment=False,
                             download_name=f"Statement_{sup.get('company','supplier')}_{date.today()}.pdf")
    except Exception:
        pass
    return html, 200, {"Content-Type": "text/html; charset=utf-8"}


# ── Supplier Ledger Excel Export ──────────────────────────────────────────────

@app.route("/api/suppliers/<sid>/ledger/excel", methods=["GET"])
def supplier_ledger_excel(sid):
    """Export supplier ledger to Excel file"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        return jsonify({"error": "openpyxl not installed"}), 500
    
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM suppliers WHERE id = ?", (sid,))
        sup_row = cursor.fetchone()
        if not sup_row:
            return "Supplier not found", 404
        sup = dict(sup_row)
        
        cursor.execute("""
            SELECT * FROM supplier_ledger_entries
            WHERE supplier_id = ? AND deleted_at IS NULL
            ORDER BY entry_date ASC, created_at ASC
        """, (sid,))
        entries = [dict(r) for r in cursor.fetchall()]
    
    entries, closing_balance = calc_running_balance(entries)
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Ledger"
    
    # Header
    ws['A1'] = "Account Statement"
    ws['A1'].font = Font(size=14, bold=True)
    ws['A2'] = sup.get('company', '')
    ws['A3'] = f"As of {date.today().strftime('%d %b %Y')}"
    
    # Column headers
    headers = ["Date", "Type", "Ref No", "Description", "Debit (₹)", "Credit (₹)", "Balance (₹)"]
    for col, hdr in enumerate(headers, 1):
        cell = ws.cell(row=5, column=col, value=hdr)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="1e3a8a", end_color="1e3a8a", fill_type="solid")
    
    # Data rows
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    for row_idx, e in enumerate(entries, start=6):
        ws.cell(row=row_idx, column=1, value=e.get('entry_date', ''))
        ws.cell(row=row_idx, column=2, value=e.get('entry_type', ''))
        ws.cell(row=row_idx, column=3, value=e.get('ref_number', ''))
        ws.cell(row=row_idx, column=4, value=e.get('description', ''))
        
        dr_val = e['amount_inr'] if e['dr_cr'] == 'DR' else 0
        cr_val = e['amount_inr'] if e['dr_cr'] == 'CR' else 0
        ws.cell(row=row_idx, column=5, value=dr_val).number_format = '#,##0.00'
        ws.cell(row=row_idx, column=6, value=cr_val).number_format = '#,##0.00'
        ws.cell(row=row_idx, column=7, value=e.get('running_balance', 0)).number_format = '#,##0.00'
        
        for col in range(1, 8):
            ws.cell(row=row_idx, column=col).border = thin_border
    
    # Totals row
    total_row = len(entries) + 6
    ws.cell(row=total_row, column=2, value="TOTAL").font = Font(bold=True)
    ws.cell(row=total_row, column=5, value=sum(e['amount_inr'] for e in entries if e['dr_cr'] == 'DR')).number_format = '#,##0.00'
    ws.cell(row=total_row, column=6, value=sum(e['amount_inr'] for e in entries if e['dr_cr'] == 'CR')).number_format = '#,##0.00'
    ws.cell(row=total_row, column=7, value=closing_balance).number_format = '#,##0.00'
    
    for col in range(1, 8):
        ws.cell(row=total_row, column=col).font = Font(bold=True)
        ws.cell(row=total_row, column=col).fill = PatternFill(start_color="dbeafe", end_color="dbeafe", fill_type="solid")
    
    # Adjust column widths
    ws.column_dimensions['A'].width = 12
    ws.column_dimensions['B'].width = 14
    ws.column_dimensions['C'].width = 12
    ws.column_dimensions['D'].width = 25
    ws.column_dimensions['E'].width = 14
    ws.column_dimensions['F'].width = 14
    ws.column_dimensions['G'].width = 14
    
    # Save to bytes
    xlsx_bytes = BytesIO()
    wb.save(xlsx_bytes)
    xlsx_bytes.seek(0)
    
    return send_file(xlsx_bytes, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                     as_attachment=True, download_name=f"Supplier_Statement_{sup.get('company','').replace(' ','_')}_{date.today()}.xlsx")


@app.route("/api/suppliers/<sid>/ageing/excel", methods=["GET"])
def supplier_ageing_excel(sid):
    """Export outstanding ageing analysis to Excel"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        return jsonify({"error": "openpyxl not installed"}), 500
    
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM suppliers WHERE id = ?", (sid,))
        sup_row = cursor.fetchone()
        if not sup_row:
            return "Supplier not found", 404
        sup = dict(sup_row)
    
    # Get outstanding data
    response = get_supplier_outstanding(sid)
    if isinstance(response, tuple):
        response_obj = response[0]
        status_code = response[1] if len(response) > 1 else 200
    else:
        response_obj = response
        status_code = response.status_code

    response_data = response_obj.get_json() if hasattr(response_obj, 'get_json') else None
    if status_code != 200 or not response_data:
        return jsonify(response_data or {"error": "Failed to retrieve outstanding data"}), status_code or 400
    buckets = response_data['buckets']
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Ageing Analysis"
    
    # Header
    ws['A1'] = f"Outstanding Ageing Analysis - {sup.get('company', '')}"
    ws['A1'].font = Font(size=12, bold=True)
    ws['A2'] = f"As of {response_data['as_of']}"
    
    # Column headers
    headers = ["Ref No", "Bill Date", "Due Date", "Days Overdue", "Outstanding (₹)", "Age Bucket"]
    for col, hdr in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col, value=hdr)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="dc2626", end_color="dc2626", fill_type="solid")
    
    row_idx = 5
    bucket_colors = {
        "0_30": "dcfce7",
        "31_60": "fef08a",
        "61_90": "fed7aa",
        "91_plus": "fecaca"
    }
    
    for bucket_key, bucket_name in [("0_30", "Not Due"), ("31_60", "1-30 Days"), 
                                     ("61_90", "31-60 Days"), ("91_plus", "61-90 Days")]:
        items = buckets.get(bucket_key, [])
        for item in items:
            ws.cell(row=row_idx, column=1, value=item.get('ref', ''))
            ws.cell(row=row_idx, column=2, value=item.get('bill_date', ''))
            ws.cell(row=row_idx, column=3, value=item.get('due_date', ''))
            ws.cell(row=row_idx, column=4, value=item.get('days_overdue', 0))
            ws.cell(row=row_idx, column=5, value=item.get('pending_inr', 0)).number_format = '#,##0.00'
            ws.cell(row=row_idx, column=6, value=item.get('age_label', ''))
            
            fill_color = bucket_colors.get(bucket_key, "ffffff")
            for col in range(1, 7):
                ws.cell(row=row_idx, column=col).fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type="solid")
            
            row_idx += 1
    
    # Adjust column widths
    ws.column_dimensions['A'].width = 12
    ws.column_dimensions['B'].width = 12
    ws.column_dimensions['C'].width = 12
    ws.column_dimensions['D'].width = 14
    ws.column_dimensions['E'].width = 16
    ws.column_dimensions['F'].width = 14
    
    xlsx_bytes = BytesIO()
    wb.save(xlsx_bytes)
    xlsx_bytes.seek(0)
    
    return send_file(xlsx_bytes, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                     as_attachment=True, download_name=f"Ageing_Analysis_{sup.get('company','').replace(' ','_')}_{date.today()}.xlsx")


# ═══════════════════════════════════════════════════════════════════════════════
# ITEMS
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/api/items", methods=["GET"])
def get_items():
    """Fetch all items"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM items ORDER BY name")
        items = [dict(row) for row in cursor.fetchall()]
    return jsonify(items)


@app.route("/api/items", methods=["POST"])
@require_permission("po_edit")
def add_item():
    """Create new item"""
    req = request.get_json(silent=True)
    if not req:
        return jsonify({"error": "Invalid JSON"}), 400
    
    item_id = str(uuid.uuid4())
    item = {
        "id": item_id,
        "name": req.get("name", "").strip(),
        "description": req.get("description", "").strip(),
        "hs_code": req.get("hs_code", "").strip(),
        "unit": req.get("unit", "PCS").strip(),
        "currency": req.get("currency", "CNY").strip(),
        "default_price_usd": _safe_float(req.get("default_price_usd"), 0.0),
    }
    
    if not item["name"]:
        return jsonify({"error": "Item name is required"}), 400
    
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO items (id, name, description, hs_code, unit, currency, default_price_usd)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (
                item["id"], item["name"], item["description"], item["hs_code"],
                item["unit"], item["currency"], item["default_price_usd"]
            ))
    except sqlite3.IntegrityError:
        return jsonify({"error": f"Item '{item['name']}' already exists"}), 409
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    
    return jsonify(item), 201


@app.route("/api/items/<iid>", methods=["PUT"])
@require_permission("po_edit")
def update_item(iid):
    """Update item"""
    req = request.get_json(silent=True)
    if not req:
        return jsonify({"error": "Invalid JSON"}), 400
    
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            
            cursor.execute("SELECT * FROM items WHERE id = ?", (iid,))
            row = cursor.fetchone()
            if not row:
                return jsonify({"error": "Not found"}), 404
            
            item = dict(row)
            
            # Update fields
            item["name"] = req.get("name", item["name"])
            item["description"] = req.get("description", item["description"])
            item["hs_code"] = req.get("hs_code", item["hs_code"])
            item["unit"] = req.get("unit", item["unit"])
            item["currency"] = req.get("currency", item.get("currency", "CNY"))
            item["default_price_usd"] = _safe_float(req.get("default_price_usd"), item["default_price_usd"])
            
            cursor.execute("""
                UPDATE items
                SET name = ?, description = ?, hs_code = ?, unit = ?, currency = ?, default_price_usd = ?
                WHERE id = ?
            """, (
                item["name"], item["description"], item["hs_code"], item["unit"],
                item["currency"], item["default_price_usd"], iid
            ))
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    
    return jsonify(item)


@app.route("/api/items/<iid>", methods=["DELETE"])
@require_permission("po_delete")
def delete_item(iid):
    """Delete item (check for usage in POs and Quotations)"""
    with get_db() as conn:
        cursor = conn.cursor()
        blockers = []
        
        # Check if item is used in any PO line item
        cursor.execute("SELECT DISTINCT po_id FROM po_items WHERE item_id = ?", (iid,))
        po_ids = [row[0] for row in cursor.fetchall()]
        if po_ids:
            cursor.execute(f"SELECT po_number FROM purchase_orders WHERE id IN ({','.join(['?']*len(po_ids))})", po_ids)
            used_pos = [row[0] for row in cursor.fetchall()]
            blockers.append(f"POs: {', '.join(used_pos)}")
        
        # O(1) lookup via quotation_refs index (no JSON scanning)
        cursor.execute("""
            SELECT q.quotation_number, q.title, q.id
            FROM quotation_refs r
            JOIN quotations q ON q.id = r.quotation_id
            WHERE r.ref_type = 'item' AND r.ref_id = ?
        """, (iid,))
        used_quotes = [
            row["quotation_number"] or row["title"] or row["id"][:8]
            for row in cursor.fetchall()
        ]
        if used_quotes:
            blockers.append(f"Quotations: {', '.join(used_quotes)}")
        
        if blockers:
            return jsonify({"error": f"Cannot delete — used in {' | '.join(blockers)}"}), 409
        
        cursor.execute("DELETE FROM items WHERE id = ?", (iid,))
    
    return jsonify({"ok": True})


@app.route("/api/items/<item_id>/stock", methods=["GET"])
def get_item_stock(item_id):
    """Dynamically calculate current stock (Total PO Qty - Total Sales Qty)."""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT COALESCE(SUM(qty), 0) FROM po_items 
            JOIN purchase_orders ON po_items.po_id = purchase_orders.id 
            WHERE item_id = ? AND purchase_orders.deleted_at IS NULL
        """, (item_id,))
        po_qty = cursor.fetchone()[0]

        cursor.execute("""
            SELECT COALESCE(SUM(qty), 0) FROM customer_invoice_items
            WHERE item_id = ?
        """, (item_id,))
        sales_qty = cursor.fetchone()[0]

    stock = po_qty - sales_qty
    return jsonify({"stock": stock, "po_qty": po_qty, "sales_qty": sales_qty})


@app.route("/api/customers/<cid>/last-sale-rate/<item_id>", methods=["GET"])
def get_last_sale_rate(cid, item_id):
    """Fetch the rate from the last invoice for this customer and item."""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT unit_price FROM customer_invoice_items cii
            JOIN customer_invoices ci ON cii.invoice_id = ci.id
            WHERE ci.customer_id = ? AND cii.item_id = ?
            ORDER BY ci.invoice_date DESC, ci.created_at DESC
            LIMIT 1
        """, (cid, item_id))
        row = cursor.fetchone()
        rate = row[0] if row else 0.0
        
    return jsonify({"last_rate": rate})


# ═══════════════════════════════════════════════════════════════════════════════
# FORWARDERS
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/api/forwarders", methods=["GET"])
def get_forwarders():
    """Fetch all forwarders"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM forwarders ORDER BY name")
        forwarders = []
        for row in cursor.fetchall():
            fw = dict(row)
            cursor.execute(
                "SELECT id, name, contact_person, phone, email, address "
                "FROM forwarder_godowns WHERE forwarder_id = ? ORDER BY sort_order",
                (fw["id"],)
            )
            godown_rows = cursor.fetchall()
            if godown_rows:
                fw["godowns"] = [
                    {
                        "id":             r["id"],
                        "label":          r["name"],        # 'name' col stores the label
                        "contact_person": r["contact_person"] or "",
                        "phone":          r["phone"] or "",
                        "email":          r["email"] or "",
                        "address":        r["address"] or "",
                    }
                    for r in godown_rows
                ]
            else:
                # Fallback: try legacy JSON blob in forwarders.godowns column
                raw = json.loads(fw.get("godowns") or "[]")
                # Normalise — old format was plain strings, new format is objects
                fw["godowns"] = [
                    g if isinstance(g, dict) else {"id": str(uuid.uuid4()), "label": g,
                                                   "contact_person": "", "phone": "",
                                                   "email": "", "address": ""}
                    for g in raw
                ]
            forwarders.append(fw)
    return jsonify(forwarders)


@app.route("/api/forwarders", methods=["POST"])
@require_permission("forwarder_edit")
def add_forwarder():
    """Create new forwarder"""
    req = request.get_json(silent=True)
    if not req:
        return jsonify({"error": "Invalid JSON"}), 400
    
    fw_id = str(uuid.uuid4())
    forwarder = {
        "id": fw_id,
        "name": req.get("name", "").strip(),
        "contact_person": req.get("contact_person", "").strip(),
        "phone": req.get("phone", "").strip(),
        "email": req.get("email", "").strip(),
        "godowns": req.get("godowns", []),
    }
    
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO forwarders (id, name, contact_person, phone, email, godowns)
                VALUES (?, ?, ?, ?, ?, ?)
            """, (
                forwarder["id"], forwarder["name"], forwarder["contact_person"],
                forwarder["phone"], forwarder["email"], json.dumps(forwarder["godowns"])
            ))
            for i, g in enumerate(forwarder["godowns"]):
                # g is a rich object: {id, label, contact_person, phone, email, address}
                gid = g.get("id") or str(uuid.uuid4())
                cursor.execute(
                    "INSERT INTO forwarder_godowns "
                    "(id, forwarder_id, name, contact_person, phone, email, address, sort_order) "
                    "VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                    (gid, forwarder["id"],
                     g.get("label", "") or g.get("name", ""),
                     g.get("contact_person", ""),
                     g.get("phone", ""),
                     g.get("email", ""),
                     g.get("address", ""),
                     i)
                )
    except sqlite3.IntegrityError:
        return jsonify({"error": f"Forwarder '{forwarder['name']}' already exists"}), 409
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    
    return jsonify(forwarder), 201


@app.route("/api/forwarders/<fid>", methods=["PUT"])
@require_permission("forwarder_edit")
def update_forwarder(fid):
    """Update forwarder"""
    req = request.get_json(silent=True)
    if not req:
        return jsonify({"error": "Invalid JSON"}), 400
    
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            
            cursor.execute("SELECT * FROM forwarders WHERE id = ?", (fid,))
            row = cursor.fetchone()
            if not row:
                return jsonify({"error": "Not found"}), 404
            
            forwarder = dict(row)
            forwarder["godowns"] = json.loads(forwarder.get("godowns") or "[]")
            
            forwarder["name"] = req.get("name", forwarder["name"])
            forwarder["contact_person"] = req.get("contact_person", forwarder.get("contact_person", ""))
            forwarder["phone"] = req.get("phone", forwarder["phone"])
            forwarder["email"] = req.get("email", forwarder.get("email", ""))
            forwarder["godowns"] = req.get("godowns", forwarder.get("godowns", []))
            
            cursor.execute("""
                UPDATE forwarders
                SET name = ?, contact_person = ?, phone = ?, email = ?, godowns = ?
                WHERE id = ?
            """, (
                forwarder["name"], forwarder["contact_person"], forwarder["phone"],
                forwarder["email"], json.dumps(forwarder["godowns"]), fid
            ))
            cursor.execute("DELETE FROM forwarder_godowns WHERE forwarder_id = ?", (fid,))
            for i, g in enumerate(forwarder["godowns"]):
                # g is a rich object: {id, label, contact_person, phone, email, address}
                gid = g.get("id") or str(uuid.uuid4())
                cursor.execute(
                    "INSERT INTO forwarder_godowns "
                    "(id, forwarder_id, name, contact_person, phone, email, address, sort_order) "
                    "VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                    (gid, fid,
                     g.get("label", "") or g.get("name", ""),
                     g.get("contact_person", ""),
                     g.get("phone", ""),
                     g.get("email", ""),
                     g.get("address", ""),
                     i)
                )
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    
    return jsonify(forwarder)


@app.route("/api/forwarders/<fid>", methods=["DELETE"])
@require_permission("forwarder_delete")
def delete_forwarder(fid):
    """Delete forwarder (check for usage in POs)"""
    with get_db() as conn:
        cursor = conn.cursor()
        
        cursor.execute("SELECT po_number FROM purchase_orders WHERE forwarder_id = ? AND deleted_at IS NULL", (fid,))
        used_pos = [row[0] for row in cursor.fetchall()]
        
        if used_pos:
            return jsonify({"error": f"Cannot delete — used in POs: {', '.join(used_pos)}"}), 409
        
        cursor.execute("DELETE FROM forwarders WHERE id = ?", (fid,))
    
    return jsonify({"ok": True})


# ═══════════════════════════════════════════════════════════════════════════════
# PURCHASE ORDERS
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/api/po", methods=["GET"])
def get_pos():
    """Fetch all purchase orders"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM purchase_orders WHERE deleted_at IS NULL ORDER BY po_number DESC")
        po_rows = cursor.fetchall()
        
        # Batch fetch all related shipments
        cursor.execute("""
            SELECT spl.po_id, s.id, s.description, s.status, s.expected_arrival
            FROM shipments s
            JOIN shipment_po_link spl ON spl.shipment_id = s.id
            WHERE s.deleted_at IS NULL
        """)
        shipments_by_po = {row["po_id"]: dict(row) for row in cursor.fetchall()}
        
        # Batch fetch all line items
        cursor.execute("SELECT * FROM po_items ORDER BY line_sequence")
        line_items_by_po = {}
        for item in cursor.fetchall():
            po_id = item["po_id"]
            if po_id not in line_items_by_po:
                line_items_by_po[po_id] = []
            line_items_by_po[po_id].append({
                "id": item["id"],
                "item_id": item["item_id"],
                "item_name": item["item_name"],
                "description": item["description"],
                "hs_code": item["hs_code"],
                "qty": item["qty"],
                "unit": item["unit"],
                "unit_price": item["unit_price"],
            })
            
        purchase_orders = []
        for row in po_rows:
            po = dict(row)
            po_id = po["id"]
            
            # Attach linked shipment info
            po["shipment"] = shipments_by_po.get(po_id)
            
            # Attach line items
            po["line_items"] = line_items_by_po.get(po_id, [])
            
            # Parse supplier snapshot
            po["supplier_snapshot"] = json.loads(po.get("supplier_snapshot") or "{}")
            
            purchase_orders.append(po)
    
    return jsonify(purchase_orders)


@app.route("/api/po", methods=["POST"])
@require_permission("po_edit")
def create_po():
    """Create new purchase order"""
    req = request.get_json(silent=True)
    if not req:
        return jsonify({"error": "Invalid JSON"}), 400
    
    po_id = str(uuid.uuid4())
    po = {
        "id": po_id,
        "po_number": req.get("po_number", ""),
        "po_date": req.get("po_date", str(date.today())),
        "supplier_id": req.get("supplier_id", ""),
        "godown_id": req.get("godown_id", ""),
        "supplier_snapshot": req.get("supplier_snapshot", {}),
        "payment_conditions": req.get("payment_conditions", ""),
        "delivery_terms": req.get("delivery_terms", "FOB"),
        "delivery_address": req.get("delivery_address", ""),
        "forwarder_id": req.get("forwarder_id", ""),
        "forwarder_name": req.get("forwarder_name", ""),
        "forwarder_contact": req.get("forwarder_contact", ""),
        "remarks": req.get("remarks", ""),
        "currency": req.get("currency", "CNY"),
        "line_items": req.get("line_items", []),
        "status": req.get("status", "Draft"),
        "created_at": str(date.today()),
        "lead_time_days": req.get("lead_time_days", 0),
        "due_date": req.get("due_date", ""),
        "lc_usd_rate": req.get("lc_usd_rate", 84.0),
        "lc_rmb_rate": req.get("lc_rmb_rate", 11.5),
        "lc_bank": req.get("lc_bank", 0.0),
        "lc_ship": req.get("lc_ship", 0.0),
        "lc_duty": req.get("lc_duty", 0.0),
        "lc_trans": req.get("lc_trans", 0.0),
        "lc_gst_duty": req.get("lc_gst_duty", 0.0),
        "lc_doc_pct": req.get("lc_doc_pct", 0.0),
    }
    
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            
            # Check if PO number already exists
            if po["po_number"].strip():
                cursor.execute("SELECT id FROM purchase_orders WHERE po_number = ?", (po["po_number"],))
                if cursor.fetchone():
                    return jsonify({"error": f"PO number '{po['po_number']}' already exists"}), 409

            # Validation: Supplier is required
            if not po["supplier_id"] or not po["supplier_id"].strip():
                return jsonify({"error": "Supplier is required"}), 400
            
            # Insert PO
            cursor.execute("""
                INSERT INTO purchase_orders (
                    id, po_number, po_date, supplier_id, supplier_snapshot,
                    payment_conditions, delivery_terms, delivery_address,
                    forwarder_id, forwarder_name, forwarder_contact, remarks,
                    currency, status, created_at, lead_time_days, due_date, godown_id,
                    lc_usd_rate, lc_rmb_rate, lc_bank, lc_ship, lc_duty, lc_trans, lc_gst_duty, lc_doc_pct
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                po["id"], po["po_number"], po["po_date"], po["supplier_id"],
                json.dumps(po["supplier_snapshot"]),
                po["payment_conditions"], po["delivery_terms"], po["delivery_address"],
                po["forwarder_id"], po["forwarder_name"], po["forwarder_contact"],
                po["remarks"], po["currency"], po["status"], po["created_at"],
                int(po["lead_time_days"]), po["due_date"], po["godown_id"],
                float(po["lc_usd_rate"]), float(po["lc_rmb_rate"]),
                float(po["lc_bank"]), float(po["lc_ship"]), float(po["lc_duty"]),
                float(po["lc_trans"]), float(po["lc_gst_duty"]), float(po["lc_doc_pct"])
            ))
            
            # Log initial status (from_status is NULL = creation event)
            _log_status_change(
                cursor,
                po["id"],
                from_status=None,
                to_status=po["status"],
                note="PO created"
            )
            
            # Insert line items
            for seq, li in enumerate(po.get("line_items", [])):
                cursor.execute("""
                    INSERT INTO po_items (id, po_id, item_id, item_name, description, hs_code, qty, unit, unit_price, line_sequence)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    str(uuid.uuid4()), po["id"], li.get("item_id", ""),
                    li.get("item_name", ""), li.get("description", ""),
                    li.get("hs_code", ""), float(li.get("qty", 0)),
                    li.get("unit", "PCS"), float(li.get("unit_price", 0)), seq
                ))
            
            # Update po_sequence in settings if the new PO number has a higher sequence.
            _sync_po_sequence(cursor, po["po_number"])

            # Auto-create shipment if PO is already Confirmed/Shipped with a forwarder
            if po.get("status") in _AUTO_SHIP_STATUSES and po.get("forwarder_id"):
                _auto_create_shipment_for_po(cursor, po_id, po)
            
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    
    return jsonify(po), 201


@app.route("/api/po/<pid>", methods=["PUT"])
@require_permission("po_edit")
def update_po(pid):
    """Update purchase order"""
    req = request.get_json(silent=True)
    if not req:
        return jsonify({"error": "Invalid JSON"}), 400
    
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            
            # Fetch existing PO
            cursor.execute("SELECT * FROM purchase_orders WHERE id = ?", (pid,))
            row = cursor.fetchone()
            if not row:
                return jsonify({"error": "Not found"}), 404
            
            po = dict(row)
            po["supplier_snapshot"] = json.loads(po.get("supplier_snapshot") or "{}")
            old_status = po["status"]   # capture before any field merge
            
            # Update fields
            update_fields = [
                "po_number", "po_date", "supplier_id", "supplier_snapshot",
                "payment_conditions", "delivery_terms", "delivery_address",
                "forwarder_id", "forwarder_name", "forwarder_contact", "remarks",
                "currency", "status", "lead_time_days", "due_date", "godown_id",
                "lc_usd_rate", "lc_rmb_rate", "lc_bank", "lc_ship", "lc_duty",
                "lc_trans", "lc_gst_duty", "lc_doc_pct"
            ]
            
            for field in update_fields:
                if field in req:
                    po[field] = req[field]
            
            # Update PO
            cursor.execute("""
                UPDATE purchase_orders
                SET po_number = ?, po_date = ?, supplier_id = ?, supplier_snapshot = ?,
                    payment_conditions = ?, delivery_terms = ?, delivery_address = ?,
                    forwarder_id = ?, forwarder_name = ?, forwarder_contact = ?, remarks = ?,
                    currency = ?, status = ?, lead_time_days = ?, due_date = ?, godown_id = ?,
                    lc_usd_rate = ?, lc_rmb_rate = ?, lc_bank = ?, lc_ship = ?, lc_duty = ?,
                    lc_trans = ?, lc_gst_duty = ?, lc_doc_pct = ?, updated_at = CURRENT_TIMESTAMP
                WHERE id = ?
            """, (
                po["po_number"], po["po_date"], po["supplier_id"],
                json.dumps(po.get("supplier_snapshot", {})),
                po["payment_conditions"], po["delivery_terms"], po["delivery_address"],
                po["forwarder_id"], po["forwarder_name"], po["forwarder_contact"],
                po["remarks"], po["currency"], po["status"],
                int(po.get("lead_time_days", 0)), po.get("due_date", ""), po.get("godown_id", ""),
                float(po.get("lc_usd_rate", 84.0)), float(po.get("lc_rmb_rate", 11.5)),
                float(po.get("lc_bank", 0)), float(po.get("lc_ship", 0)),
                float(po.get("lc_duty", 0)), float(po.get("lc_trans", 0)),
                float(po.get("lc_gst_duty", 0)), float(po.get("lc_doc_pct", 0)), pid
            ))
            
            # Log status change if status actually changed
            if "status" in req and req["status"] != old_status:
                note = req.get("status_note")  # frontend can pass an optional note
                _log_status_change(cursor, pid, old_status, po["status"], note)
                # Auto-sync linked shipment status
                _sync_shipment_from_po(cursor, pid, po["status"])
                # Auto-create shipment if PO now has a forwarder and no shipment yet
                _auto_create_shipment_for_po(cursor, pid, po)
                # Auto-create Supplier Ledger entry (INVOICE on Confirmed, DEBIT_NOTE on Cancelled)
                try:
                    _auto_ledger_on_status_change(cursor, pid, po["status"], po)
                except Exception as le:
                    print(f"[supplier-books] auto-ledger error (non-fatal): {le}")
            elif req.get("status_note"):
                # Save a note-only log entry even when status has not changed
                _log_status_change(cursor, pid, old_status, old_status, req["status_note"], force=True)

            # Auto-create shipment if forwarder was just assigned (even with no status change)
            if "forwarder_id" in req and req.get("forwarder_id"):
                _auto_create_shipment_for_po(cursor, pid, po)
            
            # Sync sequence in case of manual renumbering
            _sync_po_sequence(cursor, po["po_number"])
            
            # Update line items if provided
            if "line_items" in req:
                # Delete existing line items
                cursor.execute("DELETE FROM po_items WHERE po_id = ?", (pid,))
                
                # Insert new line items
                for seq, li in enumerate(req["line_items"]):
                    cursor.execute("""
                        INSERT INTO po_items (id, po_id, item_id, item_name, description, hs_code, qty, unit, unit_price, line_sequence)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """, (
                        str(uuid.uuid4()), pid, li.get("item_id", ""),
                        li.get("item_name", ""), li.get("description", ""),
                        li.get("hs_code", ""), float(li.get("qty", 0)),
                        li.get("unit", "PCS"), float(li.get("unit_price", 0)), seq
                    ))
            
            # Fetch updated line items
            cursor2 = conn.cursor()
            cursor2.execute("SELECT * FROM po_items WHERE po_id = ? ORDER BY line_sequence", (pid,))
            po["line_items"] = [
                {
                    "id": item["id"],
                    "item_id": item["item_id"],
                    "item_name": item["item_name"],
                    "description": item["description"],
                    "hs_code": item["hs_code"],
                    "qty": item["qty"],
                    "unit": item["unit"],
                    "unit_price": item["unit_price"],
                }
                for item in cursor2.fetchall()
            ]
            
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    
    return jsonify(po)


@app.route("/api/po/<pid>", methods=["DELETE"])
@require_permission("po_delete")
def delete_po(pid):
    """Soft-delete a purchase order (sets deleted_at; data is preserved).

    Use DELETE /api/po/<pid>/hard to permanently destroy a PO and its files.
    """
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM purchase_orders WHERE id = ? AND deleted_at IS NULL", (pid,))
        if not cursor.fetchone():
            return jsonify({"error": "PO not found"}), 404
        cursor.execute(
            "UPDATE purchase_orders SET deleted_at = CURRENT_TIMESTAMP WHERE id = ?",
            (pid,)
        )
    return jsonify({"ok": True})


@app.route("/api/po/<pid>/restore", methods=["POST"])
@require_permission("po_edit")
def restore_po(pid):
    """Restore a soft-deleted purchase order"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM purchase_orders WHERE id = ? AND deleted_at IS NOT NULL", (pid,))
        if not cursor.fetchone():
            return jsonify({"error": "No deleted PO found with that id"}), 404
        cursor.execute(
            "UPDATE purchase_orders SET deleted_at = NULL WHERE id = ?",
            (pid,)
        )
    return jsonify({"ok": True})


@app.route('/api/po/<pid>/timeline', methods=['GET'])
def get_po_timeline(pid):
    """
    Returns all status log entries for one PO, newest first.
    Example URL: GET /api/po/17/timeline
    """
    with get_db() as db:
        rows = db.execute('''
            SELECT id, from_status, to_status, changed_at, note
            FROM   po_status_log
            WHERE  po_id = ?
            ORDER  BY changed_at DESC
        ''', [pid]).fetchall()

    # Convert rows to plain dictionaries so Flask can turn them into JSON
    entries = [dict(r) for r in rows]
    return jsonify({'timeline': entries})



@app.route("/api/analytics/suppliers", methods=["GET"])
def supplier_analytics():
    """
    Aggregate PO data per supplier:
    - Total PO count, total order value (sum of line items)
    - Average promised lead time (lead_time_days column)
    - Actual lead time (days from po_date to Received status log entry)
    - On-time delivery % (actual lead time <= promised lead time, for Received POs)
    - Total value over last 12 months (month-by-month)
    """
    try:
        with get_db() as conn:
            cursor = conn.cursor()

            # ── Per-supplier PO totals ────────────────────────────────────────
            cursor.execute("""
                SELECT
                    s.id            AS supplier_id,
                    s.name          AS supplier_name,
                    s.company       AS company,
                    COUNT(po.id)    AS total_pos,
                    SUM(CASE WHEN po.status NOT IN ('Cancelled') THEN 1 ELSE 0 END)
                                    AS active_pos,
                    SUM(CASE WHEN po.status = 'Shipped' THEN 1 ELSE 0 END)
                                    AS shipped_pos,
                    SUM(CASE WHEN po.status = 'In Transit' THEN 1 ELSE 0 END)
                                    AS in_transit_pos,
                    AVG(CASE WHEN po.lead_time_days > 0 THEN po.lead_time_days END)
                                    AS avg_promised_days,
                    SUM(COALESCE(pv.value_cny, 0))
                                    AS total_value_cny
                FROM suppliers s
                LEFT JOIN purchase_orders po
                    ON po.supplier_id = s.id AND po.deleted_at IS NULL
                LEFT JOIN (
                    SELECT po_id, SUM(qty * unit_price) AS value_cny
                    FROM po_items GROUP BY po_id
                ) pv ON pv.po_id = po.id
                GROUP BY s.id
                ORDER BY total_value_cny DESC
            """)
            supplier_rows = [dict(r) for r in cursor.fetchall()]

            # ── On-time delivery calculation (using status log) ───────────────
            cursor.execute("""
                SELECT
                    po.supplier_id,
                    po.po_date,
                    po.lead_time_days,
                    COALESCE(rlog.received_at, po.updated_at) AS received_at
                FROM purchase_orders po
                LEFT JOIN (
                    SELECT po_id, MIN(changed_at) as received_at
                    FROM po_status_log
                    WHERE to_status = 'Received'
                    GROUP BY po_id
                ) rlog ON rlog.po_id = po.id
                WHERE po.deleted_at IS NULL AND po.status = 'Received'
            """)
            delivery_rows = cursor.fetchall()

            # Compute per-supplier on-time stats
            from datetime import datetime
            delivery_stats = {}  # supplier_id -> {on_time: int, total: int, actual_days_sum: int}
            for dr in delivery_rows:
                sid       = dr["supplier_id"]
                po_date   = dr["po_date"]
                promised  = dr["lead_time_days"]
                recv_str  = dr["received_at"]
                if not po_date or not recv_str:
                    continue
                try:
                    d0      = datetime.strptime(po_date[:10], "%Y-%m-%d")
                    d1      = datetime.strptime(recv_str[:10], "%Y-%m-%d")
                    actual  = (d1 - d0).days
                except ValueError:
                    continue
                if sid not in delivery_stats:
                    delivery_stats[sid] = {"on_time": 0, "total": 0, "actual_days_sum": 0}
                delivery_stats[sid]["total"]           += 1
                delivery_stats[sid]["actual_days_sum"] += actual
                if actual <= promised:
                    delivery_stats[sid]["on_time"] += 1

            # ── Monthly spend per supplier (last 12 months) ───────────────────
            cursor.execute("""
                SELECT
                    po.supplier_id,
                    substr(po.po_date, 1, 7) AS ym,
                    SUM(COALESCE(pv.value_cny, 0)) AS monthly_value
                FROM purchase_orders po
                LEFT JOIN (
                    SELECT po_id, SUM(qty * unit_price) AS value_cny
                    FROM po_items GROUP BY po_id
                ) pv ON pv.po_id = po.id
                WHERE po.deleted_at IS NULL
                  AND po.po_date >= date('now', '-12 months')
                GROUP BY po.supplier_id, ym
                ORDER BY ym ASC
            """)
            monthly_raw = cursor.fetchall()
            monthly_by_supplier = {}
            for mr in monthly_raw:
                sid = mr["supplier_id"]
                monthly_by_supplier.setdefault(sid, []).append({
                    "month": mr["ym"], "value": mr["monthly_value"] or 0
                })

            # ── Merge everything ──────────────────────────────────────────────
            result = []
            for s in supplier_rows:
                sid   = s["supplier_id"]
                stats = delivery_stats.get(sid, {})
                total_d = stats.get("total", 0)
                on_t    = stats.get("on_time", 0)
                actual_days_sum = stats.get("actual_days_sum", 0)
                result.append({
                    **s,
                    "avg_actual_days":   round(actual_days_sum / total_d, 1) if total_d else None,
                    "on_time_pct":       round(on_t / total_d * 100, 1)              if total_d else None,
                    "delivered_pos":     total_d,
                    "shipped_pos":       s.get("shipped_pos", 0) or 0,
                    "in_transit_pos":    s.get("in_transit_pos", 0) or 0,
                    "monthly_spend":     monthly_by_supplier.get(sid, []),
                })

    except Exception as e:
        return jsonify({"error": str(e)}), 400

    return jsonify(result)


@app.route("/api/po/deleted", methods=["GET"])
def list_deleted_pos():
    """List soft-deleted purchase orders"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT id, po_number, po_date, supplier_id, status, deleted_at "
            "FROM purchase_orders WHERE deleted_at IS NOT NULL ORDER BY deleted_at DESC"
        )
        return jsonify([dict(r) for r in cursor.fetchall()])


@app.route("/api/po/<pid>/hard", methods=["DELETE"])
@require_permission("po_delete")
def hard_delete_po(pid):
    """Permanently delete a PO and all its attachments / payment files.

    Only works on already-soft-deleted POs so there is always a two-step
    confirmation before data is irrecoverably lost.
    """
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM purchase_orders WHERE id = ? AND deleted_at IS NOT NULL", (pid,))
        if not cursor.fetchone():
            return jsonify({"error": "PO must be soft-deleted first"}), 409

        # Remove physical attachment files
        attach_folder = os.path.join(ATTACH_DIR, pid)
        if os.path.isdir(attach_folder):
                    shutil.rmtree(attach_folder, ignore_errors=True)

        cursor.execute("DELETE FROM purchase_orders WHERE id = ?", (pid,))
    return jsonify({"ok": True})


# ═══════════════════════════════════════════════════════════════════════════════
# SETTINGS
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/api/settings", methods=["GET"])
def get_settings():
    """Fetch all settings"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT key, value FROM settings")
        settings = {row["key"]: row["value"] for row in cursor.fetchall()}
    return jsonify(settings)


@app.route("/api/settings", methods=["POST"])
@require_permission("admin_rbac")
def save_settings():
    """Update settings"""
    req = request.get_json(silent=True)
    if not req:
        return jsonify({"error": "Invalid JSON"}), 400
    
    with get_db() as conn:
        cursor = conn.cursor()
        for key, value in req.items():
            cursor.execute(
                "INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)",
                (key, str(value))
            )
        
        # Return updated settings
        cursor.execute("SELECT key, value FROM settings")
        settings = {row["key"]: row["value"] for row in cursor.fetchall()}
    
    return jsonify(settings)


# ═══════════════════════════════════════════════════════════════════════════════
# NEXT PO NUMBER (Auto-Numbering)
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/api/next-po-number", methods=["GET"])
def next_po_number():
    """Predict the next PO number based on current settings and existing POs.

    Does NOT persist the sequence to the database. This ensures that opening
    the "New PO" dialog without saving doesn't waste a PO number. The sequence
    is only advanced when a PO is actually saved in `create_po`.
    """
    with get_db() as conn:
        cursor = conn.cursor()

        # Ensure default settings exist
        init_default_settings(conn)

        cursor.execute("SELECT key, value FROM settings WHERE key IN ('po_prefix', 'po_sequence')")
        cfg = {row["key"]: row["value"] for row in cursor.fetchall()}

        prefix = (cfg.get("po_prefix") or "PO").strip()
        seq = int(cfg.get("po_sequence") or "0")
        year = date.today().year

        # Advance past any gaps or existing POs
        while True:
            seq += 1
            po_number = f"{prefix}-{year}-{str(seq).zfill(3)}"
            cursor.execute("SELECT id FROM purchase_orders WHERE po_number = ?", (po_number,))
            if not cursor.fetchone():
                break

    return jsonify({"po_number": po_number, "sequence": seq})


@app.route("/api/next-invoice-number", methods=["GET"])
def next_invoice_number():
    """Predict the next Invoice number based on current settings and existing invoices."""
    with get_db() as conn:
        cursor = conn.cursor()
        init_default_settings(conn)
        cursor.execute("SELECT key, value FROM settings WHERE key IN ('inv_prefix', 'inv_sequence')")
        cfg = {row["key"]: row["value"] for row in cursor.fetchall()}

        prefix = (cfg.get("inv_prefix") or "INV").strip()
        seq = int(cfg.get("inv_sequence") or "0")
        year = date.today().year

        while True:
            seq += 1
            inv_number = f"{prefix}-{year}-{str(seq).zfill(3)}"
            cursor.execute("SELECT id FROM customer_invoices WHERE invoice_no = ?", (inv_number,))
            if not cursor.fetchone():
                break

    return jsonify({"inv_number": inv_number, "sequence": seq})

# ═══════════════════════════════════════════════════════════════════════════════
# EXPORT / IMPORT
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/api/export", methods=["GET"])
def export_data():
    """
    📤 EXPORT ENDPOINT
    Read all data from SQLite → Convert to JSON → Return downloadable backup
    """
    bundle = {}
    
    with get_db() as conn:
        cursor = conn.cursor()
        
        # Export suppliers
        cursor.execute("SELECT * FROM suppliers ORDER BY name")
        bundle["suppliers"] = [dict(row) for row in cursor.fetchall()]
        
        # Export items
        cursor.execute("SELECT * FROM items ORDER BY name")
        bundle["items"] = [dict(row) for row in cursor.fetchall()]
        
        # Export forwarders
        cursor.execute("SELECT * FROM forwarders ORDER BY name")
        forwarders = []
        for row in cursor.fetchall():
            fw = dict(row)
            cursor2 = conn.cursor()
            cursor2.execute(
                "SELECT id, name, contact_person, phone, email, address "
                "FROM forwarder_godowns WHERE forwarder_id = ? ORDER BY sort_order",
                (fw["id"],)
            )
            gd_rows = cursor2.fetchall()
            if gd_rows:
                fw["godowns"] = [
                    {
                        "id":             r["id"],
                        "label":          r["name"],
                        "contact_person": r["contact_person"] or "",
                        "phone":          r["phone"] or "",
                        "email":          r["email"] or "",
                        "address":        r["address"] or "",
                    }
                    for r in gd_rows
                ]
            else:
                fw["godowns"] = json.loads(fw.get("godowns") or "[]")
            forwarders.append(fw)
        bundle["forwarders"] = forwarders
        
        # Export purchase orders with line items (including deleted for full backup)
        cursor.execute("SELECT * FROM purchase_orders ORDER BY po_date DESC")
        purchase_orders = []
        for po_row in cursor.fetchall():
            po = dict(po_row)
            po["supplier_snapshot"] = json.loads(po.get("supplier_snapshot") or "{}")
            
            # Fetch line items — use a separate cursor to avoid clobbering the outer result set
            cursor2 = conn.cursor()
            cursor2.execute("SELECT * FROM po_items WHERE po_id = ? ORDER BY line_sequence", (po["id"],))
            po["line_items"] = [
                {
                    "item_id": item["item_id"],
                    "item_name": item["item_name"],
                    "description": item["description"],
                    "hs_code": item["hs_code"],
                    "qty": item["qty"],
                    "unit": item["unit"],
                    "unit_price": item["unit_price"],
                }
                for item in cursor2.fetchall()
            ]
            purchase_orders.append(po)
        
        bundle["purchase_orders"] = purchase_orders
        
        # Export customers and customer ledger entries
        cursor.execute("SELECT * FROM customers ORDER BY name")
        bundle["customers"] = [dict(row) for row in cursor.fetchall()]

        cursor.execute("SELECT * FROM customer_ledger_entries ORDER BY entry_date")
        bundle["customer_ledger_entries"] = [dict(row) for row in cursor.fetchall()]

        # Export suppliers' ledger entries
        cursor.execute("SELECT * FROM supplier_ledger_entries ORDER BY entry_date")
        bundle["supplier_ledger_entries"] = [dict(row) for row in cursor.fetchall()]

        # Export quotations
        cursor.execute("SELECT * FROM quotations")
        quotations = []
        for row in cursor.fetchall():
            q = dict(row)
            q["details"] = json.loads(q.get("details") or "{}")
            quotations.append(q)
        bundle["quotations"] = quotations
        
        # Export shipments with linked PO IDs
        cursor.execute("SELECT * FROM shipments")
        shipments = []
        for s_row in cursor.fetchall():
            s = dict(s_row)
            s["notes"] = json.loads(s.get("notes") or "[]")
            # Fetch linked PO IDs
            cursor2 = conn.cursor()
            cursor2.execute("SELECT po_id FROM shipment_po_link WHERE shipment_id = ?", (s["id"],))
            s["po_ids"] = [r["po_id"] for r in cursor2.fetchall()]
            shipments.append(s)
        bundle["shipments"] = shipments
        
        # Export settings
        cursor.execute("SELECT key, value FROM settings")
        settings = {row["key"]: row["value"] for row in cursor.fetchall()}
        bundle["settings"] = settings
    
    bundle["export_date"] = str(date.today())
    bundle["version"] = "2.0 (SQLite)"
    
    json_bytes = json.dumps(bundle, indent=2, ensure_ascii=False).encode("utf-8")
    
    zip_buf = BytesIO()
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        # Add database backup
        zf.writestr('database_backup.json', json_bytes)
        
        # Add physical attachments
        if os.path.exists(ATTACH_DIR):
            for root, dirs, files in os.walk(ATTACH_DIR):
                for file in files:
                    file_path = os.path.join(root, file)
                    # Create relative path starting with 'attachments/'
                    rel_path = os.path.relpath(file_path, ATTACH_DIR)
                    arcname = os.path.join('attachments', rel_path)
                    zf.write(file_path, arcname)
                    
    zip_buf.seek(0)
    
    return send_file(
        zip_buf,
        mimetype="application/zip",
        as_attachment=True,
        download_name=f"portal_full_backup_{date.today().strftime('%Y%m%d')}.zip"
    )


@app.route("/api/import", methods=["POST"])
@require_permission("admin_rbac")
def import_data():
    """
    📥 IMPORT ENDPOINT WITH TRANSACTION & ROLLBACK
    Upload JSON file → Validate → Replace existing DB data
    On error: ROLLBACK entire transaction (no partial import)
    """
    if "file" not in request.files:
        return jsonify({"error": "No file"}), 400
    
    file = request.files["file"]
    file_bytes = file.read()
    
    try:
        if file.filename.lower().endswith('.zip'):
            with zipfile.ZipFile(BytesIO(file_bytes), 'r') as zf:
                # Read the JSON database backup
                if 'database_backup.json' not in zf.namelist():
                    return jsonify({"error": "ZIP file does not contain database_backup.json"}), 400
                bundle = json.loads(zf.read('database_backup.json').decode("utf-8"))
                
                # Extract attachments to the ATTACH_DIR
                for item in zf.namelist():
                    if item.startswith('attachments/') and not item.endswith('/'):
                        # Calculate the relative path within attachments/
                        rel_path = item[len('attachments/'):]
                        if not rel_path:
                            continue
                        target_path = os.path.join(ATTACH_DIR, rel_path)
                        # Ensure the target directory exists
                        os.makedirs(os.path.dirname(target_path), exist_ok=True)
                        # Write the file directly
                        with open(target_path, 'wb') as f_out:
                            f_out.write(zf.read(item))
        else:
            # Fallback for old .json uploads
            bundle = json.loads(file_bytes.decode("utf-8"))
    except Exception as e:
        return jsonify({"error": f"Invalid backup file: {e}"}), 400
    
    # Validate structure
    required_keys = ["suppliers", "items", "purchase_orders", "customers"]
    for key in required_keys:
        if key not in bundle:
            return jsonify({"error": f"Missing required key: {key}"}), 400
    
    # ── 🔐 BACKUP BEFORE IMPORT (must run OUTSIDE any transaction) ──────────
    # VACUUM INTO cannot be called inside a transaction; SQLite will raise
    # "cannot VACUUM from within a transaction".  We do the backup first on
    # a separate connection so the import transaction starts clean.
    backup_dir = os.path.join(os.path.dirname(__file__), "data", "po")
    try:
        for fname in os.listdir(backup_dir):
            if fname.startswith("backup_before_import_") and fname.endswith(".db"):
                try:
                    os.remove(os.path.join(backup_dir, fname))
                    print(f"🗑️ Removed old backup: {fname}")
                except Exception:
                    pass
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        backup_db_path = os.path.join(backup_dir, f"backup_before_import_{timestamp}.db")
        print(f"💾 Creating backup: {backup_db_path}")
        _bconn = get_connection()
        try:
            _bconn.execute(f"VACUUM INTO '{backup_db_path}'")
            print("✅ Backup created successfully")
        finally:
            _bconn.close()
    except Exception as backup_err:
        print(f"⚠️ Backup skipped (non-fatal): {backup_err}")
        # The uploaded JSON file itself serves as an implicit backup.

    # Start import transaction
    with get_db() as conn:
        try:
            cursor = conn.cursor()
            
            # ⚠️ CLEAR EXISTING DATA (within transaction)
            # Order matters due to FOREIGN KEY constraints
            cursor.execute("DELETE FROM shipment_po_link")
            cursor.execute("DELETE FROM shipments")
            cursor.execute("DELETE FROM po_status_log")
            cursor.execute("DELETE FROM po_attachments")
            cursor.execute("DELETE FROM po_payments")
            cursor.execute("DELETE FROM supplier_ledger_entries")
            cursor.execute("DELETE FROM customer_ledger_entries")
            cursor.execute("DELETE FROM customers")
            cursor.execute("DELETE FROM forwarder_godowns")
            cursor.execute("DELETE FROM po_items")
            cursor.execute("DELETE FROM purchase_orders")
            cursor.execute("DELETE FROM quotation_refs")
            cursor.execute("DELETE FROM quotation_supplier_rows")
            cursor.execute("DELETE FROM quotation_line_items")
            cursor.execute("DELETE FROM quotations")
            cursor.execute("DELETE FROM forwarders")
            cursor.execute("DELETE FROM items")
            cursor.execute("DELETE FROM suppliers")
            
            imported_stats = {
                "suppliers": 0,
                "items": 0,
                "forwarders": 0,
                "customers": 0,
                "customer_ledger_entries": 0,
                "supplier_ledger_entries": 0,
                "purchase_orders": 0,
                "po_items": 0,
                "quotations": 0,
                "shipments": 0,
            }
            
            # ── IMPORT SUPPLIERS ──────────────────────────────────────────
            for supplier in bundle.get("suppliers", []):
                cursor.execute("""
                    INSERT INTO suppliers (id, name, company, address, country, email, phone, wechat, bank_name, bank_account, swift_code)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    supplier.get("id", str(uuid.uuid4())),
                    supplier.get("name", ""),
                    supplier.get("company", ""),
                    supplier.get("address", ""),
                    supplier.get("country", "China"),
                    supplier.get("email", ""),
                    supplier.get("phone", ""),
                    supplier.get("wechat", ""),
                    supplier.get("bank_name", ""),
                    supplier.get("bank_account", ""),
                    supplier.get("swift_code", ""),
                ))
                imported_stats["suppliers"] += 1
            
            # ── IMPORT ITEMS ──────────────────────────────────────────────
            for item in bundle.get("items", []):
                cursor.execute("""
                    INSERT INTO items (id, name, description, hs_code, unit, currency, default_price_usd)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                """, (
                    item.get("id", str(uuid.uuid4())),
                    item.get("name", ""),
                    item.get("description", ""),
                    item.get("hs_code", ""),
                    item.get("unit", "PCS"),
                    item.get("currency", "CNY"),
                    float(item.get("default_price_usd", 0)),
                ))
                imported_stats["items"] += 1
            
            # ── IMPORT FORWARDERS ─────────────────────────────────────────
            for fwd in bundle.get("forwarders", []):
                fwd_id = fwd.get("id", str(uuid.uuid4()))
                godowns = fwd.get("godowns", [])
                cursor.execute("""
                    INSERT INTO forwarders (id, name, contact_person, phone, email, godowns)
                    VALUES (?, ?, ?, ?, ?, ?)
                """, (
                    fwd_id,
                    fwd.get("name", ""),
                    fwd.get("contact_person", ""),
                    fwd.get("phone", ""),
                    fwd.get("email", ""),
                    json.dumps(godowns),
                ))
                for i, g in enumerate(godowns):
                    if isinstance(g, dict):
                        gid = g.get("id") or str(uuid.uuid4())
                        cursor.execute(
                            "INSERT INTO forwarder_godowns "
                            "(id, forwarder_id, name, contact_person, phone, email, address, sort_order) "
                            "VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                            (gid, fwd_id,
                             g.get("label", "") or g.get("name", ""),
                             g.get("contact_person", ""),
                             g.get("phone", ""),
                             g.get("email", ""),
                             g.get("address", ""),
                             i)
                        )
                    else:
                        # Legacy: plain string godown name
                        cursor.execute(
                            "INSERT INTO forwarder_godowns "
                            "(id, forwarder_id, name, contact_person, phone, email, address, sort_order) "
                            "VALUES (?, ?, ?, '', '', '', '', ?)",
                            (str(uuid.uuid4()), fwd_id, str(g), i)
                        )
                imported_stats["forwarders"] += 1

            # ── IMPORT CUSTOMERS ─────────────────────────────────────────
            for customer in bundle.get("customers", []):
                cursor.execute("""
                    INSERT INTO customers (id, name, company, address, city, state, pincode, country,
                                           gstin, email, phone, credit_limit, credit_days, active, created_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    customer.get("id", str(uuid.uuid4())),
                    customer.get("name", ""),
                    customer.get("company", ""),
                    customer.get("address", ""),
                    customer.get("city", ""),
                    customer.get("state", ""),
                    customer.get("pincode", ""),
                    customer.get("country", "India"),
                    customer.get("gstin", ""),
                    customer.get("email", ""),
                    customer.get("phone", ""),
                    float(customer.get("credit_limit", 0.0)),
                    int(customer.get("credit_days", 30)),
                    int(customer.get("active", 1)),
                    customer.get("created_at", datetime.now().isoformat()),
                ))
                imported_stats["customers"] += 1

            # ── IMPORT CUSTOMER LEDGER ENTRIES ─────────────────────────────
            for entry in bundle.get("customer_ledger_entries", []):
                cursor.execute("""
                    INSERT INTO customer_ledger_entries (id, customer_id, entry_type, entry_date, ref_number,
                                                         description, amount_inr, amount_usd, usd_rate, dr_cr,
                                                         payment_mode, bank_ref, due_date, notes, deleted_at,
                                                         created_at, created_by)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    entry.get("id", str(uuid.uuid4())),
                    entry.get("customer_id", ""),
                    entry.get("entry_type", ""),
                    entry.get("entry_date", str(date.today())),
                    entry.get("ref_number", ""),
                    entry.get("description", ""),
                    float(entry.get("amount_inr", 0.0)),
                    float(entry.get("amount_usd", 0.0)),
                    float(entry.get("usd_rate", 84.0)),
                    entry.get("dr_cr", "DR"),
                    entry.get("payment_mode", ""),
                    entry.get("bank_ref", ""),
                    entry.get("due_date"),
                    entry.get("notes", ""),
                    entry.get("deleted_at"),
                    entry.get("created_at", datetime.now().isoformat()),
                    entry.get("created_by", "User"),
                ))
                imported_stats["customer_ledger_entries"] += 1

            # ── IMPORT PURCHASE ORDERS ────────────────────────────────────
            for po in bundle.get("purchase_orders", []):
                po_id = po.get("id", str(uuid.uuid4()))
                
                cursor.execute("""
                    INSERT INTO purchase_orders (
                        id, po_number, po_date, supplier_id, supplier_snapshot,
                        payment_conditions, delivery_terms, delivery_address,
                        forwarder_id, forwarder_name, forwarder_contact, remarks,
                        currency, status, created_at, lead_time_days, due_date, godown_id,
                        lc_usd_rate, lc_rmb_rate, lc_bank, lc_ship, lc_duty,
                        lc_trans, lc_gst_duty, lc_doc_pct
                    )
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    po_id,
                    po.get("po_number", ""),
                    po.get("po_date", str(date.today())),
                    po.get("supplier_id", ""),
                    json.dumps(po.get("supplier_snapshot", {})),
                    po.get("payment_conditions", ""),
                    po.get("delivery_terms", "FOB"),
                    po.get("delivery_address", ""),
                    po.get("forwarder_id", ""),
                    po.get("forwarder_name", ""),
                    po.get("forwarder_contact", ""),
                    po.get("remarks", ""),
                    po.get("currency", "USD"),
                    po.get("status", "Draft"),
                    po.get("created_at", str(date.today())),
                    int(po.get("lead_time_days", 0)),
                    po.get("due_date", ""),
                    po.get("godown_id", ""),          # ← was missing; caused godown to be lost on import
                    float(po.get("lc_usd_rate", 84.0)),
                    float(po.get("lc_rmb_rate", 11.5)),
                    float(po.get("lc_bank", 0.0)),
                    float(po.get("lc_ship", 0.0)),
                    float(po.get("lc_duty", 0.0)),
                    float(po.get("lc_trans", 0.0)),
                    float(po.get("lc_gst_duty", 0.0)),
                    float(po.get("lc_doc_pct", 0.0)),
                ))
                imported_stats["purchase_orders"] += 1
                
                # ── IMPORT LINE ITEMS ──────────────────────────────────────
                for seq, li in enumerate(po.get("line_items", [])):
                    cursor.execute("""
                        INSERT INTO po_items (id, po_id, item_id, item_name, description, hs_code, qty, unit, unit_price, line_sequence)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """, (
                        str(uuid.uuid4()),
                        po_id,
                        li.get("item_id", ""),
                        li.get("item_name", ""),
                        li.get("description", ""),
                        li.get("hs_code", ""),
                        float(li.get("qty", 0)),
                        li.get("unit", "PCS"),
                        float(li.get("unit_price", 0)),
                        seq,
                    ))
                    imported_stats["po_items"] += 1

            # ── IMPORT SUPPLIER LEDGER ENTRIES ─────────────────────────────
            for entry in bundle.get("supplier_ledger_entries", []):
                cursor.execute("""
                    INSERT INTO supplier_ledger_entries (id, supplier_id, po_id, entry_type, entry_date, ref_number,
                                                         description, amount_usd, amount_inr, usd_rate, dr_cr,
                                                         payment_mode, bank_ref, attachment_id, notes, deleted_at,
                                                         created_at, created_by)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    entry.get("id", str(uuid.uuid4())),
                    entry.get("supplier_id", ""),
                    entry.get("po_id"),
                    entry.get("entry_type", ""),
                    entry.get("entry_date", str(date.today())),
                    entry.get("ref_number", ""),
                    entry.get("description", ""),
                    float(entry.get("amount_usd", 0.0)),
                    float(entry.get("amount_inr", 0.0)),
                    float(entry.get("usd_rate", 84.0)),
                    entry.get("dr_cr", "CR"),
                    entry.get("payment_mode", ""),
                    entry.get("bank_ref", ""),
                    entry.get("attachment_id"),
                    entry.get("notes", ""),
                    entry.get("deleted_at"),
                    entry.get("created_at", datetime.now().isoformat()),
                    entry.get("created_by", "User"),
                ))
                imported_stats["supplier_ledger_entries"] += 1

            # ── IMPORT QUOTATIONS ─────────────────────────────────────────
            for quotation in bundle.get("quotations", []):
                q_id = quotation.get("id", str(uuid.uuid4()))
                q_number = quotation.get("quotation_number", quotation.get("quotation_no", ""))
                q_title = quotation.get("title", "")
                q_currency = quotation.get("currency", "CNY")
                q_status = quotation.get("status", "Open")
                q_awarded = quotation.get("awarded_to", "")
                q_custref = quotation.get("customer_ref", "")
                q_notes = quotation.get("notes", "")
                q_date = quotation.get("date", "")
                q_details = json.dumps(quotation.get("details", quotation))
                q_created = quotation.get("created_at", datetime.now().isoformat())
                q_updated = quotation.get("updated_at", datetime.now().isoformat())
                cursor.execute("""
                    INSERT INTO quotations (id, quotation_number, title, currency, status,
                                            awarded_to, customer_ref, notes, date, details,
                                            created_at, updated_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    q_id, q_number, q_title, q_currency, q_status,
                    q_awarded, q_custref, q_notes, q_date, q_details,
                    q_created, q_updated,
                ))
                _sync_quotation_refs(cursor, q_id, q_details)
                _sync_quotation_line_items(cursor, q_id, q_details)
                imported_stats["quotations"] += 1

            # ── IMPORT SHIPMENTS ──────────────────────────────────────────
            for ship in bundle.get("shipments", []):
                ship_id = ship.get("id", str(uuid.uuid4()))
                cursor.execute("""
                    INSERT INTO shipments (
                        id, forwarder_id, booking_ref, departure_date, expected_arrival,
                        actual_arrival, status, description, notes, created_at, updated_at, deleted_at
                    )
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    ship_id,
                    ship.get("forwarder_id", ""),
                    ship.get("booking_ref", ""),
                    ship.get("departure_date", ""),
                    ship.get("expected_arrival", ""),
                    ship.get("actual_arrival"),
                    ship.get("status", "With Forwarder"),
                    ship.get("description", ""),
                    json.dumps(ship.get("notes", [])),
                    ship.get("created_at", datetime.now().isoformat()),
                    ship.get("updated_at", datetime.now().isoformat()),
                    ship.get("deleted_at")
                ))
                
                # Link POs
                for po_id in ship.get("po_ids", []):
                    cursor.execute("""
                        INSERT OR IGNORE INTO shipment_po_link (id, shipment_id, po_id)
                        VALUES (?, ?, ?)
                    """, (str(uuid.uuid4()), ship_id, po_id))
                
                imported_stats["shipments"] += 1
            
            # ── IMPORT SETTINGS ───────────────────────────────────────────
            if "settings" in bundle and isinstance(bundle["settings"], dict):
                init_default_settings(conn)
                for key, value in bundle["settings"].items():
                    cursor.execute(
                        "INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)",
                        (key, str(value))
                    )
            
            # ✅ COMMIT TRANSACTION (only if all succeeded)
            
            return jsonify({
                "ok": True,
                "imported": imported_stats,
                "message": "✅ Import completed successfully with transaction backup"
            }), 200
            
        except Exception as e:
                # ❌ ROLLBACK on error (automatic with context manager)
                conn.rollback()
                return jsonify({"error": f"Import failed - rolled back: {str(e)}"}), 400


@app.route("/api/reset", methods=["POST"])
@require_permission("admin_rbac")
def reset_data():
    """Wipe all app data and restore default settings."""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("DELETE FROM shipment_po_link")
        cursor.execute("DELETE FROM shipments")
        cursor.execute("DELETE FROM po_status_log")
        cursor.execute("DELETE FROM po_attachments")
        cursor.execute("DELETE FROM po_payments")
        cursor.execute("DELETE FROM supplier_ledger_entries")
        cursor.execute("DELETE FROM customer_ledger_entries")
        cursor.execute("DELETE FROM customers")
        cursor.execute("DELETE FROM forwarder_godowns")
        cursor.execute("DELETE FROM po_items")
        cursor.execute("DELETE FROM purchase_orders")
        cursor.execute("DELETE FROM quotation_refs")
        cursor.execute("DELETE FROM quotation_supplier_rows")
        cursor.execute("DELETE FROM quotation_line_items")
        cursor.execute("DELETE FROM quotations")
        cursor.execute("DELETE FROM forwarders")
        cursor.execute("DELETE FROM items")
        cursor.execute("DELETE FROM suppliers")
        cursor.execute("DELETE FROM settings")
        init_default_settings(conn)

    # Clear uploaded attachment files too
    if os.path.isdir(ATTACH_DIR):
        for name in os.listdir(ATTACH_DIR):
            path = os.path.join(ATTACH_DIR, name)
            if os.path.isdir(path):
                shutil.rmtree(path, ignore_errors=True)
            else:
                try:
                    os.remove(path)
                except OSError:
                    pass

    return jsonify({"ok": True, "message": "All app data wiped and default settings restored."}), 200

# ═══════════════════════════════════════════════════════════════════════════════
# QUOTATION ANALYZER API
# ═══════════════════════════════════════════════════════════════════════════════

def _sync_quotation_refs(cursor, quotation_id, details_json):
    """Rebuild quotation_refs rows for *quotation_id* from the JSON payload.

    Called inside an open transaction whenever a quotation is created or
    updated.  The DELETE + INSERT approach is a simple upsert that handles
    adds, removals, and changes in one pass.
    """
    cursor.execute("DELETE FROM quotation_refs WHERE quotation_id = ?", (quotation_id,))
    try:
        details = json.loads(details_json) if isinstance(details_json, str) else (details_json or {})
    except Exception:
        details = {}
    seen = set()
    for li in details.get("line_items", []):
        item_id = li.get("item_id", "")
        if item_id:
            key = ("item", item_id)
            if key not in seen:
                cursor.execute(
                    "INSERT OR IGNORE INTO quotation_refs (quotation_id, ref_type, ref_id) VALUES (?, ?, ?)",
                    (quotation_id, "item", item_id)
                )
                seen.add(key)
        for sr in li.get("supplier_rows", []):
            sup_id = sr.get("supplier_id", "")
            if sup_id:
                key = ("supplier", sup_id)
                if key not in seen:
                    cursor.execute(
                        "INSERT OR IGNORE INTO quotation_refs (quotation_id, ref_type, ref_id) VALUES (?, ?, ?)",
                        (quotation_id, "supplier", sup_id)
                    )
                    seen.add(key)



def _sync_quotation_line_items(cursor, quotation_id, details_json):
    """Rebuild normalised quotation_line_items and quotation_supplier_rows.

    Called inside an open transaction on every quotation write so the
    structured tables stay in sync with the details JSON blob.
    """
    cursor.execute("DELETE FROM quotation_supplier_rows WHERE quotation_id = ?", (quotation_id,))
    cursor.execute("DELETE FROM quotation_line_items WHERE quotation_id = ?", (quotation_id,))
    try:
        details = json.loads(details_json) if isinstance(details_json, str) else (details_json or {})
    except Exception:
        details = {}

    for seq, li in enumerate(details.get("line_items", [])):
        li_id = str(uuid.uuid4())
        cursor.execute(
            "INSERT INTO quotation_line_items "
            "(id, quotation_id, item_id, item_name, qty, unit, description, "
            " selected_supplier_id, selected_supplier_name, line_sequence) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (
                li_id, quotation_id,
                li.get("item_id") or None,
                li.get("item_name", ""),
                float(li.get("qty", 1)),
                li.get("unit", "PCS"),
                li.get("description", ""),
                li.get("selected") or None,
                li.get("selected_supplier") or None,
                seq,
            )
        )
        for sr in li.get("supplier_rows", []):
            cursor.execute(
                "INSERT INTO quotation_supplier_rows "
                "(id, line_item_id, quotation_id, supplier_id, supplier_name, price) "
                "VALUES (?, ?, ?, ?, ?, ?)",
                (
                    str(uuid.uuid4()), li_id, quotation_id,
                    sr.get("supplier_id") or None,
                    sr.get("supplier_name", ""),
                    float(sr.get("price", 0)),
                )
            )


def _quotation_to_dict(row):
    """Convert a quotation DB row to a dict, merging stored details"""
    q = dict(row)
    details = {}
    if q.get("details"):
        try:
            details = json.loads(q["details"])
        except Exception:
            pass
    # Merge details into top-level, but DB columns take priority for indexed fields
    merged = {**details, **{k: v for k, v in q.items() if v is not None and k != "details"}}
    # Ensure quotation_no is available (alias)
    if "quotation_no" not in merged and "quotation_number" in merged:
        merged["quotation_no"] = merged["quotation_number"]
    return merged


@app.route("/api/quotations", methods=["GET"])
def list_quotations():
    """Get all quotations"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM quotations ORDER BY created_at DESC")
        return jsonify([_quotation_to_dict(r) for r in cursor.fetchall()])


@app.route("/api/quotations", methods=["POST"])
@require_permission("po_edit")
def create_quotation():
    """Create a new quotation"""
    data = request.get_json(force=True)
    qid = str(uuid.uuid4())
    q_number = data.get("quotation_no", "")
    title = data.get("title", "")
    currency = data.get("currency", "CNY")
    customer_ref = data.get("customer_ref", "")
    notes = data.get("notes", "")
    q_date = data.get("date", str(date.today()))
    now = datetime.now().isoformat()

    # Store full payload as JSON details
    details = json.dumps(data)

    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO quotations (id, quotation_number, title, currency, status,
                                    customer_ref, notes, date, details, created_at, updated_at)
            VALUES (?, ?, ?, ?, 'Open', ?, ?, ?, ?, ?, ?)
        """, (qid, q_number, title, currency, customer_ref, notes, q_date, details, now, now))
        _sync_quotation_refs(cursor, qid, details)
        _sync_quotation_line_items(cursor, qid, details)
        cursor.execute("SELECT * FROM quotations WHERE id = ?", (qid,))
        return jsonify(_quotation_to_dict(cursor.fetchone())), 201


@app.route("/api/quotations/<qid>", methods=["GET"])
def get_quotation(qid):
    """Get a single quotation by ID"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM quotations WHERE id = ?", (qid,))
        row = cursor.fetchone()
        if not row:
            return jsonify({"error": "Quotation not found"}), 404
        return jsonify(_quotation_to_dict(row))


@app.route("/api/quotations/<qid>", methods=["PUT"])
@require_permission("po_edit")
def update_quotation(qid):
    """Update an existing quotation"""
    data = request.get_json(force=True)
    now = datetime.now().isoformat()

    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM quotations WHERE id = ?", (qid,))
        existing = cursor.fetchone()
        if not existing:
            return jsonify({"error": "Quotation not found"}), 404

        q_number = data.get("quotation_no", existing["quotation_number"])
        title = data.get("title", existing["title"])
        currency = data.get("currency", existing["currency"])
        customer_ref = data.get("customer_ref", existing["customer_ref"] or "")
        notes = data.get("notes", existing["notes"] or "")
        q_date = data.get("date", existing["date"])

        # Preserve status and awarded_to from existing unless explicitly provided
        status = data.get("status", existing["status"])
        awarded_to = data.get("awarded_to", existing["awarded_to"])

        details = json.dumps(data)

        cursor.execute("""
            UPDATE quotations SET quotation_number=?, title=?, currency=?, status=?,
                awarded_to=?, customer_ref=?, notes=?, date=?, details=?, updated_at=?
            WHERE id=?
        """, (q_number, title, currency, status, awarded_to, customer_ref,
              notes, q_date, details, now, qid))
        _sync_quotation_refs(cursor, qid, details)
        _sync_quotation_line_items(cursor, qid, details)
        cursor.execute("SELECT * FROM quotations WHERE id = ?", (qid,))
        return jsonify(_quotation_to_dict(cursor.fetchone()))


@app.route("/api/quotations/<qid>", methods=["DELETE"])
@require_permission("po_delete")
def delete_quotation(qid):
    """Delete a quotation"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("DELETE FROM quotations WHERE id = ?", (qid,))
    return jsonify({"ok": True})


@app.route("/api/quotations/<qid>/award", methods=["POST"])
@require_permission("po_edit")
def award_quotation(qid):
    """Award a quotation to a supplier"""
    data = request.get_json(force=True)
    supplier_id = data.get("supplier_id", "")
    now = datetime.now().isoformat()

    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM quotations WHERE id = ?", (qid,))
        existing = cursor.fetchone()
        if not existing:
            return jsonify({"error": "Quotation not found"}), 404

        # Update details JSON with awarded info
        details = {}
        try:
            details = json.loads(existing["details"] or "{}")
        except Exception:
            pass
        details["status"] = "Awarded"
        details["awarded_to"] = supplier_id

        cursor.execute("""
            UPDATE quotations SET status='Awarded', awarded_to=?, details=?, updated_at=?
            WHERE id=?
        """, (supplier_id, json.dumps(details), now, qid))

        cursor.execute("SELECT * FROM quotations WHERE id = ?", (qid,))
        q = _quotation_to_dict(cursor.fetchone())

        # Build po_seed for auto-populating PO form
        line_items_data = q.get("line_items", [])
        po_seed = {
            "supplier_id": supplier_id,
            "currency": q.get("currency", "CNY"),
            "line_items": line_items_data,
        }

        return jsonify({**q, "po_seed": po_seed})


@app.route("/api/quotations/<qid>/reopen", methods=["POST"])
@require_permission("po_edit")
def reopen_quotation(qid):
    """Reopen an awarded quotation"""
    now = datetime.now().isoformat()
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM quotations WHERE id = ?", (qid,))
        existing = cursor.fetchone()
        if not existing:
            return jsonify({"error": "Quotation not found"}), 404

        details = {}
        try:
            details = json.loads(existing["details"] or "{}")
        except Exception:
            pass
        details["status"] = "Open"
        details["awarded_to"] = None

        cursor.execute("""
            UPDATE quotations SET status='Open', awarded_to=NULL, details=?, updated_at=?
            WHERE id=?
        """, (json.dumps(details), now, qid))
        cursor.execute("SELECT * FROM quotations WHERE id = ?", (qid,))
        return jsonify(_quotation_to_dict(cursor.fetchone()))


@app.route("/api/quotations/<qid>/duplicate", methods=["POST"])
@require_permission("po_edit")
def duplicate_quotation(qid):
    """Duplicate a quotation with cleared prices"""
    now = datetime.now().isoformat()
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM quotations WHERE id = ?", (qid,))
        existing = cursor.fetchone()
        if not existing:
            return jsonify({"error": "Quotation not found"}), 404

        details = {}
        try:
            details = json.loads(existing["details"] or "{}")
        except Exception:
            pass

        # Generate new ID and number
        new_id = str(uuid.uuid4())
        old_num = existing["quotation_number"] or ""
        new_num = old_num + "-COPY" if old_num else f"QT-{date.today().year}-COPY"

        # Clear prices in line items
        for li in details.get("line_items", []):
            for sr in li.get("supplier_rows", []):
                sr["price"] = 0
            li["selected"] = ""
            li["selected_supplier"] = ""

        # Reset status
        details["status"] = "Open"
        details["awarded_to"] = None
        details["quotation_no"] = new_num

        new_details_json = json.dumps(details)
        cursor.execute("""
            INSERT INTO quotations (id, quotation_number, title, currency, status,
                                    customer_ref, notes, date, details, created_at, updated_at)
            VALUES (?, ?, ?, ?, 'Open', ?, ?, ?, ?, ?, ?)
        """, (new_id, new_num, details.get("title", existing["title"]),
              existing["currency"], existing["customer_ref"] or "",
              existing["notes"] or "", str(date.today()), new_details_json, now, now))
        _sync_quotation_refs(cursor, new_id, new_details_json)
        _sync_quotation_line_items(cursor, new_id, new_details_json)

        cursor.execute("SELECT * FROM quotations WHERE id = ?", (new_id,))
        return jsonify(_quotation_to_dict(cursor.fetchone())), 201


def _safe_float(val, default=0.0):
    try:
        if val is None or val == "": return default
        return float(val)
    except (ValueError, TypeError):
        return default

def _safe_int(val, default=0):
    try:
        if val is None or val == "": return default
        return int(float(val))
    except (ValueError, TypeError):
        return default

@app.route("/api/po/<pid>/print", methods=["GET"])
def print_po(pid):
    """Generate a printable HTML view of a Purchase Order"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM purchase_orders WHERE id = ?", (pid,))
        po_row = cursor.fetchone()
        if not po_row:
            return "PO not found", 404
        po = dict(po_row)
        
        # Fetch supplier
        cursor.execute("SELECT * FROM suppliers WHERE id = ?", (po["supplier_id"],))
        sup_row = cursor.fetchone()
        sup = dict(sup_row) if sup_row else {}
        
        # Fetch line items
        cursor.execute("SELECT * FROM po_items WHERE po_id = ? ORDER BY line_sequence", (pid,))
        line_items = [dict(item) for item in cursor.fetchall()]
        
        # Fetch settings
        cursor.execute("SELECT key, value FROM settings")
        settings = {row["key"]: row["value"] for row in cursor.fetchall()}

    total = sum(_safe_float(li.get("qty", 0)) * _safe_float(li.get("unit_price", 0)) for li in line_items)
    total_qty = sum(_safe_float(li.get("qty", 0)) for li in line_items)
    cur = po.get("currency", "USD")
    fmt_total = f"{total:,.2f}"

    item_rows = ""
    for i, li in enumerate(line_items):
        amt = _safe_float(li.get("qty", 0)) * _safe_float(li.get("unit_price", 0))
        item_rows += f"""<tr>
          <td style="text-align:center;color:#374151">{i+1}</td>
          <td style="font-weight:600;color:#111;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;max-width:0">{_e(li.get('item_name',''))}</td>
          <td style="color:#4b5563;font-size:11px;white-space:nowrap">{_e(li.get('description','—'))}</td>
          <td style="text-align:center;color:#4b5563">{_e(li.get('hs_code','—'))}</td>
          <td style="text-align:center;font-weight:600">{li.get('qty',0)}</td>
          <td style="text-align:center;color:#4b5563">{_e(li.get('unit','PCS'))}</td>
          <td style="text-align:right">{_e(cur)} {li.get('unit_price',0):,.2f}</td>
          <td style="text-align:right;font-weight:700;color:#1e3a8a">{_e(cur)} {amt:,.2f}</td>
        </tr>"""

    buyer_lines = []
    if settings.get("company_name"): buyer_lines.append(f'<div style="font-size:16px;font-weight:700;color:#111;margin-bottom:3px">{_e(settings["company_name"])}</div>')
    if settings.get("company_address"):
        addr_html = _e(settings["company_address"]).replace("\n", "<br>")
        buyer_lines.append(f'<div style="color:#4b5563">{addr_html}</div>')
    if settings.get("company_phone"): buyer_lines.append(f'<div style="color:#4b5563">Ph: {_e(settings["company_phone"])}</div>')
    if settings.get("company_email"): buyer_lines.append(f'<div style="color:#4b5563">Email: {_e(settings["company_email"])}</div>')
    if settings.get("company_gstin"): buyer_lines.append(f'<div style="color:#4b5563">GSTIN: {_e(settings["company_gstin"])}</div>')
    buyer_html = "".join(buyer_lines) or '<div style="color:#9ca3af">Not configured.</div>'

    sup_lines = []
    sup_lines.append(f'<div style="font-size:14px;font-weight:700;color:#111;margin-bottom:3px">{_e(sup.get("company","—"))}</div>')
    if sup.get("name"): sup_lines.append(f'<div style="color:#4b5563">Attn: {_e(sup["name"])}</div>')
    if sup.get("address"): sup_lines.append(f'<div style="color:#4b5563">{_e(sup["address"])}</div>')
    if sup.get("country"): sup_lines.append(f'<div style="color:#4b5563">{_e(sup["country"])}</div>')
    if sup.get("email"): sup_lines.append(f'<div style="color:#4b5563">Email: {_e(sup["email"])}</div>')
    if sup.get("phone"): sup_lines.append(f'<div style="color:#4b5563">Ph/WA: {_e(sup["phone"])}</div>')
    if sup.get("wechat"): sup_lines.append(f'<div style="color:#4b5563">WeChat: {_e(sup["wechat"])}</div>')
    supplier_html = "".join(sup_lines)

    bank_section = ""
    if sup.get("bank_name") or sup.get("bank_account"):
        swift = f' &nbsp;|&nbsp; SWIFT: {_e(sup["swift_code"])}' if sup.get("swift_code") else ""
        acc = f' &nbsp;|&nbsp; A/C: {_e(sup["bank_account"])}' if sup.get("bank_account") else ""
        bank_section = f"""
        <div style="margin-top:10px;padding:10px 16px;background:#f8faff;border:1px solid #dbeafe;border-radius:6px;font-size:11px">
          <div style="font-weight:700;text-transform:uppercase;letter-spacing:.05em;font-size:9px;color:#1e3a8a;margin-bottom:4px">Bank Details (Supplier) · 银行详情</div>
          <div style="color:#1e40af">{_e(sup.get('bank_name',''))}{acc}{swift}</div>
        </div>"""

    status_color = {"Draft":"#6b7280","Sent":"#2563eb","Confirmed":"#16a34a","Shipped":"#7c3aed","In Transit":"#f97316","Received":"#059669","Cancelled":"#dc2626"}.get(po.get("status"), "#6b7280")

    # Payment Breakdown
    pay_breakdown_html = ""
    pc = po.get("payment_conditions") or ""
    pay_match = re.search(r"(\d+)%\s*Advance.*?(\d+)%", pc, re.I)
    if pay_match and total > 0:
        pct1 = int(pay_match.group(1))
        pct2 = int(pay_match.group(2))
        adv = total * pct1 / 100.0
        rem = total * pct2 / 100.0
        lbl2 = "against B/L (Bill of Lading)" if "bl" in pc.lower() else "on Shipment"
        pay_breakdown_html = f"""
          <div style="padding:7px 12px;background:#eff6ff;border:1px solid #bfdbfe;border-radius:5px;font-size:10px;line-height:1.8">
            <div style="font-size:8px;font-weight:700;letter-spacing:.1em;text-transform:uppercase;color:#1d4ed8;margin-bottom:3px">Payment Schedule · 付款计划</div>
            <div style="display:flex;justify-content:space-between;gap:8px;flex-wrap:nowrap"><span style="color:#374151;white-space:nowrap">① {pct1}% Advance (Proforma Invoice)</span><span style="font-weight:700;color:#1d4ed8;white-space:nowrap">{_e(cur)} {adv:,.2f}</span></div>
            <div style="display:flex;justify-content:space-between;gap:8px;flex-wrap:nowrap"><span style="color:#374151;white-space:nowrap">② {pct2}% {lbl2}</span><span style="font-weight:700;color:#1d4ed8;white-space:nowrap">{_e(cur)} {rem:,.2f}</span></div>
          </div>"""

    lead_time_html = ""
    lt = po.get("lead_time_days")
    dd = po.get("due_date")
    if lt or dd:
        lt_str = f"{lt} Days" if lt else "—"
        dd_str = f'<div style="font-size:11px;color:#1d4ed8;margin-top:1px;font-weight:600">Due: {_e(dd)}</div>' if dd else ""
        lead_time_html = f"""
        <div>
          <div style="font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#9ca3af;margin-bottom:3px">Lead Time · 交货期</div>
          <div style="font-weight:600;color:#111;font-size:14px">{_e(lt_str)}</div>
          {dd_str}
        </div>"""

    delivery_logistics = ""
    if po.get("forwarder_name") or po.get("forwarder_contact") or po.get("delivery_address"):
        fwd_name = po.get("forwarder_name", "—")
        fwd_contact = f' <span style="font-size:11px;color:#6b7280;font-weight:400">{_e(po["forwarder_contact"])}</span>' if po.get("forwarder_contact") else ""
        fwd_html = f"""<div style="flex:0 0 auto">
          <div style="font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#9ca3af;margin-bottom:3px">Freight Forwarder · 货运代理</div>
          <div style="display:inline;font-weight:600;color:#111">{_e(fwd_name)}</div>{fwd_contact}
        </div>""" if po.get("forwarder_name") or po.get("forwarder_contact") else ""
        
        addr_html = _e(po.get('delivery_address','')).replace("\n", "<br>")
        addr_html = f"""<div style="flex:1">
          <div style="font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#9ca3af;margin-bottom:3px">📍 Delivery Address · 交货地址</div>
          <div style="font-weight:600;color:#111;font-size:12px;line-height:1.6;background:#fff;border:1px solid #fed7aa;border-radius:5px;padding:8px 12px;margin-top:4px">{addr_html}</div>
        </div>""" if po.get("delivery_address") else ""
        
        delivery_logistics = f"""
        <div style="display:flex;gap:16px;align-items:flex-start;margin-top:6px;padding-top:6px;border-top:1px solid #fed7aa">
          {fwd_html}
          {addr_html}
        </div>"""

    remarks_html = ""
    if po.get("remarks"):
        remarks_html = f'<div style="margin-top:8px;padding:6px 10px;background:#fffbeb;border:1px solid #fde68a;border-left:4px solid #f59e0b;border-radius:0 6px 6px 0;font-size:11px;color:#92400e"><strong style="font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#78350f;margin-right:4px">REMARKS / SPECIAL INSTRUCTIONS · 备注:</strong>{_e(po["remarks"])}</div>'

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>PO {_e(po.get('po_number',''))}</title>
<link href="https://fonts.googleapis.com/css2?family=Noto+Sans+SC:wght@400;600;700&display=swap" rel="stylesheet">
<style>
  *{{box-sizing:border-box;margin:0;padding:0}}
  body{{font-family:'Noto Sans SC','Segoe UI',Arial,sans-serif;font-size:12px;color:#111;background:#fff;padding:32px 36px;max-width:920px;margin:0 auto}}
  .po-header{{display:flex;justify-content:space-between;align-items:flex-start;padding-bottom:12px;border-bottom:3px solid #1e3a8a;margin-bottom:14px}}
  .po-main-title{{font-size:28px;font-weight:700;color:#1e3a8a;letter-spacing:-.02em;line-height:1}}
  .po-subtitle{{font-size:11px;color:#6b7280;margin-top:3px;letter-spacing:.06em;text-transform:uppercase}}
  .po-header-right{{text-align:right;font-size:12px}}
  .po-meta-row{{display:flex;gap:6px;align-items:center;justify-content:flex-end;margin-bottom:4px;color:#374151}}
  .po-meta-row strong{{color:#111;min-width:60px;text-align:right}}
  .status-pill{{display:inline-block;padding:3px 12px;border-radius:20px;font-size:10px;font-weight:700;letter-spacing:.07em;text-transform:uppercase;color:#fff;background:{status_color};margin-top:8px}}
  .parties-grid{{display:grid;grid-template-columns:1fr 1fr;gap:0;border:1px solid #e5e7eb;border-radius:8px;overflow:hidden;margin-bottom:10px}}
  .party-box{{padding:10px 14px}}
  .party-box:first-child{{border-right:1px solid #e5e7eb;background:#f9fafb}}
  .party-box:last-child{{background:#f0f6ff}}
  .party-label{{font-size:9px;font-weight:700;letter-spacing:.1em;text-transform:uppercase;color:#6b7280;margin-bottom:8px;display:flex;align-items:center;gap:5px}}
  .party-label::before{{content:'';display:block;width:10px;height:2px;background:#1e3a8a;border-radius:1px}}
  .delivery-box{{background:#fff8f0;border:1.5px solid #f97316;border-radius:8px;padding:8px 14px;margin-bottom:10px}}
  .delivery-box-title{{font-size:9px;font-weight:700;letter-spacing:.1em;text-transform:uppercase;color:#c2410c;margin-bottom:6px;display:flex;align-items:center;gap:6px}}
  .delivery-box-title::before{{content:'📦';font-size:13px}}
  .delivery-grid{{display:grid;grid-template-columns:1fr 1fr 1.4fr 1fr;gap:8px 14px;font-size:11px}}
  table{{width:100%;border-collapse:collapse;margin-bottom:0}}
  thead tr{{background:#1e3a8a}}
  thead th{{padding:7px 10px;color:#fff;font-size:10px;font-weight:600;letter-spacing:.07em;text-transform:uppercase;text-align:left}}
  tbody tr:nth-child(even){{background:#f8faff}}
  tbody td{{padding:6px 10px;border-bottom:1px solid #e5e7eb;font-size:11px;vertical-align:top}}
  tfoot tr{{background:#1e3a8a}}
  tfoot td{{padding:7px 10px;color:#fff;font-weight:700;font-size:13px}}
  .table-wrap{{border:1px solid #d1d5db;border-radius:8px;overflow:hidden;margin-bottom:8px}}
  .sig-grid{{display:grid;grid-template-columns:1fr 1fr;gap:48px;margin-top:16px}}
  .sig-box{{border-top:2px solid #d1d5db;padding-top:8px}}
  .sig-label{{font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#9ca3af;margin-bottom:4px}}
  .sig-name{{font-size:12px;font-weight:700;color:#111}}
  .sig-sub{{font-size:10px;color:#6b7280;margin-top:2px}}
  .footer-note{{margin-top:12px;padding-top:8px;border-top:1px solid #e5e7eb;font-size:9px;color:#9ca3af;text-align:center;letter-spacing:.04em}}
  @media print{{
    body{{padding:10px 14px}}
    @page{{margin:6mm 8mm;size:A4}}
  }}
</style>
</head>
<body>

<div class="po-header">
  <div class="po-header-left">
    <div class="po-main-title">PURCHASE ORDER</div>
    <div class="po-subtitle">采购订单</div>
  </div>
  <div class="po-header-right">
    <div class="po-meta-row"><strong>PO No:</strong>&nbsp;{_e(po.get('po_number',''))}</div>
    <div class="po-meta-row"><strong>Date:</strong>&nbsp;{_e(po.get('po_date',''))}</div>
    <div class="po-meta-row"><strong>Currency:</strong>&nbsp;{_e(cur)}</div>
    <div><span class="status-pill">{_e(po.get('status',''))}</span></div>
  </div>
</div>

<div class="parties-grid">
  <div class="party-box">
    <div class="party-label">Buyer / Importer · 买方</div>
    {buyer_html}
  </div>
  <div class="party-box">
    <div class="party-label">Supplier / Vendor · 供应商</div>
    {supplier_html}
  </div>
</div>

<div class="delivery-box">
  <div class="delivery-box-title">Delivery &amp; Logistics Details · 交货与物流详情</div>
  <div class="delivery-grid">
    <div>
      <div style="font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#9ca3af;margin-bottom:3px">Delivery Terms · 交货条款</div>
      <div style="font-weight:600;color:#111;font-size:14px">{_e(po.get('delivery_terms','—'))}</div>
    </div>
    {lead_time_html}
    <div>
      <div style="font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#9ca3af;margin-bottom:3px">Payment Conditions · 付款条件</div>
      <div style="font-weight:600;color:#111;font-size:13px">{_e(pc or '—')}</div>
    </div>
    <div>{pay_breakdown_html}</div>
  </div>
  {delivery_logistics}
  {remarks_html}
</div>

<div class="table-wrap">
  <table>
    <thead><tr>
      <th style="width:3%;text-align:center">#</th>
      <th style="width:30%">Item Name<br><span style="font-weight:400;opacity:.75;font-size:9px">品名</span></th>
      <th style="width:10%">Description<br><span style="font-weight:400;opacity:.75;font-size:9px">描述</span></th>
      <th style="width:10%;text-align:center">HS Code</th>
      <th style="width:7%;text-align:center">Qty<br><span style="font-weight:400;opacity:.75;font-size:9px">数量</span></th>
      <th style="width:7%;text-align:center">Unit<br><span style="font-weight:400;opacity:.75;font-size:9px">单位</span></th>
      <th style="width:15%;text-align:right">Unit Price<br><span style="font-weight:400;opacity:.75;font-size:9px">单价</span></th>
      <th style="width:15%;text-align:right">Amount<br><span style="font-weight:400;opacity:.75;font-size:9px">金额</span></th>
    </tr></thead>
    <tbody>{item_rows}</tbody>
    <tfoot>
      <tr>
        <td colspan="4" style="text-align:right;font-size:11px;text-transform:uppercase;padding-right:14px;color:#e0e7ff">Total Quantity Required</td>
        <td style="text-align:center;font-size:13px;color:#fff">{total_qty:g}</td>
        <td colspan="2" style="text-align:right;font-size:11px;text-transform:uppercase;padding-right:14px;color:#e0e7ff">Grand Total · 总金额</td>
        <td style="text-align:right;font-size:14px">{_e(cur)} {fmt_total}</td>
      </tr>
    </tfoot>
  </table>
</div>
{bank_section}

<div class="sig-grid">
  <div class="sig-box">
    <div class="sig-label">Authorised Signatory — Buyer · 买方授权签字</div>
    <div class="sig-name">{_e(settings.get('company_name','_______________________'))}</div>
    <div class="sig-sub">Signature &amp; Stamp · 签名盖章</div>
  </div>
  <div class="sig-box">
    <div class="sig-label">Accepted — Supplier · 供应商确认</div>
    <div class="sig-name">{_e(sup.get('company','_______________________'))}</div>
    <div class="sig-sub">Signature &amp; Stamp · 签名盖章</div>
  </div>
</div>

<div class="footer-note">This is a computer-generated Purchase Order · {date.today().strftime('%d %b %Y')}</div>
</body></html>"""
    return html, 200, {"Content-Type": "text/html; charset=utf-8"}


@app.route("/api/quotations/<qid>/print", methods=["GET"])
def print_quotation(qid):
    """Generate a printable HTML view of a quotation matching the client-side design"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM quotations WHERE id = ?", (qid,))
        row = cursor.fetchone()
        if not row:
            return "Quotation not found", 404
        cursor.execute("SELECT key, value FROM settings")
        settings = {r["key"]: r["value"] for r in cursor.fetchall()}

    q = _quotation_to_dict(row)
    cur = q.get("currency", "CNY")
    cur_sym = "¥" if cur == "CNY" else "$"
    line_items = q.get("line_items", [])

    # Get all unique supplier names
    sup_names = []
    for li in line_items:
        for sr in li.get("supplier_rows", []):
            nm = sr.get("supplier_name", "")
            if nm and nm not in sup_names:
                sup_names.append(nm)

    sup_headers = "".join(f'<th class="th-sup">{_e(nm)}</th>' for nm in sup_names)

    item_rows = ""
    for i, li in enumerate(line_items):
        best_price = min((sr.get('price', 0) for sr in li.get('supplier_rows', []) if sr.get('price', 0) > 0), default=0)
        
        sup_cells = ""
        for nm in sup_names:
            sr_match = next((sr for sr in li.get("supplier_rows",[]) if sr.get("supplier_name") == nm), None)
            p = sr_match["price"] if sr_match else 0
            is_best = p > 0 and p == best_price
            style = ' style="background:#f0fdf4;color:#15803d;font-weight:700"' if is_best else ''
            sup_cells += f'<td class="td-sup"{style}>{cur_sym}{p:,.2f}</td>'

        best_cell = f'<td class="best-cell">{cur_sym}{best_price:,.2f}</td>' if best_price > 0 else '<td>—</td>'
        
        item_rows += f"""<tr>
          <td style="text-align:center;color:#64748b">{i+1}</td>
          <td><div class="item-name">{_e(li.get('item_name',''))}</div></td>
          <td><div class="item-desc">{_e(li.get('description',''))}</div></td>
          <td style="text-align:center">{li.get('qty',1)}</td>
          <td style="text-align:center">{_e(li.get('unit','PCS'))}</td>
          {sup_cells}
          {best_cell}
        </tr>"""

    html = f"""<!DOCTYPE html>
<html><head>
<meta charset="UTF-8">
<title>Comparison {_e(q.get('quotation_no',''))}</title>
<link href="https://fonts.googleapis.com/css2?family=Outfit:wght@400;600;800&family=Inter:wght@400;500;600;700&display=swap" rel="stylesheet">
<style>
  @page {{ size: A4 landscape; margin: 10mm 12mm; }}
  *{{box-sizing:border-box;margin:0;padding:0}}
  body{{font-family:'Inter',sans-serif;font-size:11px;color:#1e293b;background:#fff;padding:0;-webkit-print-color-adjust:exact;print-color-adjust:exact}}
  .doc-header{{display:flex;justify-content:space-between;align-items:flex-end;margin-bottom:20px;padding:25px 30px;padding-bottom:12px;border-bottom:2px solid #3b82f6}}
  .doc-title{{font-family:'Outfit',sans-serif;font-size:26px;font-weight:800;color:#1e3a8a}}
  .doc-subtitle{{font-family:'Outfit',sans-serif;font-size:11px;font-weight:600;color:#3b82f6;text-transform:uppercase}}
  .doc-meta{{text-align:right;font-size:11px;color:#64748b;display:grid;grid-template-columns:auto auto;gap:4px 20px}}
  .doc-meta span{{color:#0f172a;font-weight:700}}
  .main{{padding:0 30px}}
  table{{width:100%;border-collapse:separate;border-spacing:0;table-layout:fixed;margin-top:10px}}
  thead th{{padding:10px 8px;background:#1e3a8a;color:#fff;font-family:'Outfit',sans-serif;font-size:9px;text-transform:uppercase;text-align:left;border-right:1px solid rgba(255,255,255,0.1)}}
  th.th-sup{{text-align:right;min-width:100px}}
  tbody td{{padding:8px 10px;border-bottom:1px solid #e2e8f0;border-right:1px solid #f1f5f9;vertical-align:top}}
  .td-sup{{text-align:right}}
  .best-cell{{background:#fef3c7!important;color:#92400e!important;font-weight:800;text-align:right}}
  .item-name{{font-weight:700;color:#0f172a}}
  .item-desc{{color:#64748b;font-size:9.5px}}
  .footer{{margin-top:30px;padding:20px 30px;border-top:1px dashed #cbd5e1;font-size:10px;color:#64748b;display:flex;justify-content:space-between}}
</style>
</head><body>
<div class="doc-header">
  <div><div class="doc-title">QUOTATION COMPARISON</div><div class="doc-subtitle">{_e(q.get('title',''))}</div></div>
  <div class="doc-meta">
    <div>No: <span>{_e(q.get('quotation_no',''))}</span></div>
    <div>Date: <span>{_e(q.get('date',''))}</span></div>
    <div>Currency: <span>{_e(cur)}</span></div>
  </div>
</div>
<div class="main">
<table>
  <thead><tr>
    <th style="width:35px;text-align:center">#</th>
    <th style="width:18%">Product</th>
    <th style="width:15%">Description</th>
    <th style="width:50px;text-align:center">Qty</th>
    <th style="width:40px;text-align:center">Unit</th>
    {sup_headers}
    <th style="width:90px;text-align:right">Best Price</th>
  </tr></thead>
  <tbody>{item_rows}</tbody>
</table>
</div>
<div class="footer">
  <div>Generated by {_e(settings.get('company_name',''))}</div>
  <div>Printed: {date.today().strftime('%d %b %Y')}</div>
</div>
</body></html>"""
    return html, 200, {"Content-Type": "text/html; charset=utf-8"}

@app.route("/api/quotations/<qid>/export-excel", methods=["GET"])
def export_quotation_excel(qid):
    """Export quotation as CSV (Excel-compatible)"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM quotations WHERE id = ?", (qid,))
        row = cursor.fetchone()
        if not row:
            return "Quotation not found", 404

    q = _quotation_to_dict(row)
    cur = q.get("currency", "CNY")
    line_items = q.get("line_items", [])

    # Get all unique supplier names
    sup_names = []
    for li in line_items:
        for sr in li.get("supplier_rows", []):
            nm = sr.get("supplier_name", "")
            if nm and nm not in sup_names:
                sup_names.append(nm)

    # Build CSV
    output = io.StringIO()
    writer = csv.writer(output)

    # Header row
    writer.writerow(["#", "Product", "Description", "Qty", "Unit"] + sup_names + ["Best Price", "Best Supplier"])

    for i, li in enumerate(line_items):
        prices = {}
        for sr in li.get("supplier_rows", []):
            prices[sr.get("supplier_name", "")] = sr.get("price", 0)
        price_vals = [prices.get(nm, 0) for nm in sup_names]
        non_zero = [(nm, p) for nm, p in zip(sup_names, price_vals) if p > 0]
        best_price = min(non_zero, key=lambda x: x[1]) if non_zero else ("", 0)
        writer.writerow([
            i + 1,
            li.get("item_name", ""),
            li.get("description", ""),
            li.get("qty", 1),
            li.get("unit", "PCS"),
        ] + price_vals + [best_price[1], best_price[0]])

    csv_bytes = output.getvalue().encode("utf-8-sig")  # BOM for Excel
    return send_file(
        BytesIO(csv_bytes),
        mimetype="text/csv",
        as_attachment=True,
        download_name=f"Quotation_{q.get('quotation_no','export')}_{date.today()}.csv"
    )


# ═══════════════════════════════════════════════════════════════════════════════
# PO REPORT (unchanged from original)
# ═══════════════════════════════════════════════════════════════════════════════


@app.route("/api/po-report/<pid>", methods=["GET"])
def po_report(pid):
    """Generate PO report"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM purchase_orders WHERE id = ?", (pid,))
        po_row = cursor.fetchone()
        
        if not po_row:
            return "PO not found", 404
        
        po = dict(po_row)
        po["supplier_snapshot"] = json.loads(po.get("supplier_snapshot") or "{}")
        
        # Fetch line items
        cursor.execute("SELECT * FROM po_items WHERE po_id = ? ORDER BY line_sequence", (pid,))
        po["line_items"] = [
            {
                "item_name": item["item_name"],
                "description": item["description"],
                "hs_code": item["hs_code"],
                "qty": item["qty"],
                "unit": item["unit"],
                "unit_price": item["unit_price"],
            }
            for item in cursor.fetchall()
        ]
        
        # Fetch settings
        cursor.execute("SELECT key, value FROM settings")
        settings = {row["key"]: row["value"] for row in cursor.fetchall()}
    
    s = po.get("supplier_snapshot", {})
    cfg = settings

    currency = po.get("currency", "USD")
    cur_sym = "¥" if currency == "CNY" else "$"

    lines = po.get("line_items", [])
    total = sum(_safe_float(li.get("qty", 0)) * _safe_float(li.get("unit_price", 0)) for li in lines)
    total_qty = sum(_safe_float(li.get("qty", 0)) for li in lines)

    status_str = po.get("status", "Draft").upper()

    rows_html = ""
    for i, li in enumerate(lines):
        item_total = _safe_float(li.get("qty", 0)) * _safe_float(li.get("unit_price", 0))
        hs = li.get('hs_code', '') or '—'
        desc = li.get('description', '') or '—'
        item_name = li.get('item_name', '') or f'Item {i+1}'
        rows_html += f"""<tr>
          <td style="text-align:center;color:#6b7280;font-size:10px">{i+1}</td>
          <td style="font-weight:600;color:#111827">{_e(item_name)}</td>
          <td style="text-align:center;color:#374151">{_e(desc)}</td>
          <td style="text-align:center;color:#374151;font-family:monospace">{_e(hs)}</td>
          <td style="text-align:center;font-weight:600;color:#111827">{li.get('qty','')}</td>
          <td style="text-align:center;color:#374151">{_e(li.get('unit',''))}</td>
          <td style="text-align:right;color:#1d4ed8;font-weight:600">{_e(currency)} {li.get('unit_price',0):,.2f}</td>
          <td style="text-align:right;font-weight:700;color:#1d4ed8">{_e(currency)} {item_total:,.2f}</td>
        </tr>"""

    po_date_fmt = ""
    try:
        po_date_fmt = datetime.strptime(po.get("po_date", ""), "%Y-%m-%d").strftime("%Y-%m-%d")
    except:
        po_date_fmt = po.get("po_date", "")

    # Build contact line for supplier
    sup_contacts = []
    if s.get('email'): sup_contacts.append(f"Email: {s.get('email')}")
    if s.get('phone'): sup_contacts.append(f"Ph/WA: {s.get('phone')}")
    sup_contact_html = "<br>".join([_e(c) for c in sup_contacts])

    buyer_lines = []
    if cfg.get('company_address'): buyer_lines.append(cfg['company_address'])
    if cfg.get('company_phone'):   buyer_lines.append(f"Ph: {cfg['company_phone']}")
    if cfg.get('company_email'):   buyer_lines.append(f"Email: {cfg['company_email']}")
    if cfg.get('company_gstin'):   buyer_lines.append(f"GSTIN: {cfg['company_gstin']}")
    buyer_detail_html = "<br>".join([_e(b) for b in buyer_lines])

    pc = po.get("payment_conditions") or ""
    payment_breakdown_html = ""
    m = re.match(r'(\d+)%\s*Advance.*?(\d+)%', pc, re.IGNORECASE)
    if m and total > 0:
        pct1 = int(m.group(1))
        pct2 = int(m.group(2))
        adv  = total * pct1 / 100
        rem  = total * pct2 / 100
        lbl2 = "against B/L (Bill of Lading)" if "BL" in pc.upper() else "on Shipment"
        payment_breakdown_html = f"""
        <div style="margin-top:6px;padding:7px 12px;background:#eff6ff;border:1px solid #bfdbfe;border-radius:5px;font-size:10px;line-height:1.8">
          <div style="font-size:8px;font-weight:700;letter-spacing:.1em;text-transform:uppercase;color:#1d4ed8;margin-bottom:3px">Payment Schedule · 付款计划</div>
          <div style="display:flex;justify-content:space-between;gap:8px;flex-wrap:nowrap"><span style="color:#374151;white-space:nowrap">&#9312; {pct1}% Advance (Proforma Invoice)</span><span style="font-weight:700;color:#1d4ed8;white-space:nowrap">{_e(currency)} {adv:,.2f}</span></div>
          <div style="display:flex;justify-content:space-between;gap:8px;flex-wrap:nowrap"><span style="color:#374151;white-space:nowrap">&#9313; {pct2}% {lbl2}</span><span style="font-weight:700;color:#1d4ed8;white-space:nowrap">{_e(currency)} {rem:,.2f}</span></div>
        </div>"""

    lead_time_days = po.get("lead_time_days", 0) or 0
    due_date_val   = po.get("due_date", "") or ""
    lead_time_html = ""
    if lead_time_days or due_date_val:
        due_line = f'<div style="font-size:10px;color:#1d4ed8;margin-top:2px;font-weight:600">Due: {_e(due_date_val)}</div>' if due_date_val else ""
        lead_time_html = f"""<div class="lg-block lg-divider">
      <div class="lg-sub">LEAD TIME · 交货期</div>
      <div class="lg-val">{lead_time_days} Days</div>
      {due_line}
    </div>"""

    # (Truncated for brevity - rest is same HTML generation as original)
    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>PO {_e(po.get('po_number',''))}</title>
<link href="https://fonts.googleapis.com/css2?family=Noto+Sans+SC:wght@400;600;700&display=swap" rel="stylesheet">
<style>
  *{{box-sizing:border-box;margin:0;padding:0}}
  body{{font-family:'Noto Sans SC','Segoe UI',Arial,sans-serif;font-size:12px;color:#111;background:#fff;padding:32px 36px;max-width:920px;margin:0 auto}}
  .po-header{{display:flex;justify-content:space-between;align-items:flex-start;padding-bottom:12px;border-bottom:3px solid #1e3a8a;margin-bottom:14px}}
  .po-main-title{{font-size:28px;font-weight:700;color:#1e3a8a;letter-spacing:-.02em;line-height:1}}
  .po-subtitle{{font-size:11px;color:#6b7280;margin-top:3px;letter-spacing:.06em;text-transform:uppercase}}
  .po-header-right{{text-align:right;font-size:12px}}
  .po-meta-row{{display:flex;gap:6px;align-items:center;justify-content:flex-end;margin-bottom:4px;color:#374151}}
  .po-meta-row strong{{color:#111;min-width:60px;text-align:right}}
  .parties-grid{{display:grid;grid-template-columns:1fr 1fr;gap:0;border:1px solid #e5e7eb;border-radius:8px;overflow:hidden;margin-bottom:10px}}
  .party-box{{padding:10px 14px}}
  .party-box:first-child{{border-right:1px solid #e5e7eb;background:#f9fafb}}
  .party-box:last-child{{background:#f0f6ff}}
  .party-label{{font-size:9px;font-weight:700;letter-spacing:.1em;text-transform:uppercase;color:#6b7280;margin-bottom:8px;display:flex;align-items:center;gap:5px}}
  .party-label::before{{content:'';display:block;width:10px;height:2px;background:#1e3a8a;border-radius:1px}}
  .delivery-box{{background:#fff8f0;border:1.5px solid #f97316;border-radius:8px;padding:8px 14px;margin-bottom:10px}}
  .delivery-box-title{{font-size:9px;font-weight:700;letter-spacing:.1em;text-transform:uppercase;color:#c2410c;margin-bottom:6px;display:flex;align-items:center;gap:6px}}
  .delivery-box-title::before{{content:'📦';font-size:13px}}
  .delivery-grid{{display:grid;grid-template-columns:1fr 1fr 1.4fr 1fr;gap:8px 14px;font-size:11px}}
  table{{width:100%;border-collapse:collapse;margin-bottom:0}}
  thead tr{{background:#1e3a8a}}
  thead th{{padding:7px 10px;color:#fff;font-size:10px;font-weight:600;letter-spacing:.07em;text-transform:uppercase;text-align:left}}
  tbody tr:nth-child(even){{background:#f8faff}}
  tbody td{{padding:6px 10px;border-bottom:1px solid #e5e7eb;font-size:11px;vertical-align:top}}
  tfoot tr{{background:#1e3a8a}}
  tfoot td{{padding:7px 10px;color:#fff;font-weight:700;font-size:13px}}
  .table-wrap{{border:1px solid #d1d5db;border-radius:8px;overflow:hidden;margin-bottom:8px}}
  .sig-grid{{display:grid;grid-template-columns:1fr 1fr;gap:48px;margin-top:16px}}
  .sig-box{{border-top:2px solid #d1d5db;padding-top:8px}}
  .sig-label{{font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#9ca3af;margin-bottom:4px}}
  .sig-name{{font-size:12px;font-weight:700;color:#111}}
  .sig-sub{{font-size:10px;color:#6b7280;margin-top:2px}}
  .footer-note{{margin-top:12px;padding-top:8px;border-top:1px solid #e5e7eb;font-size:9px;color:#9ca3af;text-align:center;letter-spacing:.04em}}
  @media print{{
    body{{padding:10px 14px}}
    @page{{margin:6mm 8mm;size:A4}}
  }}
</style>
</head>
<body>

<div class="po-header">
  <div class="po-header-left">
    <div class="po-main-title">PURCHASE ORDER</div>
    <div class="po-subtitle">采购订单</div>
  </div>
  <div class="po-header-right">
    <div class="po-meta-row"><strong>PO No:</strong>&nbsp;{_e(po.get('po_number',''))}</div>
    <div class="po-meta-row"><strong>Date:</strong>&nbsp;{_e(po_date_fmt)}</div>
    <div class="po-meta-row"><strong>Currency:</strong>&nbsp;{_e(currency)}</div>
    <div><span style="display:inline-block;padding:3px 12px;border-radius:20px;font-size:10px;font-weight:700;letter-spacing:.07em;text-transform:uppercase;color:#fff;background:#6b7280;margin-top:8px">{_e(status_str)}</span></div>
  </div>
</div>

<div class="parties-grid">
  <div class="party-box">
    <div class="party-label">Buyer / Importer · 买方</div>
    <div style="font-size:14px;font-weight:700;color:#111;margin-bottom:3px">{_e(cfg.get('company_name',''))}</div>
    <div style="color:#4b5563">{buyer_detail_html}</div>
  </div>
  <div class="party-box">
    <div class="party-label">Supplier / Vendor · 供应商</div>
    <div style="font-size:14px;font-weight:700;color:#111;margin-bottom:3px">{_e(s.get('company','—'))}</div>
    <div style="color:#4b5563">{sup_contact_html}</div>
  </div>
</div>

<div class="delivery-box">
  <div class="delivery-box-title">Delivery &amp; Logistics Details · 交货与物流详情</div>
  <div class="delivery-grid">
    <div>
      <div style="font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#9ca3af;margin-bottom:3px">Delivery Terms · 交货条款</div>
      <div style="font-weight:600;color:#111;font-size:14px">{_e(po.get('delivery_terms','—'))}</div>
    </div>
    {lead_time_html}
    <div>
      <div style="font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#9ca3af;margin-bottom:3px">Payment Conditions · 付款条件</div>
      <div style="font-weight:600;color:#111;font-size:13px">{_e(pc or '—')}</div>
    </div>
    <div>{payment_breakdown_html}</div>
  </div>
</div>

<div class="table-wrap">
  <table>
    <thead><tr>
      <th style="width:3%;text-align:center">#</th>
      <th style="width:30%">Item Name<br><span style="font-weight:400;opacity:.75;font-size:9px">品名</span></th>
      <th style="width:10%">Description<br><span style="font-weight:400;opacity:.75;font-size:9px">描述</span></th>
      <th style="width:10%;text-align:center">HS Code</th>
      <th style="width:7%;text-align:center">Qty<br><span style="font-weight:400;opacity:.75;font-size:9px">数量</span></th>
      <th style="width:7%;text-align:center">Unit<br><span style="font-weight:400;opacity:.75;font-size:9px">单位</span></th>
      <th style="width:15%;text-align:right">Unit Price<br><span style="font-weight:400;opacity:.75;font-size:9px">单价</span></th>
      <th style="width:15%;text-align:right">Amount<br><span style="font-weight:400;opacity:.75;font-size:9px">金额</span></th>
    </tr></thead>
    <tbody>{rows_html}</tbody>
    <tfoot>
      <tr>
        <td colspan="4" style="text-align:right;font-size:11px;text-transform:uppercase;padding-right:14px;color:#e0e7ff">Total Quantity Required</td>
        <td style="text-align:center;font-size:13px;color:#fff">{total_qty:g}</td>
        <td colspan="2" style="text-align:right;font-size:11px;text-transform:uppercase;padding-right:14px;color:#e0e7ff">Grand Total · 总金额</td>
        <td style="text-align:right;font-size:14px">{_e(currency)} {total:,.2f}</td>
      </tr>
    </tfoot>
  </table>
</div>

<div class="sig-grid">
  <div class="sig-box">
    <div class="sig-label">Authorised Signatory — Buyer · 买方授权签字</div>
    <div class="sig-name">{_e(cfg.get('company_name','_______________________'))}</div>
    <div class="sig-sub">Signature &amp; Stamp · 签名盖章</div>
  </div>
  <div class="sig-box">
    <div class="sig-label">Accepted — Supplier · 供应商确认</div>
    <div class="sig-name">{_e(s.get('company','_______________________'))}</div>
    <div class="sig-sub">Signature &amp; Stamp · 签名盖章</div>
  </div>
</div>

<div class="footer-note">This is a computer-generated Purchase Order · {date.today().strftime('%d %b %Y')}</div>
</body></html>'''

    try:
        if WEASYPRINT_AVAILABLE:
            pdf_bytes = HTML(string=html, base_url=None).write_pdf()
            return send_file(
                BytesIO(pdf_bytes),
                mimetype="application/pdf",
                as_attachment=False,
                download_name=f"PO_{po.get('po_number','export')}_{date.today()}.pdf"
            )
    except Exception:
        pass
    
    return html, 200, {"Content-Type": "text/html; charset=utf-8"}


# ═══════════════════════════════════════════════════════════════════════════════
# LANDED COST REPORT
# ═══════════════════════════════════════════════════════════════════════════════

def _build_lc_html(po, settings, currency, c, rate, bank, ship, duty, trans, gst_duty, lc_cur, doc_pct):
    cur_sym = "¥" if currency == "CNY" else "$"
    company_name = settings.get("company_name", "")
    total_qty = sum(it.get('qty', 0) for it in c.get('items', []))

    # ── Build HTML Report ────────────────────────────────────────────────────

    # Item rows for main table
    item_rows = ""
    for i, it in enumerate(c.get("items", [])):
        item_rows += f"""<tr>
          <td style="text-align:center;color:#6b7280;font-size:9px">{i+1}</td>
          <td style="font-weight:600;color:#111827;font-size:10px">{_e(it.get('name',''))}</td>
          <td style="text-align:center;font-weight:600;font-size:10px">{it.get('qty',0):,.0f}</td>
          <td style="text-align:right;color:#1d4ed8;font-weight:600;font-size:10px">{cur_sym}{it.get('unitPrice',0):,.2f}</td>
          <td style="text-align:right;font-weight:600;font-size:10px">{fmt_inr(it.get('item_inr',0))}</td>
          <td style="text-align:center;color:#6b7280;font-size:10px">{it.get('share',0):.1f}%</td>
          <td style="text-align:right;color:#059669;font-weight:600;font-size:10px">{fmt_inr(it.get('addl_share',0))}</td>
          <td style="text-align:right;font-weight:700;color:#7c3aed;font-size:10px">{fmt_inr(it.get('total_item',0))}</td>
          <td style="text-align:right;font-weight:700;color:#dc2626;font-size:10px">{fmt_inr(it.get('per_unit',0))}</td>
        </tr>"""

    # Cost breakdown rows for overhead table
    overhead_items = [
        ("Banking Charges", bank, "bank_s"),
        ("Shipping / Freight", ship, "ship_s"),
        ("Customs Duty", duty, "duty_s"),
        ("Transport / Clearing", trans, "trans_s"),
        ("GST on Duty", gst_duty, "gst_s"),
        (f"Documentation ({doc_pct}%)", c.get("doc_inr", 0), "doc_s"),
    ]
    overhead_rows = ""
    for label, total_val, share_key in overhead_items:
        if total_val > 0:
            overhead_rows += f"""<tr>
              <td style="color:#374151;font-weight:500;font-size:10px">{_e(label)}</td>
              <td style="text-align:right;font-weight:600;color:#111;font-size:10px">{fmt_inr(total_val)}</td>"""
            for it in c.get("items", []):
                overhead_rows += f'<td style="text-align:right;color:#6b7280;font-size:10px">{fmt_inr(it.get(share_key,0))}</td>'
            overhead_rows += "</tr>"

    # Item headers for overhead table
    item_headers = ""
    for it in c.get("items", []):
        item_headers += f'<th style="text-align:right;min-width:70px;font-size:9px">{it.get("name","")}</th>'

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Landing Cost — PO {_e(po.get('po_number',''))}</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap" rel="stylesheet">
<style>
  *{{box-sizing:border-box;margin:0;padding:0}}
  body{{font-family:'Inter','Segoe UI',Arial,sans-serif;font-size:11px;color:#111;background:#f8fafc;padding:0;margin:0}}
  .page{{max-width:100%;margin:0 auto;background:#fff;min-height:100vh;padding:20px 24px}}

  /* Header */
  .report-header{{display:grid;grid-template-columns:1fr;gap:12px;padding-bottom:14px;border-bottom:3px solid #1e3a8a;margin-bottom:16px}}
  .report-title{{font-size:22px;font-weight:800;color:#1e3a8a;letter-spacing:-.03em;line-height:1}}
  .report-subtitle{{font-size:10px;color:#6b7280;margin-top:2px;letter-spacing:.06em;text-transform:uppercase}}
  .report-meta{{display:grid;grid-template-columns:1fr 1fr;gap:6px;font-size:10px;color:#374151}}
  .report-meta strong{{color:#111}}

  /* Summary Cards */
  .summary-grid{{display:grid;grid-template-columns:repeat(2,1fr);gap:12px;margin-bottom:24px}}
  .summary-card{{background:linear-gradient(135deg,#f8faff,#eff6ff);border:1px solid #dbeafe;border-radius:8px;padding:12px 14px;text-align:center}}
  .summary-card.highlight{{background:linear-gradient(135deg,#faf5ff,#f3e8ff);border-color:#e9d5ff}}
  .summary-card.grand{{background:linear-gradient(135deg,#ecfdf5,#d1fae5);border-color:#a7f3d0}}
  .summary-card.rate{{background:linear-gradient(135deg,#fffbeb,#fef3c7);border-color:#fde68a}}
  .sc-label{{font-size:8px;font-weight:700;letter-spacing:.1em;text-transform:uppercase;color:#6b7280;margin-bottom:4px}}
  .sc-value{{font-size:18px;font-weight:800;color:#1e3a8a}}
  .summary-card.highlight .sc-value{{color:#7c3aed}}
  .summary-card.grand .sc-value{{color:#059669}}
  .summary-card.rate .sc-value{{color:#d97706;font-size:16px}}

  /* Tables */
  .section-title{{font-size:13px;font-weight:700;color:#1e3a8a;margin:16px 0 8px;display:flex;align-items:center;gap:8px}}
  .section-title::before{{content:'';display:block;width:4px;height:16px;background:#1e3a8a;border-radius:2px}}
  .table-wrap{{border:1px solid #d1d5db;border-radius:6px;overflow:hidden;margin-bottom:14px;font-size:10px}}
  table{{width:100%;border-collapse:collapse}}
  thead tr{{background:#1e3a8a}}
  thead th{{padding:6px 8px;color:#fff;font-size:9px;font-weight:600;letter-spacing:.05em;text-transform:uppercase;text-align:left;white-space:nowrap}}
  tbody tr:nth-child(even){{background:#f8faff}}
  tbody tr:nth-child(odd){{background:#fff}}
  tbody td{{padding:5px 8px;border-bottom:1px solid #e5e7eb;font-size:10px;vertical-align:middle}}
  tfoot tr{{background:#1e3a8a}}
  tfoot td{{padding:6px 8px;color:#fff;font-weight:700;font-size:10px}}

  /* Overhead breakdown */
  .overhead-wrap{{border:1px solid #e9d5ff;border-radius:6px;overflow:hidden;margin-bottom:14px;font-size:10px}}
  .overhead-wrap thead tr{{background:#7c3aed}}
  .overhead-wrap tfoot tr{{background:#7c3aed}}

  /* Footer */
  .report-footer{{margin-top:14px;padding-top:8px;border-top:1px solid #e5e7eb;font-size:8px;color:#9ca3af;text-align:center;letter-spacing:.04em}}

  @media print{{
    body{{background:#fff;padding:0}}
    .page{{padding:14px 18px;box-shadow:none}}
    @page{{margin:10mm;size:A4 portrait}}
    .table-wrap{{page-break-inside:avoid}}
    .overhead-wrap{{page-break-inside:avoid}}
    .summary-grid{{page-break-inside:avoid}}
    thead{{display:table-header-group}}
    tr{{page-break-inside:avoid}}
  }}
</style>
</head>
<body>
<div class="page">

  <!-- HEADER -->
  <div class="report-header">
    <div>
      <div class="report-title">LANDED COST ANALYSIS</div>
      <div class="report-subtitle">Import Costing Breakdown Report</div>
    </div>
    <div class="report-meta">
      <div><strong>PO No:</strong> {_e(po.get('po_number',''))}</div>
      <div><strong>Date:</strong> {po.get('po_date','')}</div>
      <div><strong>Currency:</strong> {currency}</div>
      <div><strong>Status:</strong> {po.get('status','Draft')}</div>
      {f'<div style="margin-top:4px;font-weight:600;color:#1e3a8a">{_e(company_name)}</div>' if company_name else ''}
    </div>
  </div>

  <!-- SUMMARY CARDS -->
  <div class="summary-grid">
    <div class="summary-card rate">
      <div class="sc-label">{lc_cur} → INR Rate</div>
      <div class="sc-value">{rate:,.2f}</div>
    </div>
    <div class="summary-card">
      <div class="sc-label">Invoice ({currency})</div>
      <div class="sc-value">{cur_sym}{c.get('inv_usd',0):,.2f}</div>
    </div>
    <div class="summary-card highlight">
      <div class="sc-label">Total Overhead (₹)</div>
      <div class="sc-value">{fmt_inr(c.get('total_addl',0))}</div>
    </div>
    <div class="summary-card grand">
      <div class="sc-label">Grand Landed Cost (₹)</div>
      <div class="sc-value">{fmt_inr(c.get('grand',0))}</div>
    </div>
  </div>

  <!-- ITEM-WISE BREAKDOWN -->
  <div class="section-title">Item-Wise Landed Cost Breakdown</div>
  <div class="table-wrap">
    <table>
      <thead><tr>
        <th style="width:3%;text-align:center">#</th>
        <th style="width:28%">Item Name</th>
        <th style="text-align:center;width:10%">Qty</th>
        <th style="text-align:right;width:16%">Unit Price</th>
        <th style="text-align:right;width:14%">INR Value</th>
        <th style="text-align:center;width:9%">Share %</th>
        <th style="text-align:right;width:14%">Overhead</th>
        <th style="text-align:right;width:14%">Total Landed</th>
        <th style="text-align:right;width:12%;background:#b91c1c">Per Unit</th>
      </tr></thead>
      <tbody>{item_rows}</tbody>
      <tfoot>
        <tr>
          <td colspan="2" style="text-align:right;font-size:10px;letter-spacing:.04em;text-transform:uppercase;color:#e0e7ff">Total Quantity Required</td>
          <td style="text-align:center;font-size:11px;color:#fff">{total_qty:g}</td>
          <td style="text-align:right;font-size:10px;letter-spacing:.04em;text-transform:uppercase;color:#e0e7ff">Grand Total</td>
          <td style="text-align:right">{fmt_inr(c.get('inv_inr',0))}</td>
          <td style="text-align:center">100%</td>
          <td style="text-align:right">{fmt_inr(c.get('total_addl',0))}</td>
          <td style="text-align:right;font-size:14px">{fmt_inr(c.get('grand',0))}</td>
          <td></td>
        </tr>
      </tfoot>
    </table>
  </div>

  <!-- OVERHEAD ALLOCATION -->
  <div class="section-title" style="color:#7c3aed">Overhead Cost Allocation by Item</div>
  <div class="overhead-wrap">
    <table>
      <thead><tr>
        <th style="min-width:160px">Cost Head</th>
        <th style="text-align:right;min-width:100px">Total (₹)</th>
        {item_headers}
      </tr></thead>
      <tbody>{overhead_rows}</tbody>
      <tfoot>
        <tr>
          <td style="font-weight:700">TOTAL OVERHEAD</td>
          <td style="text-align:right;font-weight:700">{fmt_inr(c.get('total_addl',0))}</td>"""

    for it in c.get("items", []):
        html += f'<td style="text-align:right;font-weight:700">{fmt_inr(it.get("addl_share",0))}</td>'

    html += f"""
        </tr>
      </tfoot>
    </table>
  </div>

  <!-- COST PARAMETERS -->
  <div style="display:grid;grid-template-columns:1fr;gap:14px;margin-top:14px">
    <div style="background:#f8fafc;border:1px solid #e2e8f0;border-radius:8px;padding:12px 16px">
      <div style="font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#64748b;margin-bottom:8px">Cost Parameters Used</div>
      <div style="display:grid;grid-template-columns:1fr 1fr;gap:4px 12px;font-size:10px">
        <div style="color:#64748b">Exchange Rate ({lc_cur}→INR):</div><div style="font-weight:600;color:#111">₹{rate:,.2f}</div>
        <div style="color:#64748b">Banking Charges:</div><div style="font-weight:600;color:#111">{fmt_inr(bank)}</div>
        <div style="color:#64748b">Shipping / Freight:</div><div style="font-weight:600;color:#111">{fmt_inr(ship)}</div>
        <div style="color:#64748b">Customs Duty:</div><div style="font-weight:600;color:#111">{fmt_inr(duty)}</div>
        <div style="color:#64748b">Transport / Clearing:</div><div style="font-weight:600;color:#111">{fmt_inr(trans)}</div>
        <div style="color:#64748b">GST on Duty:</div><div style="font-weight:600;color:#111">{fmt_inr(gst_duty)}</div>
        <div style="color:#64748b">Documentation ({doc_pct}%):</div><div style="font-weight:600;color:#111">{fmt_inr(c.get('doc_inr',0))}</div>
      </div>
    </div>
    <div style="background:#ecfdf5;border:1px solid #a7f3d0;border-radius:8px;padding:12px 16px;display:flex;flex-direction:column;justify-content:center;align-items:center">
      <div style="font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#059669;margin-bottom:4px">Total Landed Cost</div>
      <div style="font-size:28px;font-weight:800;color:#059669">{fmt_inr(c.get('grand',0))}</div>
      <div style="font-size:10px;color:#6b7280;margin-top:2px">Invoice: {fmt_inr(c.get('inv_inr',0))} + Overhead: {fmt_inr(c.get('total_addl',0))}</div>
    </div>
  </div>

  <div class="report-footer">
    This is a computer-generated Landed Cost Report · {_e(company_name)} · Generated on {date.today().strftime('%d %b %Y')}
  </div>

</div>
</body>
</html>"""


    return html

@app.route("/api/lc-report/<pid>", methods=["POST"])
@require_permission("po_edit")
def lc_report(pid):
    """Generate landed cost report"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM purchase_orders WHERE id = ?", (pid,))
        po_row = cursor.fetchone()
        
        if not po_row:
            return "PO not found", 404
        
        po = dict(po_row)
        
        # Fetch line items
        cursor.execute("SELECT * FROM po_items WHERE po_id = ? ORDER BY line_sequence", (pid,))
        po["line_items"] = [dict(item) for item in cursor.fetchall()]
        
        # Fetch settings
        cursor.execute("SELECT key, value FROM settings")
        settings = {row["key"]: row["value"] for row in cursor.fetchall()}

    req = request.get_json(force=True)
    currency = po.get("currency", "USD")
    lc_cur = "RMB" if currency == "CNY" else "USD"
    default_rate = 11.5 if lc_cur == "RMB" else 84.0
    rate = _safe_float(req.get("rate"), po.get("lc_rmb_rate" if lc_cur == "RMB" else "lc_usd_rate", default_rate))
    bank = _safe_float(req.get("bank"), po.get("lc_bank", 0))
    ship = _safe_float(req.get("ship"), po.get("lc_ship", 0))
    duty = _safe_float(req.get("duty"), po.get("lc_duty", 0))
    trans = _safe_float(req.get("trans"), po.get("lc_trans", 0))
    gst_duty = _safe_float(req.get("gst_duty"), po.get("lc_gst_duty", 0))
    doc_pct = _safe_float(req.get("doc_pct"), po.get("lc_doc_pct", 0))

    items = []
    for li in po.get("line_items", []):
        items.append({
            "name": li.get("item_name", ""),
            "qty": float(li.get("qty", 0)),
            "unitPrice": float(li.get("unit_price", 0)),
        })

    c = calc_landed(items, rate, bank, ship, duty, trans, gst_duty, doc_pct)

    # ── Build HTML Report ────────────────────────────────────────────────────
    html = _build_lc_html(po, settings, currency, c, rate, bank, ship, duty, trans, gst_duty, lc_cur, doc_pct)
    return html, 200, {"Content-Type": "text/html; charset=utf-8"}


@app.route("/api/lc-report/<pid>/pdf", methods=["POST"])
@require_permission("po_edit")
def lc_report_pdf(pid):
    """Generate and download landed cost report as PDF (A4 Portrait)"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM purchase_orders WHERE id = ?", (pid,))
        po_row = cursor.fetchone()
        if not po_row: return "PO not found", 404
        po = dict(po_row)
        cursor.execute("SELECT * FROM po_items WHERE po_id = ? ORDER BY line_sequence", (pid,))
        po["line_items"] = [dict(item) for item in cursor.fetchall()]
        cursor.execute("SELECT key, value FROM settings")
        settings = {row["key"]: row["value"] for row in cursor.fetchall()}

    req = request.get_json(force=True)
    currency = po.get("currency", "USD")
    lc_cur = "RMB" if currency == "CNY" else "USD"
    default_rate = 11.5 if lc_cur == "RMB" else 84.0
    rate = _safe_float(req.get("rate"), po.get("lc_rmb_rate" if lc_cur == "RMB" else "lc_usd_rate", default_rate))
    bank = _safe_float(req.get("bank"), po.get("lc_bank", 0))
    ship = _safe_float(req.get("ship"), po.get("lc_ship", 0))
    duty = _safe_float(req.get("duty"), po.get("lc_duty", 0))
    trans = _safe_float(req.get("trans"), po.get("lc_trans", 0))
    gst_duty = _safe_float(req.get("gst_duty"), po.get("lc_gst_duty", 0))
    doc_pct = _safe_float(req.get("doc_pct"), po.get("lc_doc_pct", 0))

    items = []
    for li in po.get("line_items", []):
        items.append({
            "name": li.get("item_name", ""),
            "qty": float(li.get("qty", 0)),
            "unitPrice": float(li.get("unit_price", 0)),
        })

    c = calc_landed(items, rate, bank, ship, duty, trans, gst_duty, doc_pct)
    html = _build_lc_html(po, settings, currency, c, rate, bank, ship, duty, trans, gst_duty, lc_cur, doc_pct)
    # Generate PDF from HTML
    try:
        if WEASYPRINT_AVAILABLE:
            po_number = po["po_number"] if po.get("po_number") else "LC_Report"
            
            pdf_bytes = HTML(string=html, base_url=None).write_pdf()
            return send_file(
                BytesIO(pdf_bytes),
                mimetype="application/pdf",
                as_attachment=True,
                download_name=f"Landed_Cost_Analysis_PO_{po_number}_{date.today().strftime('%d_%b_%Y')}.pdf"
            )
        else:
            return jsonify({"error": "PDF generation not available. WeasyPrint is not installed."}), 500
    except Exception as e:
        return jsonify({"error": f"PDF generation failed: {str(e)}"}), 500



# ═══════════════════════════════════════════════════════════════════════════════
# ATTACHMENTS
# ═══════════════════════════════════════════════════════════════════════════════

ALLOWED_EXTENSIONS = {
    'pdf', 'png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp',
    'doc', 'docx', 'xls', 'xlsx', 'csv', 'txt', 'zip', 'rar',
}
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB


def _attach_dir(pid):
    """Get/create per-PO attachment directory (physical files still on disk)"""
    d = os.path.join(ATTACH_DIR, pid)
    os.makedirs(d, exist_ok=True)
    return d


def _migrate_attach_json(pid, conn):
    """One-time migration: import legacy _meta.json into po_attachments table."""
    meta_path = os.path.join(_attach_dir(pid), "_meta.json")
    if not os.path.exists(meta_path):
        return
    try:
        with open(meta_path, "r", encoding="utf-8") as f:
            entries = json.load(f)
        cursor = conn.cursor()
        for e in entries:
            cursor.execute(
                "INSERT OR IGNORE INTO po_attachments "
                "(id, po_id, filename, original, label, mime, uploaded_at) "
                "VALUES (?, ?, ?, ?, ?, ?, ?)",
                (e["id"], pid, e["filename"], e.get("original", e["filename"]),
                 e.get("label", e["filename"]),
                 e.get("mime", "application/octet-stream"),
                 e.get("uploaded_at", datetime.now().strftime("%Y-%m-%d %H:%M")))
            )
        os.rename(meta_path, meta_path + ".migrated")
    except Exception as ex:
        print(f"[attach-migrate] {pid}: {ex}")


@app.route("/api/po/<pid>/attachments", methods=["GET"])
def list_attachments(pid):
    """List all attachments for a PO"""
    with get_db() as conn:
        _migrate_attach_json(pid, conn)
        cursor = conn.cursor()
        cursor.execute(
            "SELECT id, filename, original, label, mime, uploaded_at "
            "FROM po_attachments WHERE po_id = ? ORDER BY uploaded_at",
            (pid,)
        )
        rows = [dict(r) for r in cursor.fetchall()]

    attach_dir = _attach_dir(pid)
    for row in rows:
        fp = os.path.join(attach_dir, row["filename"])
        row["size"] = os.path.getsize(fp) if os.path.exists(fp) else 0
    return jsonify(rows)


@app.route("/api/po/<pid>/attachments", methods=["POST"])
@require_permission("po_edit")
def upload_attachment(pid):
    """Upload a file attachment to a PO"""
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["file"]
    if not file.filename:
        return jsonify({"error": "Empty filename"}), 400

    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        return jsonify({"error": f"File type .{ext} not allowed"}), 400

    label = request.form.get("label", file.filename.rsplit(".", 1)[0])
    aid = str(uuid.uuid4())
    safe_name = f"{aid}.{ext}"

    dest = os.path.join(_attach_dir(pid), safe_name)
    file.save(dest)

    if os.path.getsize(dest) > MAX_FILE_SIZE:
        os.remove(dest)
        return jsonify({"error": "File too large (max 20 MB)"}), 400

    mime_type = mimetypes.guess_type(file.filename)[0] or "application/octet-stream"
    uploaded_at = datetime.now().strftime("%Y-%m-%d %H:%M")

    with get_db() as conn:
        _migrate_attach_json(pid, conn)
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO po_attachments "
            "(id, po_id, filename, original, label, mime, uploaded_at) "
            "VALUES (?, ?, ?, ?, ?, ?, ?)",
            (aid, pid, safe_name, file.filename, label, mime_type, uploaded_at)
        )
        cursor.execute("UPDATE purchase_orders SET attach_count = attach_count + 1 WHERE id = ?", (pid,))

    return jsonify({"ok": True, "id": aid, "filename": safe_name})


@app.route("/api/po/<pid>/attachments/<aid>", methods=["GET"])
def download_attachment(pid, aid):
    """Download/view an attachment"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT filename, original, mime FROM po_attachments "
            "WHERE id = ? AND po_id = ?",
            (aid, pid)
        )
        row = cursor.fetchone()

    if not row:
        return jsonify({"error": "Attachment not found"}), 404

    fp = os.path.join(_attach_dir(pid), row["filename"])
    if not os.path.exists(fp):
        return jsonify({"error": "File missing from disk"}), 404

    return send_file(fp, download_name=row["original"] or row["filename"],
                     mimetype=row["mime"] or "application/octet-stream")


@app.route("/api/po/<pid>/attachments/<aid>", methods=["DELETE"])
@require_permission("po_delete")
def delete_attachment(pid, aid):
    """Delete an attachment"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT filename FROM po_attachments WHERE id = ? AND po_id = ?",
            (aid, pid)
        )
        row = cursor.fetchone()
        if not row:
            return jsonify({"error": "Attachment not found"}), 404

        fp = os.path.join(_attach_dir(pid), row["filename"])
        if os.path.exists(fp):
            os.remove(fp)

        cursor.execute("DELETE FROM po_attachments WHERE id = ?", (aid,))
        cursor.execute(
            "UPDATE purchase_orders SET attach_count = CASE WHEN attach_count > 0 THEN attach_count - 1 ELSE 0 END WHERE id = ?",
            (pid,)
        )

    return jsonify({"ok": True})


@app.route("/api/po/<pid>/attachments/<aid>/rename", methods=["POST"])
@require_permission("po_edit")
def rename_attachment(pid, aid):
    """Rename an attachment label"""
    req = request.get_json(silent=True) or {}
    new_label = req.get("label", "").strip()
    if not new_label:
        return jsonify({"error": "Label cannot be empty"}), 400

    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT id FROM po_attachments WHERE id = ? AND po_id = ?",
            (aid, pid)
        )
        if not cursor.fetchone():
            return jsonify({"error": "Attachment not found"}), 404
        cursor.execute(
            "UPDATE po_attachments SET label = ? WHERE id = ?",
            (new_label, aid)
        )

    return jsonify({"ok": True})


# ═══════════════════════════════════════════════════════════════════════════════
# PAYMENT PROOF
# ═══════════════════════════════════════════════════════════════════════════════

def _payment_dir(pid):
    """Get/create payment proof directory for a PO (physical file storage)"""
    d = os.path.join(ATTACH_DIR, pid, "_payment")
    os.makedirs(d, exist_ok=True)
    return d


def _migrate_payment_json(pid, conn):
    """One-time migration: import legacy _payment_meta.json into po_payments."""
    meta_path = os.path.join(_payment_dir(pid), "_payment_meta.json")
    if not os.path.exists(meta_path):
        return
    try:
        with open(meta_path, "r", encoding="utf-8") as f:
            m = json.load(f)
        cursor = conn.cursor()
        cursor.execute(
            "INSERT OR IGNORE INTO po_payments "
            "(po_id, filename, original, uploaded_at, confirmed, confirmed_at) "
            "VALUES (?, ?, ?, ?, ?, ?)",
            (pid, m.get("filename", ""),
             m.get("original", m.get("filename", "")),
             m.get("uploaded_at", datetime.now().strftime("%Y-%m-%d %H:%M")),
             1 if m.get("confirmed") else 0,
             m.get("confirmed_at"))
        )
        os.rename(meta_path, meta_path + ".migrated")
    except Exception as ex:
        print(f"[payment-migrate] {pid}: {ex}")


@app.route("/api/po/<pid>/payment", methods=["GET"])
def get_payment_status(pid):
    """Check payment proof status"""
    with get_db() as conn:
        _migrate_payment_json(pid, conn)
        cursor = conn.cursor()
        cursor.execute(
            "SELECT filename, original, uploaded_at, confirmed, confirmed_at "
            "FROM po_payments WHERE po_id = ?",
            (pid,)
        )
        row = cursor.fetchone()

    if not row:
        return jsonify({})

    meta = dict(row)
    meta["confirmed"] = bool(meta["confirmed"])
    fp = os.path.join(_payment_dir(pid), meta["filename"])
    try:
        meta["size"] = os.path.getsize(fp) if os.path.exists(fp) else 0
    except (OSError, IOError) as e:
        meta["size"] = 0
        meta["error"] = f"Could not determine file size: {str(e)}"
    return jsonify(meta)


@app.route("/api/po/<pid>/payment", methods=["POST"])
@require_permission("po_edit")
def upload_payment_proof(pid):
    """Upload payment proof PDF"""
    with get_db() as conn:
        _migrate_payment_json(pid, conn)
        cursor = conn.cursor()
        cursor.execute(
            "SELECT confirmed FROM po_payments WHERE po_id = ?", (pid,)
        )
        existing = cursor.fetchone()
        if existing and existing["confirmed"]:
            return jsonify({"error": "Payment proof already confirmed and locked"}), 409

        if "file" not in request.files:
            return jsonify({"error": "No file provided"}), 400
        file = request.files["file"]
        if not file.filename:
            return jsonify({"error": "Empty filename"}), 400
        if not file.filename.lower().endswith(".pdf"):
            return jsonify({"error": "Only PDF files allowed for payment proof"}), 400

        cursor.execute("SELECT po_number FROM purchase_orders WHERE id = ?", (pid,))
        po_row = cursor.fetchone()
        po_num = po_row["po_number"] if po_row else pid

        safe_po = re.sub(r'[^\w\-]', '_', po_num)
        dest_name = f"Payment_{safe_po}.pdf"
        dest_path = os.path.join(_payment_dir(pid), dest_name)
        file.save(dest_path)

        if os.path.getsize(dest_path) > MAX_FILE_SIZE:
            os.remove(dest_path)
            return jsonify({"error": "File too large (max 20 MB)"}), 400

        uploaded_at = datetime.now().strftime("%Y-%m-%d %H:%M")
        cursor.execute(
            "INSERT OR REPLACE INTO po_payments "
            "(po_id, filename, original, uploaded_at, confirmed, confirmed_at) "
            "VALUES (?, ?, ?, ?, 0, NULL)",
            (pid, dest_name, file.filename, uploaded_at)
        )

    return jsonify({"ok": True, "filename": dest_name})


@app.route("/api/po/<pid>/payment/confirm", methods=["POST"])
@require_permission("po_edit")
def confirm_payment_proof(pid):
    """Confirm and lock payment proof (permanent)"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT filename, confirmed FROM po_payments WHERE po_id = ?", (pid,)
        )
        row = cursor.fetchone()
        if not row:
            return jsonify({"error": "No payment proof uploaded yet"}), 400
        if row["confirmed"]:
            return jsonify({"error": "Already confirmed"}), 409

        confirmed_at = datetime.now().strftime("%Y-%m-%d %H:%M")
        cursor.execute(
            "UPDATE po_payments SET confirmed = 1, confirmed_at = ? WHERE po_id = ?",
            (confirmed_at, pid)
        )
        cursor.execute(
            "SELECT filename, original, uploaded_at, confirmed, confirmed_at "
            "FROM po_payments WHERE po_id = ?",
            (pid,)
        )
        updated = dict(cursor.fetchone())

    updated["confirmed"] = bool(updated["confirmed"])
    fp = os.path.join(_payment_dir(pid), updated["filename"])
    updated["size"] = os.path.getsize(fp) if os.path.exists(fp) else 0
    return jsonify(updated)


@app.route("/api/po/<pid>/payment/download", methods=["GET"])
def download_payment_proof(pid):
    """Download the payment proof PDF"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT filename FROM po_payments WHERE po_id = ?", (pid,)
        )
        row = cursor.fetchone()

    if not row:
        return jsonify({"error": "No payment proof found"}), 404

    fp = os.path.join(_payment_dir(pid), row["filename"])
    if not os.path.exists(fp):
        return jsonify({"error": "File missing from disk"}), 404

    return send_file(fp, download_name=row["filename"], mimetype="application/pdf")

@app.route("/api/scan-invoice", methods=["POST"])
@require_permission("po_edit")
def scan_invoice():
    """
    AI INVOICE SCANNER
    Takes an image/pdf, calls Gemini, and extracts structured data.
    Matches items against master data.
    """
    api_key = request.form.get("api_key")
    model = "gemini-2.5-flash"
    
    if not api_key:
        return jsonify({"error": "Gemini API Key is required"}), 400
    
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    file = request.files["file"]
    file_data = file.read()
    mime_type = mimetypes.guess_type(file.filename)[0] or "application/octet-stream"
    file_ext = file.filename.lower()

    # 1. Fetch Master Data for matching
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id, name, description, hs_code, unit, currency, default_price_usd FROM items")
        master_items = [dict(row) for row in cursor.fetchall()]
        
        cursor.execute("SELECT id, name, company FROM suppliers")
        master_suppliers = [dict(row) for row in cursor.fetchall()]

    is_document = False
    extracted_text = ""

    if file_ext.endswith(('.xlsx', '.xls')):
        is_document = True
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(file_data), data_only=True)
            for sheet in wb.worksheets:
                extracted_text += f"\nSheet: {sheet.title}\n"
                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell) for cell in row if cell is not None]
                    if row_data:
                        extracted_text += " | ".join(row_data) + "\n"
        except Exception as e:
            return jsonify({"error": f"Failed to parse Excel file: {str(e)}"}), 400
    elif file_ext.endswith(('.docx', '.doc')):
        is_document = True
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_data))
            for para in doc.paragraphs:
                if para.text.strip(): extracted_text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        extracted_text += " | ".join(row_data) + "\n"
        except Exception as e:
            return jsonify({"error": f"Failed to parse Word file: {str(e)}"}), 400
    elif file_ext.endswith('.csv'):
        is_document = True
        try:
            text_data = file_data.decode('utf-8', errors='replace')
            reader = csv.reader(io.StringIO(text_data))
            for row in reader:
                extracted_text += " | ".join([str(cell).strip() for cell in row if str(cell).strip()]) + "\n"
        except Exception as e:
            return jsonify({"error": f"Failed to parse CSV file: {str(e)}"}), 400

    # 2. Build Prompt
    base_prompt = f"""
    MASTER ITEMS LIST:
    {json.dumps(master_items)}
    
    MASTER SUPPLIERS LIST:
    {json.dumps(master_suppliers)}

    INSTRUCTIONS:
    1. Identify the Supplier. If the supplier matches one in the MASTER SUPPLIERS LIST (match by name or company), return its "id".
    2. Extract the Date and Currency.
    3. Extract all Line Items. For each item, you must try extremely hard to find a match in the MASTER ITEMS LIST. 
       MATCHING LOGIC:
       - First, try to match the master item's "name" with the invoice item's "description".
       - If no direct match, look for keywords (model numbers, sizes, colors) across both name and description fields to find the closest match.
       - Maximize the number of matched items with the available master list.
       If matched, return the "item_id" from the master list. 
       Always return the item_name, description, qty, unit, and unit_price as found in the invoice.
    4. Return ONLY the JSON object, no other text or formatting.

    JSON SCHEMA:
    {{
      "supplier_id": "...", 
      "supplier_name": "...",
      "date": "YYYY-MM-DD",
      "currency": "USD/CNY",
      "items": [
        {{
          "item_id": "...", 
          "item_name": "...", 
          "description": "...", 
          "qty": 0, 
          "unit": "PCS", 
          "unit_price": 0
        }}
      ]
    }}
    """

    if is_document:
        prompt = f"""
    You are an expert invoice processing assistant. 
    Analyze the following extracted text from an invoice document and extract the data in a structured JSON format.
    
    EXTRACTED INVOICE TEXT:
    {extracted_text}
    """ + base_prompt
    else:
        prompt = f"""
    You are an expert invoice processing assistant. 
    Analyze the attached invoice image/pdf and extract the data in a structured JSON format.
    """ + base_prompt

    try:
        if is_document:
            res = _call_gemini(api_key, model, prompt, file_data=None, mime_type=None)
        else:
            res = _call_gemini(api_key, model, prompt, file_data, mime_type)
        
        # Robustly extract text from Gemini response
        if 'error' in res:
            return jsonify({"error": res['error'].get('message', 'Unknown API Error')}), 400
            
        if not res.get('candidates') or not res['candidates'][0].get('content'):
            # Check if it was blocked
            finish_reason = res.get('candidates', [{}])[0].get('finishReason', 'UNKNOWN')
            return jsonify({"error": f"Gemini returned no results. Reason: {finish_reason}"}), 400
            
        text = res['candidates'][0]['content']['parts'][0]['text']
        
        # Clean up Markdown formatting if present
        text = re.sub(r'```json\s*|\s*```', '', text).strip()
        
        extracted_data = json.loads(text)
        
        # Validate extracted data structure
        if not isinstance(extracted_data, dict):
            return jsonify({"error": "Invalid JSON structure from AI scanner"}), 400
        if 'items' not in extracted_data:
            extracted_data['items'] = []
        if not isinstance(extracted_data.get('items'), list):
            return jsonify({"error": "Items field must be a list"}), 400
        
        # --- Python Logical Matching Fallback ---
        # A deterministic, token-based matching algorithm to ensure high accuracy
        def normalize_str(s):
            if not s: return set()
            s = str(s).lower()
            # Keep alphanumeric and periods (for decimals like 2.5mm), replace rest with space
            s = re.sub(r'[^a-z0-9.]', ' ', s)
            return set(s.split())

        valid_items = []
        for inv_item in extracted_data.get('items', []):
            # Validate item structure
            if not isinstance(inv_item, dict):
                continue
            inv_desc_tokens = normalize_str(inv_item.get('description', ''))
            inv_name_tokens = normalize_str(inv_item.get('item_name', ''))
            
            best_match = None
            best_score = 0
            
            for m_item in master_items:
                m_tokens = normalize_str(m_item.get('name', '')) | normalize_str(m_item.get('description', ''))
                
                score = 0
                
                # Check description tokens (Highest priority)
                desc_intersection = inv_desc_tokens.intersection(m_tokens)
                for token in desc_intersection:
                    if any(c.isdigit() for c in token):
                        score += 6  # Digits + Description = Highest Priority
                    else:
                        score += 2  # Words in Description
                        
                # Check name tokens (Lower priority)
                name_intersection = inv_name_tokens.intersection(m_tokens)
                for token in name_intersection:
                    if token not in desc_intersection: # Avoid double counting
                        if any(c.isdigit() for c in token):
                            score += 3
                        else:
                            score += 1
                        
                if score > best_score:
                    best_score = score
                    best_match = m_item
            
            # If we found a strong match, update the item_id
            if best_match and best_score >= 3:
                inv_item['item_id'] = best_match['id']
            elif best_match and best_score > 0 and not inv_item.get('item_id'):
                inv_item['item_id'] = best_match['id']
                
            # Ignore items that do not have a valid master match
            if inv_item.get('item_id') and any(m.get('id') == inv_item.get('item_id') for m in master_items):
                valid_items.append(inv_item)

        extracted_data['items'] = valid_items
        # ----------------------------------------
        
        return jsonify(extracted_data)
        
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 400


# ═══════════════════════════════════════════════════════════════════════════════
# FORWARDER DASHBOARD
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/forwarder-dashboard")
@require_permission("forwarder_dashboard")
def forwarder_dashboard():
    """Render the Forwarder Dashboard page"""
    return render_template("forwarder_dashboard.html")


# ── HELPERS ───────────────────────────────────────────────────────────────────

def _get_shipment_with_details(cursor, shipment_id):
    """Return a fully enriched shipment dict (forwarder name + linked POs)."""
    cursor.execute("""
        SELECT s.*, f.name AS forwarder_name
        FROM shipments s
        LEFT JOIN forwarders f ON f.id = s.forwarder_id
        WHERE s.id = ? AND s.deleted_at IS NULL
    """, (shipment_id,))
    row = cursor.fetchone()
    if not row:
        return None
    ship = dict(row)
    ship["notes"] = json.loads(ship.get("notes") or "[]")

    # Fetch linked POs
    cursor.execute("""
        SELECT po.id, po.po_number, po.status, po.supplier_snapshot
        FROM shipment_po_link spl
        JOIN purchase_orders po ON po.id = spl.po_id
        WHERE spl.shipment_id = ? AND po.deleted_at IS NULL
        ORDER BY po.po_number
    """, (shipment_id,))
    ship["linked_pos"] = [dict(r) for r in cursor.fetchall()]
    return ship


# ── GET /api/shipments/active-items ───────────────────────────────────────────

@app.route("/api/shipments/active-items", methods=["GET"])
def get_active_shipments_items():
    """Return all items for all active shipments, with unit prices calculated in INR."""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT value FROM settings WHERE key = 'default_usd_rate'")
        usd_rate_row = cursor.fetchone()
        usd_rate = float(usd_rate_row['value']) if usd_rate_row else 84.0
        
        cursor.execute("SELECT value FROM settings WHERE key = 'default_rmb_rate'")
        rmb_rate_row = cursor.fetchone()
        rmb_rate = float(rmb_rate_row['value']) if rmb_rate_row else 11.5
        
        cursor.execute("""
            SELECT 
                spl.shipment_id,
                pi.item_name,
                pi.qty,
                pi.unit_price,
                po.currency,
                po.status,
                po.po_number
            FROM shipment_po_link spl
            JOIN purchase_orders po ON po.id = spl.po_id
            JOIN po_items pi ON pi.po_id = po.id
            JOIN shipments s ON s.id = spl.shipment_id
            WHERE po.deleted_at IS NULL AND s.deleted_at IS NULL
        """)
        rows = cursor.fetchall()
        
        result = {}
        for row in rows:
            sid = row['shipment_id']
            if sid not in result:
                result[sid] = []
            
            # Calculate price in INR based on currency
            currency = row['currency'] or 'USD'
            price = float(row['unit_price'] or 0)
            if currency == 'USD':
                price_inr = round(price * usd_rate, 2)
            elif currency == 'CNY':
                price_inr = round(price * rmb_rate, 2)
            else:
                price_inr = price # Fallback
                
            result[sid].append({
                "item_name": row['item_name'],
                "qty": row['qty'],
                "unit_price_inr": price_inr,
                "status": row['status'],
                "po_number": row['po_number']
            })
            
    return jsonify(result)

# ── GET /api/shipments/<sid>/items ────────────────────────────────────────────

@app.route("/api/shipments/<sid>/items", methods=["GET"])
def get_shipment_items(sid):
    """Fetch all items from all POs linked to a shipment, with unit prices calculated in INR."""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT value FROM settings WHERE key = 'default_usd_rate'")
        usd_rate_row = cursor.fetchone()
        usd_rate = float(usd_rate_row['value']) if usd_rate_row else 84.0
        
        cursor.execute("SELECT value FROM settings WHERE key = 'default_rmb_rate'")
        rmb_rate_row = cursor.fetchone()
        rmb_rate = float(rmb_rate_row['value']) if rmb_rate_row else 11.5

        cursor.execute("""
            SELECT 
                pi.item_name,
                pi.qty,
                pi.unit_price,
                po.currency,
                po.status,
                po.po_number
            FROM shipment_po_link spl
            JOIN purchase_orders po ON po.id = spl.po_id
            JOIN po_items pi ON pi.po_id = po.id
            WHERE spl.shipment_id = ? AND po.deleted_at IS NULL
        """, (sid,))
        rows = cursor.fetchall()
        items = []
        for row in rows:
            currency = row['currency'] or 'USD'
            price = float(row['unit_price'] or 0)
            if currency == 'USD':
                price_inr = round(price * usd_rate, 2)
            elif currency == 'CNY':
                price_inr = round(price * rmb_rate, 2)
            else:
                price_inr = price # Fallback

            items.append({
                "item_name": row['item_name'],
                "qty": row['qty'],
                "unit_price_inr": price_inr,
                "status": row['status'],
                "po_number": row['po_number']
            })
            
    return jsonify(items)

# ── GET /api/shipments ────────────────────────────────────────────────────────

@app.route("/api/shipments", methods=["GET"])
def get_shipments():
    """Return all active (non-deleted) shipments with forwarder name and linked POs."""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT s.*, f.name AS forwarder_name
            FROM shipments s
            LEFT JOIN forwarders f ON f.id = s.forwarder_id
            WHERE s.deleted_at IS NULL
            ORDER BY s.expected_arrival ASC
        """)
        rows = cursor.fetchall()
        shipments = []
        for row in rows:
            ship = dict(row)
            ship["notes"] = json.loads(ship.get("notes") or "[]")
            # Linked POs — just id + po_number for the table view
            cursor2 = conn.cursor()
            cursor2.execute("""
                SELECT po.id, po.po_number, po.status
                FROM shipment_po_link spl
                JOIN purchase_orders po ON po.id = spl.po_id
                WHERE spl.shipment_id = ? AND po.deleted_at IS NULL
                ORDER BY po.po_number
            """, (ship["id"],))
            ship["linked_pos"] = [dict(r) for r in cursor2.fetchall()]
            shipments.append(ship)
    return jsonify(shipments)


# ── POST /api/shipments ───────────────────────────────────────────────────────

@app.route("/api/shipments", methods=["POST"])
@require_permission("forwarder_edit")
def create_shipment():
    """Create a new shipment and optionally link POs."""
    req = request.get_json(silent=True)
    if not req:
        return jsonify({"error": "Invalid JSON"}), 400

    if not req.get("forwarder_id"):
        return jsonify({"error": "forwarder_id is required"}), 400
    if not req.get("departure_date"):
        return jsonify({"error": "departure_date is required"}), 400
    if not req.get("expected_arrival"):
        return jsonify({"error": "expected_arrival is required"}), 400

    ship_id = str(uuid.uuid4())
    now = datetime.utcnow().isoformat(timespec="seconds")

    # Seed initial note if provided
    initial_note = req.get("notes", "").strip()
    notes_list = []
    if initial_note:
        notes_list.append({"date": str(date.today()), "text": initial_note})

    try:
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO shipments
                    (id, forwarder_id, booking_ref, departure_date, expected_arrival,
                     actual_arrival, status, description, notes, created_at, updated_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                ship_id,
                req["forwarder_id"],
                req.get("booking_ref", "").strip() or None,
                req["departure_date"],
                req["expected_arrival"],
                req.get("actual_arrival") or None,
                req.get("status", "Shipped"),
                req.get("description", "").strip() or None,
                json.dumps(notes_list),
                now, now,
            ))

            # Link POs if provided
            for po_id in req.get("po_ids", []):
                if po_id:
                    cursor.execute("""
                        INSERT OR IGNORE INTO shipment_po_link (id, shipment_id, po_id)
                        VALUES (?, ?, ?)
                    """, (str(uuid.uuid4()), ship_id, po_id))

        with get_db() as conn:
            cursor = conn.cursor()
            ship = _get_shipment_with_details(cursor, ship_id)
    except Exception as e:
        return jsonify({"error": str(e)}), 400

    return jsonify(ship), 201


# ── PUT /api/shipments/<id> ───────────────────────────────────────────────────

@app.route("/api/shipments/<sid>", methods=["PUT"])
@require_permission("forwarder_edit")
def update_shipment(sid):
    """Update shipment fields."""
    req = request.get_json(silent=True)
    if not req:
        return jsonify({"error": "Invalid JSON"}), 400

    try:
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM shipments WHERE id = ? AND deleted_at IS NULL", (sid,))
            row = cursor.fetchone()
            if not row:
                return jsonify({"error": "Shipment not found"}), 404

            ship = dict(row)
            old_status = ship["status"]
            old_fwd_id = ship["forwarder_id"]

            # Merge updatable fields
            fields = ["forwarder_id", "booking_ref", "departure_date",
                      "expected_arrival", "actual_arrival", "status",
                      "description"]
            for f in fields:
                if f in req:
                    ship[f] = req[f] if req[f] != "" else None

            # Auto-set actual_arrival if status changed to Delivered
            if ship["status"] == "Delivered" and not ship["actual_arrival"]:
                ship["actual_arrival"] = str(date.today())

            # Append note if provided
            if req.get("notes"):
                notes = json.loads(ship.get("notes") or "[]")
                notes.append({"date": str(date.today()), "text": req["notes"], "author": "User"})
                ship["notes"] = json.dumps(notes)

            cursor.execute("""
                UPDATE shipments
                SET forwarder_id=?, booking_ref=?, departure_date=?, expected_arrival=?,
                    actual_arrival=?, status=?, description=?, notes=?, updated_at=CURRENT_TIMESTAMP
                WHERE id=?
            """, (
                ship["forwarder_id"], ship["booking_ref"], ship["departure_date"],
                ship["expected_arrival"], ship["actual_arrival"], ship["status"],
                ship["description"], ship.get("notes") or "[]", sid,
            ))

            # Update linked POs if po_ids provided
            if "po_ids" in req:
                cursor.execute("DELETE FROM shipment_po_link WHERE shipment_id = ?", (sid,))
                for po_id in req["po_ids"]:
                    if po_id:
                        cursor.execute("""
                            INSERT OR IGNORE INTO shipment_po_link (id, shipment_id, po_id)
                            VALUES (?, ?, ?)
                        """, (str(uuid.uuid4()), sid, po_id))

            # TWO-WAY SYNC: If status or forwarder changed, propagate to ALL linked POs
            if ship["status"] != old_status or ship["forwarder_id"] != old_fwd_id:
                # Get all currently linked POs
                cursor.execute("SELECT po_id FROM shipment_po_link WHERE shipment_id = ?", (sid,))
                linked_po_ids = [r[0] for r in cursor.fetchall()]
                
                if linked_po_ids:
                    # 1. Update forwarder info if forwarder changed
                    if ship["forwarder_id"] != old_fwd_id:
                        cursor.execute("SELECT name, contact_person FROM forwarders WHERE id = ?", (ship["forwarder_id"],))
                        fwd_row = cursor.fetchone()
                        if fwd_row:
                            f_name = fwd_row["name"]
                            f_contact = fwd_row["contact_person"]
                            cursor.execute(f"""
                                UPDATE purchase_orders 
                                SET forwarder_id = ?, forwarder_name = ?, forwarder_contact = ?
                                WHERE id IN ({','.join(['?']*len(linked_po_ids))})
                            """, (ship["forwarder_id"], f_name, f_contact, *linked_po_ids))

                    # 2. Update status if status changed
                    if ship["status"] != old_status:
                        new_po_status = None
                        if ship["status"] == "Shipped": new_po_status = "Shipped"
                        elif ship["status"] == "In Transit": new_po_status = "In Transit"
                        elif ship["status"] == "Delivered": new_po_status = "Received"
                        
                        if new_po_status:
                            for p_id in linked_po_ids:
                                # Get current PO status for logging
                                cursor.execute("SELECT status FROM purchase_orders WHERE id = ?", (p_id,))
                                po_row = cursor.fetchone()
                                if po_row and po_row["status"] != new_po_status:
                                    _log_status_change(cursor, p_id, po_row["status"], new_po_status, 
                                                       f"Status auto-updated via Shipment Dashboard (Shipment status: {ship['status']})")
                                    cursor.execute("UPDATE purchase_orders SET status = ? WHERE id = ?", (new_po_status, p_id))

            ship = _get_shipment_with_details(cursor, sid)
    except Exception as e:
        return jsonify({"error": str(e)}), 400

    return jsonify(ship)


# ── DELETE /api/shipments/<id> ────────────────────────────────────────────────

@app.route("/api/shipments/<sid>", methods=["DELETE"])
@require_permission("forwarder_delete")
def delete_shipment(sid):
    """Soft-delete a shipment."""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM shipments WHERE id = ? AND deleted_at IS NULL", (sid,))
        if not cursor.fetchone():
            return jsonify({"error": "Not found"}), 404
        cursor.execute(
            "UPDATE shipments SET deleted_at = CURRENT_TIMESTAMP WHERE id = ?", (sid,)
        )
    return jsonify({"ok": True})


# ── POST /api/shipments/<id>/notes ────────────────────────────────────────────

@app.route("/api/shipments/<sid>/notes", methods=["POST"])
@require_permission("forwarder_edit")
def add_shipment_note(sid):
    """Append a timestamped note to the shipment's JSON notes log."""
    req = request.get_json(silent=True)
    if not req or not req.get("text", "").strip():
        return jsonify({"error": "Note text is required"}), 400

    try:
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT notes FROM shipments WHERE id = ? AND deleted_at IS NULL", (sid,))
            row = cursor.fetchone()
            if not row:
                return jsonify({"error": "Shipment not found"}), 404

            notes = json.loads(row["notes"] or "[]")
            notes.append({
                "date": str(date.today()),
                "text": req["text"].strip(),
                "author": "User"
            })
            cursor.execute(
                "UPDATE shipments SET notes=?, updated_at=CURRENT_TIMESTAMP WHERE id=?",
                (json.dumps(notes), sid)
            )
        with get_db() as conn:
            cursor = conn.cursor()
            ship = _get_shipment_with_details(cursor, sid)
    except Exception as e:
        return jsonify({"error": str(e)}), 400

    return jsonify(ship)


# ── POST /api/shipments/<id>/link-po ─────────────────────────────────────────

@app.route("/api/shipments/<sid>/link-po", methods=["POST"])
@require_permission("forwarder_edit")
def link_po(sid):
    """Link a purchase order to a shipment."""
    req = request.get_json(silent=True)
    if not req or not req.get("po_id"):
        return jsonify({"error": "po_id is required"}), 400

    try:
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id FROM shipments WHERE id = ? AND deleted_at IS NULL", (sid,))
            if not cursor.fetchone():
                return jsonify({"error": "Shipment not found"}), 404
            cursor.execute("SELECT id FROM purchase_orders WHERE id = ? AND deleted_at IS NULL", (req["po_id"],))
            if not cursor.fetchone():
                return jsonify({"error": "PO not found"}), 404

            cursor.execute("""
                INSERT OR IGNORE INTO shipment_po_link (id, shipment_id, po_id)
                VALUES (?, ?, ?)
            """, (str(uuid.uuid4()), sid, req["po_id"]))

        with get_db() as conn:
            cursor = conn.cursor()
            ship = _get_shipment_with_details(cursor, sid)
    except Exception as e:
        return jsonify({"error": str(e)}), 400

    return jsonify(ship)


# ── POST /api/shipments/<id>/unlink-po ───────────────────────────────────────

@app.route("/api/shipments/<sid>/unlink-po", methods=["POST"])
@require_permission("forwarder_edit")
def unlink_po(sid):
    """Remove a PO link from a shipment."""
    req = request.get_json(silent=True)
    if not req or not req.get("po_id"):
        return jsonify({"error": "po_id is required"}), 400

    po_id = req["po_id"]

    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute(
            "DELETE FROM shipment_po_link WHERE shipment_id=? AND po_id=?",
            (sid, po_id)
        )
    return jsonify({"ok": True})


# ── GET /api/dashboard-stats ──────────────────────────────────────────────────

@app.route("/api/dashboard-stats", methods=["GET"])
def get_dashboard_stats():
    """Return KPI counts for the Forwarder Dashboard cards."""
    today = str(date.today())
    # today+7
    from datetime import timedelta
    week_end = str(date.today() + timedelta(days=7))
    # first day of current month
    month_start = date.today().replace(day=1).isoformat()

    with get_db() as conn:
        cursor = conn.cursor()

        cursor.execute("""
            SELECT COUNT(*) FROM shipments
            WHERE status NOT IN ('Delivered') AND deleted_at IS NULL
        """)
        in_transit = cursor.fetchone()[0]

        cursor.execute("""
            SELECT COUNT(*) FROM shipments
            WHERE expected_arrival BETWEEN ? AND ?
            AND status NOT IN ('Delivered') AND deleted_at IS NULL
        """, (today, week_end))
        arriving_this_week = cursor.fetchone()[0]

        cursor.execute("""
            SELECT COUNT(*) FROM shipments
            WHERE expected_arrival < ? AND status NOT IN ('Delivered') AND deleted_at IS NULL
        """, (today,))
        overdue = cursor.fetchone()[0]

        cursor.execute("""
            SELECT COUNT(*) FROM shipments
            WHERE status IN ('Arrived','Under Clearance') AND deleted_at IS NULL
        """)
        under_clearance = cursor.fetchone()[0]

        cursor.execute("""
            SELECT COUNT(*) FROM shipments
            WHERE status = 'Delivered'
            AND COALESCE(actual_arrival, substr(updated_at, 1, 10)) >= ?
            AND deleted_at IS NULL
        """, (month_start,))

        delivered_this_month = cursor.fetchone()[0]

        # Forwarder performance stats
        cursor.execute("""
            SELECT
                f.id, f.name,
                COUNT(CASE WHEN s.status != 'Delivered' AND s.deleted_at IS NULL THEN 1 END) AS active_shipments,
                COUNT(CASE WHEN s.deleted_at IS NULL THEN 1 END) AS total_shipments,
                ROUND(AVG(CASE WHEN s.actual_arrival IS NOT NULL
                    THEN CAST(julianday(s.actual_arrival) - julianday(s.departure_date) AS REAL)
                    END), 1) AS avg_transit_days,
                COUNT(CASE WHEN s.actual_arrival IS NOT NULL
                    AND s.actual_arrival <= s.expected_arrival
                    AND s.deleted_at IS NULL THEN 1 END) AS on_time_count,
                COUNT(CASE WHEN s.actual_arrival IS NOT NULL
                    AND s.deleted_at IS NULL THEN 1 END) AS delivered_count,
                COUNT(CASE WHEN s.expected_arrival < ?
                    AND s.status != 'Delivered'
                    AND s.deleted_at IS NULL THEN 1 END) AS current_delays
            FROM forwarders f
            LEFT JOIN shipments s ON s.forwarder_id = f.id
            WHERE f.active = 1
            GROUP BY f.id, f.name
            ORDER BY f.name
        """, (today,))
        fwd_rows = cursor.fetchall()
        forwarder_stats = []
        for r in fwd_rows:
            d = dict(r)
            delivered = d["delivered_count"] or 0
            on_time = d["on_time_count"] or 0
            d["on_time_pct"] = round((on_time / delivered * 100), 1) if delivered > 0 else None
            forwarder_stats.append(d)

    return jsonify({
        "in_transit": in_transit,
        "arriving_this_week": arriving_this_week,
        "overdue": overdue,
        "under_clearance": under_clearance,
        "delivered_this_month": delivered_this_month,
        "forwarder_stats": forwarder_stats,
    })


# ═══════════════════════════════════════════════════════════════════════════════
# CUSTOMER MASTER API
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/api/customers", methods=["GET"])
def get_customers():
    """Fetch all active customers with their net outstanding balance."""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM customers WHERE active = 1 ORDER BY company, name")
        customers = [dict(row) for row in cursor.fetchall()]
        
        cursor.execute("""
            SELECT customer_id, 
                   SUM(CASE WHEN dr_cr = 'DR' THEN amount_inr ELSE 0 END) as total_dr,
                   SUM(CASE WHEN dr_cr = 'CR' THEN amount_inr ELSE 0 END) as total_cr
            FROM customer_ledger_entries
            WHERE deleted_at IS NULL
            GROUP BY customer_id
        """)
        balances = {row["customer_id"]: dict(row) for row in cursor.fetchall()}
        
        for c in customers:
            cid = c["id"]
            agg = balances.get(cid, {"total_dr": 0, "total_cr": 0})
            total_dr = agg["total_dr"] or 0
            total_cr = agg["total_cr"] or 0
            c["outstanding_balance"] = round(total_dr - total_cr, 2)
            
    return jsonify(customers)


@app.route("/api/customers", methods=["POST"])
@require_permission("customer_edit")
def add_customer():
    """Create new customer"""
    req = request.get_json(silent=True) or {}
    name = req.get("name", "").strip()
    company = req.get("company", "").strip()
    
    if not name or not company:
        return jsonify({"error": "Customer name and company are required"}), 400
        
    cid = str(uuid.uuid4())
    customer = {
        "id": cid,
        "name": name,
        "company": company,
        "address": req.get("address", "").strip(),
        "city": req.get("city", "").strip(),
        "state": req.get("state", "").strip(),
        "pincode": req.get("pincode", "").strip(),
        "country": req.get("country", "India").strip(),
        "gstin": req.get("gstin", "").strip(),
        "email": req.get("email", "").strip(),
        "phone": req.get("phone", "").strip(),
        "credit_limit": _safe_float(req.get("credit_limit"), 0.0),
        "credit_days": _safe_int(req.get("credit_days"), 30),
        "active": 1
    }
    
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO customers (id, name, company, address, city, state, pincode, country, gstin, email, phone, credit_limit, credit_days, active)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 1)
            """, (
                customer["id"], customer["name"], customer["company"], customer["address"],
                customer["city"], customer["state"], customer["pincode"], customer["country"],
                customer["gstin"], customer["email"], customer["phone"], customer["credit_limit"],
                customer["credit_days"]
            ))
            
            opening_bal = _safe_float(req.get("opening_balance"), 0.0)
            if opening_bal > 0:
                eid = str(uuid.uuid4())
                today = str(date.today())
                cursor.execute("""
                    INSERT INTO customer_ledger_entries (id, customer_id, entry_type, entry_date, ref_number, description, amount_inr, dr_cr, created_by)
                    VALUES (?, ?, 'OPENING', ?, 'OB', 'Opening Balance', ?, 'DR', 'System')
                """, (eid, cid, today, opening_bal))
                
    except sqlite3.IntegrityError:
        return jsonify({"error": f"Customer with name '{name}' and company '{company}' already exists"}), 409
    except Exception as e:
        return jsonify({"error": str(e)}), 400
        
    return jsonify(customer), 201


@app.route("/api/customers/<cid>", methods=["PUT"])
@require_permission("customer_edit")
def update_customer(cid):
    """Update customer"""
    req = request.get_json(silent=True) or {}
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM customers WHERE id = ?", (cid,))
            row = cursor.fetchone()
            if not row:
                return jsonify({"error": "Customer not found"}), 404
                
            customer = dict(row)
            fields = ["name", "company", "address", "city", "state", "pincode", "country", "gstin", "email", "phone", "credit_limit", "credit_days", "active"]
            for f in fields:
                if f in req:
                    if f == "credit_limit":
                        customer[f] = _safe_float(req[f], 0.0)
                    elif f == "credit_days":
                        customer[f] = _safe_int(req[f], 30)
                    elif f == "active":
                        customer[f] = int(req[f])
                    else:
                        customer[f] = req[f].strip()
                        
            cursor.execute("""
                UPDATE customers
                SET name = ?, company = ?, address = ?, city = ?, state = ?, pincode = ?, country = ?, gstin = ?, email = ?, phone = ?, credit_limit = ?, credit_days = ?, active = ?
                WHERE id = ?
            """, (
                customer["name"], customer["company"], customer["address"], customer["city"],
                customer["state"], customer["pincode"], customer["country"], customer["gstin"],
                customer["email"], customer["phone"], customer["credit_limit"], customer["credit_days"],
                customer["active"], cid
            ))
    except sqlite3.IntegrityError:
        return jsonify({"error": "Customer with this name and company already exists"}), 409
    except Exception as e:
        return jsonify({"error": str(e)}), 400
        
    return jsonify(customer)


@app.route("/api/customers/<cid>", methods=["DELETE"])
@require_permission("customer_delete")
def delete_customer(cid):
    """Soft-delete customer."""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM customers WHERE id = ?", (cid,))
        if not cursor.fetchone():
            return jsonify({"error": "Customer not found"}), 404
            
        cursor.execute("""
            SELECT dr_cr, amount_inr FROM customer_ledger_entries
            WHERE customer_id = ? AND deleted_at IS NULL
        """, (cid,))
        entries = cursor.fetchall()
        total_dr = sum(e["amount_inr"] for e in entries if e["dr_cr"] == "DR")
        total_cr = sum(e["amount_inr"] for e in entries if e["dr_cr"] == "CR")
        outstanding = round(total_dr - total_cr, 2)
        
        if abs(outstanding) > 0.01:
            return jsonify({"error": f"Cannot delete customer with non-zero outstanding balance: {fmt_inr(outstanding)}"}), 400
            
        cursor.execute("UPDATE customers SET active = 0 WHERE id = ?", (cid,))
        
    return jsonify({"ok": True})


# ═══════════════════════════════════════════════════════════════════════════════
# CUSTOMER LEDGER API
# ═══════════════════════════════════════════════════════════════════════════════

def calc_customer_running_balance(entries):
    balance = 0.0
    for e in entries:
        if e['dr_cr'] == 'DR':
            balance += e['amount_inr']
        elif e['dr_cr'] == 'CR':
            balance -= e['amount_inr']
        e['running_balance'] = round(balance, 2)
        e['balance_suffix'] = 'Dr' if balance >= 0 else 'Cr'
    return entries, round(balance, 2)

@app.route("/api/customers/<cid>/invoices/<invoice_no>/pdf", methods=["GET"])
def get_customer_invoice_pdf(cid, invoice_no):
    """Generate and return a PDF for a specific customer invoice."""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM customers WHERE id = ?", (cid,))
        customer = cursor.fetchone()
        if not customer:
            return jsonify({"error": "Customer not found"}), 404
            
        cursor.execute("SELECT * FROM customer_invoices WHERE customer_id = ? AND invoice_no = ?", (cid, invoice_no))
        invoice_row = cursor.fetchone()
        if not invoice_row:
            return jsonify({"error": "Invoice not found"}), 404
        invoice = dict(invoice_row)
            
        cursor.execute("""
            SELECT ci.*, i.name as item_name 
            FROM customer_invoice_items ci
            LEFT JOIN items i ON ci.item_id = i.id
            WHERE ci.invoice_id = ?
        """, (invoice["id"],))
        items = [dict(i) for i in cursor.fetchall()]
        
        cursor.execute("SELECT key, value FROM settings")
        settings = {row["key"]: row["value"] for row in cursor.fetchall()}
        
    try:
        from flask import render_template_string
        
        # Build Seller HTML
        seller_lines = []
        if settings.get("company_name"): seller_lines.append(f'<div style="font-size:16px;font-weight:700;color:#111;margin-bottom:3px">{_e(settings["company_name"])}</div>')
        if settings.get("company_address"):
            addr_html = _e(settings["company_address"]).replace("\n", "<br>")
            seller_lines.append(f'<div style="color:#4b5563">{addr_html}</div>')
        if settings.get("company_phone"): seller_lines.append(f'<div style="color:#4b5563">Ph: {_e(settings["company_phone"])}</div>')
        if settings.get("company_email"): seller_lines.append(f'<div style="color:#4b5563">Email: {_e(settings["company_email"])}</div>')
        if settings.get("company_gstin"): seller_lines.append(f'<div style="color:#4b5563">GSTIN: {_e(settings["company_gstin"])}</div>')
        seller_html = "".join(seller_lines) or '<div style="color:#9ca3af">Company details not configured.</div>'

        # Build Customer HTML
        cust = dict(customer)
        cust_lines = []
        cust_lines.append(f'<div style="font-size:14px;font-weight:700;color:#111;margin-bottom:3px">{_e(cust.get("company") or cust.get("name") or "—")}</div>')
        if cust.get("name") and cust.get("name") != cust.get("company"): cust_lines.append(f'<div style="color:#4b5563">Attn: {_e(cust["name"])}</div>')
        if cust.get("address"): cust_lines.append(f'<div style="color:#4b5563">{_e(cust["address"])}</div>')
        if cust.get("email"): cust_lines.append(f'<div style="color:#4b5563">Email: {_e(cust["email"])}</div>')
        if cust.get("phone"): cust_lines.append(f'<div style="color:#4b5563">Ph/WA: {_e(cust["phone"])}</div>')
        if cust.get("gstin"): cust_lines.append(f'<div style="color:#4b5563">GSTIN: {_e(cust["gstin"])}</div>')
        customer_html = "".join(cust_lines)

        # Build Items Table
        item_rows = ""
        for i, li in enumerate(items):
            item_name = _e(li.get('item_name') or 'Custom Item')
            description = _e(li.get('description') or '-')
                
            item_rows += f"""<tr>
              <td style="text-align:center;color:#374151">{i+1}</td>
              <td style="font-weight:600;color:#111;">{item_name}</td>
              <td style="color:#4b5563;">{description}</td>
              <td style="text-align:center;font-weight:600">{li['qty']}</td>
              <td style="text-align:right">{li['unit_price']:,.2f}</td>
              <td style="text-align:right;font-weight:700;color:#1e3a8a">{li['total_price']:,.2f}</td>
            </tr>"""

        invoice_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>Invoice {_e(invoice['invoice_no'])}</title>
<style>
  *{{box-sizing:border-box;margin:0;padding:0}}
  body{{font-family:'Segoe UI',Arial,sans-serif;font-size:12px;color:#111;background:#fff;padding:32px 36px;max-width:920px;margin:0 auto}}
  .po-header{{display:flex;justify-content:space-between;align-items:flex-start;padding-bottom:12px;border-bottom:3px solid #1e3a8a;margin-bottom:14px}}
  .po-main-title{{font-size:28px;font-weight:700;color:#1e3a8a;letter-spacing:-.02em;line-height:1}}
  .po-subtitle{{font-size:11px;color:#6b7280;margin-top:3px;letter-spacing:.06em;text-transform:uppercase}}
  .po-header-right{{text-align:right;font-size:12px}}
  .po-meta-row{{display:flex;gap:6px;align-items:center;justify-content:flex-end;margin-bottom:4px;color:#374151}}
  .po-meta-row strong{{color:#111;min-width:60px;text-align:right}}
  .parties-grid{{display:grid;grid-template-columns:1fr 1fr;gap:0;border:1px solid #e5e7eb;border-radius:8px;overflow:hidden;margin-bottom:14px}}
  .party-box{{padding:10px 14px}}
  .party-box:first-child{{border-right:1px solid #e5e7eb;background:#f9fafb}}
  .party-box:last-child{{background:#f0f6ff}}
  .party-label{{font-size:9px;font-weight:700;letter-spacing:.1em;text-transform:uppercase;color:#6b7280;margin-bottom:8px;display:flex;align-items:center;gap:5px}}
  .party-label::before{{content:'';display:block;width:10px;height:2px;background:#1e3a8a;border-radius:1px}}
  table{{width:100%;border-collapse:collapse;margin-bottom:0}}
  thead tr{{background:#1e3a8a}}
  thead th{{padding:7px 10px;color:#fff;font-size:10px;font-weight:600;letter-spacing:.07em;text-transform:uppercase;text-align:left}}
  tbody tr:nth-child(even){{background:#f8faff}}
  tbody td{{padding:6px 10px;border-bottom:1px solid #e5e7eb;font-size:11px;vertical-align:top}}
  tfoot tr{{background:#1e3a8a}}
  tfoot td{{padding:7px 10px;color:#fff;font-weight:700;font-size:13px}}
  .table-wrap{{border:1px solid #d1d5db;border-radius:8px;overflow:hidden;margin-bottom:16px}}
  .totals-table{{width:300px;margin-left:auto;border-collapse:collapse}}
  .totals-table td{{padding:6px 8px;border:none;font-size:12px}}
  .totals-table tr.grand-total{{border-top:2px solid #1e3a8a;font-weight:700;font-size:15px;color:#1e3a8a}}
  @media print{{
    body{{padding:10px 14px}}
    @page{{margin:6mm 8mm;size:A4}}
  }}
</style>
</head>
<body onload="window.print()">

<div class="po-header">
  <div class="po-header-left">
    <div class="po-main-title">SALES INVOICE</div>
    <div class="po-subtitle">Original for Recipient</div>
  </div>
  <div class="po-header-right">
    <div class="po-meta-row"><strong>Invoice No:</strong>&nbsp;{_e(invoice['invoice_no'])}</div>
    <div class="po-meta-row"><strong>Date:</strong>&nbsp;{_e(invoice['invoice_date'])}</div>
    {f'<div class="po-meta-row"><strong>Due Date:</strong>&nbsp;{_e(invoice["due_date"])}</div>' if invoice['due_date'] else ''}
  </div>
</div>

<div class="parties-grid">
  <div class="party-box">
    <div class="party-label">Billed By (Seller)</div>
    {seller_html}
  </div>
  <div class="party-box">
    <div class="party-label">Billed To (Customer)</div>
    {customer_html}
  </div>
</div>

<div class="table-wrap">
  <table>
    <thead>
      <tr>
        <th style="width:40px;text-align:center">#</th>
        <th style="width:250px">Item Name</th>
        <th>Description</th>
        <th style="width:80px;text-align:center">Qty</th>
        <th style="width:100px;text-align:right">Rate (INR)</th>
        <th style="width:120px;text-align:right">Amount (INR)</th>
      </tr>
    </thead>
    <tbody>
      {item_rows}
    </tbody>
  </table>
</div>

<table class="totals-table">
  <tr>
    <td>Subtotal</td>
    <td style="text-align:right">{invoice['subtotal']:,.2f}</td>
  </tr>"""

        if invoice['discount_amount'] > 0:
            invoice_html += f"""
  <tr>
    <td>Discount</td>
    <td style="text-align:right">- {invoice['discount_amount']:,.2f}</td>
  </tr>"""
        elif invoice['discount_amount'] < 0:
            invoice_html += f"""
  <tr>
    <td>Freight / Transport</td>
    <td style="text-align:right">+ {abs(invoice['discount_amount']):,.2f}</td>
  </tr>"""

        if invoice['tax_amount'] > 0:
            invoice_html += f"""
  <tr>
    <td>Tax ({invoice['tax_rate']}%)</td>
    <td style="text-align:right">{invoice['tax_amount']:,.2f}</td>
  </tr>"""

        invoice_html += f"""
  <tr class="grand-total">
    <td>Grand Total</td>
    <td style="text-align:right">INR {invoice['grand_total']:,.2f}</td>
  </tr>
</table>

{f'<div style="margin-top:20px;padding:10px 14px;background:#f8faff;border:1px solid #dbeafe;border-radius:6px;font-size:11px"><div style="font-weight:700;text-transform:uppercase;letter-spacing:.05em;font-size:9px;color:#1e3a8a;margin-bottom:4px">Notes / Reference</div><div style="color:#1e40af">{_e(invoice["description"])}</div></div>' if invoice.get("description") else ''}

</body>
</html>"""
        
        return invoice_html
    except Exception as e:
        return jsonify({"error": f"Failed to generate Invoice View: {str(e)}"}), 500

@app.route("/api/customers/<cid>/ledger", methods=["GET"])
def get_customer_ledger(cid):
    """Fetch all active customer ledger entries with running balance."""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM customers WHERE id = ?", (cid,))
        if not cursor.fetchone():
            return jsonify({"error": "Customer not found"}), 404
            
        cursor.execute("""
            SELECT * FROM customer_ledger_entries
            WHERE customer_id = ? AND deleted_at IS NULL
            ORDER BY entry_date ASC, created_at ASC
        """, (cid,))
        entries = [dict(r) for r in cursor.fetchall()]
        
    entries, closing_balance = calc_customer_running_balance(entries)
    balance_suffix = "Dr" if closing_balance >= 0 else "Cr"
    
    return jsonify({
        "entries": entries,
        "closing_balance": closing_balance,
        "balance_suffix": balance_suffix
    })

@app.route("/api/customers/<cid>/ledger", methods=["POST"])
@require_permission("customer_edit")
def add_customer_ledger_entry(cid):
    """Add new customer ledger entry (voucher)."""
    req = request.get_json(silent=True) or {}
    VALID_TYPES = {"INVOICE", "PAYMENT", "DEBIT_NOTE", "CREDIT_NOTE", "ADJUSTMENT", "OPENING"}
    entry_type = req.get("entry_type", "").upper()
    if entry_type not in VALID_TYPES:
        return jsonify({"error": f"entry_type must be one of {VALID_TYPES}"}), 400
        
    if entry_type == "INVOICE":
        dr_cr = "DR"
        amount_inr = _safe_float(req.get("amount_inr"), 0) # Fallback if not calculated
    elif entry_type == "PAYMENT" or entry_type == "CREDIT_NOTE":
        dr_cr = "CR"
        amount_inr = _safe_float(req.get("amount_inr"), 0)
    elif entry_type == "DEBIT_NOTE":
        dr_cr = "DR"
        amount_inr = _safe_float(req.get("amount_inr"), 0)
    else:
        dr_cr = req.get("dr_cr", "").upper()
        if dr_cr not in ("DR", "CR"):
            return jsonify({"error": "dr_cr must be DR or CR for ADJUSTMENT/OPENING"}), 400
        amount_inr = _safe_float(req.get("amount_inr"), 0)
            
    if entry_type != "INVOICE" and amount_inr <= 0:
        return jsonify({"error": "amount_inr must be positive and non-zero"}), 400
        
    entry_date = req.get("entry_date", str(date.today()))
    due_date = req.get("due_date") or None
    
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("BEGIN TRANSACTION")
            
            cursor.execute("SELECT credit_days FROM customers WHERE id = ?", (cid,))
            cust_row = cursor.fetchone()
            if not cust_row:
                cursor.execute("ROLLBACK")
                return jsonify({"error": "Customer not found"}), 404
                
            credit_days = cust_row["credit_days"]
            if not due_date and entry_type == "INVOICE":
                from datetime import timedelta
                try:
                    base = date.fromisoformat(entry_date)
                    due_date = (base + timedelta(days=credit_days)).isoformat()
                except:
                    due_date = entry_date
                    
            eid = str(uuid.uuid4())
            ref_number = req.get("ref_number", "")
            description = req.get("description", "")
            created_by = req.get("created_by", "User")
            
            # If it's an invoice, perform atomic insertion into both customer_invoices and customer_ledger_entries
            if entry_type == "INVOICE":
                invoice_id = str(uuid.uuid4())
                
                if ref_number:
                    cursor.execute("SELECT id FROM customer_invoices WHERE invoice_no = ?", (ref_number,))
                    if cursor.fetchone():
                        cursor.execute("ROLLBACK")
                        return jsonify({"error": f"Invoice number '{ref_number}' already exists. Please refresh to get a new number."}), 409
                    inv_no = ref_number
                    _sync_inv_sequence(cursor, inv_no)
                else:
                    cursor.execute("SELECT key, value FROM settings WHERE key IN ('inv_prefix', 'inv_sequence')")
                    cfg = {row["key"]: row["value"] for row in cursor.fetchall()}
                    prefix = (cfg.get("inv_prefix") or "INV").strip()
                    seq = int(cfg.get("inv_sequence") or "0")
                    year = date.today().year
                    while True:
                        seq += 1
                        inv_no = f"{prefix}-{year}-{str(seq).zfill(3)}"
                        cursor.execute("SELECT id FROM customer_invoices WHERE invoice_no = ?", (inv_no,))
                        if not cursor.fetchone():
                            break
                    _sync_inv_sequence(cursor, inv_no)
                    
                # We need to make sure ledger entry also uses this generated invoice number as reference
                ref_number = inv_no
                
                items = req.get("items", [])
                if not items:
                    cursor.execute("ROLLBACK")
                    return jsonify({"error": "Invoice must contain at least one item."}), 400
                
                subtotal = 0.0
                prepared_items = []
                for item in items:
                    qty = _safe_float(item.get("qty"), 1.0)
                    price = _safe_float(item.get("unit_price"), 0.0)
                    line_total = round(qty * price, 2)
                    subtotal += line_total
                    prepared_items.append((
                        str(uuid.uuid4()), invoice_id, item.get("item_id"), item.get("description", ""),
                        qty, price, line_total
                    ))
                
                tax_rate = _safe_float(req.get("tax_rate"), 0.0)
                discount_amount = _safe_float(req.get("discount_amount"), 0.0)
                tax_amount = round((subtotal - discount_amount) * (tax_rate / 100.0), 2)
                if tax_amount < 0: tax_amount = 0.0
                grand_total = round(subtotal - discount_amount + tax_amount, 2)
                amount_inr = grand_total
                
                if amount_inr <= 0:
                    cursor.execute("ROLLBACK")
                    return jsonify({"error": "Invoice grand total must be positive."}), 400
                
                cursor.execute("""
                    INSERT INTO customer_invoices 
                        (id, customer_id, invoice_no, invoice_date, due_date, 
                         subtotal, tax_rate, tax_amount, discount_amount, grand_total, 
                         description, created_by)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    invoice_id, cid, inv_no, entry_date, due_date, 
                    subtotal, tax_rate, tax_amount, discount_amount, grand_total, 
                    description, created_by
                ))

                for item_tuple in prepared_items:
                    cursor.execute("""
                        INSERT INTO customer_invoice_items 
                            (id, invoice_id, item_id, description, qty, unit_price, total_price)
                        VALUES (?, ?, ?, ?, ?, ?, ?)
                    """, item_tuple)
                
                # USD calculations
                usd_rate = _safe_float(req.get("usd_rate"), 84.0)
                amount_usd = round(amount_inr / usd_rate, 2)
                
                # The ledger entry uses the invoice_id (or inv_no) as reference
                cursor.execute("""
                    INSERT INTO customer_ledger_entries
                        (id, customer_id, entry_type, entry_date, ref_number, description,
                         amount_inr, amount_usd, usd_rate, dr_cr, payment_mode, bank_ref,
                         due_date, notes, created_by)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    eid, cid, entry_type, entry_date, inv_no, description,
                    amount_inr, amount_usd, usd_rate, dr_cr,
                    req.get("payment_mode", ""), req.get("bank_ref", ""),
                    due_date, req.get("notes", ""), created_by
                ))
            else:
                usd_rate = _safe_float(req.get("usd_rate"), 84.0)
                amount_usd = round(amount_inr / usd_rate, 2)
                
                cursor.execute("""
                    INSERT INTO customer_ledger_entries
                        (id, customer_id, entry_type, entry_date, ref_number, description,
                         amount_inr, amount_usd, usd_rate, dr_cr, payment_mode, bank_ref,
                         due_date, notes, created_by)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    eid, cid, entry_type, entry_date, ref_number, description,
                    amount_inr, amount_usd, usd_rate, dr_cr,
                    req.get("payment_mode", ""), req.get("bank_ref", ""),
                    due_date, req.get("notes", ""), created_by
                ))
            
            conn.commit()
    except Exception as e:
        return jsonify({"error": f"Database error: {str(e)}"}), 500
        
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM customer_ledger_entries WHERE id = ?", (eid,))
        row = dict(cursor.fetchone())
        
    return jsonify(row), 201


def is_older_than_7_days(created_str):
    try:
        if ' ' in created_str:
            dt = datetime.strptime(created_str, "%Y-%m-%d %H:%M:%S")
        else:
            dt = datetime.strptime(created_str, "%Y-%m-%d")
        return (datetime.utcnow() - dt).days > 7
    except:
        return False

@app.route("/api/invoices/ledger/<eid>", methods=["GET"])
@require_permission("customer_books")
def get_invoice_by_ledger_id(eid):
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM customer_ledger_entries WHERE id = ?", (eid,))
        ledger = cursor.fetchone()
        if not ledger or ledger["entry_type"] != "INVOICE":
            return jsonify({"error": "Invoice ledger entry not found"}), 404
            
        cursor.execute("SELECT * FROM customer_invoices WHERE invoice_no = ? AND customer_id = ?", (ledger["ref_number"], ledger["customer_id"]))
        invoice = cursor.fetchone()
        if not invoice:
            return jsonify({"error": "Invoice details not found"}), 404
            
        cursor.execute("SELECT * FROM customer_invoice_items WHERE invoice_id = ?", (invoice["id"],))
        items = [dict(i) for i in cursor.fetchall()]
        
        return jsonify({
            "ledger": dict(ledger),
            "invoice": dict(invoice),
            "items": items
        })

@app.route("/api/ledger/customer/<eid>", methods=["PUT"])
@require_permission("customer_edit")
def update_customer_ledger_entry(eid):
    """Edit customer ledger entry. Locked if older than 7 days. Also handles invoice updates."""
    req = request.get_json(silent=True) or {}
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM customer_ledger_entries WHERE id = ? AND deleted_at IS NULL", (eid,))
        row = cursor.fetchone()
        if not row:
            return jsonify({"error": "Ledger entry not found"}), 404
            
        created_str = row["created_at"] or row["entry_date"]
        if is_older_than_7_days(created_str):
            return jsonify({"error": "Entries older than 7 days are locked and cannot be edited"}), 400
            
        entry = dict(row)
        
        if entry["entry_type"] == "INVOICE":
            # Handle full invoice update
            cursor.execute("SELECT * FROM customer_invoices WHERE invoice_no = ? AND customer_id = ?", (entry["ref_number"], entry["customer_id"]))
            inv = cursor.fetchone()
            if not inv:
                return jsonify({"error": "Underlying invoice record not found"}), 404
                
            items = req.get("items", [])
            if not items:
                return jsonify({"error": "Invoice must contain at least one item."}), 400
                
            subtotal = 0.0
            prepared_items = []
            invoice_id = inv["id"]
            
            for item in items:
                qty = _safe_float(item.get("qty"), 1.0)
                price = _safe_float(item.get("unit_price"), 0.0)
                line_total = round(qty * price, 2)
                subtotal += line_total
                prepared_items.append((
                    str(uuid.uuid4()), invoice_id, item.get("item_id"), item.get("description", ""),
                    qty, price, line_total
                ))
                
            tax_rate = _safe_float(req.get("tax_rate"), 0.0)
            discount_amount = _safe_float(req.get("discount_amount"), 0.0)
            tax_amount = round((subtotal - discount_amount) * (tax_rate / 100.0), 2)
            if tax_amount < 0: tax_amount = 0.0
            grand_total = round(subtotal - discount_amount + tax_amount, 2)
            amount_inr = grand_total
            
            if amount_inr <= 0:
                return jsonify({"error": "Invoice grand total must be positive."}), 400
                
            entry_date = req.get("entry_date", entry["entry_date"])
            due_date = req.get("due_date", entry["due_date"])
            description = req.get("description", entry["description"])
            
            cursor.execute("BEGIN TRANSACTION")
            try:
                # Update invoice
                cursor.execute("""
                    UPDATE customer_invoices 
                    SET invoice_date = ?, due_date = ?, subtotal = ?, tax_rate = ?, 
                        tax_amount = ?, discount_amount = ?, grand_total = ?, description = ?
                    WHERE id = ?
                """, (entry_date, due_date, subtotal, tax_rate, tax_amount, discount_amount, grand_total, description, invoice_id))
                
                # Update items
                cursor.execute("DELETE FROM customer_invoice_items WHERE invoice_id = ?", (invoice_id,))
                cursor.executemany("""
                    INSERT INTO customer_invoice_items 
                        (id, invoice_id, item_id, description, qty, unit_price, total_price)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                """, prepared_items)
                
                # Update ledger entry
                cursor.execute("""
                    UPDATE customer_ledger_entries
                    SET entry_date = ?, due_date = ?, description = ?, amount_inr = ?
                    WHERE id = ?
                """, (entry_date, due_date, description, amount_inr, eid))
                
                cursor.execute("COMMIT")
            except Exception as e:
                cursor.execute("ROLLBACK")
                return jsonify({"error": str(e)}), 500
                
            cursor.execute("SELECT * FROM customer_ledger_entries WHERE id = ?", (eid,))
            updated = dict(cursor.fetchone())
            return jsonify(updated)
            
        else:
            updated_amount_inr = False
            updated_amount_usd = False
            updated_usd_rate = False

            for k in ["entry_date", "ref_number", "description", "amount_inr", "amount_usd", "usd_rate", "dr_cr", "payment_mode", "bank_ref", "due_date", "notes"]:
                if k in req:
                    if k == "amount_inr":
                        amount_inr = _safe_float(req[k], None)
                        if amount_inr is None or amount_inr <= 0:
                            return jsonify({"error": "amount_inr must be positive and non-zero"}), 400
                        entry[k] = amount_inr
                        updated_amount_inr = True
                    elif k == "amount_usd":
                        amount_usd = _safe_float(req[k], None)
                        if amount_usd is None or amount_usd <= 0:
                            return jsonify({"error": "amount_usd must be positive and non-zero"}), 400
                        entry[k] = amount_usd
                        updated_amount_usd = True
                    elif k == "usd_rate":
                        usd_rate = _safe_float(req[k], None)
                        if usd_rate is None or usd_rate <= 0:
                            return jsonify({"error": "usd_rate must be positive and non-zero"}), 400
                        entry[k] = usd_rate
                        updated_usd_rate = True
                    elif k == "dr_cr":
                        if req[k].upper() not in ("DR", "CR"):
                            return jsonify({"error": "dr_cr must be DR or CR"}), 400
                        entry[k] = req[k].upper()
                    else:
                        entry[k] = req[k]

            if updated_amount_inr or updated_usd_rate:
                entry["amount_usd"] = round(entry["amount_inr"] / entry["usd_rate"], 4)
            elif updated_amount_usd:
                entry["amount_inr"] = round(entry["amount_usd"] * entry["usd_rate"], 2)

            cursor.execute("""
                UPDATE customer_ledger_entries
                SET entry_date = ?, ref_number = ?, description = ?, amount_inr = ?,
                    amount_usd = ?, usd_rate = ?, dr_cr = ?, payment_mode = ?,
                    bank_ref = ?, due_date = ?, notes = ?
                WHERE id = ?
            """, (
                entry["entry_date"], entry["ref_number"], entry["description"], entry["amount_inr"],
                entry["amount_usd"], entry["usd_rate"], entry["dr_cr"], entry["payment_mode"],
                entry["bank_ref"], entry["due_date"], entry["notes"], eid
            ))
            
            cursor.execute("SELECT * FROM customer_ledger_entries WHERE id = ?", (eid,))
            updated = dict(cursor.fetchone())
            
        return jsonify(updated)

@app.route("/api/ledger/customer/<eid>", methods=["DELETE"])
@require_permission("customer_delete")
def delete_customer_ledger_entry(eid):
    """Soft-delete a customer ledger entry, and remove associated invoice."""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM customer_ledger_entries WHERE id = ? AND deleted_at IS NULL", (eid,))
        row = cursor.fetchone()
        if not row:
            return jsonify({"error": "Ledger entry not found"}), 404
            
        cursor.execute("BEGIN TRANSACTION")
        try:
            cursor.execute("UPDATE customer_ledger_entries SET deleted_at = CURRENT_TIMESTAMP WHERE id = ?", (eid,))
            
            if row["entry_type"] == "INVOICE":
                cursor.execute("DELETE FROM customer_invoices WHERE invoice_no = ? AND customer_id = ?", (row["ref_number"], row["customer_id"]))
                
            cursor.execute("COMMIT")
        except Exception as e:
            cursor.execute("ROLLBACK")
            return jsonify({"error": str(e)}), 500
        
    return jsonify({"ok": True})


@app.route("/api/ledger/supplier/<eid>", methods=["PUT"])
@require_permission("supplier_edit")
def update_supplier_ledger_entry_v2(eid):
    """Edit supplier ledger entry. Locked if older than 7 days."""
    req = request.get_json(silent=True) or {}
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM supplier_ledger_entries WHERE id = ? AND deleted_at IS NULL", (eid,))
        row = cursor.fetchone()
        if not row:
            return jsonify({"error": "Ledger entry not found"}), 404
            
        created_str = row["created_at"] or row["entry_date"]
        if is_older_than_7_days(created_str):
            return jsonify({"error": "Entries older than 7 days are locked and cannot be edited"}), 400
            
        entry = dict(row)
        updated_amount_usd = False
        updated_amount_inr = False

        for k in ["entry_date", "ref_number", "description", "amount_usd", "amount_cny", "amount_inr", "usd_rate", "cny_rate", "dr_cr", "payment_mode", "bank_ref", "notes"]:
            if k in req:
                if k in ("amount_usd", "amount_cny"):
                    amount = _safe_float(req[k], None)
                    if amount is None or amount <= 0:
                        return jsonify({"error": "Amount must be positive and non-zero"}), 400
                    entry["amount_usd"] = amount
                    updated_amount_usd = True
                elif k == "amount_inr":
                    amount = _safe_float(req[k], None)
                    if amount is None or amount <= 0:
                        return jsonify({"error": "amount_inr must be positive and non-zero"}), 400
                    entry["amount_inr"] = amount
                    updated_amount_inr = True
                elif k in ("usd_rate", "cny_rate"):
                    rate = _safe_float(req[k], None)
                    if rate is None or rate <= 0:
                        return jsonify({"error": "usd_rate must be positive and non-zero"}), 400
                    entry["usd_rate"] = rate
                elif k == "dr_cr":
                    if req[k].upper() not in ("DR", "CR"):
                        return jsonify({"error": "dr_cr must be DR or CR"}), 400
                    entry[k] = req[k].upper()
                else:
                    entry[k] = req[k]

        if updated_amount_usd:
            entry["amount_inr"] = round(entry["amount_usd"] * entry["usd_rate"], 2)
        elif updated_amount_inr:
            entry["amount_usd"] = round(entry["amount_inr"] / entry["usd_rate"], 4)
        
        cursor.execute("""
            UPDATE supplier_ledger_entries
            SET entry_date = ?, ref_number = ?, description = ?, amount_usd = ?,
                amount_inr = ?, usd_rate = ?, dr_cr = ?, payment_mode = ?,
                bank_ref = ?, notes = ?
            WHERE id = ?
        """, (
            entry["entry_date"], entry["ref_number"], entry["description"], entry["amount_usd"],
            entry["amount_inr"], entry["usd_rate"], entry["dr_cr"], entry["payment_mode"],
            entry["bank_ref"], entry["notes"], eid
        ))
        
        cursor.execute("SELECT * FROM supplier_ledger_entries WHERE id = ?", (eid,))
        updated = dict(cursor.fetchone())
        
    return jsonify(updated)

@app.route("/api/ledger/supplier/<eid>", methods=["DELETE"])
@require_permission("supplier_delete")
def delete_supplier_ledger_entry_v2(eid):
    """Soft-delete supplier ledger entry."""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM supplier_ledger_entries WHERE id = ? AND deleted_at IS NULL", (eid,))
        if not cursor.fetchone():
            return jsonify({"error": "Ledger entry not found"}), 404
            
        cursor.execute("UPDATE supplier_ledger_entries SET deleted_at = CURRENT_TIMESTAMP WHERE id = ?", (eid,))
        
    return jsonify({"ok": True})


# ═══════════════════════════════════════════════════════════════════════════════
# OUTSTANDING & SUMMARY API (FIFO AGEING)
# ═══════════════════════════════════════════════════════════════════════════════

def calculate_customer_fifo_ageing(entries, credit_days_default=30):
    from datetime import date, timedelta
    invoices = []
    payments = []
    
    for e in entries:
        if e.get('deleted_at') is not None: continue
        amt = _safe_float(e.get('amount_inr'), 0.0)
        
        due_date = e.get('due_date')
        if not due_date and e['entry_type'] in ('INVOICE', 'Sales Invoice', 'OPENING'):
            try:
                base = date.fromisoformat(e['entry_date'])
                due_date = (base + timedelta(days=credit_days_default)).isoformat()
            except:
                due_date = e['entry_date']
                
        if e['dr_cr'] == 'DR':
            invoices.append({
                'id': e['id'], 'date': e['entry_date'], 'due_date': due_date,
                'ref': e['ref_number'] or '', 'entry_type': e['entry_type'],
                'description': e['description'] or '', 'amount': amt, 'remaining': amt
            })
        elif e['dr_cr'] == 'CR':
            payments.append(amt)
            
    # Apply payments
    for pay_amt in payments:
        rem_pay = pay_amt
        for inv in invoices:
            if inv['remaining'] > 0:
                if rem_pay >= inv['remaining']:
                    rem_pay -= inv['remaining']
                    inv['remaining'] = 0.0
                else:
                    inv['remaining'] -= rem_pay
                    rem_pay = 0.0
                    break
                    
    today = date.today()
    summary = {
        'total_due': 0.0, 'not_due': 0.0, 'days_1_30': 0.0, 'days_31_60': 0.0,
        'days_61_90': 0.0, 'days_90_plus': 0.0, 'overdue_total': 0.0, 'unpaid_invoices': [],
        'total_advance': 0.0
    }
    
    total_invoices = sum(inv['amount'] for inv in invoices)
    total_payments = sum(payments)
    if total_payments > total_invoices:
        summary['total_advance'] = round(total_payments - total_invoices, 2)
    
    for inv in invoices:
        if inv['remaining'] <= 0.01: continue
        rem = round(inv['remaining'], 2)
        summary['total_due'] += rem
        
        due = inv['due_date']
        if not due:
            summary['not_due'] += rem
            inv['days_overdue'] = 0
            inv['bucket'] = 'not_due'
        else:
            try:
                due_dt = date.fromisoformat(due)
                diff_days = (today - due_dt).days
            except:
                diff_days = 0
                
            inv['days_overdue'] = diff_days
            
            if diff_days <= 0:
                summary['not_due'] += rem
                inv['bucket'] = 'not_due'
            elif 1 <= diff_days <= 30:
                summary['days_1_30'] += rem
                summary['overdue_total'] += rem
                inv['bucket'] = '1-30'
            elif 31 <= diff_days <= 60:
                summary['days_31_60'] += rem
                summary['overdue_total'] += rem
                inv['bucket'] = '31-60'
            elif 61 <= diff_days <= 90:
                summary['days_61_90'] += rem
                summary['overdue_total'] += rem
                inv['bucket'] = '61-90'
            else:
                summary['days_90_plus'] += rem
                summary['overdue_total'] += rem
                inv['bucket'] = '90+'
                
        inv['remaining'] = rem
        summary['unpaid_invoices'].append(inv)
        
    for k in ['total_due', 'not_due', 'days_1_30', 'days_31_60', 'days_61_90', 'days_90_plus', 'overdue_total']:
        summary[k] = round(summary[k], 2)
        
    return summary


def calculate_supplier_fifo_ageing(entries, credit_days_default=30):
    from datetime import date, timedelta
    invoices = []
    payments = []
    
    for e in entries:
        if e.get('deleted_at') is not None: continue
        amt = _safe_float(e.get('amount_usd'), 0.0)
        
        due_date = e.get('due_date')
        if not due_date and e['entry_type'] in ('INVOICE', 'Purchase Invoice', 'OPENING'):
            try:
                base = date.fromisoformat(e['entry_date'])
                due_date = (base + timedelta(days=credit_days_default)).isoformat()
            except:
                due_date = e['entry_date']
                
        if e['dr_cr'] == 'CR':
            invoices.append({
                'id': e['id'], 'date': e['entry_date'], 'due_date': due_date,
                'ref': e['ref_number'] or '', 'entry_type': e['entry_type'],
                'description': e['description'] or '', 'amount': amt, 'remaining': amt,
                'po_id': e.get('po_id')
            })
        elif e['dr_cr'] == 'DR':
            payments.append(amt)
            
    for pay_amt in payments:
        rem_pay = pay_amt
        for inv in invoices:
            if inv['remaining'] > 0:
                if rem_pay >= inv['remaining']:
                    rem_pay -= inv['remaining']
                    inv['remaining'] = 0.0
                else:
                    inv['remaining'] -= rem_pay
                    rem_pay = 0.0
                    break
                    
    today = date.today()
    summary = {
        'total_due': 0.0, 'not_due': 0.0, 'days_1_30': 0.0, 'days_31_60': 0.0,
        'days_61_90': 0.0, 'days_90_plus': 0.0, 'overdue_total': 0.0, 'unpaid_invoices': [],
        'total_advance': 0.0
    }
    
    total_invoices = sum(inv['amount'] for inv in invoices)
    total_payments = sum(payments)
    if total_payments > total_invoices:
        summary['total_advance'] = round(total_payments - total_invoices, 2)
    
    for inv in invoices:
        if inv['remaining'] <= 0.01: continue
        rem = round(inv['remaining'], 2)
        summary['total_due'] += rem
        
        due = inv['due_date']
        if not due:
            summary['not_due'] += rem
            inv['days_overdue'] = 0
            inv['bucket'] = 'not_due'
        else:
            try:
                due_dt = date.fromisoformat(due)
                diff_days = (today - due_dt).days
            except:
                diff_days = 0
                
            inv['days_overdue'] = diff_days
            
            if diff_days <= 0:
                summary['not_due'] += rem
                inv['bucket'] = 'not_due'
            elif 1 <= diff_days <= 30:
                summary['days_1_30'] += rem
                summary['overdue_total'] += rem
                inv['bucket'] = '1-30'
            elif 31 <= diff_days <= 60:
                summary['days_31_60'] += rem
                summary['overdue_total'] += rem
                inv['bucket'] = '31-60'
            elif 61 <= diff_days <= 90:
                summary['days_61_90'] += rem
                summary['overdue_total'] += rem
                inv['bucket'] = '61-90'
            else:
                summary['days_90_plus'] += rem
                summary['overdue_total'] += rem
                inv['bucket'] = '90+'
                
        inv['remaining'] = rem
        inv['pending_usd'] = rem
        inv['bill_date'] = inv['date']
        inv['age_label'] = f"{inv['days_overdue']} Days" if inv['days_overdue'] > 0 else "Not Due"
        summary['unpaid_invoices'].append(inv)
        
    for k in ['total_due', 'not_due', 'days_1_30', 'days_31_60', 'days_61_90', 'days_90_plus', 'overdue_total']:
        summary[k] = round(summary[k], 2)
        
    return summary


@app.route("/api/customers/outstanding", methods=["GET"])
def get_all_customers_outstanding():
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id, name, company, credit_days FROM customers WHERE active = 1")
        customers = [dict(r) for r in cursor.fetchall()]
        
        parties_list = []
        totals = { "total_due": 0.0, "not_due": 0.0, "days_1_30": 0.0, "days_31_60": 0.0, "days_61_90": 0.0, "days_90_plus": 0.0 }
        
        for c in customers:
            cursor.execute("""
                SELECT * FROM customer_ledger_entries
                WHERE customer_id = ? AND deleted_at IS NULL
                ORDER BY entry_date ASC, created_at ASC
            """, (c["id"],))
            entries = [dict(r) for r in cursor.fetchall()]
            if not entries: continue
                
            ageing = calculate_customer_fifo_ageing(entries, credit_days_default=c["credit_days"] or 30)
            if ageing["total_due"] > 0 or ageing.get("total_advance", 0) > 0:
                party_data = {
                    "id": c["id"], "name": c["name"], "company": c["company"],
                    "total_due": ageing["total_due"], "not_due": ageing["not_due"],
                    "days_1_30": ageing["days_1_30"], "days_31_60": ageing["days_31_60"],
                    "days_61_90": ageing["days_61_90"], "days_90_plus": ageing["days_90_plus"],
                    "total_advance": ageing.get("total_advance", 0.0)
                }
                parties_list.append(party_data)
                for k in totals:
                    if k in ageing: totals[k] += ageing[k]
                    
        for k in totals: totals[k] = round(totals[k], 2)
            
    return jsonify({ "parties": parties_list, "totals": totals, "as_of": str(date.today()) })


@app.route("/api/suppliers/outstanding", methods=["GET"])
def get_all_suppliers_outstanding():
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id, name, company FROM suppliers")
        suppliers = [dict(r) for r in cursor.fetchall()]

        cursor.execute("SELECT value FROM settings WHERE key = 'default_rmb_rate'")
        r2 = cursor.fetchone()
        
        parties_list = []
        totals = { "total_due": 0.0, "not_due": 0.0, "days_1_30": 0.0, "days_31_60": 0.0, "days_61_90": 0.0, "days_90_plus": 0.0 }
        
        for s in suppliers:
            cursor.execute("SELECT credit_days FROM supplier_payment_terms WHERE supplier_id = ?", (s["id"],))
            terms = cursor.fetchone()
            credit_days = terms["credit_days"] if terms else 30
            
            cursor.execute("""
                SELECT * FROM supplier_ledger_entries
                WHERE supplier_id = ? AND deleted_at IS NULL
                ORDER BY entry_date ASC, created_at ASC
            """, (s["id"],))
            entries = [dict(r) for r in cursor.fetchall()]
            if not entries: continue
                
            ageing = calculate_supplier_fifo_ageing(entries, credit_days_default=credit_days)
            if ageing["total_due"] > 0 or ageing.get("total_advance", 0) > 0:
                party_data = {
                    "id": s["id"], "name": s["name"], "company": s["company"],
                    "total_due": ageing["total_due"], "not_due": ageing["not_due"],
                    "days_1_30": ageing["days_1_30"], "days_31_60": ageing["days_31_60"],
                    "days_61_90": ageing["days_61_90"], "days_90_plus": ageing["days_90_plus"],
                    "total_advance": ageing.get("total_advance", 0.0)
                }
                parties_list.append(party_data)
                for k in totals: 
                    if k in ageing: totals[k] += ageing[k]
                    
        for k in totals: totals[k] = round(totals[k], 2)
            
    return jsonify({ "parties": parties_list, "totals": totals, "as_of": str(date.today()) })


@app.route("/api/customers/<cid>/outstanding", methods=["GET"])
def get_single_customer_outstanding(cid):
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT name, company, credit_days FROM customers WHERE id = ?", (cid,))
        cust = cursor.fetchone()
        if not cust: return jsonify({"error": "Customer not found"}), 404
            
        cursor.execute("""
            SELECT * FROM customer_ledger_entries
            WHERE customer_id = ? AND deleted_at IS NULL
            ORDER BY entry_date ASC, created_at ASC
        """, (cid,))
        entries = [dict(r) for r in cursor.fetchall()]
        
    credit_days = cust["credit_days"] or 30
    ageing = calculate_customer_fifo_ageing(entries, credit_days_default=credit_days)
    
    buckets = {
        "0_30": [inv for inv in ageing["unpaid_invoices"] if inv["bucket"] == "1-30" or inv["bucket"] == "not_due"],
        "31_60": [inv for inv in ageing["unpaid_invoices"] if inv["bucket"] == "31-60"],
        "61_90": [inv for inv in ageing["unpaid_invoices"] if inv["bucket"] == "61-90"],
        "91_plus": [inv for inv in ageing["unpaid_invoices"] if inv["bucket"] == "90+"]
    }
    
    return jsonify({
        "buckets": buckets, "as_of": str(date.today()),
        "summary": { k: ageing[k] for k in ['total_due','not_due','days_1_30','days_31_60','days_61_90','days_90_plus'] },
        "total_advance": ageing.get('total_advance', 0.0)
    })


@app.route("/api/suppliers/<sid>/outstanding/fifo", methods=["GET"])
def get_single_supplier_outstanding_fifo(sid):
    """FIFO based itemized outstanding for a specific supplier."""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT name, company FROM suppliers WHERE id = ?", (sid,))
        sup = cursor.fetchone()
        if not sup: return jsonify({"error": "Supplier not found"}), 404
            
        cursor.execute("SELECT credit_days FROM supplier_payment_terms WHERE supplier_id = ?", (sid,))
        terms = cursor.fetchone()
        credit_days = terms["credit_days"] if terms else 30
        
        cursor.execute("""
            SELECT * FROM supplier_ledger_entries
            WHERE supplier_id = ? AND deleted_at IS NULL
            ORDER BY entry_date ASC, created_at ASC
        """, (sid,))
        entries = [dict(r) for r in cursor.fetchall()]
        
    ageing = calculate_supplier_fifo_ageing(entries, credit_days_default=credit_days)
    
    buckets = {
        "0_30": [inv for inv in ageing["unpaid_invoices"] if inv["bucket"] == "1-30" or inv["bucket"] == "not_due"],
        "31_60": [inv for inv in ageing["unpaid_invoices"] if inv["bucket"] == "31-60"],
        "61_90": [inv for inv in ageing["unpaid_invoices"] if inv["bucket"] == "61-90"],
        "91_plus": [inv for inv in ageing["unpaid_invoices"] if inv["bucket"] == "90+"]
    }
    
    return jsonify({
        "buckets": buckets, "as_of": str(date.today()),
        "summary": { k: ageing[k] for k in ['total_due','not_due','days_1_30','days_31_60','days_61_90','days_90_plus'] },
        "total_advance": ageing.get('total_advance', 0.0)
    })


@app.route("/api/customers/summary", methods=["GET"])
def get_customer_summary():
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id, name, company, credit_days FROM customers WHERE active = 1")
        customers = [dict(r) for r in cursor.fetchall()]
        
        total_receivable = 0.0
        customers_with_balance = 0
        total_overdue = 0.0
        weighted_overdue_days_sum = 0.0
        total_overdue_invoices_amount = 0.0
        party_summaries = []
        
        for c in customers:
            cursor.execute("""
                SELECT * FROM customer_ledger_entries
                WHERE customer_id = ? AND deleted_at IS NULL
                ORDER BY entry_date ASC, created_at ASC
            """, (c["id"],))
            entries = [dict(r) for r in cursor.fetchall()]
            
            total_dr = sum(e["amount_inr"] for e in entries if e["dr_cr"] == "DR")
            total_cr = sum(e["amount_inr"] for e in entries if e["dr_cr"] == "CR")
            balance = round(total_dr - total_cr, 2)
            
            if balance != 0:
                customers_with_balance += 1
                total_receivable += balance
                
            ageing = calculate_customer_fifo_ageing(entries, credit_days_default=c["credit_days"] or 30)
            
            for inv in ageing["unpaid_invoices"]:
                if inv["days_overdue"] > 0:
                    weighted_overdue_days_sum += inv["days_overdue"] * inv["remaining"]
                    total_overdue_invoices_amount += inv["remaining"]
                    
            total_overdue += ageing["overdue_total"]
            party_summaries.append({
                "id": c["id"], "name": c["name"], "company": c["company"],
                "balance": balance, "suffix": "Dr" if balance >= 0 else "Cr",
                "ageing": { k: ageing[k] for k in ['total_due','not_due','days_1_30','days_31_60','days_61_90','days_90_plus'] }
            })
            
        avg_collection_days = 0
        if total_overdue_invoices_amount > 0:
            avg_collection_days = round(weighted_overdue_days_sum / total_overdue_invoices_amount, 1)
            
    return jsonify({
        "total_receivable": round(total_receivable, 2),
        "customers_with_balance": customers_with_balance,
        "overdue_amount": round(total_overdue, 2),
        "avg_collection_days": avg_collection_days,
        "parties": party_summaries
    })


@app.route("/api/suppliers/summary", methods=["GET"])
def get_supplier_summary_v2():
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id, name, company FROM suppliers")
        suppliers = [dict(r) for r in cursor.fetchall()]

        cursor.execute("SELECT value FROM settings WHERE key = 'default_rmb_rate'")
        r2 = cursor.fetchone()
        
        cursor.execute("SELECT value FROM settings WHERE key = 'default_rmb_rate'")
        rate_row = cursor.fetchone()
        cny_rate = _safe_float(rate_row['value'] if rate_row else None, 11.5)
        
        total_payable_usd = 0.0
        total_payable_inr = 0.0
        suppliers_with_balance = 0
        total_overdue_usd = 0.0
        total_overdue_inr = 0.0
        weighted_overdue_days_sum = 0.0
        total_overdue_invoices_amount = 0.0
        party_summaries = []
        
        for s in suppliers:
            cursor.execute("SELECT credit_days FROM supplier_payment_terms WHERE supplier_id = ?", (s["id"],))
            terms = cursor.fetchone()
            credit_days = terms["credit_days"] if terms else 30
            
            cursor.execute("""
                SELECT * FROM supplier_ledger_entries
                WHERE supplier_id = ? AND deleted_at IS NULL
                ORDER BY entry_date ASC, created_at ASC
            """, (s["id"],))
            entries = [dict(r) for r in cursor.fetchall()]
            
            total_cr = sum(e["amount_usd"] for e in entries if e["dr_cr"] == "CR")
            total_dr = sum(e["amount_usd"] for e in entries if e["dr_cr"] == "DR")
            balance_usd = round(total_cr - total_dr, 2)
            balance_inr = round(balance_usd * cny_rate, 2)
            
            if balance_usd != 0:
                suppliers_with_balance += 1
                total_payable_usd += balance_usd
                total_payable_inr += balance_inr
                
            ageing = calculate_supplier_fifo_ageing(entries, credit_days_default=credit_days)
            
            for inv in ageing["unpaid_invoices"]:
                if inv["days_overdue"] > 0:
                    weighted_overdue_days_sum += inv["days_overdue"] * inv["remaining"]
                    total_overdue_invoices_amount += inv["remaining"]
                    
            total_overdue_usd += ageing["overdue_total"]
            total_overdue_inr += ageing["overdue_total"] * cny_rate
            party_summaries.append({
                "id": s["id"], "name": s["name"], "company": s["company"],
                "balance_usd": balance_usd, "balance_inr": balance_inr,
                "suffix": "Cr" if balance_usd >= 0 else "Dr",
                "ageing": { k: ageing[k] for k in ['total_due','not_due','days_1_30','days_31_60','days_61_90','days_90_plus'] }
            })
            
        avg_payment_days = 0
        if total_overdue_invoices_amount > 0:
            avg_payment_days = round(weighted_overdue_days_sum / total_overdue_invoices_amount, 1)
            
    return jsonify({
        "total_payable_usd": round(total_payable_usd, 2),
        "total_payable_inr": round(total_payable_inr, 2),
        "suppliers_with_balance": suppliers_with_balance,
        "overdue_amount_usd": round(total_overdue_usd, 2),
        "overdue_amount_inr": round(total_overdue_inr, 2),
        "avg_payment_days": avg_payment_days,
        "parties": party_summaries
    })


@app.route("/api/customers/<cid>/ledger/pdf", methods=["GET"])
def customer_ledger_pdf(cid):
    """Generate a formal Account Statement PDF for a customer."""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM customers WHERE id = ?", (cid,))
        cust_row = cursor.fetchone()
        if not cust_row:
            return "Customer not found", 404
        cust = dict(cust_row)
        cursor.execute("""
            SELECT * FROM customer_ledger_entries
            WHERE customer_id = ? AND deleted_at IS NULL
            ORDER BY entry_date ASC, created_at ASC
        """, (cid,))
        entries = [dict(r) for r in cursor.fetchall()]
        cursor.execute("SELECT key, value FROM settings")
        settings = {r["key"]: r["value"] for r in cursor.fetchall()}

    entries, closing_balance = calc_customer_running_balance(entries)
    balance_suffix = "Dr" if closing_balance >= 0 else "Cr"
    company_name = settings.get("company_name", "")

    rows_html = ""
    for e in entries:
        dr_amt = f"₹ {e['amount_inr']:,.2f}" if e["dr_cr"] == "DR" else ""
        cr_amt = f"₹ {e['amount_inr']:,.2f}" if e["dr_cr"] == "CR" else ""
        bal = e.get("running_balance", 0)
        bal_sfx = e.get("balance_suffix", "Dr")
        row_color = {"INVOICE": "#e8f4fd", "PAYMENT": "#e8fdf0", "DEBIT_NOTE": "#fff3e0",
                     "CREDIT_NOTE": "#fce4ec", "ADJUSTMENT": "#f3f4f6", "OPENING": "#f9fafb"}.get(e["entry_type"], "#fff")
        rows_html += f'''<tr style="background:{row_color}">
          <td>{_e(e['entry_date'])}</td>
          <td><span style="font-size:9px;padding:2px 6px;border-radius:3px;background:#ddd">{_e(e['entry_type'])}</span></td>
          <td>{_e(e['ref_number'] or '')}</td>
          <td style="max-width:200px;font-size:10px">{_e(e['description'] or '')}</td>
          <td style="text-align:right;color:#1565c0">{_e(dr_amt)}</td>
          <td style="text-align:right;color:#c62828">{_e(cr_amt)}</td>
          <td style="text-align:right;font-weight:700">₹ {abs(bal):,.2f} {_e(bal_sfx)}</td>
        </tr>'''

    total_dr = sum(e["amount_inr"] for e in entries if e["dr_cr"] == "DR")
    total_cr = sum(e["amount_inr"] for e in entries if e["dr_cr"] == "CR")

    html = f'''<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8">
<title>Account Statement — {_e(cust.get('company',''))}</title>
<style>
  body{{font-family:'Segoe UI',Arial,sans-serif;font-size:11px;color:#111;margin:0;padding:24px 32px}}
  .hdr{{border-bottom:3px solid #1e3a8a;padding-bottom:12px;margin-bottom:16px;display:flex;justify-content:space-between;align-items:flex-start}}
  .co-name{{font-size:20px;font-weight:800;color:#1e3a8a}}.co-sub{{font-size:9px;color:#6b7280;margin-top:2px}}
  .stmt-title{{font-size:14px;font-weight:700;color:#111;text-align:right}}.stmt-meta{{font-size:10px;color:#6b7280;text-align:right}}
  .parties{{display:grid;grid-template-columns:1fr 1fr;gap:0;border:1px solid #e5e7eb;border-radius:6px;overflow:hidden;margin-bottom:14px}}
  .party{{padding:10px 14px}}.party:first-child{{background:#f9fafb;border-right:1px solid #e5e7eb}}
  .party-label{{font-size:8px;font-weight:700;letter-spacing:.1em;color:#6b7280;text-transform:uppercase;margin-bottom:6px}}
  table{{width:100%;border-collapse:collapse;font-size:10px}}
  thead th{{padding:7px 10px;background:#1e3a8a;color:#fff;font-size:9px;text-align:left;font-weight:600;text-transform:uppercase;letter-spacing:.05em}}
  tbody td{{padding:6px 10px;border-bottom:1px solid #e5e7eb;vertical-align:middle}}
  tfoot td{{padding:7px 10px;background:#1e3a8a;color:#fff;font-weight:700;font-size:11px}}
  .closing{{margin-top:16px;padding:12px 16px;border-radius:8px;border:2px solid #1e3a8a;background:#eff6ff;text-align:right}}
  .closing-lbl{{font-size:10px;color:#6b7280;font-weight:700;text-transform:uppercase;letter-spacing:.06em}}
  .closing-val{{font-size:22px;font-weight:900;color:#1e3a8a;margin-top:4px}}
  .footer{{margin-top:16px;font-size:8px;color:#9ca3af;text-align:center;border-top:1px solid #e5e7eb;padding-top:8px}}
  @media print{{@page{{margin:10mm;size:A4}}body{{padding:0}}}}
</style></head><body>
<div class="hdr">
  <div><div class="co-name">{_e(company_name or 'Purchase Order')}</div>
  <div class="co-sub">Account Statement</div></div>
  <div><div class="stmt-title">CUSTOMER ACCOUNT STATEMENT</div>
  <div class="stmt-meta">As of {date.today().strftime('%d %b %Y')}</div></div>
</div>
<div class="parties">
  <div class="party"><div class="party-label">Statement For</div>
  <div style="font-weight:700;font-size:13px">{_e(cust.get('company',''))}</div>
  <div style="color:#4b5563">{_e(cust.get('name',''))}</div>
  <div style="color:#4b5563">{_e(cust.get('city',''))}, {_e(cust.get('state',''))}</div>
  <div style="color:#4b5563">GSTIN: {_e(cust.get('gstin',''))}</div></div>
  <div class="party"><div class="party-label">Prepared By</div>
  <div style="font-weight:700">{_e(company_name)}</div></div>
</div>
<table><thead><tr>
  <th style="width:80px">Date</th><th style="width:90px">Type</th>
  <th style="width:110px">Reference</th><th>Narration</th>
  <th style="width:100px;text-align:right">Dr (INR)</th>
  <th style="width:100px;text-align:right">Cr (INR)</th>
  <th style="width:110px;text-align:right">Balance</th>
</tr></thead>
<tbody>{rows_html}</tbody>
<tfoot><tr>
  <td colspan="4" style="text-align:right">TOTAL</td>
  <td style="text-align:right">₹ {total_dr:,.2f}</td>
  <td style="text-align:right">₹ {total_cr:,.2f}</td>
  <td style="text-align:right">₹ {abs(closing_balance):,.2f} {_e(balance_suffix)}</td>
</tr></tfoot></table>
<div class="closing">
  <div class="closing-lbl">Closing Balance</div>
  <div class="closing-val">₹ {abs(closing_balance):,.2f} <span style="font-size:14px">{_e(balance_suffix)}</span></div>
</div>
<div class="footer">Computer-generated Account Statement · {_e(company_name)} · {date.today().strftime('%d %b %Y')}</div>
</body></html>'''

    try:
        if WEASYPRINT_AVAILABLE:
            pdf_bytes = HTML(string=html, base_url=None).write_pdf()
            return send_file(BytesIO(pdf_bytes), mimetype="application/pdf", as_attachment=False,
                             download_name=f"Statement_{cust.get('company','customer')}_{date.today()}.pdf")
    except Exception:
        pass
    return html, 200, {"Content-Type": "text/html; charset=utf-8"}


@app.route("/customer-book")
@require_permission("customer_books")
def customer_book():
    return render_template("customer_book.html")


@app.route("/supplier-book")
@require_permission("supplier_books")
def supplier_book():
    return render_template("supplier_book.html")



# ── ADMIN RBAC ──────────────────────────────────────────────────────────────

@app.route("/admin/rbac")
@require_permission("admin_rbac")
def admin_rbac():
    return render_template("admin_roles.html")

@app.route("/api/admin/roles", methods=["GET"])
@require_permission("admin_rbac")
def get_roles():
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM roles")
        roles = [dict(r) for r in cursor.fetchall()]
        for r in roles:
            cursor.execute("SELECT permission_id FROM role_permissions WHERE role_id = ?", (r["id"],))
            r["permissions"] = [p["permission_id"] for p in cursor.fetchall()]
    return jsonify(roles)

@app.route("/api/admin/permissions", methods=["GET"])
@require_permission("admin_rbac")
def get_permissions():
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM permissions")
        perms = [dict(r) for r in cursor.fetchall()]
    return jsonify(perms)

@app.route("/api/admin/users", methods=["GET"])
@require_permission("admin_rbac")
def get_users():
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id, username, role_id, created_at FROM users")
        users = [dict(r) for r in cursor.fetchall()]
    return jsonify(users)

@app.route("/api/admin/roles", methods=["POST"])
@require_permission("admin_rbac")
def save_role():
    req = request.get_json()
    role_id = req.get("id") or str(uuid.uuid4())
    name = req.get("name")
    perms = req.get("permissions", [])
    
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM roles WHERE id = ?", (role_id,))
        if cursor.fetchone():
            cursor.execute("UPDATE roles SET name = ?, description = ? WHERE id = ?", (name, req.get("description", ""), role_id))
        else:
            cursor.execute("INSERT INTO roles (id, name, description) VALUES (?, ?, ?)", (role_id, name, req.get("description", "")))
        
        cursor.execute("DELETE FROM role_permissions WHERE role_id = ?", (role_id,))
        for p in perms:
            cursor.execute("INSERT INTO role_permissions (role_id, permission_id) VALUES (?, ?)", (role_id, p))
    return jsonify({"success": True, "id": role_id})

@app.route("/api/admin/users", methods=["POST"])
@require_permission("admin_rbac")
def save_user():
    req = request.get_json()
    user_id = req.get("id") or str(uuid.uuid4())
    username = req.get("username")
    role_id = req.get("role_id")
    password = req.get("password")
    
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM users WHERE id = ?", (user_id,))
        exists = cursor.fetchone()
        
        if password:
            hashed = hashlib.sha256(password.encode()).hexdigest()
            if exists:
                cursor.execute("UPDATE users SET username = ?, password_hash = ?, role_id = ? WHERE id = ?", (username, hashed, role_id, user_id))
            else:
                cursor.execute("INSERT INTO users (id, username, password_hash, role_id) VALUES (?, ?, ?, ?)", (user_id, username, hashed, role_id))
        else:
            if exists:
                cursor.execute("UPDATE users SET username = ?, role_id = ? WHERE id = ?", (username, role_id, user_id))
            else:
                return jsonify({"error": "Password required for new user"}), 400
                
    return jsonify({"success": True, "id": user_id})

# ── GOOGLE DRIVE BACKUP (RCLONE) ──────────────────────────────────────────────
import threading
import subprocess
import re
import time
import os

backup_state = {
    "running": False,
    "lastStatus": None,
    "lastError": None,
    "lastTimestamp": None,
    "stage": None,
    "overallPct": 0
}

def parse_rclone_stats(line):
    result = {}
    
    # Transferred:   28.808 MiB / 206.701 MiB, 13%, 14.400 MiB/s, ETA 12s
    transfer_match = re.search(r'Transferred:\s+([\d.]+)\s*(\w+)\s*/\s*([\d.]+)\s*(\w+),\s*(\d+)%', line)
    if transfer_match:
        result['percentage'] = int(transfer_match.group(5))
        result['uploadedBytesLabel'] = f"{transfer_match.group(1)} {transfer_match.group(2)}"
        result['totalBytesLabel'] = f"{transfer_match.group(3)} {transfer_match.group(4)}"
        
    speed_match = re.search(r'([\d.]+)\s*(\w+)/s', line)
    if speed_match:
        result['speedLabel'] = f"{speed_match.group(1)} {speed_match.group(2)}/s"
        
    eta_match = re.search(r'ETA\s+(\S+)', line)
    if eta_match:
        eta_str = eta_match.group(1)
        seconds = 0
        h = re.search(r'(\d+)h', eta_str)
        m = re.search(r'(\d+)m', eta_str)
        s = re.search(r'(\d+)s', eta_str)
        if h: seconds += int(h.group(1)) * 3600
        if m: seconds += int(m.group(1)) * 60
        if s: seconds += int(s.group(1))
        result['etaSeconds'] = seconds
        
    # Transferred:   5 / 100, 5%
    file_match = re.search(r'Transferred:\s+(\d+)\s*/\s*(\d+),', line)
    if file_match:
        result['processedFiles'] = int(file_match.group(1))
        result['totalFiles'] = int(file_match.group(2))
        
    return result

def run_rclone_with_progress(args, update_state_cb):
    try:
        rclone_path = os.path.join(os.path.dirname(__file__), 'rclone.exe')
        if not os.path.exists(rclone_path):
            rclone_path = 'rclone' # fallback to PATH
            
        proc = subprocess.Popen(
            [rclone_path] + args,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            bufsize=1
        )
        
        last_stats = {}
        for line in iter(proc.stdout.readline, ''):
            if not line: break
            clean_line = re.sub(r'\x1b\[[0-9;]*m', '', line).strip()
            if not clean_line: continue
            
            stats = parse_rclone_stats(clean_line)
            if stats:
                last_stats.update(stats)
                update_state_cb(last_stats)
                
        proc.wait()
        if proc.returncode != 0:
            raise Exception(f"Rclone exited with code {proc.returncode}. Is rclone configured properly with the 'gdrive' remote?")
    except FileNotFoundError:
        raise Exception("Failed to start rclone: 'rclone' executable not found in PATH. Please install rclone.")
    except Exception as e:
        raise Exception(f"Failed to start rclone: {str(e)}")

def perform_backup():
    base_dir = os.path.dirname(__file__)
    db_path = os.path.join(base_dir, "data", "po", "database.db")
    uploads_path = os.path.join(base_dir, "data", "po", "attachments")
    
    def update_stage_1(stats):
        pct = stats.get('percentage', 0)
        overall = min(5, int(pct * 0.05))
        backup_state.update(stats)
        backup_state.update({"overallPct": overall, "stageLabel": "Stage 1 / 2"})

    def update_stage_2(stats):
        pct = stats.get('percentage', 0)
        overall = 5 + int(pct * 0.95)
        backup_state.update(stats)
        backup_state.update({"overallPct": overall, "stageLabel": "Stage 2 / 2"})

    try:
        backup_state['stage'] = "Backing up database..."
        update_stage_1({"percentage": 0})
        # Try to backup DB
        run_rclone_with_progress([
            '--stats', '1s', '--stats-one-line', 'copy', db_path, 'gdrive:PurchaseOrderBackups/db'
        ], update_stage_1)
        
        backup_state['stage'] = "Uploading attachments..."
        update_stage_2({"percentage": 0})
        # Try to backup Attachments
        if os.path.exists(uploads_path):
            run_rclone_with_progress([
                '--stats', '1s', '--stats-one-line', 'copy', uploads_path, 'gdrive:PurchaseOrderBackups/uploads'
            ], update_stage_2)
            
        backup_state['running'] = False
        backup_state['lastStatus'] = 'success'
        backup_state['lastTimestamp'] = int(time.time() * 1000)
        backup_state['overallPct'] = 100
        backup_state['stage'] = 'Backup completed successfully!'
    except Exception as e:
        backup_state['running'] = False
        backup_state['lastStatus'] = 'failed'
        backup_state['lastError'] = str(e)
        backup_state['overallPct'] = 0

@app.route("/api/admin/backup", methods=["POST"])
@require_permission("admin_rbac")
def trigger_backup():
    if backup_state.get("running"):
        return jsonify({"success": True, "message": "Backup already in progress", "running": True})
        
    backup_state.update({
        "running": True, "lastStatus": "running", "lastError": None,
        "lastTimestamp": None, "stage": "Starting...", "overallPct": 0
    })
    
    t = threading.Thread(target=perform_backup)
    t.daemon = True
    t.start()
    
    return jsonify({"success": True, "message": "Backup started.", "running": True})

@app.route("/api/admin/backup/status", methods=["GET"])
@require_permission("admin_rbac")
def backup_status():
    return jsonify({
        "success": True,
        "running": backup_state.get("running"),
        "status": backup_state.get("lastStatus"),
        "error": backup_state.get("lastError"),
        "timestamp": backup_state.get("lastTimestamp"),
        "stage": backup_state.get("stage"),
        "stageLabel": backup_state.get("stageLabel"),
        "overallPct": backup_state.get("overallPct"),
        "processedFiles": backup_state.get("processedFiles"),
        "totalFiles": backup_state.get("totalFiles"),
        "uploadedBytesLabel": backup_state.get("uploadedBytesLabel"),
        "totalBytesLabel": backup_state.get("totalBytesLabel"),
        "speedLabel": backup_state.get("speedLabel"),
        "etaSeconds": backup_state.get("etaSeconds")
    })

if __name__ == "__main__":
    app.run(debug=True, port=8090, host="0.0.0.0")

