"""
Import Portal — Unified Flask App (SQLite Version)
Replaces JSON storage with SQLite for performance & safety
Port: 5005
"""

import csv
import os, json, uuid, shutil, base64, re, html, urllib.error, urllib.request, mimetypes, sqlite3, time, traceback
from datetime import date, datetime
from flask import Flask, render_template, request, jsonify, send_file
import io
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    WEASYPRINT_AVAILABLE = False
from io import BytesIO

from database import get_db, get_connection, init_db, init_default_settings

app = Flask(__name__, template_folder="templates")
app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024  # 20 MB

# Initialize database on startup
init_db()

DATA_DIR   = os.path.join(os.path.dirname(__file__), "data", "po")
ATTACH_DIR = os.path.join(DATA_DIR, "attachments")
os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(ATTACH_DIR, exist_ok=True)

GEMINI_URL = (
    "https://generativelanguage.googleapis.com/v1beta/"
    "models/{model}:generateContent?key={key}"
)


# ── HELPER FUNCTIONS ──────────────────────────────────────────────────────────

def _call_gemini(api_key, model, prompt, file_data=None, mime_type=None):
    """Call Gemini API using urllib with retry logic for 503 errors"""
    url = GEMINI_URL.format(model=model, key=api_key)
    
    contents = []
    part_prompt = {"text": prompt}
    
    if file_data and mime_type:
        part_file = {
            "inline_data": {
                "mime_type": mime_type,
                "data": base64.b64encode(file_data).decode("utf-8")
            }
        }
        contents.append({"parts": [part_prompt, part_file]})
    else:
        contents.append({"parts": [part_prompt]})

    data = json.dumps({"contents": contents}).encode("utf-8")
    
    max_retries = 3
    retry_delay = 2 # seconds
    
    for attempt in range(max_retries):
        req = urllib.request.Request(
            url,
            data=data,
            headers={"Content-Type": "application/json"}
        )
        
        try:
            with urllib.request.urlopen(req) as response:
                res_data = json.loads(response.read().decode("utf-8"))
                return res_data
        except urllib.error.HTTPError as e:
            error_body = e.read().decode("utf-8")
            # If 429 (Too Many Requests) or 503 (Service Unavailable) and we have retries left, wait and try again
            if e.code in (429, 503) and attempt < max_retries - 1:
                # For 429, we might want to wait longer
                wait_time = retry_delay * (attempt + 1)
                if e.code == 429: wait_time *= 2 
                time.sleep(wait_time)
                continue
            
            # User friendly error message for common codes
            msg = f"Gemini API Error: {e.code}"
            try:
                err_json = json.loads(error_body)
                msg = err_json.get("error", {}).get("message", error_body)
                if e.code == 429:
                    msg = "AI Rate Limit / Quota exceeded. Please wait 60 seconds and try again."
            except: pass
            raise Exception(msg)
        except Exception as e:
            raise Exception(f"Failed to call Gemini: {str(e)}")

def _e(val):
    if val is None: return ""
    return html.escape(str(val))

def fmt_inr(n):
    return "₹" + f"{int(round(float(n))):,}"


def _sync_po_sequence(cursor, po_number):
    """Update po_sequence in settings if the given PO number has a higher sequence.
    This ensures the auto-numbering stays in sync with the latest created or edited PO.
    """
    if not po_number:
        return
    seq_match = re.search(r'(\d+)$', po_number)
    if seq_match:
        new_seq = int(seq_match.group(1))
        cursor.execute("SELECT value FROM settings WHERE key = 'po_sequence'")
        row = cursor.fetchone()
        current_seq = int(row[0]) if row and row[0].isdigit() else 0

        if new_seq > current_seq:
            cursor.execute(
                "INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)",
                ("po_sequence", str(new_seq))
            )

def _log_status_change(cursor, po_id, from_status, to_status, note=None, force=False):
    """Insert one row into po_status_log. Call within an open transaction."""
    import uuid as _uuid
    if from_status == to_status and not force:
        return  # no change, nothing to log
    cursor.execute(
        """INSERT INTO po_status_log (id, po_id, from_status, to_status, note)
           VALUES (?, ?, ?, ?, ?)""",
        (str(_uuid.uuid4()), po_id, from_status, to_status, note)
    )

def calc_landed(items, rate, bank, ship, duty, trans, gst_duty=0, doc_pct=0):
    # Filter out items with zero or negative qty before any calculation
    items = [it for it in items if float(it.get("qty", 0)) > 0]
    if not items:
        return {"items": [], "inv_usd": 0, "inv_inr": 0, "total_addl": 0,
                "grand": 0, "bank": bank, "ship": ship, "duty": duty,
                "trans": trans, "gst_duty": gst_duty, "doc_pct": doc_pct,
                "doc_inr": 0, "rate": rate}

    inv_usd    = sum(float(it["qty"]) * float(it.get("unitPrice", it.get("unit_price", 0))) for it in items)
    inv_inr    = inv_usd * rate
    doc_inr    = inv_inr * (doc_pct / 100.0)
    total_addl = bank + ship + duty + trans + gst_duty + doc_inr
    grand      = inv_inr + total_addl
    result = []
    for it in items:
        qty = float(it.get("qty", 0))
        price = float(it.get("unitPrice", it.get("unit_price", 0)))
        item_inr   = qty * price * rate
        share      = (item_inr / inv_inr) if inv_inr > 0 else 0
        addl_share = total_addl * share
        total_item = item_inr + addl_share
        per_unit   = (total_item / qty) if qty > 0 else 0
        result.append({
            **it,
            "item_inr":   round(item_inr, 2),
            "share":      round(share * 100, 2),
            "addl_share": round(addl_share, 2),
            "total_item": round(total_item, 2),
            "per_unit":   round(per_unit, 2),
            "bank_s":     round(bank * share, 2),
            "ship_s":     round(ship * share, 2),
            "duty_s":     round(duty * share, 2),
            "trans_s":    round(trans * share, 2),
            "gst_s":      round(gst_duty * share, 2),
            "doc_s":      round(doc_inr * share, 2),
        })
    return {
        "items":      result,
        "inv_usd":    round(inv_usd, 2),
        "inv_inr":    round(inv_inr, 2),
        "total_addl": round(total_addl, 2),
        "grand":      round(grand, 2),
        "bank":       round(bank, 2),
        "ship":       round(ship, 2),
        "duty":       round(duty, 2),
        "trans":      round(trans, 2),
        "gst_duty":   round(gst_duty, 2),
        "doc_pct":    round(doc_pct, 2),
        "doc_inr":    round(doc_inr, 2),
        "rate":       rate,
    }


# ── ROUTES ────────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("po_index.html")


# ═══════════════════════════════════════════════════════════════════════════════
# SUPPLIERS
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/api/suppliers", methods=["GET"])
def get_suppliers():
    """Fetch all suppliers"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM suppliers ORDER BY name")
        suppliers = [dict(row) for row in cursor.fetchall()]
    return jsonify(suppliers)


@app.route("/api/suppliers", methods=["POST"])
def add_supplier():
    """Create new supplier"""
    req = request.get_json(silent=True)
    if not req:
        return jsonify({"error": "Invalid JSON"}), 400
    
    s_id = str(uuid.uuid4())
    supplier = {
        "id": s_id,
        "name": req.get("name", "").strip(),
        "company": req.get("company", "").strip(),
        "address": req.get("address", "").strip(),
        "country": req.get("country", "China").strip(),
        "email": req.get("email", "").strip(),
        "phone": req.get("phone", "").strip(),
        "wechat": req.get("wechat", "").strip(),
        "bank_name": req.get("bank_name", "").strip(),
        "bank_account": req.get("bank_account", "").strip(),
        "swift_code": req.get("swift_code", "").strip(),
    }
    
    if not supplier["name"]:
        return jsonify({"error": "Supplier name is required"}), 400
    
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO suppliers (id, name, company, address, country, email, phone, wechat, bank_name, bank_account, swift_code)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                supplier["id"],
                supplier["name"],
                supplier["company"],
                supplier["address"],
                supplier["country"],
                supplier["email"],
                supplier["phone"],
                supplier["wechat"],
                supplier["bank_name"],
                supplier["bank_account"],
                supplier["swift_code"],
            ))
    except sqlite3.IntegrityError:
        return jsonify({"error": f"Supplier '{supplier['name']}' already exists"}), 409
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    
    return jsonify(supplier), 201


@app.route("/api/suppliers/<sid>", methods=["PUT"])
def update_supplier(sid):
    """Update supplier"""
    req = request.get_json(silent=True)
    if not req:
        return jsonify({"error": "Invalid JSON"}), 400
    
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            
            # Fetch existing supplier
            cursor.execute("SELECT * FROM suppliers WHERE id = ?", (sid,))
            row = cursor.fetchone()
            if not row:
                return jsonify({"error": "Not found"}), 404
            
            supplier = dict(row)
            
            # Update fields
            for k in ["name", "company", "address", "country", "email", "phone", "wechat", "bank_name", "bank_account", "swift_code"]:
                if k in req:
                    supplier[k] = req[k]
            
            cursor.execute("""
                UPDATE suppliers
                SET name = ?, company = ?, address = ?, country = ?, email = ?, phone = ?, wechat = ?, bank_name = ?, bank_account = ?, swift_code = ?
                WHERE id = ?
            """, (
                supplier["name"], supplier["company"], supplier["address"], supplier["country"],
                supplier["email"], supplier["phone"], supplier["wechat"], supplier["bank_name"],
                supplier["bank_account"], supplier["swift_code"], sid
            ))
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    
    return jsonify(supplier)


@app.route("/api/suppliers/<sid>", methods=["DELETE"])
def delete_supplier(sid):
    """Delete supplier (check for usage in POs and Quotations)"""
    with get_db() as conn:
        cursor = conn.cursor()
        blockers = []
        
        # Check if supplier is used in any PO
        cursor.execute("SELECT po_number FROM purchase_orders WHERE supplier_id = ? AND deleted_at IS NULL", (sid,))
        used_pos = [row[0] for row in cursor.fetchall()]
        if used_pos:
            blockers.append(f"POs: {', '.join(used_pos)}")
        
        # O(1) lookup via quotation_refs index (no JSON scanning)
        cursor.execute("""
            SELECT q.quotation_number, q.title, q.id
            FROM quotation_refs r
            JOIN quotations q ON q.id = r.quotation_id
            WHERE r.ref_type = 'supplier' AND r.ref_id = ?
        """, (sid,))
        used_quotes = [
            row["quotation_number"] or row["title"] or row["id"][:8]
            for row in cursor.fetchall()
        ]
        if used_quotes:
            blockers.append(f"Quotations: {', '.join(used_quotes)}")
        
        if blockers:
            return jsonify({"error": f"Cannot delete — used in {' | '.join(blockers)}"}), 409
        
        cursor.execute("DELETE FROM suppliers WHERE id = ?", (sid,))
    
    return jsonify({"ok": True})


# ═══════════════════════════════════════════════════════════════════════════════
# ITEMS
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/api/items", methods=["GET"])
def get_items():
    """Fetch all items"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM items ORDER BY name")
        items = [dict(row) for row in cursor.fetchall()]
    return jsonify(items)


@app.route("/api/items", methods=["POST"])
def add_item():
    """Create new item"""
    req = request.get_json(silent=True)
    if not req:
        return jsonify({"error": "Invalid JSON"}), 400
    
    item_id = str(uuid.uuid4())
    item = {
        "id": item_id,
        "name": req.get("name", "").strip(),
        "description": req.get("description", "").strip(),
        "hs_code": req.get("hs_code", "").strip(),
        "unit": req.get("unit", "PCS").strip(),
        "currency": req.get("currency", "CNY").strip(),
        "default_price_usd": float(req.get("default_price_usd", 0)),
    }
    
    if not item["name"]:
        return jsonify({"error": "Item name is required"}), 400
    
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO items (id, name, description, hs_code, unit, currency, default_price_usd)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (
                item["id"], item["name"], item["description"], item["hs_code"],
                item["unit"], item["currency"], item["default_price_usd"]
            ))
    except sqlite3.IntegrityError:
        return jsonify({"error": f"Item '{item['name']}' already exists"}), 409
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    
    return jsonify(item), 201


@app.route("/api/items/<iid>", methods=["PUT"])
def update_item(iid):
    """Update item"""
    req = request.get_json(silent=True)
    if not req:
        return jsonify({"error": "Invalid JSON"}), 400
    
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            
            cursor.execute("SELECT * FROM items WHERE id = ?", (iid,))
            row = cursor.fetchone()
            if not row:
                return jsonify({"error": "Not found"}), 404
            
            item = dict(row)
            
            # Update fields
            item["name"] = req.get("name", item["name"])
            item["description"] = req.get("description", item["description"])
            item["hs_code"] = req.get("hs_code", item["hs_code"])
            item["unit"] = req.get("unit", item["unit"])
            item["currency"] = req.get("currency", item.get("currency", "CNY"))
            item["default_price_usd"] = float(req.get("default_price_usd", item["default_price_usd"]))
            
            cursor.execute("""
                UPDATE items
                SET name = ?, description = ?, hs_code = ?, unit = ?, currency = ?, default_price_usd = ?
                WHERE id = ?
            """, (
                item["name"], item["description"], item["hs_code"], item["unit"],
                item["currency"], item["default_price_usd"], iid
            ))
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    
    return jsonify(item)


@app.route("/api/items/<iid>", methods=["DELETE"])
def delete_item(iid):
    """Delete item (check for usage in POs and Quotations)"""
    with get_db() as conn:
        cursor = conn.cursor()
        blockers = []
        
        # Check if item is used in any PO line item
        cursor.execute("SELECT DISTINCT po_id FROM po_items WHERE item_id = ?", (iid,))
        po_ids = [row[0] for row in cursor.fetchall()]
        if po_ids:
            cursor.execute(f"SELECT po_number FROM purchase_orders WHERE id IN ({','.join(['?']*len(po_ids))})", po_ids)
            used_pos = [row[0] for row in cursor.fetchall()]
            blockers.append(f"POs: {', '.join(used_pos)}")
        
        # O(1) lookup via quotation_refs index (no JSON scanning)
        cursor.execute("""
            SELECT q.quotation_number, q.title, q.id
            FROM quotation_refs r
            JOIN quotations q ON q.id = r.quotation_id
            WHERE r.ref_type = 'item' AND r.ref_id = ?
        """, (iid,))
        used_quotes = [
            row["quotation_number"] or row["title"] or row["id"][:8]
            for row in cursor.fetchall()
        ]
        if used_quotes:
            blockers.append(f"Quotations: {', '.join(used_quotes)}")
        
        if blockers:
            return jsonify({"error": f"Cannot delete — used in {' | '.join(blockers)}"}), 409
        
        cursor.execute("DELETE FROM items WHERE id = ?", (iid,))
    
    return jsonify({"ok": True})


# ═══════════════════════════════════════════════════════════════════════════════
# FORWARDERS
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/api/forwarders", methods=["GET"])
def get_forwarders():
    """Fetch all forwarders"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM forwarders ORDER BY name")
        forwarders = []
        for row in cursor.fetchall():
            fw = dict(row)
            cursor.execute(
                "SELECT id, name, contact_person, phone, email, address "
                "FROM forwarder_godowns WHERE forwarder_id = ? ORDER BY sort_order",
                (fw["id"],)
            )
            godown_rows = cursor.fetchall()
            if godown_rows:
                fw["godowns"] = [
                    {
                        "id":             r["id"],
                        "label":          r["name"],        # 'name' col stores the label
                        "contact_person": r["contact_person"] or "",
                        "phone":          r["phone"] or "",
                        "email":          r["email"] or "",
                        "address":        r["address"] or "",
                    }
                    for r in godown_rows
                ]
            else:
                # Fallback: try legacy JSON blob in forwarders.godowns column
                raw = json.loads(fw.get("godowns") or "[]")
                # Normalise — old format was plain strings, new format is objects
                fw["godowns"] = [
                    g if isinstance(g, dict) else {"id": str(uuid.uuid4()), "label": g,
                                                   "contact_person": "", "phone": "",
                                                   "email": "", "address": ""}
                    for g in raw
                ]
            forwarders.append(fw)
    return jsonify(forwarders)


@app.route("/api/forwarders", methods=["POST"])
def add_forwarder():
    """Create new forwarder"""
    req = request.get_json(silent=True)
    if not req:
        return jsonify({"error": "Invalid JSON"}), 400
    
    fw_id = str(uuid.uuid4())
    forwarder = {
        "id": fw_id,
        "name": req.get("name", "").strip(),
        "contact_person": req.get("contact_person", "").strip(),
        "phone": req.get("phone", "").strip(),
        "email": req.get("email", "").strip(),
        "godowns": req.get("godowns", []),
    }
    
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO forwarders (id, name, contact_person, phone, email, godowns)
                VALUES (?, ?, ?, ?, ?, ?)
            """, (
                forwarder["id"], forwarder["name"], forwarder["contact_person"],
                forwarder["phone"], forwarder["email"], json.dumps(forwarder["godowns"])
            ))
            for i, g in enumerate(forwarder["godowns"]):
                # g is a rich object: {id, label, contact_person, phone, email, address}
                gid = g.get("id") or str(uuid.uuid4())
                cursor.execute(
                    "INSERT INTO forwarder_godowns "
                    "(id, forwarder_id, name, contact_person, phone, email, address, sort_order) "
                    "VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                    (gid, forwarder["id"],
                     g.get("label", "") or g.get("name", ""),
                     g.get("contact_person", ""),
                     g.get("phone", ""),
                     g.get("email", ""),
                     g.get("address", ""),
                     i)
                )
    except sqlite3.IntegrityError:
        return jsonify({"error": f"Forwarder '{forwarder['name']}' already exists"}), 409
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    
    return jsonify(forwarder), 201


@app.route("/api/forwarders/<fid>", methods=["PUT"])
def update_forwarder(fid):
    """Update forwarder"""
    req = request.get_json(silent=True)
    if not req:
        return jsonify({"error": "Invalid JSON"}), 400
    
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            
            cursor.execute("SELECT * FROM forwarders WHERE id = ?", (fid,))
            row = cursor.fetchone()
            if not row:
                return jsonify({"error": "Not found"}), 404
            
            forwarder = dict(row)
            forwarder["godowns"] = json.loads(forwarder.get("godowns") or "[]")
            
            forwarder["name"] = req.get("name", forwarder["name"])
            forwarder["contact_person"] = req.get("contact_person", forwarder.get("contact_person", ""))
            forwarder["phone"] = req.get("phone", forwarder["phone"])
            forwarder["email"] = req.get("email", forwarder.get("email", ""))
            forwarder["godowns"] = req.get("godowns", forwarder.get("godowns", []))
            
            cursor.execute("""
                UPDATE forwarders
                SET name = ?, contact_person = ?, phone = ?, email = ?, godowns = ?
                WHERE id = ?
            """, (
                forwarder["name"], forwarder["contact_person"], forwarder["phone"],
                forwarder["email"], json.dumps(forwarder["godowns"]), fid
            ))
            cursor.execute("DELETE FROM forwarder_godowns WHERE forwarder_id = ?", (fid,))
            for i, g in enumerate(forwarder["godowns"]):
                # g is a rich object: {id, label, contact_person, phone, email, address}
                gid = g.get("id") or str(uuid.uuid4())
                cursor.execute(
                    "INSERT INTO forwarder_godowns "
                    "(id, forwarder_id, name, contact_person, phone, email, address, sort_order) "
                    "VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                    (gid, fid,
                     g.get("label", "") or g.get("name", ""),
                     g.get("contact_person", ""),
                     g.get("phone", ""),
                     g.get("email", ""),
                     g.get("address", ""),
                     i)
                )
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    
    return jsonify(forwarder)


@app.route("/api/forwarders/<fid>", methods=["DELETE"])
def delete_forwarder(fid):
    """Delete forwarder (check for usage in POs)"""
    with get_db() as conn:
        cursor = conn.cursor()
        
        cursor.execute("SELECT po_number FROM purchase_orders WHERE forwarder_id = ? AND deleted_at IS NULL", (fid,))
        used_pos = [row[0] for row in cursor.fetchall()]
        
        if used_pos:
            return jsonify({"error": f"Cannot delete — used in POs: {', '.join(used_pos)}"}), 409
        
        cursor.execute("DELETE FROM forwarders WHERE id = ?", (fid,))
    
    return jsonify({"ok": True})


# ═══════════════════════════════════════════════════════════════════════════════
# PURCHASE ORDERS
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/api/po", methods=["GET"])
def get_pos():
    """Fetch all purchase orders"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM purchase_orders WHERE deleted_at IS NULL ORDER BY po_number DESC")
        
        purchase_orders = []
        for row in cursor.fetchall():
            po = dict(row)
            
            # Fetch line items for this PO
            cursor2 = conn.cursor()
            cursor2.execute("SELECT * FROM po_items WHERE po_id = ? ORDER BY line_sequence", (po["id"],))
            po["line_items"] = [
                {
                    "id": item["id"],
                    "item_id": item["item_id"],
                    "item_name": item["item_name"],
                    "description": item["description"],
                    "hs_code": item["hs_code"],
                    "qty": item["qty"],
                    "unit": item["unit"],
                    "unit_price": item["unit_price"],
                }
                for item in cursor2.fetchall()
            ]
            
            # Parse supplier snapshot
            po["supplier_snapshot"] = json.loads(po.get("supplier_snapshot") or "{}")
            
            purchase_orders.append(po)
    
    return jsonify(purchase_orders)


@app.route("/api/po", methods=["POST"])
def create_po():
    """Create new purchase order"""
    req = request.get_json(silent=True)
    if not req:
        return jsonify({"error": "Invalid JSON"}), 400
    
    po_id = str(uuid.uuid4())
    po = {
        "id": po_id,
        "po_number": req.get("po_number", ""),
        "po_date": req.get("po_date", str(date.today())),
        "supplier_id": req.get("supplier_id", ""),
        "godown_id": req.get("godown_id", ""),
        "supplier_snapshot": req.get("supplier_snapshot", {}),
        "payment_conditions": req.get("payment_conditions", ""),
        "delivery_terms": req.get("delivery_terms", "FOB"),
        "delivery_address": req.get("delivery_address", ""),
        "forwarder_id": req.get("forwarder_id", ""),
        "forwarder_name": req.get("forwarder_name", ""),
        "forwarder_contact": req.get("forwarder_contact", ""),
        "remarks": req.get("remarks", ""),
        "currency": req.get("currency", "CNY"),
        "line_items": req.get("line_items", []),
        "status": req.get("status", "Draft"),
        "created_at": str(date.today()),
        "lead_time_days": req.get("lead_time_days", 0),
        "due_date": req.get("due_date", ""),
        "lc_usd_rate": req.get("lc_usd_rate", 84.0),
        "lc_rmb_rate": req.get("lc_rmb_rate", 11.5),
        "lc_bank": req.get("lc_bank", 0.0),
        "lc_ship": req.get("lc_ship", 0.0),
        "lc_duty": req.get("lc_duty", 0.0),
        "lc_trans": req.get("lc_trans", 0.0),
        "lc_gst_duty": req.get("lc_gst_duty", 0.0),
        "lc_doc_pct": req.get("lc_doc_pct", 0.0),
    }
    
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            
            # Check if PO number already exists
            if po["po_number"].strip():
                cursor.execute("SELECT id FROM purchase_orders WHERE po_number = ?", (po["po_number"],))
                if cursor.fetchone():
                    return jsonify({"error": f"PO number '{po['po_number']}' already exists"}), 409

            # Validation: Supplier is required
            if not po["supplier_id"] or not po["supplier_id"].strip():
                return jsonify({"error": "Supplier is required"}), 400
            
            # Insert PO
            cursor.execute("""
                INSERT INTO purchase_orders (
                    id, po_number, po_date, supplier_id, supplier_snapshot,
                    payment_conditions, delivery_terms, delivery_address,
                    forwarder_id, forwarder_name, forwarder_contact, remarks,
                    currency, status, created_at, lead_time_days, due_date, godown_id,
                    lc_usd_rate, lc_rmb_rate, lc_bank, lc_ship, lc_duty, lc_trans, lc_gst_duty, lc_doc_pct
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                po["id"], po["po_number"], po["po_date"], po["supplier_id"],
                json.dumps(po["supplier_snapshot"]),
                po["payment_conditions"], po["delivery_terms"], po["delivery_address"],
                po["forwarder_id"], po["forwarder_name"], po["forwarder_contact"],
                po["remarks"], po["currency"], po["status"], po["created_at"],
                int(po["lead_time_days"]), po["due_date"], po["godown_id"],
                float(po["lc_usd_rate"]), float(po["lc_rmb_rate"]),
                float(po["lc_bank"]), float(po["lc_ship"]), float(po["lc_duty"]),
                float(po["lc_trans"]), float(po["lc_gst_duty"]), float(po["lc_doc_pct"])
            ))
            
            # Log initial status (from_status is NULL = creation event)
            _log_status_change(
                cursor,
                po["id"],
                from_status=None,
                to_status=po["status"],
                note="PO created"
            )
            
            # Insert line items
            for seq, li in enumerate(po.get("line_items", [])):
                cursor.execute("""
                    INSERT INTO po_items (id, po_id, item_id, item_name, description, hs_code, qty, unit, unit_price, line_sequence)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    str(uuid.uuid4()), po["id"], li.get("item_id", ""),
                    li.get("item_name", ""), li.get("description", ""),
                    li.get("hs_code", ""), float(li.get("qty", 0)),
                    li.get("unit", "PCS"), float(li.get("unit_price", 0)), seq
                ))
            
            # Update po_sequence in settings if the new PO number has a higher sequence.
            _sync_po_sequence(cursor, po["po_number"])
            
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    
    return jsonify(po), 201


@app.route("/api/po/<pid>", methods=["PUT"])
def update_po(pid):
    """Update purchase order"""
    req = request.get_json(silent=True)
    if not req:
        return jsonify({"error": "Invalid JSON"}), 400
    
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            
            # Fetch existing PO
            cursor.execute("SELECT * FROM purchase_orders WHERE id = ?", (pid,))
            row = cursor.fetchone()
            if not row:
                return jsonify({"error": "Not found"}), 404
            
            po = dict(row)
            po["supplier_snapshot"] = json.loads(po.get("supplier_snapshot") or "{}")
            old_status = po["status"]   # capture before any field merge
            
            # Update fields
            update_fields = [
                "po_number", "po_date", "supplier_id", "supplier_snapshot",
                "payment_conditions", "delivery_terms", "delivery_address",
                "forwarder_id", "forwarder_name", "forwarder_contact", "remarks",
                "currency", "status", "lead_time_days", "due_date", "godown_id",
                "lc_usd_rate", "lc_rmb_rate", "lc_bank", "lc_ship", "lc_duty",
                "lc_trans", "lc_gst_duty", "lc_doc_pct"
            ]
            
            for field in update_fields:
                if field in req:
                    po[field] = req[field]
            
            # Update PO
            cursor.execute("""
                UPDATE purchase_orders
                SET po_number = ?, po_date = ?, supplier_id = ?, supplier_snapshot = ?,
                    payment_conditions = ?, delivery_terms = ?, delivery_address = ?,
                    forwarder_id = ?, forwarder_name = ?, forwarder_contact = ?, remarks = ?,
                    currency = ?, status = ?, lead_time_days = ?, due_date = ?, godown_id = ?,
                    lc_usd_rate = ?, lc_rmb_rate = ?, lc_bank = ?, lc_ship = ?, lc_duty = ?,
                    lc_trans = ?, lc_gst_duty = ?, lc_doc_pct = ?, updated_at = CURRENT_TIMESTAMP
                WHERE id = ?
            """, (
                po["po_number"], po["po_date"], po["supplier_id"],
                json.dumps(po.get("supplier_snapshot", {})),
                po["payment_conditions"], po["delivery_terms"], po["delivery_address"],
                po["forwarder_id"], po["forwarder_name"], po["forwarder_contact"],
                po["remarks"], po["currency"], po["status"],
                int(po.get("lead_time_days", 0)), po.get("due_date", ""), po.get("godown_id", ""),
                float(po.get("lc_usd_rate", 84.0)), float(po.get("lc_rmb_rate", 11.5)),
                float(po.get("lc_bank", 0)), float(po.get("lc_ship", 0)),
                float(po.get("lc_duty", 0)), float(po.get("lc_trans", 0)),
                float(po.get("lc_gst_duty", 0)), float(po.get("lc_doc_pct", 0)), pid
            ))
            
            # Log status change if status actually changed
            if "status" in req and req["status"] != old_status:
                note = req.get("status_note")  # frontend can pass an optional note
                _log_status_change(cursor, pid, old_status, po["status"], note)
            elif req.get("status_note"):
                # Save a note-only log entry even when status has not changed
                _log_status_change(cursor, pid, old_status, old_status, req["status_note"], force=True)
            
            # Sync sequence in case of manual renumbering
            _sync_po_sequence(cursor, po["po_number"])
            
            # Update line items if provided
            if "line_items" in req:
                # Delete existing line items
                cursor.execute("DELETE FROM po_items WHERE po_id = ?", (pid,))
                
                # Insert new line items
                for seq, li in enumerate(req["line_items"]):
                    cursor.execute("""
                        INSERT INTO po_items (id, po_id, item_id, item_name, description, hs_code, qty, unit, unit_price, line_sequence)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """, (
                        str(uuid.uuid4()), pid, li.get("item_id", ""),
                        li.get("item_name", ""), li.get("description", ""),
                        li.get("hs_code", ""), float(li.get("qty", 0)),
                        li.get("unit", "PCS"), float(li.get("unit_price", 0)), seq
                    ))
            
            # Fetch updated line items
            cursor2 = conn.cursor()
            cursor2.execute("SELECT * FROM po_items WHERE po_id = ? ORDER BY line_sequence", (pid,))
            po["line_items"] = [
                {
                    "id": item["id"],
                    "item_id": item["item_id"],
                    "item_name": item["item_name"],
                    "description": item["description"],
                    "hs_code": item["hs_code"],
                    "qty": item["qty"],
                    "unit": item["unit"],
                    "unit_price": item["unit_price"],
                }
                for item in cursor2.fetchall()
            ]
            
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    
    return jsonify(po)


@app.route("/api/po/<pid>", methods=["DELETE"])
def delete_po(pid):
    """Soft-delete a purchase order (sets deleted_at; data is preserved).

    Use DELETE /api/po/<pid>/hard to permanently destroy a PO and its files.
    """
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM purchase_orders WHERE id = ? AND deleted_at IS NULL", (pid,))
        if not cursor.fetchone():
            return jsonify({"error": "PO not found"}), 404
        cursor.execute(
            "UPDATE purchase_orders SET deleted_at = CURRENT_TIMESTAMP WHERE id = ?",
            (pid,)
        )
    return jsonify({"ok": True})


@app.route("/api/po/<pid>/restore", methods=["POST"])
def restore_po(pid):
    """Restore a soft-deleted purchase order"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM purchase_orders WHERE id = ? AND deleted_at IS NOT NULL", (pid,))
        if not cursor.fetchone():
            return jsonify({"error": "No deleted PO found with that id"}), 404
        cursor.execute(
            "UPDATE purchase_orders SET deleted_at = NULL WHERE id = ?",
            (pid,)
        )
    return jsonify({"ok": True})


@app.route('/api/po/<pid>/timeline', methods=['GET'])
def get_po_timeline(pid):
    """
    Returns all status log entries for one PO, newest first.
    Example URL: GET /api/po/17/timeline
    """
    with get_db() as db:
        rows = db.execute('''
            SELECT id, from_status, to_status, changed_at, note
            FROM   po_status_log
            WHERE  po_id = ?
            ORDER  BY changed_at DESC
        ''', [pid]).fetchall()

    # Convert rows to plain dictionaries so Flask can turn them into JSON
    entries = [dict(r) for r in rows]
    return jsonify({'timeline': entries})



@app.route("/api/analytics/suppliers", methods=["GET"])
def supplier_analytics():
    """
    Aggregate PO data per supplier:
    - Total PO count, total order value (sum of line items)
    - Average promised lead time (lead_time_days column)
    - Actual lead time (days from po_date to Received status log entry)
    - On-time delivery % (actual lead time <= promised lead time, for Received POs)
    - Total value over last 12 months (month-by-month)
    """
    try:
        with get_db() as conn:
            cursor = conn.cursor()

            # ── Per-supplier PO totals ────────────────────────────────────────
            cursor.execute("""
                SELECT
                    s.id            AS supplier_id,
                    s.name          AS supplier_name,
                    s.company       AS company,
                    COUNT(po.id)    AS total_pos,
                    SUM(CASE WHEN po.status NOT IN ('Cancelled') THEN 1 ELSE 0 END)
                                    AS active_pos,
                    SUM(CASE WHEN po.status = 'Shipped' THEN 1 ELSE 0 END)
                                    AS shipped_pos,
                    SUM(CASE WHEN po.status = 'In Transit' THEN 1 ELSE 0 END)
                                    AS in_transit_pos,
                    AVG(CASE WHEN po.lead_time_days > 0 THEN po.lead_time_days END)
                                    AS avg_promised_days,
                    SUM(COALESCE(pv.value_cny, 0))
                                    AS total_value_cny
                FROM suppliers s
                LEFT JOIN purchase_orders po
                    ON po.supplier_id = s.id AND po.deleted_at IS NULL
                LEFT JOIN (
                    SELECT po_id, SUM(qty * unit_price) AS value_cny
                    FROM po_items GROUP BY po_id
                ) pv ON pv.po_id = po.id
                GROUP BY s.id
                ORDER BY total_value_cny DESC
            """)
            supplier_rows = [dict(r) for r in cursor.fetchall()]

            # ── On-time delivery calculation (using status log) ───────────────
            cursor.execute("""
                SELECT
                    po.supplier_id,
                    po.po_date,
                    po.lead_time_days,
                    COALESCE(rlog.received_at, po.updated_at) AS received_at
                FROM purchase_orders po
                LEFT JOIN (
                    SELECT po_id, MIN(changed_at) as received_at
                    FROM po_status_log
                    WHERE to_status = 'Received'
                    GROUP BY po_id
                ) rlog ON rlog.po_id = po.id
                WHERE po.deleted_at IS NULL AND po.status = 'Received'
            """)
            delivery_rows = cursor.fetchall()

            # Compute per-supplier on-time stats
            from datetime import datetime
            delivery_stats = {}  # supplier_id -> {on_time: int, total: int, actual_days_sum: int}
            for dr in delivery_rows:
                sid       = dr["supplier_id"]
                po_date   = dr["po_date"]
                promised  = dr["lead_time_days"]
                recv_str  = dr["received_at"]
                if not po_date or not recv_str:
                    continue
                try:
                    d0      = datetime.strptime(po_date[:10], "%Y-%m-%d")
                    d1      = datetime.strptime(recv_str[:10], "%Y-%m-%d")
                    actual  = (d1 - d0).days
                except ValueError:
                    continue
                if sid not in delivery_stats:
                    delivery_stats[sid] = {"on_time": 0, "total": 0, "actual_days_sum": 0}
                delivery_stats[sid]["total"]           += 1
                delivery_stats[sid]["actual_days_sum"] += actual
                if actual <= promised:
                    delivery_stats[sid]["on_time"] += 1

            # ── Monthly spend per supplier (last 12 months) ───────────────────
            cursor.execute("""
                SELECT
                    po.supplier_id,
                    substr(po.po_date, 1, 7) AS ym,
                    SUM(COALESCE(pv.value_cny, 0)) AS monthly_value
                FROM purchase_orders po
                LEFT JOIN (
                    SELECT po_id, SUM(qty * unit_price) AS value_cny
                    FROM po_items GROUP BY po_id
                ) pv ON pv.po_id = po.id
                WHERE po.deleted_at IS NULL
                  AND po.po_date >= date('now', '-12 months')
                GROUP BY po.supplier_id, ym
                ORDER BY ym ASC
            """)
            monthly_raw = cursor.fetchall()
            monthly_by_supplier = {}
            for mr in monthly_raw:
                sid = mr["supplier_id"]
                monthly_by_supplier.setdefault(sid, []).append({
                    "month": mr["ym"], "value": mr["monthly_value"] or 0
                })

            # ── Merge everything ──────────────────────────────────────────────
            result = []
            for s in supplier_rows:
                sid   = s["supplier_id"]
                stats = delivery_stats.get(sid, {})
                total_d = stats.get("total", 0)
                on_t    = stats.get("on_time", 0)
                actual_days_sum = stats.get("actual_days_sum", 0)
                result.append({
                    **s,
                    "avg_actual_days":   round(actual_days_sum / total_d, 1) if total_d else None,
                    "on_time_pct":       round(on_t / total_d * 100, 1)              if total_d else None,
                    "delivered_pos":     total_d,
                    "shipped_pos":       s.get("shipped_pos", 0) or 0,
                    "in_transit_pos":    s.get("in_transit_pos", 0) or 0,
                    "monthly_spend":     monthly_by_supplier.get(sid, []),
                })

    except Exception as e:
        return jsonify({"error": str(e)}), 400

    return jsonify(result)


@app.route("/api/po/deleted", methods=["GET"])
def list_deleted_pos():
    """List soft-deleted purchase orders"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT id, po_number, po_date, supplier_id, status, deleted_at "
            "FROM purchase_orders WHERE deleted_at IS NOT NULL ORDER BY deleted_at DESC"
        )
        return jsonify([dict(r) for r in cursor.fetchall()])


@app.route("/api/po/<pid>/hard", methods=["DELETE"])
def hard_delete_po(pid):
    """Permanently delete a PO and all its attachments / payment files.

    Only works on already-soft-deleted POs so there is always a two-step
    confirmation before data is irrecoverably lost.
    """
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM purchase_orders WHERE id = ? AND deleted_at IS NOT NULL", (pid,))
        if not cursor.fetchone():
            return jsonify({"error": "PO must be soft-deleted first"}), 409

        # Remove physical attachment files
        attach_folder = os.path.join(ATTACH_DIR, pid)
        if os.path.isdir(attach_folder):
                    shutil.rmtree(attach_folder, ignore_errors=True)

        cursor.execute("DELETE FROM purchase_orders WHERE id = ?", (pid,))
    return jsonify({"ok": True})


# ═══════════════════════════════════════════════════════════════════════════════
# SETTINGS
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/api/settings", methods=["GET"])
def get_settings():
    """Fetch all settings"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT key, value FROM settings")
        settings = {row["key"]: row["value"] for row in cursor.fetchall()}
    return jsonify(settings)


@app.route("/api/settings", methods=["POST"])
def save_settings():
    """Update settings"""
    req = request.get_json(silent=True)
    if not req:
        return jsonify({"error": "Invalid JSON"}), 400
    
    with get_db() as conn:
        cursor = conn.cursor()
        for key, value in req.items():
            cursor.execute(
                "INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)",
                (key, str(value))
            )
        
        # Return updated settings
        cursor.execute("SELECT key, value FROM settings")
        settings = {row["key"]: row["value"] for row in cursor.fetchall()}
    
    return jsonify(settings)


# ═══════════════════════════════════════════════════════════════════════════════
# NEXT PO NUMBER (Auto-Numbering)
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/api/next-po-number", methods=["GET"])
def next_po_number():
    """Predict the next PO number based on current settings and existing POs.

    Does NOT persist the sequence to the database. This ensures that opening
    the "New PO" dialog without saving doesn't waste a PO number. The sequence
    is only advanced when a PO is actually saved in `create_po`.
    """
    with get_db() as conn:
        cursor = conn.cursor()

        # Ensure default settings exist
        init_default_settings(conn)

        cursor.execute("SELECT key, value FROM settings WHERE key IN ('po_prefix', 'po_sequence')")
        cfg = {row["key"]: row["value"] for row in cursor.fetchall()}

        prefix = (cfg.get("po_prefix") or "PO").strip()
        seq = int(cfg.get("po_sequence") or "0")
        year = date.today().year

        # Advance past any gaps or existing POs
        while True:
            seq += 1
            po_number = f"{prefix}-{year}-{str(seq).zfill(3)}"
            cursor.execute("SELECT id FROM purchase_orders WHERE po_number = ?", (po_number,))
            if not cursor.fetchone():
                break

    return jsonify({"po_number": po_number, "sequence": seq})


# ═══════════════════════════════════════════════════════════════════════════════
# EXPORT / IMPORT
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/api/export", methods=["GET"])
def export_data():
    """
    📤 EXPORT ENDPOINT
    Read all data from SQLite → Convert to JSON → Return downloadable backup
    """
    bundle = {}
    
    with get_db() as conn:
        cursor = conn.cursor()
        
        # Export suppliers
        cursor.execute("SELECT * FROM suppliers ORDER BY name")
        bundle["suppliers"] = [dict(row) for row in cursor.fetchall()]
        
        # Export items
        cursor.execute("SELECT * FROM items ORDER BY name")
        bundle["items"] = [dict(row) for row in cursor.fetchall()]
        
        # Export forwarders
        cursor.execute("SELECT * FROM forwarders ORDER BY name")
        forwarders = []
        for row in cursor.fetchall():
            fw = dict(row)
            cursor2 = conn.cursor()
            cursor2.execute(
                "SELECT id, name, contact_person, phone, email, address "
                "FROM forwarder_godowns WHERE forwarder_id = ? ORDER BY sort_order",
                (fw["id"],)
            )
            gd_rows = cursor2.fetchall()
            if gd_rows:
                fw["godowns"] = [
                    {
                        "id":             r["id"],
                        "label":          r["name"],
                        "contact_person": r["contact_person"] or "",
                        "phone":          r["phone"] or "",
                        "email":          r["email"] or "",
                        "address":        r["address"] or "",
                    }
                    for r in gd_rows
                ]
            else:
                fw["godowns"] = json.loads(fw.get("godowns") or "[]")
            forwarders.append(fw)
        bundle["forwarders"] = forwarders
        
        # Export purchase orders with line items
        cursor.execute("SELECT * FROM purchase_orders WHERE deleted_at IS NULL ORDER BY po_date DESC")
        purchase_orders = []
        for po_row in cursor.fetchall():
            po = dict(po_row)
            po["supplier_snapshot"] = json.loads(po.get("supplier_snapshot") or "{}")
            
            # Fetch line items — use a separate cursor to avoid clobbering the outer result set
            cursor2 = conn.cursor()
            cursor2.execute("SELECT * FROM po_items WHERE po_id = ? ORDER BY line_sequence", (po["id"],))
            po["line_items"] = [
                {
                    "item_id": item["item_id"],
                    "item_name": item["item_name"],
                    "description": item["description"],
                    "hs_code": item["hs_code"],
                    "qty": item["qty"],
                    "unit": item["unit"],
                    "unit_price": item["unit_price"],
                }
                for item in cursor2.fetchall()
            ]
            purchase_orders.append(po)
        
        bundle["purchase_orders"] = purchase_orders
        
        # Export quotations
        cursor.execute("SELECT * FROM quotations")
        quotations = []
        for row in cursor.fetchall():
            q = dict(row)
            q["details"] = json.loads(q.get("details") or "{}")
            quotations.append(q)
        bundle["quotations"] = quotations
        
        # Export settings
        cursor.execute("SELECT key, value FROM settings")
        settings = {row["key"]: row["value"] for row in cursor.fetchall()}
        bundle["settings"] = settings
    
    bundle["export_date"] = str(date.today())
    bundle["version"] = "2.0 (SQLite)"
    
    json_bytes = json.dumps(bundle, indent=2, ensure_ascii=False).encode("utf-8")
    buf = BytesIO(json_bytes)
    buf.seek(0)
    
    return send_file(
        buf,
        mimetype="application/json",
        as_attachment=True,
        download_name=f"portal_backup_{date.today().strftime('%Y%m%d')}.json"
    )


@app.route("/api/import", methods=["POST"])
def import_data():
    """
    📥 IMPORT ENDPOINT WITH TRANSACTION & ROLLBACK
    Upload JSON file → Validate → Replace existing DB data
    On error: ROLLBACK entire transaction (no partial import)
    """
    if "file" not in request.files:
        return jsonify({"error": "No file"}), 400
    
    file = request.files["file"]
    
    try:
        bundle = json.loads(file.read().decode("utf-8"))
    except Exception as e:
        return jsonify({"error": f"Invalid JSON: {e}"}), 400
    
    # Validate structure
    required_keys = ["suppliers", "items", "purchase_orders"]
    for key in required_keys:
        if key not in bundle:
            return jsonify({"error": f"Missing required key: {key}"}), 400
    
    # ── 🔐 BACKUP BEFORE IMPORT (must run OUTSIDE any transaction) ──────────
    # VACUUM INTO cannot be called inside a transaction; SQLite will raise
    # "cannot VACUUM from within a transaction".  We do the backup first on
    # a separate connection so the import transaction starts clean.
    backup_dir = os.path.join(os.path.dirname(__file__), "data", "po")
    try:
        for fname in os.listdir(backup_dir):
            if fname.startswith("backup_before_import_") and fname.endswith(".db"):
                try:
                    os.remove(os.path.join(backup_dir, fname))
                    print(f"🗑️ Removed old backup: {fname}")
                except Exception:
                    pass
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        backup_db_path = os.path.join(backup_dir, f"backup_before_import_{timestamp}.db")
        print(f"💾 Creating backup: {backup_db_path}")
        _bconn = get_connection()
        try:
            _bconn.execute(f"VACUUM INTO '{backup_db_path}'")
            print("✅ Backup created successfully")
        finally:
            _bconn.close()
    except Exception as backup_err:
        print(f"⚠️ Backup skipped (non-fatal): {backup_err}")
        # The uploaded JSON file itself serves as an implicit backup.

    # Start import transaction
    with get_db() as conn:
        try:
            cursor = conn.cursor()
            
            # ⚠️ CLEAR EXISTING DATA (within transaction)
            cursor.execute("DELETE FROM po_items")
            cursor.execute("DELETE FROM purchase_orders")
            cursor.execute("DELETE FROM quotations")
            cursor.execute("DELETE FROM forwarders")
            cursor.execute("DELETE FROM items")
            cursor.execute("DELETE FROM suppliers")
            
            imported_stats = {
                "suppliers": 0,
                "items": 0,
                "forwarders": 0,
                "purchase_orders": 0,
                "po_items": 0,
                "quotations": 0,
            }
            
            # ── IMPORT SUPPLIERS ──────────────────────────────────────────
            for supplier in bundle.get("suppliers", []):
                cursor.execute("""
                    INSERT INTO suppliers (id, name, company, address, country, email, phone, wechat, bank_name, bank_account, swift_code)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    supplier.get("id", str(uuid.uuid4())),
                    supplier.get("name", ""),
                    supplier.get("company", ""),
                    supplier.get("address", ""),
                    supplier.get("country", "China"),
                    supplier.get("email", ""),
                    supplier.get("phone", ""),
                    supplier.get("wechat", ""),
                    supplier.get("bank_name", ""),
                    supplier.get("bank_account", ""),
                    supplier.get("swift_code", ""),
                ))
                imported_stats["suppliers"] += 1
            
            # ── IMPORT ITEMS ──────────────────────────────────────────────
            for item in bundle.get("items", []):
                cursor.execute("""
                    INSERT INTO items (id, name, description, hs_code, unit, currency, default_price_usd)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                """, (
                    item.get("id", str(uuid.uuid4())),
                    item.get("name", ""),
                    item.get("description", ""),
                    item.get("hs_code", ""),
                    item.get("unit", "PCS"),
                    item.get("currency", "CNY"),
                    float(item.get("default_price_usd", 0)),
                ))
                imported_stats["items"] += 1
            
            # ── IMPORT FORWARDERS ─────────────────────────────────────────
            for fwd in bundle.get("forwarders", []):
                fwd_id = fwd.get("id", str(uuid.uuid4()))
                godowns = fwd.get("godowns", [])
                cursor.execute("""
                    INSERT INTO forwarders (id, name, contact_person, phone, email, godowns)
                    VALUES (?, ?, ?, ?, ?, ?)
                """, (
                    fwd_id,
                    fwd.get("name", ""),
                    fwd.get("contact_person", ""),
                    fwd.get("phone", ""),
                    fwd.get("email", ""),
                    json.dumps(godowns),
                ))
                for i, g in enumerate(godowns):
                    if isinstance(g, dict):
                        gid = g.get("id") or str(uuid.uuid4())
                        cursor.execute(
                            "INSERT INTO forwarder_godowns "
                            "(id, forwarder_id, name, contact_person, phone, email, address, sort_order) "
                            "VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                            (gid, fwd_id,
                             g.get("label", "") or g.get("name", ""),
                             g.get("contact_person", ""),
                             g.get("phone", ""),
                             g.get("email", ""),
                             g.get("address", ""),
                             i)
                        )
                    else:
                        # Legacy: plain string godown name
                        cursor.execute(
                            "INSERT INTO forwarder_godowns "
                            "(id, forwarder_id, name, contact_person, phone, email, address, sort_order) "
                            "VALUES (?, ?, ?, '', '', '', '', ?)",
                            (str(uuid.uuid4()), fwd_id, str(g), i)
                        )
                imported_stats["forwarders"] += 1
            
            # ── IMPORT PURCHASE ORDERS ────────────────────────────────────
            for po in bundle.get("purchase_orders", []):
                po_id = po.get("id", str(uuid.uuid4()))
                
                cursor.execute("""
                    INSERT INTO purchase_orders (
                        id, po_number, po_date, supplier_id, supplier_snapshot,
                        payment_conditions, delivery_terms, delivery_address,
                        forwarder_id, forwarder_name, forwarder_contact, remarks,
                        currency, status, created_at, lead_time_days, due_date, godown_id,
                        lc_usd_rate, lc_rmb_rate, lc_bank, lc_ship, lc_duty,
                        lc_trans, lc_gst_duty, lc_doc_pct
                    )
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    po_id,
                    po.get("po_number", ""),
                    po.get("po_date", str(date.today())),
                    po.get("supplier_id", ""),
                    json.dumps(po.get("supplier_snapshot", {})),
                    po.get("payment_conditions", ""),
                    po.get("delivery_terms", "FOB"),
                    po.get("delivery_address", ""),
                    po.get("forwarder_id", ""),
                    po.get("forwarder_name", ""),
                    po.get("forwarder_contact", ""),
                    po.get("remarks", ""),
                    po.get("currency", "USD"),
                    po.get("status", "Draft"),
                    po.get("created_at", str(date.today())),
                    int(po.get("lead_time_days", 0)),
                    po.get("due_date", ""),
                    po.get("godown_id", ""),          # ← was missing; caused godown to be lost on import
                    float(po.get("lc_usd_rate", 84.0)),
                    float(po.get("lc_rmb_rate", 11.5)),
                    float(po.get("lc_bank", 0.0)),
                    float(po.get("lc_ship", 0.0)),
                    float(po.get("lc_duty", 0.0)),
                    float(po.get("lc_trans", 0.0)),
                    float(po.get("lc_gst_duty", 0.0)),
                    float(po.get("lc_doc_pct", 0.0)),
                ))
                imported_stats["purchase_orders"] += 1
                
                # ── IMPORT LINE ITEMS ──────────────────────────────────────
                for seq, li in enumerate(po.get("line_items", [])):
                    cursor.execute("""
                        INSERT INTO po_items (id, po_id, item_id, item_name, description, hs_code, qty, unit, unit_price, line_sequence)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """, (
                        str(uuid.uuid4()),
                        po_id,
                        li.get("item_id", ""),
                        li.get("item_name", ""),
                        li.get("description", ""),
                        li.get("hs_code", ""),
                        float(li.get("qty", 0)),
                        li.get("unit", "PCS"),
                        float(li.get("unit_price", 0)),
                        seq,
                    ))
                    imported_stats["po_items"] += 1
            
            # ── IMPORT QUOTATIONS ─────────────────────────────────────────
            for quotation in bundle.get("quotations", []):
                q_id = quotation.get("id", str(uuid.uuid4()))
                q_number = quotation.get("quotation_number", quotation.get("quotation_no", ""))
                q_title = quotation.get("title", "")
                q_currency = quotation.get("currency", "CNY")
                q_status = quotation.get("status", "Open")
                q_awarded = quotation.get("awarded_to", "")
                q_custref = quotation.get("customer_ref", "")
                q_notes = quotation.get("notes", "")
                q_date = quotation.get("date", "")
                q_details = json.dumps(quotation.get("details", quotation))
                q_created = quotation.get("created_at", datetime.now().isoformat())
                q_updated = quotation.get("updated_at", datetime.now().isoformat())
                cursor.execute("""
                    INSERT INTO quotations (id, quotation_number, title, currency, status,
                                            awarded_to, customer_ref, notes, date, details,
                                            created_at, updated_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    q_id, q_number, q_title, q_currency, q_status,
                    q_awarded, q_custref, q_notes, q_date, q_details,
                    q_created, q_updated,
                ))
                _sync_quotation_refs(cursor, q_id, q_details)
                _sync_quotation_line_items(cursor, q_id, q_details)
                imported_stats["quotations"] += 1
            
            # ── IMPORT SETTINGS ───────────────────────────────────────────
            if "settings" in bundle and isinstance(bundle["settings"], dict):
                init_default_settings(conn)
                for key, value in bundle["settings"].items():
                    cursor.execute(
                        "INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)",
                        (key, str(value))
                    )
            
            # ✅ COMMIT TRANSACTION (only if all succeeded)
            
            return jsonify({
                "ok": True,
                "imported": imported_stats,
                "message": "✅ Import completed successfully with transaction backup"
            }), 200
            
        except Exception as e:
            # ❌ ROLLBACK on error (automatic with context manager)
            conn.rollback()
            return jsonify({"error": f"Import failed - rolled back: {str(e)}"}), 400


# ═══════════════════════════════════════════════════════════════════════════════
# QUOTATION ANALYZER API
# ═══════════════════════════════════════════════════════════════════════════════

def _sync_quotation_refs(cursor, quotation_id, details_json):
    """Rebuild quotation_refs rows for *quotation_id* from the JSON payload.

    Called inside an open transaction whenever a quotation is created or
    updated.  The DELETE + INSERT approach is a simple upsert that handles
    adds, removals, and changes in one pass.
    """
    cursor.execute("DELETE FROM quotation_refs WHERE quotation_id = ?", (quotation_id,))
    try:
        details = json.loads(details_json) if isinstance(details_json, str) else (details_json or {})
    except Exception:
        details = {}
    seen = set()
    for li in details.get("line_items", []):
        item_id = li.get("item_id", "")
        if item_id:
            key = ("item", item_id)
            if key not in seen:
                cursor.execute(
                    "INSERT OR IGNORE INTO quotation_refs (quotation_id, ref_type, ref_id) VALUES (?, ?, ?)",
                    (quotation_id, "item", item_id)
                )
                seen.add(key)
        for sr in li.get("supplier_rows", []):
            sup_id = sr.get("supplier_id", "")
            if sup_id:
                key = ("supplier", sup_id)
                if key not in seen:
                    cursor.execute(
                        "INSERT OR IGNORE INTO quotation_refs (quotation_id, ref_type, ref_id) VALUES (?, ?, ?)",
                        (quotation_id, "supplier", sup_id)
                    )
                    seen.add(key)



def _sync_quotation_line_items(cursor, quotation_id, details_json):
    """Rebuild normalised quotation_line_items and quotation_supplier_rows.

    Called inside an open transaction on every quotation write so the
    structured tables stay in sync with the details JSON blob.
    """
    cursor.execute("DELETE FROM quotation_supplier_rows WHERE quotation_id = ?", (quotation_id,))
    cursor.execute("DELETE FROM quotation_line_items WHERE quotation_id = ?", (quotation_id,))
    try:
        details = json.loads(details_json) if isinstance(details_json, str) else (details_json or {})
    except Exception:
        details = {}

    for seq, li in enumerate(details.get("line_items", [])):
        li_id = str(uuid.uuid4())
        cursor.execute(
            "INSERT INTO quotation_line_items "
            "(id, quotation_id, item_id, item_name, qty, unit, description, "
            " selected_supplier_id, selected_supplier_name, line_sequence) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (
                li_id, quotation_id,
                li.get("item_id") or None,
                li.get("item_name", ""),
                float(li.get("qty", 1)),
                li.get("unit", "PCS"),
                li.get("description", ""),
                li.get("selected") or None,
                li.get("selected_supplier") or None,
                seq,
            )
        )
        for sr in li.get("supplier_rows", []):
            cursor.execute(
                "INSERT INTO quotation_supplier_rows "
                "(id, line_item_id, quotation_id, supplier_id, supplier_name, price) "
                "VALUES (?, ?, ?, ?, ?, ?)",
                (
                    str(uuid.uuid4()), li_id, quotation_id,
                    sr.get("supplier_id") or None,
                    sr.get("supplier_name", ""),
                    float(sr.get("price", 0)),
                )
            )


def _quotation_to_dict(row):
    """Convert a quotation DB row to a dict, merging stored details"""
    q = dict(row)
    details = {}
    if q.get("details"):
        try:
            details = json.loads(q["details"])
        except Exception:
            pass
    # Merge details into top-level, but DB columns take priority for indexed fields
    merged = {**details, **{k: v for k, v in q.items() if v is not None and k != "details"}}
    # Ensure quotation_no is available (alias)
    if "quotation_no" not in merged and "quotation_number" in merged:
        merged["quotation_no"] = merged["quotation_number"]
    return merged


@app.route("/api/quotations", methods=["GET"])
def list_quotations():
    """Get all quotations"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM quotations ORDER BY created_at DESC")
        return jsonify([_quotation_to_dict(r) for r in cursor.fetchall()])


@app.route("/api/quotations", methods=["POST"])
def create_quotation():
    """Create a new quotation"""
    data = request.get_json(force=True)
    qid = str(uuid.uuid4())
    q_number = data.get("quotation_no", "")
    title = data.get("title", "")
    currency = data.get("currency", "CNY")
    customer_ref = data.get("customer_ref", "")
    notes = data.get("notes", "")
    q_date = data.get("date", str(date.today()))
    now = datetime.now().isoformat()

    # Store full payload as JSON details
    details = json.dumps(data)

    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO quotations (id, quotation_number, title, currency, status,
                                    customer_ref, notes, date, details, created_at, updated_at)
            VALUES (?, ?, ?, ?, 'Open', ?, ?, ?, ?, ?, ?)
        """, (qid, q_number, title, currency, customer_ref, notes, q_date, details, now, now))
        _sync_quotation_refs(cursor, qid, details)
        _sync_quotation_line_items(cursor, qid, details)
        cursor.execute("SELECT * FROM quotations WHERE id = ?", (qid,))
        return jsonify(_quotation_to_dict(cursor.fetchone())), 201


@app.route("/api/quotations/<qid>", methods=["GET"])
def get_quotation(qid):
    """Get a single quotation by ID"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM quotations WHERE id = ?", (qid,))
        row = cursor.fetchone()
        if not row:
            return jsonify({"error": "Quotation not found"}), 404
        return jsonify(_quotation_to_dict(row))


@app.route("/api/quotations/<qid>", methods=["PUT"])
def update_quotation(qid):
    """Update an existing quotation"""
    data = request.get_json(force=True)
    now = datetime.now().isoformat()

    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM quotations WHERE id = ?", (qid,))
        existing = cursor.fetchone()
        if not existing:
            return jsonify({"error": "Quotation not found"}), 404

        q_number = data.get("quotation_no", existing["quotation_number"])
        title = data.get("title", existing["title"])
        currency = data.get("currency", existing["currency"])
        customer_ref = data.get("customer_ref", existing["customer_ref"] or "")
        notes = data.get("notes", existing["notes"] or "")
        q_date = data.get("date", existing["date"])

        # Preserve status and awarded_to from existing unless explicitly provided
        status = data.get("status", existing["status"])
        awarded_to = data.get("awarded_to", existing["awarded_to"])

        details = json.dumps(data)

        cursor.execute("""
            UPDATE quotations SET quotation_number=?, title=?, currency=?, status=?,
                awarded_to=?, customer_ref=?, notes=?, date=?, details=?, updated_at=?
            WHERE id=?
        """, (q_number, title, currency, status, awarded_to, customer_ref,
              notes, q_date, details, now, qid))
        _sync_quotation_refs(cursor, qid, details)
        _sync_quotation_line_items(cursor, qid, details)
        cursor.execute("SELECT * FROM quotations WHERE id = ?", (qid,))
        return jsonify(_quotation_to_dict(cursor.fetchone()))


@app.route("/api/quotations/<qid>", methods=["DELETE"])
def delete_quotation(qid):
    """Delete a quotation"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("DELETE FROM quotations WHERE id = ?", (qid,))
    return jsonify({"ok": True})


@app.route("/api/quotations/<qid>/award", methods=["POST"])
def award_quotation(qid):
    """Award a quotation to a supplier"""
    data = request.get_json(force=True)
    supplier_id = data.get("supplier_id", "")
    now = datetime.now().isoformat()

    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM quotations WHERE id = ?", (qid,))
        existing = cursor.fetchone()
        if not existing:
            return jsonify({"error": "Quotation not found"}), 404

        # Update details JSON with awarded info
        details = {}
        try:
            details = json.loads(existing["details"] or "{}")
        except Exception:
            pass
        details["status"] = "Awarded"
        details["awarded_to"] = supplier_id

        cursor.execute("""
            UPDATE quotations SET status='Awarded', awarded_to=?, details=?, updated_at=?
            WHERE id=?
        """, (supplier_id, json.dumps(details), now, qid))

        cursor.execute("SELECT * FROM quotations WHERE id = ?", (qid,))
        q = _quotation_to_dict(cursor.fetchone())

        # Build po_seed for auto-populating PO form
        line_items_data = q.get("line_items", [])
        po_seed = {
            "supplier_id": supplier_id,
            "currency": q.get("currency", "CNY"),
            "line_items": line_items_data,
        }

        return jsonify({**q, "po_seed": po_seed})


@app.route("/api/quotations/<qid>/reopen", methods=["POST"])
def reopen_quotation(qid):
    """Reopen an awarded quotation"""
    now = datetime.now().isoformat()
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM quotations WHERE id = ?", (qid,))
        existing = cursor.fetchone()
        if not existing:
            return jsonify({"error": "Quotation not found"}), 404

        details = {}
        try:
            details = json.loads(existing["details"] or "{}")
        except Exception:
            pass
        details["status"] = "Open"
        details["awarded_to"] = None

        cursor.execute("""
            UPDATE quotations SET status='Open', awarded_to=NULL, details=?, updated_at=?
            WHERE id=?
        """, (json.dumps(details), now, qid))
        cursor.execute("SELECT * FROM quotations WHERE id = ?", (qid,))
        return jsonify(_quotation_to_dict(cursor.fetchone()))


@app.route("/api/quotations/<qid>/duplicate", methods=["POST"])
def duplicate_quotation(qid):
    """Duplicate a quotation with cleared prices"""
    now = datetime.now().isoformat()
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM quotations WHERE id = ?", (qid,))
        existing = cursor.fetchone()
        if not existing:
            return jsonify({"error": "Quotation not found"}), 404

        details = {}
        try:
            details = json.loads(existing["details"] or "{}")
        except Exception:
            pass

        # Generate new ID and number
        new_id = str(uuid.uuid4())
        old_num = existing["quotation_number"] or ""
        new_num = old_num + "-COPY" if old_num else f"QT-{date.today().year}-COPY"

        # Clear prices in line items
        for li in details.get("line_items", []):
            for sr in li.get("supplier_rows", []):
                sr["price"] = 0
            li["selected"] = ""
            li["selected_supplier"] = ""

        # Reset status
        details["status"] = "Open"
        details["awarded_to"] = None
        details["quotation_no"] = new_num

        new_details_json = json.dumps(details)
        cursor.execute("""
            INSERT INTO quotations (id, quotation_number, title, currency, status,
                                    customer_ref, notes, date, details, created_at, updated_at)
            VALUES (?, ?, ?, ?, 'Open', ?, ?, ?, ?, ?, ?)
        """, (new_id, new_num, details.get("title", existing["title"]),
              existing["currency"], existing["customer_ref"] or "",
              existing["notes"] or "", str(date.today()), new_details_json, now, now))
        _sync_quotation_refs(cursor, new_id, new_details_json)
        _sync_quotation_line_items(cursor, new_id, new_details_json)

        cursor.execute("SELECT * FROM quotations WHERE id = ?", (new_id,))
        return jsonify(_quotation_to_dict(cursor.fetchone())), 201


def _safe_float(val, default=0.0):
    try:
        if val is None or val == "": return default
        return float(val)
    except (ValueError, TypeError):
        return default

@app.route("/api/po/<pid>/print", methods=["GET"])
def print_po(pid):
    """Generate a printable HTML view of a Purchase Order"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM purchase_orders WHERE id = ?", (pid,))
        po_row = cursor.fetchone()
        if not po_row:
            return "PO not found", 404
        po = dict(po_row)
        
        # Fetch supplier
        cursor.execute("SELECT * FROM suppliers WHERE id = ?", (po["supplier_id"],))
        sup_row = cursor.fetchone()
        sup = dict(sup_row) if sup_row else {}
        
        # Fetch line items
        cursor.execute("SELECT * FROM po_items WHERE po_id = ? ORDER BY line_sequence", (pid,))
        line_items = [dict(item) for item in cursor.fetchall()]
        
        # Fetch settings
        cursor.execute("SELECT key, value FROM settings")
        settings = {row["key"]: row["value"] for row in cursor.fetchall()}

    total = sum(_safe_float(li.get("qty", 0)) * _safe_float(li.get("unit_price", 0)) for li in line_items)
    total_qty = sum(_safe_float(li.get("qty", 0)) for li in line_items)
    cur = po.get("currency", "USD")
    fmt_total = f"{total:,.2f}"

    item_rows = ""
    for i, li in enumerate(line_items):
        amt = _safe_float(li.get("qty", 0)) * _safe_float(li.get("unit_price", 0))
        item_rows += f"""<tr>
          <td style="text-align:center;color:#374151">{i+1}</td>
          <td style="font-weight:600;color:#111;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;max-width:0">{_e(li.get('item_name',''))}</td>
          <td style="color:#4b5563;font-size:11px;white-space:nowrap">{_e(li.get('description','—'))}</td>
          <td style="text-align:center;color:#4b5563">{_e(li.get('hs_code','—'))}</td>
          <td style="text-align:center;font-weight:600">{li.get('qty',0)}</td>
          <td style="text-align:center;color:#4b5563">{_e(li.get('unit','PCS'))}</td>
          <td style="text-align:right">{_e(cur)} {li.get('unit_price',0):,.2f}</td>
          <td style="text-align:right;font-weight:700;color:#1e3a8a">{_e(cur)} {amt:,.2f}</td>
        </tr>"""

    buyer_lines = []
    if settings.get("company_name"): buyer_lines.append(f'<div style="font-size:16px;font-weight:700;color:#111;margin-bottom:3px">{_e(settings["company_name"])}</div>')
    if settings.get("company_address"):
        addr_html = _e(settings["company_address"]).replace("\n", "<br>")
        buyer_lines.append(f'<div style="color:#4b5563">{addr_html}</div>')
    if settings.get("company_phone"): buyer_lines.append(f'<div style="color:#4b5563">Ph: {_e(settings["company_phone"])}</div>')
    if settings.get("company_email"): buyer_lines.append(f'<div style="color:#4b5563">Email: {_e(settings["company_email"])}</div>')
    if settings.get("company_gstin"): buyer_lines.append(f'<div style="color:#4b5563">GSTIN: {_e(settings["company_gstin"])}</div>')
    buyer_html = "".join(buyer_lines) or '<div style="color:#9ca3af">Not configured.</div>'

    sup_lines = []
    sup_lines.append(f'<div style="font-size:14px;font-weight:700;color:#111;margin-bottom:3px">{_e(sup.get("company","—"))}</div>')
    if sup.get("name"): sup_lines.append(f'<div style="color:#4b5563">Attn: {_e(sup["name"])}</div>')
    if sup.get("address"): sup_lines.append(f'<div style="color:#4b5563">{_e(sup["address"])}</div>')
    if sup.get("country"): sup_lines.append(f'<div style="color:#4b5563">{_e(sup["country"])}</div>')
    if sup.get("email"): sup_lines.append(f'<div style="color:#4b5563">Email: {_e(sup["email"])}</div>')
    if sup.get("phone"): sup_lines.append(f'<div style="color:#4b5563">Ph/WA: {_e(sup["phone"])}</div>')
    if sup.get("wechat"): sup_lines.append(f'<div style="color:#4b5563">WeChat: {_e(sup["wechat"])}</div>')
    supplier_html = "".join(sup_lines)

    bank_section = ""
    if sup.get("bank_name") or sup.get("bank_account"):
        swift = f' &nbsp;|&nbsp; SWIFT: {_e(sup["swift_code"])}' if sup.get("swift_code") else ""
        acc = f' &nbsp;|&nbsp; A/C: {_e(sup["bank_account"])}' if sup.get("bank_account") else ""
        bank_section = f"""
        <div style="margin-top:10px;padding:10px 16px;background:#f8faff;border:1px solid #dbeafe;border-radius:6px;font-size:11px">
          <div style="font-weight:700;text-transform:uppercase;letter-spacing:.05em;font-size:9px;color:#1e3a8a;margin-bottom:4px">Bank Details (Supplier) · 银行详情</div>
          <div style="color:#1e40af">{_e(sup.get('bank_name',''))}{acc}{swift}</div>
        </div>"""

    status_color = {"Draft":"#6b7280","Sent":"#2563eb","Confirmed":"#16a34a","Shipped":"#7c3aed","In Transit":"#f97316","Received":"#059669","Cancelled":"#dc2626"}.get(po.get("status"), "#6b7280")

    # Payment Breakdown
    pay_breakdown_html = ""
    pc = po.get("payment_conditions") or ""
    pay_match = re.search(r"(\d+)%\s*Advance.*?(\d+)%", pc, re.I)
    if pay_match and total > 0:
        pct1 = int(pay_match.group(1))
        pct2 = int(pay_match.group(2))
        adv = total * pct1 / 100.0
        rem = total * pct2 / 100.0
        lbl2 = "against B/L (Bill of Lading)" if "bl" in pc.lower() else "on Shipment"
        pay_breakdown_html = f"""
          <div style="padding:7px 12px;background:#eff6ff;border:1px solid #bfdbfe;border-radius:5px;font-size:10px;line-height:1.8">
            <div style="font-size:8px;font-weight:700;letter-spacing:.1em;text-transform:uppercase;color:#1d4ed8;margin-bottom:3px">Payment Schedule · 付款计划</div>
            <div style="display:flex;justify-content:space-between;gap:8px;flex-wrap:nowrap"><span style="color:#374151;white-space:nowrap">① {pct1}% Advance (Proforma Invoice)</span><span style="font-weight:700;color:#1d4ed8;white-space:nowrap">{_e(cur)} {adv:,.2f}</span></div>
            <div style="display:flex;justify-content:space-between;gap:8px;flex-wrap:nowrap"><span style="color:#374151;white-space:nowrap">② {pct2}% {lbl2}</span><span style="font-weight:700;color:#1d4ed8;white-space:nowrap">{_e(cur)} {rem:,.2f}</span></div>
          </div>"""

    lead_time_html = ""
    lt = po.get("lead_time_days")
    dd = po.get("due_date")
    if lt or dd:
        lt_str = f"{lt} Days" if lt else "—"
        dd_str = f'<div style="font-size:11px;color:#1d4ed8;margin-top:1px;font-weight:600">Due: {_e(dd)}</div>' if dd else ""
        lead_time_html = f"""
        <div>
          <div style="font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#9ca3af;margin-bottom:3px">Lead Time · 交货期</div>
          <div style="font-weight:600;color:#111;font-size:14px">{_e(lt_str)}</div>
          {dd_str}
        </div>"""

    delivery_logistics = ""
    if po.get("forwarder_name") or po.get("forwarder_contact") or po.get("delivery_address"):
        fwd_name = po.get("forwarder_name", "—")
        fwd_contact = f' <span style="font-size:11px;color:#6b7280;font-weight:400">{_e(po["forwarder_contact"])}</span>' if po.get("forwarder_contact") else ""
        fwd_html = f"""<div style="flex:0 0 auto">
          <div style="font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#9ca3af;margin-bottom:3px">Freight Forwarder · 货运代理</div>
          <div style="display:inline;font-weight:600;color:#111">{_e(fwd_name)}</div>{fwd_contact}
        </div>""" if po.get("forwarder_name") or po.get("forwarder_contact") else ""
        
        addr_html = _e(po.get('delivery_address','')).replace("\n", "<br>")
        addr_html = f"""<div style="flex:1">
          <div style="font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#9ca3af;margin-bottom:3px">📍 Delivery Address · 交货地址</div>
          <div style="font-weight:600;color:#111;font-size:12px;line-height:1.6;background:#fff;border:1px solid #fed7aa;border-radius:5px;padding:8px 12px;margin-top:4px">{addr_html}</div>
        </div>""" if po.get("delivery_address") else ""
        
        delivery_logistics = f"""
        <div style="display:flex;gap:16px;align-items:flex-start;margin-top:6px;padding-top:6px;border-top:1px solid #fed7aa">
          {fwd_html}
          {addr_html}
        </div>"""

    remarks_html = ""
    if po.get("remarks"):
        remarks_html = f'<div style="margin-top:8px;padding:6px 10px;background:#fffbeb;border:1px solid #fde68a;border-left:4px solid #f59e0b;border-radius:0 6px 6px 0;font-size:11px;color:#92400e"><strong style="font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#78350f;margin-right:4px">REMARKS / SPECIAL INSTRUCTIONS · 备注:</strong>{_e(po["remarks"])}</div>'

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>PO {_e(po.get('po_number',''))}</title>
<link href="https://fonts.googleapis.com/css2?family=Noto+Sans+SC:wght@400;600;700&display=swap" rel="stylesheet">
<style>
  *{{box-sizing:border-box;margin:0;padding:0}}
  body{{font-family:'Noto Sans SC','Segoe UI',Arial,sans-serif;font-size:12px;color:#111;background:#fff;padding:32px 36px;max-width:920px;margin:0 auto}}
  .po-header{{display:flex;justify-content:space-between;align-items:flex-start;padding-bottom:12px;border-bottom:3px solid #1e3a8a;margin-bottom:14px}}
  .po-main-title{{font-size:28px;font-weight:700;color:#1e3a8a;letter-spacing:-.02em;line-height:1}}
  .po-subtitle{{font-size:11px;color:#6b7280;margin-top:3px;letter-spacing:.06em;text-transform:uppercase}}
  .po-header-right{{text-align:right;font-size:12px}}
  .po-meta-row{{display:flex;gap:6px;align-items:center;justify-content:flex-end;margin-bottom:4px;color:#374151}}
  .po-meta-row strong{{color:#111;min-width:60px;text-align:right}}
  .status-pill{{display:inline-block;padding:3px 12px;border-radius:20px;font-size:10px;font-weight:700;letter-spacing:.07em;text-transform:uppercase;color:#fff;background:{status_color};margin-top:8px}}
  .parties-grid{{display:grid;grid-template-columns:1fr 1fr;gap:0;border:1px solid #e5e7eb;border-radius:8px;overflow:hidden;margin-bottom:10px}}
  .party-box{{padding:10px 14px}}
  .party-box:first-child{{border-right:1px solid #e5e7eb;background:#f9fafb}}
  .party-box:last-child{{background:#f0f6ff}}
  .party-label{{font-size:9px;font-weight:700;letter-spacing:.1em;text-transform:uppercase;color:#6b7280;margin-bottom:8px;display:flex;align-items:center;gap:5px}}
  .party-label::before{{content:'';display:block;width:10px;height:2px;background:#1e3a8a;border-radius:1px}}
  .delivery-box{{background:#fff8f0;border:1.5px solid #f97316;border-radius:8px;padding:8px 14px;margin-bottom:10px}}
  .delivery-box-title{{font-size:9px;font-weight:700;letter-spacing:.1em;text-transform:uppercase;color:#c2410c;margin-bottom:6px;display:flex;align-items:center;gap:6px}}
  .delivery-box-title::before{{content:'📦';font-size:13px}}
  .delivery-grid{{display:grid;grid-template-columns:1fr 1fr 1.4fr 1fr;gap:8px 14px;font-size:11px}}
  table{{width:100%;border-collapse:collapse;margin-bottom:0}}
  thead tr{{background:#1e3a8a}}
  thead th{{padding:7px 10px;color:#fff;font-size:10px;font-weight:600;letter-spacing:.07em;text-transform:uppercase;text-align:left}}
  tbody tr:nth-child(even){{background:#f8faff}}
  tbody td{{padding:6px 10px;border-bottom:1px solid #e5e7eb;font-size:11px;vertical-align:top}}
  tfoot tr{{background:#1e3a8a}}
  tfoot td{{padding:7px 10px;color:#fff;font-weight:700;font-size:13px}}
  .table-wrap{{border:1px solid #d1d5db;border-radius:8px;overflow:hidden;margin-bottom:8px}}
  .sig-grid{{display:grid;grid-template-columns:1fr 1fr;gap:48px;margin-top:16px}}
  .sig-box{{border-top:2px solid #d1d5db;padding-top:8px}}
  .sig-label{{font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#9ca3af;margin-bottom:4px}}
  .sig-name{{font-size:12px;font-weight:700;color:#111}}
  .sig-sub{{font-size:10px;color:#6b7280;margin-top:2px}}
  .footer-note{{margin-top:12px;padding-top:8px;border-top:1px solid #e5e7eb;font-size:9px;color:#9ca3af;text-align:center;letter-spacing:.04em}}
  @media print{{
    body{{padding:10px 14px}}
    @page{{margin:6mm 8mm;size:A4}}
  }}
</style>
</head>
<body>

<div class="po-header">
  <div class="po-header-left">
    <div class="po-main-title">PURCHASE ORDER</div>
    <div class="po-subtitle">采购订单</div>
  </div>
  <div class="po-header-right">
    <div class="po-meta-row"><strong>PO No:</strong>&nbsp;{_e(po.get('po_number',''))}</div>
    <div class="po-meta-row"><strong>Date:</strong>&nbsp;{_e(po.get('po_date',''))}</div>
    <div class="po-meta-row"><strong>Currency:</strong>&nbsp;{_e(cur)}</div>
    <div><span class="status-pill">{_e(po.get('status',''))}</span></div>
  </div>
</div>

<div class="parties-grid">
  <div class="party-box">
    <div class="party-label">Buyer / Importer · 买方</div>
    {buyer_html}
  </div>
  <div class="party-box">
    <div class="party-label">Supplier / Vendor · 供应商</div>
    {supplier_html}
  </div>
</div>

<div class="delivery-box">
  <div class="delivery-box-title">Delivery &amp; Logistics Details · 交货与物流详情</div>
  <div class="delivery-grid">
    <div>
      <div style="font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#9ca3af;margin-bottom:3px">Delivery Terms · 交货条款</div>
      <div style="font-weight:600;color:#111;font-size:14px">{_e(po.get('delivery_terms','—'))}</div>
    </div>
    {lead_time_html}
    <div>
      <div style="font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#9ca3af;margin-bottom:3px">Payment Conditions · 付款条件</div>
      <div style="font-weight:600;color:#111;font-size:13px">{_e(pc or '—')}</div>
    </div>
    <div>{pay_breakdown_html}</div>
  </div>
  {delivery_logistics}
  {remarks_html}
</div>

<div class="table-wrap">
  <table>
    <thead><tr>
      <th style="width:3%;text-align:center">#</th>
      <th style="width:30%">Item Name<br><span style="font-weight:400;opacity:.75;font-size:9px">品名</span></th>
      <th style="width:10%">Description<br><span style="font-weight:400;opacity:.75;font-size:9px">描述</span></th>
      <th style="width:10%;text-align:center">HS Code</th>
      <th style="width:7%;text-align:center">Qty<br><span style="font-weight:400;opacity:.75;font-size:9px">数量</span></th>
      <th style="width:7%;text-align:center">Unit<br><span style="font-weight:400;opacity:.75;font-size:9px">单位</span></th>
      <th style="width:15%;text-align:right">Unit Price<br><span style="font-weight:400;opacity:.75;font-size:9px">单价</span></th>
      <th style="width:15%;text-align:right">Amount<br><span style="font-weight:400;opacity:.75;font-size:9px">金额</span></th>
    </tr></thead>
    <tbody>{item_rows}</tbody>
    <tfoot>
      <tr>
        <td colspan="4" style="text-align:right;font-size:11px;text-transform:uppercase;padding-right:14px;color:#e0e7ff">Total Quantity Required</td>
        <td style="text-align:center;font-size:13px;color:#fff">{total_qty:g}</td>
        <td colspan="2" style="text-align:right;font-size:11px;text-transform:uppercase;padding-right:14px;color:#e0e7ff">Grand Total · 总金额</td>
        <td style="text-align:right;font-size:14px">{_e(cur)} {fmt_total}</td>
      </tr>
    </tfoot>
  </table>
</div>
{bank_section}

<div class="sig-grid">
  <div class="sig-box">
    <div class="sig-label">Authorised Signatory — Buyer · 买方授权签字</div>
    <div class="sig-name">{_e(settings.get('company_name','_______________________'))}</div>
    <div class="sig-sub">Signature &amp; Stamp · 签名盖章</div>
  </div>
  <div class="sig-box">
    <div class="sig-label">Accepted — Supplier · 供应商确认</div>
    <div class="sig-name">{_e(sup.get('company','_______________________'))}</div>
    <div class="sig-sub">Signature &amp; Stamp · 签名盖章</div>
  </div>
</div>

<div class="footer-note">This is a computer-generated Purchase Order · {date.today().strftime('%d %b %Y')}</div>
</body></html>"""
    return html, 200, {"Content-Type": "text/html; charset=utf-8"}


@app.route("/api/quotations/<qid>/print", methods=["GET"])
def print_quotation(qid):
    """Generate a printable HTML view of a quotation matching the client-side design"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM quotations WHERE id = ?", (qid,))
        row = cursor.fetchone()
        if not row:
            return "Quotation not found", 404
        cursor.execute("SELECT key, value FROM settings")
        settings = {r["key"]: r["value"] for r in cursor.fetchall()}

    q = _quotation_to_dict(row)
    cur = q.get("currency", "CNY")
    cur_sym = "¥" if cur == "CNY" else "$"
    line_items = q.get("line_items", [])

    # Get all unique supplier names
    sup_names = []
    for li in line_items:
        for sr in li.get("supplier_rows", []):
            nm = sr.get("supplier_name", "")
            if nm and nm not in sup_names:
                sup_names.append(nm)

    sup_headers = "".join(f'<th class="th-sup">{_e(nm)}</th>' for nm in sup_names)

    item_rows = ""
    for i, li in enumerate(line_items):
        best_price = min((sr.get('price', 0) for sr in li.get('supplier_rows', []) if sr.get('price', 0) > 0), default=0)
        
        sup_cells = ""
        for nm in sup_names:
            sr_match = next((sr for sr in li.get("supplier_rows",[]) if sr.get("supplier_name") == nm), None)
            p = sr_match["price"] if sr_match else 0
            is_best = p > 0 and p == best_price
            style = ' style="background:#f0fdf4;color:#15803d;font-weight:700"' if is_best else ''
            sup_cells += f'<td class="td-sup"{style}>{cur_sym}{p:,.2f}</td>'

        best_cell = f'<td class="best-cell">{cur_sym}{best_price:,.2f}</td>' if best_price > 0 else '<td>—</td>'
        
        item_rows += f"""<tr>
          <td style="text-align:center;color:#64748b">{i+1}</td>
          <td><div class="item-name">{_e(li.get('item_name',''))}</div></td>
          <td><div class="item-desc">{_e(li.get('description',''))}</div></td>
          <td style="text-align:center">{li.get('qty',1)}</td>
          <td style="text-align:center">{_e(li.get('unit','PCS'))}</td>
          {sup_cells}
          {best_cell}
        </tr>"""

    html = f"""<!DOCTYPE html>
<html><head>
<meta charset="UTF-8">
<title>Comparison {_e(q.get('quotation_no',''))}</title>
<link href="https://fonts.googleapis.com/css2?family=Outfit:wght@400;600;800&family=Inter:wght@400;500;600;700&display=swap" rel="stylesheet">
<style>
  @page {{ size: A4 landscape; margin: 10mm 12mm; }}
  *{{box-sizing:border-box;margin:0;padding:0}}
  body{{font-family:'Inter',sans-serif;font-size:11px;color:#1e293b;background:#fff;padding:0;-webkit-print-color-adjust:exact;print-color-adjust:exact}}
  .doc-header{{display:flex;justify-content:space-between;align-items:flex-end;margin-bottom:20px;padding:25px 30px;padding-bottom:12px;border-bottom:2px solid #3b82f6}}
  .doc-title{{font-family:'Outfit',sans-serif;font-size:26px;font-weight:800;color:#1e3a8a}}
  .doc-subtitle{{font-family:'Outfit',sans-serif;font-size:11px;font-weight:600;color:#3b82f6;text-transform:uppercase}}
  .doc-meta{{text-align:right;font-size:11px;color:#64748b;display:grid;grid-template-columns:auto auto;gap:4px 20px}}
  .doc-meta span{{color:#0f172a;font-weight:700}}
  .main{{padding:0 30px}}
  table{{width:100%;border-collapse:separate;border-spacing:0;table-layout:fixed;margin-top:10px}}
  thead th{{padding:10px 8px;background:#1e3a8a;color:#fff;font-family:'Outfit',sans-serif;font-size:9px;text-transform:uppercase;text-align:left;border-right:1px solid rgba(255,255,255,0.1)}}
  th.th-sup{{text-align:right;min-width:100px}}
  tbody td{{padding:8px 10px;border-bottom:1px solid #e2e8f0;border-right:1px solid #f1f5f9;vertical-align:top}}
  .td-sup{{text-align:right}}
  .best-cell{{background:#fef3c7!important;color:#92400e!important;font-weight:800;text-align:right}}
  .item-name{{font-weight:700;color:#0f172a}}
  .item-desc{{color:#64748b;font-size:9.5px}}
  .footer{{margin-top:30px;padding:20px 30px;border-top:1px dashed #cbd5e1;font-size:10px;color:#64748b;display:flex;justify-content:space-between}}
</style>
</head><body>
<div class="doc-header">
  <div><div class="doc-title">QUOTATION COMPARISON</div><div class="doc-subtitle">{_e(q.get('title',''))}</div></div>
  <div class="doc-meta">
    <div>No: <span>{_e(q.get('quotation_no',''))}</span></div>
    <div>Date: <span>{_e(q.get('date',''))}</span></div>
    <div>Currency: <span>{_e(cur)}</span></div>
  </div>
</div>
<div class="main">
<table>
  <thead><tr>
    <th style="width:35px;text-align:center">#</th>
    <th style="width:18%">Product</th>
    <th style="width:15%">Description</th>
    <th style="width:50px;text-align:center">Qty</th>
    <th style="width:40px;text-align:center">Unit</th>
    {sup_headers}
    <th style="width:90px;text-align:right">Best Price</th>
  </tr></thead>
  <tbody>{item_rows}</tbody>
</table>
</div>
<div class="footer">
  <div>Generated by {_e(settings.get('company_name',''))}</div>
  <div>Printed: {date.today().strftime('%d %b %Y')}</div>
</div>
</body></html>"""
    return html, 200, {"Content-Type": "text/html; charset=utf-8"}

@app.route("/api/quotations/<qid>/export-excel", methods=["GET"])
def export_quotation_excel(qid):
    """Export quotation as CSV (Excel-compatible)"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM quotations WHERE id = ?", (qid,))
        row = cursor.fetchone()
        if not row:
            return "Quotation not found", 404

    q = _quotation_to_dict(row)
    cur = q.get("currency", "CNY")
    line_items = q.get("line_items", [])

    # Get all unique supplier names
    sup_names = []
    for li in line_items:
        for sr in li.get("supplier_rows", []):
            nm = sr.get("supplier_name", "")
            if nm and nm not in sup_names:
                sup_names.append(nm)

    # Build CSV
    output = io.StringIO()
    writer = csv.writer(output)

    # Header row
    writer.writerow(["#", "Product", "Description", "Qty", "Unit"] + sup_names + ["Best Price", "Best Supplier"])

    for i, li in enumerate(line_items):
        prices = {}
        for sr in li.get("supplier_rows", []):
            prices[sr.get("supplier_name", "")] = sr.get("price", 0)
        price_vals = [prices.get(nm, 0) for nm in sup_names]
        non_zero = [(nm, p) for nm, p in zip(sup_names, price_vals) if p > 0]
        best_price = min(non_zero, key=lambda x: x[1]) if non_zero else ("", 0)
        writer.writerow([
            i + 1,
            li.get("item_name", ""),
            li.get("description", ""),
            li.get("qty", 1),
            li.get("unit", "PCS"),
        ] + price_vals + [best_price[1], best_price[0]])

    csv_bytes = output.getvalue().encode("utf-8-sig")  # BOM for Excel
    return send_file(
        BytesIO(csv_bytes),
        mimetype="text/csv",
        as_attachment=True,
        download_name=f"Quotation_{q.get('quotation_no','export')}_{date.today()}.csv"
    )


# ═══════════════════════════════════════════════════════════════════════════════
# PO REPORT (unchanged from original)
# ═══════════════════════════════════════════════════════════════════════════════


@app.route("/api/po-report/<pid>", methods=["GET"])
def po_report(pid):
    """Generate PO report"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM purchase_orders WHERE id = ?", (pid,))
        po_row = cursor.fetchone()
        
        if not po_row:
            return "PO not found", 404
        
        po = dict(po_row)
        po["supplier_snapshot"] = json.loads(po.get("supplier_snapshot") or "{}")
        
        # Fetch line items
        cursor.execute("SELECT * FROM po_items WHERE po_id = ? ORDER BY line_sequence", (pid,))
        po["line_items"] = [
            {
                "item_name": item["item_name"],
                "description": item["description"],
                "hs_code": item["hs_code"],
                "qty": item["qty"],
                "unit": item["unit"],
                "unit_price": item["unit_price"],
            }
            for item in cursor.fetchall()
        ]
        
        # Fetch settings
        cursor.execute("SELECT key, value FROM settings")
        settings = {row["key"]: row["value"] for row in cursor.fetchall()}
    
    s = po.get("supplier_snapshot", {})
    cfg = settings

    currency = po.get("currency", "USD")
    cur_sym = "¥" if currency == "CNY" else "$"

    lines = po.get("line_items", [])
    total = sum(_safe_float(li.get("qty", 0)) * _safe_float(li.get("unit_price", 0)) for li in lines)
    total_qty = sum(_safe_float(li.get("qty", 0)) for li in lines)

    status_str = po.get("status", "Draft").upper()

    rows_html = ""
    for i, li in enumerate(lines):
        item_total = _safe_float(li.get("qty", 0)) * _safe_float(li.get("unit_price", 0))
        hs = li.get('hs_code', '') or '—'
        desc = li.get('description', '') or '—'
        item_name = li.get('item_name', '') or f'Item {i+1}'
        rows_html += f"""<tr>
          <td style="text-align:center;color:#6b7280;font-size:10px">{i+1}</td>
          <td style="font-weight:600;color:#111827">{_e(item_name)}</td>
          <td style="text-align:center;color:#374151">{_e(desc)}</td>
          <td style="text-align:center;color:#374151;font-family:monospace">{_e(hs)}</td>
          <td style="text-align:center;font-weight:600;color:#111827">{li.get('qty','')}</td>
          <td style="text-align:center;color:#374151">{_e(li.get('unit',''))}</td>
          <td style="text-align:right;color:#1d4ed8;font-weight:600">{_e(currency)} {li.get('unit_price',0):,.2f}</td>
          <td style="text-align:right;font-weight:700;color:#1d4ed8">{_e(currency)} {item_total:,.2f}</td>
        </tr>"""

    po_date_fmt = ""
    try:
        po_date_fmt = datetime.strptime(po.get("po_date", ""), "%Y-%m-%d").strftime("%Y-%m-%d")
    except:
        po_date_fmt = po.get("po_date", "")

    # Build contact line for supplier
    sup_contacts = []
    if s.get('email'): sup_contacts.append(f"Email: {s.get('email')}")
    if s.get('phone'): sup_contacts.append(f"Ph/WA: {s.get('phone')}")
    sup_contact_html = "<br>".join([_e(c) for c in sup_contacts])

    buyer_lines = []
    if cfg.get('company_address'): buyer_lines.append(cfg['company_address'])
    if cfg.get('company_phone'):   buyer_lines.append(f"Ph: {cfg['company_phone']}")
    if cfg.get('company_email'):   buyer_lines.append(f"Email: {cfg['company_email']}")
    if cfg.get('company_gstin'):   buyer_lines.append(f"GSTIN: {cfg['company_gstin']}")
    buyer_detail_html = "<br>".join([_e(b) for b in buyer_lines])

    pc = po.get("payment_conditions") or ""
    payment_breakdown_html = ""
    m = re.match(r'(\d+)%\s*Advance.*?(\d+)%', pc, re.IGNORECASE)
    if m and total > 0:
        pct1 = int(m.group(1))
        pct2 = int(m.group(2))
        adv  = total * pct1 / 100
        rem  = total * pct2 / 100
        lbl2 = "against B/L (Bill of Lading)" if "BL" in pc.upper() else "on Shipment"
        payment_breakdown_html = f"""
        <div style="margin-top:6px;padding:7px 12px;background:#eff6ff;border:1px solid #bfdbfe;border-radius:5px;font-size:10px;line-height:1.8">
          <div style="font-size:8px;font-weight:700;letter-spacing:.1em;text-transform:uppercase;color:#1d4ed8;margin-bottom:3px">Payment Schedule · 付款计划</div>
          <div style="display:flex;justify-content:space-between;gap:8px;flex-wrap:nowrap"><span style="color:#374151;white-space:nowrap">&#9312; {pct1}% Advance (Proforma Invoice)</span><span style="font-weight:700;color:#1d4ed8;white-space:nowrap">{_e(currency)} {adv:,.2f}</span></div>
          <div style="display:flex;justify-content:space-between;gap:8px;flex-wrap:nowrap"><span style="color:#374151;white-space:nowrap">&#9313; {pct2}% {lbl2}</span><span style="font-weight:700;color:#1d4ed8;white-space:nowrap">{_e(currency)} {rem:,.2f}</span></div>
        </div>"""

    lead_time_days = po.get("lead_time_days", 0) or 0
    due_date_val   = po.get("due_date", "") or ""
    lead_time_html = ""
    if lead_time_days or due_date_val:
        due_line = f'<div style="font-size:10px;color:#1d4ed8;margin-top:2px;font-weight:600">Due: {_e(due_date_val)}</div>' if due_date_val else ""
        lead_time_html = f"""<div class="lg-block lg-divider">
      <div class="lg-sub">LEAD TIME · 交货期</div>
      <div class="lg-val">{lead_time_days} Days</div>
      {due_line}
    </div>"""

    # (Truncated for brevity - rest is same HTML generation as original)
    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>PO {_e(po.get('po_number',''))}</title>
<link href="https://fonts.googleapis.com/css2?family=Noto+Sans+SC:wght@400;600;700&display=swap" rel="stylesheet">
<style>
  *{{box-sizing:border-box;margin:0;padding:0}}
  body{{font-family:'Noto Sans SC','Segoe UI',Arial,sans-serif;font-size:12px;color:#111;background:#fff;padding:32px 36px;max-width:920px;margin:0 auto}}
  .po-header{{display:flex;justify-content:space-between;align-items:flex-start;padding-bottom:12px;border-bottom:3px solid #1e3a8a;margin-bottom:14px}}
  .po-main-title{{font-size:28px;font-weight:700;color:#1e3a8a;letter-spacing:-.02em;line-height:1}}
  .po-subtitle{{font-size:11px;color:#6b7280;margin-top:3px;letter-spacing:.06em;text-transform:uppercase}}
  .po-header-right{{text-align:right;font-size:12px}}
  .po-meta-row{{display:flex;gap:6px;align-items:center;justify-content:flex-end;margin-bottom:4px;color:#374151}}
  .po-meta-row strong{{color:#111;min-width:60px;text-align:right}}
  .parties-grid{{display:grid;grid-template-columns:1fr 1fr;gap:0;border:1px solid #e5e7eb;border-radius:8px;overflow:hidden;margin-bottom:10px}}
  .party-box{{padding:10px 14px}}
  .party-box:first-child{{border-right:1px solid #e5e7eb;background:#f9fafb}}
  .party-box:last-child{{background:#f0f6ff}}
  .party-label{{font-size:9px;font-weight:700;letter-spacing:.1em;text-transform:uppercase;color:#6b7280;margin-bottom:8px;display:flex;align-items:center;gap:5px}}
  .party-label::before{{content:'';display:block;width:10px;height:2px;background:#1e3a8a;border-radius:1px}}
  .delivery-box{{background:#fff8f0;border:1.5px solid #f97316;border-radius:8px;padding:8px 14px;margin-bottom:10px}}
  .delivery-box-title{{font-size:9px;font-weight:700;letter-spacing:.1em;text-transform:uppercase;color:#c2410c;margin-bottom:6px;display:flex;align-items:center;gap:6px}}
  .delivery-box-title::before{{content:'📦';font-size:13px}}
  .delivery-grid{{display:grid;grid-template-columns:1fr 1fr 1.4fr 1fr;gap:8px 14px;font-size:11px}}
  table{{width:100%;border-collapse:collapse;margin-bottom:0}}
  thead tr{{background:#1e3a8a}}
  thead th{{padding:7px 10px;color:#fff;font-size:10px;font-weight:600;letter-spacing:.07em;text-transform:uppercase;text-align:left}}
  tbody tr:nth-child(even){{background:#f8faff}}
  tbody td{{padding:6px 10px;border-bottom:1px solid #e5e7eb;font-size:11px;vertical-align:top}}
  tfoot tr{{background:#1e3a8a}}
  tfoot td{{padding:7px 10px;color:#fff;font-weight:700;font-size:13px}}
  .table-wrap{{border:1px solid #d1d5db;border-radius:8px;overflow:hidden;margin-bottom:8px}}
  .sig-grid{{display:grid;grid-template-columns:1fr 1fr;gap:48px;margin-top:16px}}
  .sig-box{{border-top:2px solid #d1d5db;padding-top:8px}}
  .sig-label{{font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#9ca3af;margin-bottom:4px}}
  .sig-name{{font-size:12px;font-weight:700;color:#111}}
  .sig-sub{{font-size:10px;color:#6b7280;margin-top:2px}}
  .footer-note{{margin-top:12px;padding-top:8px;border-top:1px solid #e5e7eb;font-size:9px;color:#9ca3af;text-align:center;letter-spacing:.04em}}
  @media print{{
    body{{padding:10px 14px}}
    @page{{margin:6mm 8mm;size:A4}}
  }}
</style>
</head>
<body>

<div class="po-header">
  <div class="po-header-left">
    <div class="po-main-title">PURCHASE ORDER</div>
    <div class="po-subtitle">采购订单</div>
  </div>
  <div class="po-header-right">
    <div class="po-meta-row"><strong>PO No:</strong>&nbsp;{_e(po.get('po_number',''))}</div>
    <div class="po-meta-row"><strong>Date:</strong>&nbsp;{_e(po_date_fmt)}</div>
    <div class="po-meta-row"><strong>Currency:</strong>&nbsp;{_e(currency)}</div>
    <div><span style="display:inline-block;padding:3px 12px;border-radius:20px;font-size:10px;font-weight:700;letter-spacing:.07em;text-transform:uppercase;color:#fff;background:#6b7280;margin-top:8px">{_e(status_str)}</span></div>
  </div>
</div>

<div class="parties-grid">
  <div class="party-box">
    <div class="party-label">Buyer / Importer · 买方</div>
    <div style="font-size:14px;font-weight:700;color:#111;margin-bottom:3px">{_e(cfg.get('company_name',''))}</div>
    <div style="color:#4b5563">{buyer_detail_html}</div>
  </div>
  <div class="party-box">
    <div class="party-label">Supplier / Vendor · 供应商</div>
    <div style="font-size:14px;font-weight:700;color:#111;margin-bottom:3px">{_e(s.get('company','—'))}</div>
    <div style="color:#4b5563">{sup_contact_html}</div>
  </div>
</div>

<div class="delivery-box">
  <div class="delivery-box-title">Delivery &amp; Logistics Details · 交货与物流详情</div>
  <div class="delivery-grid">
    <div>
      <div style="font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#9ca3af;margin-bottom:3px">Delivery Terms · 交货条款</div>
      <div style="font-weight:600;color:#111;font-size:14px">{_e(po.get('delivery_terms','—'))}</div>
    </div>
    {lead_time_html}
    <div>
      <div style="font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#9ca3af;margin-bottom:3px">Payment Conditions · 付款条件</div>
      <div style="font-weight:600;color:#111;font-size:13px">{_e(pc or '—')}</div>
    </div>
    <div>{payment_breakdown_html}</div>
  </div>
</div>

<div class="table-wrap">
  <table>
    <thead><tr>
      <th style="width:3%;text-align:center">#</th>
      <th style="width:30%">Item Name<br><span style="font-weight:400;opacity:.75;font-size:9px">品名</span></th>
      <th style="width:10%">Description<br><span style="font-weight:400;opacity:.75;font-size:9px">描述</span></th>
      <th style="width:10%;text-align:center">HS Code</th>
      <th style="width:7%;text-align:center">Qty<br><span style="font-weight:400;opacity:.75;font-size:9px">数量</span></th>
      <th style="width:7%;text-align:center">Unit<br><span style="font-weight:400;opacity:.75;font-size:9px">单位</span></th>
      <th style="width:15%;text-align:right">Unit Price<br><span style="font-weight:400;opacity:.75;font-size:9px">单价</span></th>
      <th style="width:15%;text-align:right">Amount<br><span style="font-weight:400;opacity:.75;font-size:9px">金额</span></th>
    </tr></thead>
    <tbody>{rows_html}</tbody>
    <tfoot>
      <tr>
        <td colspan="4" style="text-align:right;font-size:11px;text-transform:uppercase;padding-right:14px;color:#e0e7ff">Total Quantity Required</td>
        <td style="text-align:center;font-size:13px;color:#fff">{total_qty:g}</td>
        <td colspan="2" style="text-align:right;font-size:11px;text-transform:uppercase;padding-right:14px;color:#e0e7ff">Grand Total · 总金额</td>
        <td style="text-align:right;font-size:14px">{_e(currency)} {total:,.2f}</td>
      </tr>
    </tfoot>
  </table>
</div>

<div class="sig-grid">
  <div class="sig-box">
    <div class="sig-label">Authorised Signatory — Buyer · 买方授权签字</div>
    <div class="sig-name">{_e(cfg.get('company_name','_______________________'))}</div>
    <div class="sig-sub">Signature &amp; Stamp · 签名盖章</div>
  </div>
  <div class="sig-box">
    <div class="sig-label">Accepted — Supplier · 供应商确认</div>
    <div class="sig-name">{_e(s.get('company','_______________________'))}</div>
    <div class="sig-sub">Signature &amp; Stamp · 签名盖章</div>
  </div>
</div>

<div class="footer-note">This is a computer-generated Purchase Order · {date.today().strftime('%d %b %Y')}</div>
</body></html>'''

    try:
        if WEASYPRINT_AVAILABLE:
            pdf_bytes = HTML(string=html, base_url=None).write_pdf()
            return send_file(
                BytesIO(pdf_bytes),
                mimetype="application/pdf",
                as_attachment=False,
                download_name=f"PO_{po.get('po_number','export')}_{date.today()}.pdf"
            )
    except Exception:
        pass
    
    return html, 200, {"Content-Type": "text/html; charset=utf-8"}


# ═══════════════════════════════════════════════════════════════════════════════
# LANDED COST REPORT
# ═══════════════════════════════════════════════════════════════════════════════

def _build_lc_html(po, settings, currency, c, rate, bank, ship, duty, trans, gst_duty, lc_cur, doc_pct):
    cur_sym = "¥" if currency == "CNY" else "$"
    company_name = settings.get("company_name", "")
    total_qty = sum(it.get('qty', 0) for it in c.get('items', []))

    # ── Build HTML Report ────────────────────────────────────────────────────

    # Item rows for main table
    item_rows = ""
    for i, it in enumerate(c.get("items", [])):
        item_rows += f"""<tr>
          <td style="text-align:center;color:#6b7280;font-size:9px">{i+1}</td>
          <td style="font-weight:600;color:#111827;font-size:10px">{_e(it.get('name',''))}</td>
          <td style="text-align:center;font-weight:600;font-size:10px">{it.get('qty',0):,.0f}</td>
          <td style="text-align:right;color:#1d4ed8;font-weight:600;font-size:10px">{cur_sym}{it.get('unitPrice',0):,.2f}</td>
          <td style="text-align:right;font-weight:600;font-size:10px">{fmt_inr(it.get('item_inr',0))}</td>
          <td style="text-align:center;color:#6b7280;font-size:10px">{it.get('share',0):.1f}%</td>
          <td style="text-align:right;color:#059669;font-weight:600;font-size:10px">{fmt_inr(it.get('addl_share',0))}</td>
          <td style="text-align:right;font-weight:700;color:#7c3aed;font-size:10px">{fmt_inr(it.get('total_item',0))}</td>
          <td style="text-align:right;font-weight:700;color:#dc2626;font-size:10px">{fmt_inr(it.get('per_unit',0))}</td>
        </tr>"""

    # Cost breakdown rows for overhead table
    overhead_items = [
        ("Banking Charges", bank, "bank_s"),
        ("Shipping / Freight", ship, "ship_s"),
        ("Customs Duty", duty, "duty_s"),
        ("Transport / Clearing", trans, "trans_s"),
        ("GST on Duty", gst_duty, "gst_s"),
        (f"Documentation ({doc_pct}%)", c.get("doc_inr", 0), "doc_s"),
    ]
    overhead_rows = ""
    for label, total_val, share_key in overhead_items:
        if total_val > 0:
            overhead_rows += f"""<tr>
              <td style="color:#374151;font-weight:500;font-size:10px">{_e(label)}</td>
              <td style="text-align:right;font-weight:600;color:#111;font-size:10px">{fmt_inr(total_val)}</td>"""
            for it in c.get("items", []):
                overhead_rows += f'<td style="text-align:right;color:#6b7280;font-size:10px">{fmt_inr(it.get(share_key,0))}</td>'
            overhead_rows += "</tr>"

    # Item headers for overhead table
    item_headers = ""
    for it in c.get("items", []):
        item_headers += f'<th style="text-align:right;min-width:70px;font-size:9px">{it.get("name","")}</th>'

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Landing Cost — PO {_e(po.get('po_number',''))}</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap" rel="stylesheet">
<style>
  *{{box-sizing:border-box;margin:0;padding:0}}
  body{{font-family:'Inter','Segoe UI',Arial,sans-serif;font-size:11px;color:#111;background:#f8fafc;padding:0;margin:0}}
  .page{{max-width:100%;margin:0 auto;background:#fff;min-height:100vh;padding:20px 24px}}

  /* Header */
  .report-header{{display:grid;grid-template-columns:1fr;gap:12px;padding-bottom:14px;border-bottom:3px solid #1e3a8a;margin-bottom:16px}}
  .report-title{{font-size:22px;font-weight:800;color:#1e3a8a;letter-spacing:-.03em;line-height:1}}
  .report-subtitle{{font-size:10px;color:#6b7280;margin-top:2px;letter-spacing:.06em;text-transform:uppercase}}
  .report-meta{{display:grid;grid-template-columns:1fr 1fr;gap:6px;font-size:10px;color:#374151}}
  .report-meta strong{{color:#111}}

  /* Summary Cards */
  .summary-grid{{display:grid;grid-template-columns:repeat(2,1fr);gap:12px;margin-bottom:24px}}
  .summary-card{{background:linear-gradient(135deg,#f8faff,#eff6ff);border:1px solid #dbeafe;border-radius:8px;padding:12px 14px;text-align:center}}
  .summary-card.highlight{{background:linear-gradient(135deg,#faf5ff,#f3e8ff);border-color:#e9d5ff}}
  .summary-card.grand{{background:linear-gradient(135deg,#ecfdf5,#d1fae5);border-color:#a7f3d0}}
  .summary-card.rate{{background:linear-gradient(135deg,#fffbeb,#fef3c7);border-color:#fde68a}}
  .sc-label{{font-size:8px;font-weight:700;letter-spacing:.1em;text-transform:uppercase;color:#6b7280;margin-bottom:4px}}
  .sc-value{{font-size:18px;font-weight:800;color:#1e3a8a}}
  .summary-card.highlight .sc-value{{color:#7c3aed}}
  .summary-card.grand .sc-value{{color:#059669}}
  .summary-card.rate .sc-value{{color:#d97706;font-size:16px}}

  /* Tables */
  .section-title{{font-size:13px;font-weight:700;color:#1e3a8a;margin:16px 0 8px;display:flex;align-items:center;gap:8px}}
  .section-title::before{{content:'';display:block;width:4px;height:16px;background:#1e3a8a;border-radius:2px}}
  .table-wrap{{border:1px solid #d1d5db;border-radius:6px;overflow:hidden;margin-bottom:14px;font-size:10px}}
  table{{width:100%;border-collapse:collapse}}
  thead tr{{background:#1e3a8a}}
  thead th{{padding:6px 8px;color:#fff;font-size:9px;font-weight:600;letter-spacing:.05em;text-transform:uppercase;text-align:left;white-space:nowrap}}
  tbody tr:nth-child(even){{background:#f8faff}}
  tbody tr:nth-child(odd){{background:#fff}}
  tbody td{{padding:5px 8px;border-bottom:1px solid #e5e7eb;font-size:10px;vertical-align:middle}}
  tfoot tr{{background:#1e3a8a}}
  tfoot td{{padding:6px 8px;color:#fff;font-weight:700;font-size:10px}}

  /* Overhead breakdown */
  .overhead-wrap{{border:1px solid #e9d5ff;border-radius:6px;overflow:hidden;margin-bottom:14px;font-size:10px}}
  .overhead-wrap thead tr{{background:#7c3aed}}
  .overhead-wrap tfoot tr{{background:#7c3aed}}

  /* Footer */
  .report-footer{{margin-top:14px;padding-top:8px;border-top:1px solid #e5e7eb;font-size:8px;color:#9ca3af;text-align:center;letter-spacing:.04em}}

  @media print{{
    body{{background:#fff;padding:0}}
    .page{{padding:14px 18px;box-shadow:none}}
    @page{{margin:10mm;size:A4 portrait}}
    .table-wrap{{page-break-inside:avoid}}
    .overhead-wrap{{page-break-inside:avoid}}
    .summary-grid{{page-break-inside:avoid}}
    thead{{display:table-header-group}}
    tr{{page-break-inside:avoid}}
  }}
</style>
</head>
<body>
<div class="page">

  <!-- HEADER -->
  <div class="report-header">
    <div>
      <div class="report-title">LANDED COST ANALYSIS</div>
      <div class="report-subtitle">Import Costing Breakdown Report</div>
    </div>
    <div class="report-meta">
      <div><strong>PO No:</strong> {_e(po.get('po_number',''))}</div>
      <div><strong>Date:</strong> {po.get('po_date','')}</div>
      <div><strong>Currency:</strong> {currency}</div>
      <div><strong>Status:</strong> {po.get('status','Draft')}</div>
      {f'<div style="margin-top:4px;font-weight:600;color:#1e3a8a">{_e(company_name)}</div>' if company_name else ''}
    </div>
  </div>

  <!-- SUMMARY CARDS -->
  <div class="summary-grid">
    <div class="summary-card rate">
      <div class="sc-label">{lc_cur} → INR Rate</div>
      <div class="sc-value">{rate:,.2f}</div>
    </div>
    <div class="summary-card">
      <div class="sc-label">Invoice ({currency})</div>
      <div class="sc-value">{cur_sym}{c.get('inv_usd',0):,.2f}</div>
    </div>
    <div class="summary-card highlight">
      <div class="sc-label">Total Overhead (₹)</div>
      <div class="sc-value">{fmt_inr(c.get('total_addl',0))}</div>
    </div>
    <div class="summary-card grand">
      <div class="sc-label">Grand Landed Cost (₹)</div>
      <div class="sc-value">{fmt_inr(c.get('grand',0))}</div>
    </div>
  </div>

  <!-- ITEM-WISE BREAKDOWN -->
  <div class="section-title">Item-Wise Landed Cost Breakdown</div>
  <div class="table-wrap">
    <table>
      <thead><tr>
        <th style="width:3%;text-align:center">#</th>
        <th style="width:28%">Item Name</th>
        <th style="text-align:center;width:10%">Qty</th>
        <th style="text-align:right;width:16%">Unit Price</th>
        <th style="text-align:right;width:14%">INR Value</th>
        <th style="text-align:center;width:9%">Share %</th>
        <th style="text-align:right;width:14%">Overhead</th>
        <th style="text-align:right;width:14%">Total Landed</th>
        <th style="text-align:right;width:12%;background:#b91c1c">Per Unit</th>
      </tr></thead>
      <tbody>{item_rows}</tbody>
      <tfoot>
        <tr>
          <td colspan="2" style="text-align:right;font-size:10px;letter-spacing:.04em;text-transform:uppercase;color:#e0e7ff">Total Quantity Required</td>
          <td style="text-align:center;font-size:11px;color:#fff">{total_qty:g}</td>
          <td style="text-align:right;font-size:10px;letter-spacing:.04em;text-transform:uppercase;color:#e0e7ff">Grand Total</td>
          <td style="text-align:right">{fmt_inr(c.get('inv_inr',0))}</td>
          <td style="text-align:center">100%</td>
          <td style="text-align:right">{fmt_inr(c.get('total_addl',0))}</td>
          <td style="text-align:right;font-size:14px">{fmt_inr(c.get('grand',0))}</td>
          <td></td>
        </tr>
      </tfoot>
    </table>
  </div>

  <!-- OVERHEAD ALLOCATION -->
  <div class="section-title" style="color:#7c3aed">Overhead Cost Allocation by Item</div>
  <div class="overhead-wrap">
    <table>
      <thead><tr>
        <th style="min-width:160px">Cost Head</th>
        <th style="text-align:right;min-width:100px">Total (₹)</th>
        {item_headers}
      </tr></thead>
      <tbody>{overhead_rows}</tbody>
      <tfoot>
        <tr>
          <td style="font-weight:700">TOTAL OVERHEAD</td>
          <td style="text-align:right;font-weight:700">{fmt_inr(c.get('total_addl',0))}</td>"""

    for it in c.get("items", []):
        html += f'<td style="text-align:right;font-weight:700">{fmt_inr(it.get("addl_share",0))}</td>'

    html += f"""
        </tr>
      </tfoot>
    </table>
  </div>

  <!-- COST PARAMETERS -->
  <div style="display:grid;grid-template-columns:1fr;gap:14px;margin-top:14px">
    <div style="background:#f8fafc;border:1px solid #e2e8f0;border-radius:8px;padding:12px 16px">
      <div style="font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#64748b;margin-bottom:8px">Cost Parameters Used</div>
      <div style="display:grid;grid-template-columns:1fr 1fr;gap:4px 12px;font-size:10px">
        <div style="color:#64748b">Exchange Rate ({lc_cur}→INR):</div><div style="font-weight:600;color:#111">₹{rate:,.2f}</div>
        <div style="color:#64748b">Banking Charges:</div><div style="font-weight:600;color:#111">{fmt_inr(bank)}</div>
        <div style="color:#64748b">Shipping / Freight:</div><div style="font-weight:600;color:#111">{fmt_inr(ship)}</div>
        <div style="color:#64748b">Customs Duty:</div><div style="font-weight:600;color:#111">{fmt_inr(duty)}</div>
        <div style="color:#64748b">Transport / Clearing:</div><div style="font-weight:600;color:#111">{fmt_inr(trans)}</div>
        <div style="color:#64748b">GST on Duty:</div><div style="font-weight:600;color:#111">{fmt_inr(gst_duty)}</div>
        <div style="color:#64748b">Documentation ({doc_pct}%):</div><div style="font-weight:600;color:#111">{fmt_inr(c.get('doc_inr',0))}</div>
      </div>
    </div>
    <div style="background:#ecfdf5;border:1px solid #a7f3d0;border-radius:8px;padding:12px 16px;display:flex;flex-direction:column;justify-content:center;align-items:center">
      <div style="font-size:9px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#059669;margin-bottom:4px">Total Landed Cost</div>
      <div style="font-size:28px;font-weight:800;color:#059669">{fmt_inr(c.get('grand',0))}</div>
      <div style="font-size:10px;color:#6b7280;margin-top:2px">Invoice: {fmt_inr(c.get('inv_inr',0))} + Overhead: {fmt_inr(c.get('total_addl',0))}</div>
    </div>
  </div>

  <div class="report-footer">
    This is a computer-generated Landed Cost Report · {_e(company_name)} · Generated on {date.today().strftime('%d %b %Y')}
  </div>

</div>
</body>
</html>"""


    return html

@app.route("/api/lc-report/<pid>", methods=["POST"])
def lc_report(pid):
    """Generate landed cost report"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM purchase_orders WHERE id = ?", (pid,))
        po_row = cursor.fetchone()
        
        if not po_row:
            return "PO not found", 404
        
        po = dict(po_row)
        
        # Fetch line items
        cursor.execute("SELECT * FROM po_items WHERE po_id = ? ORDER BY line_sequence", (pid,))
        po["line_items"] = [dict(item) for item in cursor.fetchall()]
        
        # Fetch settings
        cursor.execute("SELECT key, value FROM settings")
        settings = {row["key"]: row["value"] for row in cursor.fetchall()}

    req = request.get_json(force=True)
    currency = po.get("currency", "USD")
    lc_cur = "RMB" if currency == "CNY" else "USD"
    rate = float(req.get("rate", po.get("lc_rmb_rate" if lc_cur == "RMB" else "lc_usd_rate", 84)))
    bank = float(req.get("bank", po.get("lc_bank", 0)))
    ship = float(req.get("ship", po.get("lc_ship", 0)))
    duty = float(req.get("duty", po.get("lc_duty", 0)))
    trans = float(req.get("trans", po.get("lc_trans", 0)))
    gst_duty = float(req.get("gst_duty", po.get("lc_gst_duty", 0)))
    doc_pct = float(req.get("doc_pct", po.get("lc_doc_pct", 0)))

    items = []
    for li in po.get("line_items", []):
        items.append({
            "name": li.get("item_name", ""),
            "qty": float(li.get("qty", 0)),
            "unitPrice": float(li.get("unit_price", 0)),
        })

    c = calc_landed(items, rate, bank, ship, duty, trans, gst_duty, doc_pct)

    # ── Build HTML Report ────────────────────────────────────────────────────
    html = _build_lc_html(po, settings, currency, c, rate, bank, ship, duty, trans, gst_duty, lc_cur, doc_pct)
    return html, 200, {"Content-Type": "text/html; charset=utf-8"}


@app.route("/api/lc-report/<pid>/pdf", methods=["POST"])
def lc_report_pdf(pid):
    """Generate and download landed cost report as PDF (A4 Portrait)"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM purchase_orders WHERE id = ?", (pid,))
        po_row = cursor.fetchone()
        if not po_row: return "PO not found", 404
        po = dict(po_row)
        cursor.execute("SELECT * FROM po_items WHERE po_id = ? ORDER BY line_sequence", (pid,))
        po["line_items"] = [dict(item) for item in cursor.fetchall()]
        cursor.execute("SELECT key, value FROM settings")
        settings = {row["key"]: row["value"] for row in cursor.fetchall()}

    req = request.get_json(force=True)
    currency = po.get("currency", "USD")
    lc_cur = "RMB" if currency == "CNY" else "USD"
    rate = float(req.get("rate", po.get("lc_rmb_rate" if lc_cur == "RMB" else "lc_usd_rate", 84)))
    bank = float(req.get("bank", po.get("lc_bank", 0)))
    ship = float(req.get("ship", po.get("lc_ship", 0)))
    duty = float(req.get("duty", po.get("lc_duty", 0)))
    trans = float(req.get("trans", po.get("lc_trans", 0)))
    gst_duty = float(req.get("gst_duty", po.get("lc_gst_duty", 0)))
    doc_pct = float(req.get("doc_pct", po.get("lc_doc_pct", 0)))

    items = []
    for li in po.get("line_items", []):
        items.append({
            "name": li.get("item_name", ""),
            "qty": float(li.get("qty", 0)),
            "unitPrice": float(li.get("unit_price", 0)),
        })

    c = calc_landed(items, rate, bank, ship, duty, trans, gst_duty, doc_pct)
    html = _build_lc_html(po, settings, currency, c, rate, bank, ship, duty, trans, gst_duty, lc_cur, doc_pct)
    # Generate PDF from HTML
    try:
        if WEASYPRINT_AVAILABLE:
            po_number = po["po_number"] if po.get("po_number") else "LC_Report"
            
            pdf_bytes = HTML(string=html, base_url=None).write_pdf()
            return send_file(
                BytesIO(pdf_bytes),
                mimetype="application/pdf",
                as_attachment=True,
                download_name=f"Landed_Cost_Analysis_PO_{po_number}_{date.today().strftime('%d_%b_%Y')}.pdf"
            )
        else:
            return jsonify({"error": "PDF generation not available. WeasyPrint is not installed."}), 500
    except Exception as e:
        return jsonify({"error": f"PDF generation failed: {str(e)}"}), 500



# ═══════════════════════════════════════════════════════════════════════════════
# ATTACHMENTS
# ═══════════════════════════════════════════════════════════════════════════════

ALLOWED_EXTENSIONS = {
    'pdf', 'png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp',
    'doc', 'docx', 'xls', 'xlsx', 'csv', 'txt', 'zip', 'rar',
}
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB


def _attach_dir(pid):
    """Get/create per-PO attachment directory (physical files still on disk)"""
    d = os.path.join(ATTACH_DIR, pid)
    os.makedirs(d, exist_ok=True)
    return d


def _migrate_attach_json(pid, conn):
    """One-time migration: import legacy _meta.json into po_attachments table."""
    meta_path = os.path.join(_attach_dir(pid), "_meta.json")
    if not os.path.exists(meta_path):
        return
    try:
        with open(meta_path, "r", encoding="utf-8") as f:
            entries = json.load(f)
        cursor = conn.cursor()
        for e in entries:
            cursor.execute(
                "INSERT OR IGNORE INTO po_attachments "
                "(id, po_id, filename, original, label, mime, uploaded_at) "
                "VALUES (?, ?, ?, ?, ?, ?, ?)",
                (e["id"], pid, e["filename"], e.get("original", e["filename"]),
                 e.get("label", e["filename"]),
                 e.get("mime", "application/octet-stream"),
                 e.get("uploaded_at", datetime.now().strftime("%Y-%m-%d %H:%M")))
            )
        os.rename(meta_path, meta_path + ".migrated")
    except Exception as ex:
        print(f"[attach-migrate] {pid}: {ex}")


@app.route("/api/po/<pid>/attachments", methods=["GET"])
def list_attachments(pid):
    """List all attachments for a PO"""
    with get_db() as conn:
        _migrate_attach_json(pid, conn)
        cursor = conn.cursor()
        cursor.execute(
            "SELECT id, filename, original, label, mime, uploaded_at "
            "FROM po_attachments WHERE po_id = ? ORDER BY uploaded_at",
            (pid,)
        )
        rows = [dict(r) for r in cursor.fetchall()]

    attach_dir = _attach_dir(pid)
    for row in rows:
        fp = os.path.join(attach_dir, row["filename"])
        row["size"] = os.path.getsize(fp) if os.path.exists(fp) else 0
    return jsonify(rows)


@app.route("/api/po/<pid>/attachments", methods=["POST"])
def upload_attachment(pid):
    """Upload a file attachment to a PO"""
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["file"]
    if not file.filename:
        return jsonify({"error": "Empty filename"}), 400

    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        return jsonify({"error": f"File type .{ext} not allowed"}), 400

    label = request.form.get("label", file.filename.rsplit(".", 1)[0])
    aid = str(uuid.uuid4())
    safe_name = f"{aid}.{ext}"

    dest = os.path.join(_attach_dir(pid), safe_name)
    file.save(dest)

    if os.path.getsize(dest) > MAX_FILE_SIZE:
        os.remove(dest)
        return jsonify({"error": "File too large (max 20 MB)"}), 400

    mime_type = mimetypes.guess_type(file.filename)[0] or "application/octet-stream"
    uploaded_at = datetime.now().strftime("%Y-%m-%d %H:%M")

    with get_db() as conn:
        _migrate_attach_json(pid, conn)
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO po_attachments "
            "(id, po_id, filename, original, label, mime, uploaded_at) "
            "VALUES (?, ?, ?, ?, ?, ?, ?)",
            (aid, pid, safe_name, file.filename, label, mime_type, uploaded_at)
        )
        cursor.execute("UPDATE purchase_orders SET attach_count = attach_count + 1 WHERE id = ?", (pid,))

    return jsonify({"ok": True, "id": aid, "filename": safe_name})


@app.route("/api/po/<pid>/attachments/<aid>", methods=["GET"])
def download_attachment(pid, aid):
    """Download/view an attachment"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT filename, original, mime FROM po_attachments "
            "WHERE id = ? AND po_id = ?",
            (aid, pid)
        )
        row = cursor.fetchone()

    if not row:
        return jsonify({"error": "Attachment not found"}), 404

    fp = os.path.join(_attach_dir(pid), row["filename"])
    if not os.path.exists(fp):
        return jsonify({"error": "File missing from disk"}), 404

    return send_file(fp, download_name=row["original"] or row["filename"],
                     mimetype=row["mime"] or "application/octet-stream")


@app.route("/api/po/<pid>/attachments/<aid>", methods=["DELETE"])
def delete_attachment(pid, aid):
    """Delete an attachment"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT filename FROM po_attachments WHERE id = ? AND po_id = ?",
            (aid, pid)
        )
        row = cursor.fetchone()
        if not row:
            return jsonify({"error": "Attachment not found"}), 404

        fp = os.path.join(_attach_dir(pid), row["filename"])
        if os.path.exists(fp):
            os.remove(fp)

        cursor.execute("DELETE FROM po_attachments WHERE id = ?", (aid,))
        cursor.execute(
            "UPDATE purchase_orders SET attach_count = CASE WHEN attach_count > 0 THEN attach_count - 1 ELSE 0 END WHERE id = ?",
            (pid,)
        )

    return jsonify({"ok": True})


@app.route("/api/po/<pid>/attachments/<aid>/rename", methods=["POST"])
def rename_attachment(pid, aid):
    """Rename an attachment label"""
    req = request.get_json(silent=True) or {}
    new_label = req.get("label", "").strip()
    if not new_label:
        return jsonify({"error": "Label cannot be empty"}), 400

    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT id FROM po_attachments WHERE id = ? AND po_id = ?",
            (aid, pid)
        )
        if not cursor.fetchone():
            return jsonify({"error": "Attachment not found"}), 404
        cursor.execute(
            "UPDATE po_attachments SET label = ? WHERE id = ?",
            (new_label, aid)
        )

    return jsonify({"ok": True})


# ═══════════════════════════════════════════════════════════════════════════════
# PAYMENT PROOF
# ═══════════════════════════════════════════════════════════════════════════════

def _payment_dir(pid):
    """Get/create payment proof directory for a PO (physical file storage)"""
    d = os.path.join(ATTACH_DIR, pid, "_payment")
    os.makedirs(d, exist_ok=True)
    return d


def _migrate_payment_json(pid, conn):
    """One-time migration: import legacy _payment_meta.json into po_payments."""
    meta_path = os.path.join(_payment_dir(pid), "_payment_meta.json")
    if not os.path.exists(meta_path):
        return
    try:
        with open(meta_path, "r", encoding="utf-8") as f:
            m = json.load(f)
        cursor = conn.cursor()
        cursor.execute(
            "INSERT OR IGNORE INTO po_payments "
            "(po_id, filename, original, uploaded_at, confirmed, confirmed_at) "
            "VALUES (?, ?, ?, ?, ?, ?)",
            (pid, m.get("filename", ""),
             m.get("original", m.get("filename", "")),
             m.get("uploaded_at", datetime.now().strftime("%Y-%m-%d %H:%M")),
             1 if m.get("confirmed") else 0,
             m.get("confirmed_at"))
        )
        os.rename(meta_path, meta_path + ".migrated")
    except Exception as ex:
        print(f"[payment-migrate] {pid}: {ex}")


@app.route("/api/po/<pid>/payment", methods=["GET"])
def get_payment_status(pid):
    """Check payment proof status"""
    with get_db() as conn:
        _migrate_payment_json(pid, conn)
        cursor = conn.cursor()
        cursor.execute(
            "SELECT filename, original, uploaded_at, confirmed, confirmed_at "
            "FROM po_payments WHERE po_id = ?",
            (pid,)
        )
        row = cursor.fetchone()

    if not row:
        return jsonify({})

    meta = dict(row)
    meta["confirmed"] = bool(meta["confirmed"])
    fp = os.path.join(_payment_dir(pid), meta["filename"])
    try:
        meta["size"] = os.path.getsize(fp) if os.path.exists(fp) else 0
    except (OSError, IOError) as e:
        meta["size"] = 0
        meta["error"] = f"Could not determine file size: {str(e)}"
    return jsonify(meta)


@app.route("/api/po/<pid>/payment", methods=["POST"])
def upload_payment_proof(pid):
    """Upload payment proof PDF"""
    with get_db() as conn:
        _migrate_payment_json(pid, conn)
        cursor = conn.cursor()
        cursor.execute(
            "SELECT confirmed FROM po_payments WHERE po_id = ?", (pid,)
        )
        existing = cursor.fetchone()
        if existing and existing["confirmed"]:
            return jsonify({"error": "Payment proof already confirmed and locked"}), 409

        if "file" not in request.files:
            return jsonify({"error": "No file provided"}), 400
        file = request.files["file"]
        if not file.filename:
            return jsonify({"error": "Empty filename"}), 400
        if not file.filename.lower().endswith(".pdf"):
            return jsonify({"error": "Only PDF files allowed for payment proof"}), 400

        cursor.execute("SELECT po_number FROM purchase_orders WHERE id = ?", (pid,))
        po_row = cursor.fetchone()
        po_num = po_row["po_number"] if po_row else pid

        safe_po = re.sub(r'[^\w\-]', '_', po_num)
        dest_name = f"Payment_{safe_po}.pdf"
        dest_path = os.path.join(_payment_dir(pid), dest_name)
        file.save(dest_path)

        if os.path.getsize(dest_path) > MAX_FILE_SIZE:
            os.remove(dest_path)
            return jsonify({"error": "File too large (max 20 MB)"}), 400

        uploaded_at = datetime.now().strftime("%Y-%m-%d %H:%M")
        cursor.execute(
            "INSERT OR REPLACE INTO po_payments "
            "(po_id, filename, original, uploaded_at, confirmed, confirmed_at) "
            "VALUES (?, ?, ?, ?, 0, NULL)",
            (pid, dest_name, file.filename, uploaded_at)
        )

    return jsonify({"ok": True, "filename": dest_name})


@app.route("/api/po/<pid>/payment/confirm", methods=["POST"])
def confirm_payment_proof(pid):
    """Confirm and lock payment proof (permanent)"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT filename, confirmed FROM po_payments WHERE po_id = ?", (pid,)
        )
        row = cursor.fetchone()
        if not row:
            return jsonify({"error": "No payment proof uploaded yet"}), 400
        if row["confirmed"]:
            return jsonify({"error": "Already confirmed"}), 409

        confirmed_at = datetime.now().strftime("%Y-%m-%d %H:%M")
        cursor.execute(
            "UPDATE po_payments SET confirmed = 1, confirmed_at = ? WHERE po_id = ?",
            (confirmed_at, pid)
        )
        cursor.execute(
            "SELECT filename, original, uploaded_at, confirmed, confirmed_at "
            "FROM po_payments WHERE po_id = ?",
            (pid,)
        )
        updated = dict(cursor.fetchone())

    updated["confirmed"] = bool(updated["confirmed"])
    fp = os.path.join(_payment_dir(pid), updated["filename"])
    updated["size"] = os.path.getsize(fp) if os.path.exists(fp) else 0
    return jsonify(updated)


@app.route("/api/po/<pid>/payment/download", methods=["GET"])
def download_payment_proof(pid):
    """Download the payment proof PDF"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT filename FROM po_payments WHERE po_id = ?", (pid,)
        )
        row = cursor.fetchone()

    if not row:
        return jsonify({"error": "No payment proof found"}), 404

    fp = os.path.join(_payment_dir(pid), row["filename"])
    if not os.path.exists(fp):
        return jsonify({"error": "File missing from disk"}), 404

    return send_file(fp, download_name=row["filename"], mimetype="application/pdf")

@app.route("/api/scan-invoice", methods=["POST"])
def scan_invoice():
    """
    AI INVOICE SCANNER
    Takes an image/pdf, calls Gemini, and extracts structured data.
    Matches items against master data.
    """
    api_key = request.form.get("api_key")
    model = "gemini-2.5-flash"
    
    if not api_key:
        return jsonify({"error": "Gemini API Key is required"}), 400
    
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    file = request.files["file"]
    file_data = file.read()
    mime_type = mimetypes.guess_type(file.filename)[0] or "application/octet-stream"
    file_ext = file.filename.lower()

    # 1. Fetch Master Data for matching
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id, name, description, hs_code, unit, currency, default_price_usd FROM items")
        master_items = [dict(row) for row in cursor.fetchall()]
        
        cursor.execute("SELECT id, name, company FROM suppliers")
        master_suppliers = [dict(row) for row in cursor.fetchall()]

    is_document = False
    extracted_text = ""

    if file_ext.endswith(('.xlsx', '.xls')):
        is_document = True
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(file_data), data_only=True)
            for sheet in wb.worksheets:
                extracted_text += f"\nSheet: {sheet.title}\n"
                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell) for cell in row if cell is not None]
                    if row_data:
                        extracted_text += " | ".join(row_data) + "\n"
        except Exception as e:
            return jsonify({"error": f"Failed to parse Excel file: {str(e)}"}), 400
    elif file_ext.endswith(('.docx', '.doc')):
        is_document = True
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_data))
            for para in doc.paragraphs:
                if para.text.strip(): extracted_text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        extracted_text += " | ".join(row_data) + "\n"
        except Exception as e:
            return jsonify({"error": f"Failed to parse Word file: {str(e)}"}), 400
    elif file_ext.endswith('.csv'):
        is_document = True
        try:
            text_data = file_data.decode('utf-8', errors='replace')
            reader = csv.reader(io.StringIO(text_data))
            for row in reader:
                extracted_text += " | ".join([str(cell).strip() for cell in row if str(cell).strip()]) + "\n"
        except Exception as e:
            return jsonify({"error": f"Failed to parse CSV file: {str(e)}"}), 400

    # 2. Build Prompt
    base_prompt = f"""
    MASTER ITEMS LIST:
    {json.dumps(master_items)}
    
    MASTER SUPPLIERS LIST:
    {json.dumps(master_suppliers)}

    INSTRUCTIONS:
    1. Identify the Supplier. If the supplier matches one in the MASTER SUPPLIERS LIST (match by name or company), return its "id".
    2. Extract the Date and Currency.
    3. Extract all Line Items. For each item, you must try extremely hard to find a match in the MASTER ITEMS LIST. 
       MATCHING LOGIC:
       - First, try to match the master item's "name" with the invoice item's "description".
       - If no direct match, look for keywords (model numbers, sizes, colors) across both name and description fields to find the closest match.
       - Maximize the number of matched items with the available master list.
       If matched, return the "item_id" from the master list. 
       Always return the item_name, description, qty, unit, and unit_price as found in the invoice.
    4. Return ONLY the JSON object, no other text or formatting.

    JSON SCHEMA:
    {{
      "supplier_id": "...", 
      "supplier_name": "...",
      "date": "YYYY-MM-DD",
      "currency": "USD/CNY",
      "items": [
        {{
          "item_id": "...", 
          "item_name": "...", 
          "description": "...", 
          "qty": 0, 
          "unit": "PCS", 
          "unit_price": 0
        }}
      ]
    }}
    """

    if is_document:
        prompt = f"""
    You are an expert invoice processing assistant. 
    Analyze the following extracted text from an invoice document and extract the data in a structured JSON format.
    
    EXTRACTED INVOICE TEXT:
    {extracted_text}
    """ + base_prompt
    else:
        prompt = f"""
    You are an expert invoice processing assistant. 
    Analyze the attached invoice image/pdf and extract the data in a structured JSON format.
    """ + base_prompt

    try:
        if is_document:
            res = _call_gemini(api_key, model, prompt, file_data=None, mime_type=None)
        else:
            res = _call_gemini(api_key, model, prompt, file_data, mime_type)
        
        # Robustly extract text from Gemini response
        if 'error' in res:
            return jsonify({"error": res['error'].get('message', 'Unknown API Error')}), 400
            
        if not res.get('candidates') or not res['candidates'][0].get('content'):
            # Check if it was blocked
            finish_reason = res.get('candidates', [{}])[0].get('finishReason', 'UNKNOWN')
            return jsonify({"error": f"Gemini returned no results. Reason: {finish_reason}"}), 400
            
        text = res['candidates'][0]['content']['parts'][0]['text']
        
        # Clean up Markdown formatting if present
        text = re.sub(r'```json\s*|\s*```', '', text).strip()
        
        extracted_data = json.loads(text)
        
        # Validate extracted data structure
        if not isinstance(extracted_data, dict):
            return jsonify({"error": "Invalid JSON structure from AI scanner"}), 400
        if 'items' not in extracted_data:
            extracted_data['items'] = []
        if not isinstance(extracted_data.get('items'), list):
            return jsonify({"error": "Items field must be a list"}), 400
        
        # --- Python Logical Matching Fallback ---
        # A deterministic, token-based matching algorithm to ensure high accuracy
        def normalize_str(s):
            if not s: return set()
            s = str(s).lower()
            # Keep alphanumeric and periods (for decimals like 2.5mm), replace rest with space
            s = re.sub(r'[^a-z0-9.]', ' ', s)
            return set(s.split())

        valid_items = []
        for inv_item in extracted_data.get('items', []):
            # Validate item structure
            if not isinstance(inv_item, dict):
                continue
            inv_desc_tokens = normalize_str(inv_item.get('description', ''))
            inv_name_tokens = normalize_str(inv_item.get('item_name', ''))
            
            best_match = None
            best_score = 0
            
            for m_item in master_items:
                m_tokens = normalize_str(m_item.get('name', '')) | normalize_str(m_item.get('description', ''))
                
                score = 0
                
                # Check description tokens (Highest priority)
                desc_intersection = inv_desc_tokens.intersection(m_tokens)
                for token in desc_intersection:
                    if any(c.isdigit() for c in token):
                        score += 6  # Digits + Description = Highest Priority
                    else:
                        score += 2  # Words in Description
                        
                # Check name tokens (Lower priority)
                name_intersection = inv_name_tokens.intersection(m_tokens)
                for token in name_intersection:
                    if token not in desc_intersection: # Avoid double counting
                        if any(c.isdigit() for c in token):
                            score += 3
                        else:
                            score += 1
                        
                if score > best_score:
                    best_score = score
                    best_match = m_item
            
            # If we found a strong match, update the item_id
            if best_match and best_score >= 3:
                inv_item['item_id'] = best_match['id']
            elif best_match and best_score > 0 and not inv_item.get('item_id'):
                inv_item['item_id'] = best_match['id']
                
            # Ignore items that do not have a valid master match
            if inv_item.get('item_id') and any(m.get('id') == inv_item.get('item_id') for m in master_items):
                valid_items.append(inv_item)

        extracted_data['items'] = valid_items
        # ----------------------------------------
        
        return jsonify(extracted_data)
        
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 400


if __name__ == "__main__":
    app.run(debug=True, port=5005, host="0.0.0.0")
